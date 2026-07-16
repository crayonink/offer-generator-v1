from fastapi import FastAPI, UploadFile, File, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, FileResponse, RedirectResponse, JSONResponse
from pydantic import BaseModel
from typing import List, Optional
import pandas as pd
import sqlite3
import shutil
import os
from datetime import datetime

from bom.pricelist_parser import parse_all as _parse_pricelist_all

# Swagger/ReDoc moved off /docs so the Antora handbook can own /docs.
app = FastAPI(docs_url="/api-docs", redoc_url="/api-redoc")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.middleware("http")
async def _no_cache_html(request, call_next):
    """Stop browsers serving a stale cached HTML page after a redeploy — the
    form pages are inline-HTML+JS, so a cached copy hides new code until a hard
    refresh. Force revalidation on every HTML response (JSON/assets untouched)."""
    response = await call_next(request)
    if response.headers.get("content-type", "").startswith("text/html"):
        response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
    return response

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# ── Persistent database (survives Railway redeploys) ────────────────────────
# Railway's container filesystem is ephemeral: the committed vlph.db ships in
# the image, so every redeploy resets it — wiping any edits made through the
# live UI. To make edits permanent, mount a Railway VOLUME and point
# VLPH_DB_PATH at a file on it (e.g. /data/vlph.db). On first run the committed
# vlph.db is copied to the volume as a SEED; after that the volume is the live
# database and the in-repo vlph.db is symlinked to it, so every module (which
# opens BASE_DIR/vlph.db or 'vlph.db') transparently uses the persistent copy.
# Set VLPH_DB_RESEED=1 for one deploy to force-overwrite the volume from the
# committed seed (e.g. to push a bulk data change). No env var = unchanged
# behaviour (uses the in-repo vlph.db, as before).
#
# STATIC reference tables (edited in code, no in-app editor, NOT derived from
# the price cascade) are refreshed from the committed seed on every deploy, so
# code-side edits to them appear live without a full reseed. component_price_
# master and its cascade-derived tables are deliberately NOT in this list — they
# stay on the volume so live UI price edits persist.
_SEED_REFRESH_TABLES = ["fabrication_ladle_mapping"]

def _refresh_seed_tables(seed_path, vol_path, tables):
    import sqlite3 as _sq
    sc = _sq.connect(seed_path); vc = _sq.connect(vol_path)
    try:
        for tbl in tables:
            ddl = sc.execute("SELECT sql FROM sqlite_master WHERE type='table' AND name=?",
                             (tbl,)).fetchone()
            if not ddl or not ddl[0]:
                continue
            ncols = len(sc.execute(f"PRAGMA table_info({tbl})").fetchall())
            rows = sc.execute(f"SELECT * FROM {tbl}").fetchall()
            vc.execute(f"DROP TABLE IF EXISTS {tbl}")
            vc.execute(ddl[0])
            if rows:
                vc.executemany(f"INSERT INTO {tbl} VALUES ({','.join(['?'] * ncols)})", rows)
        vc.commit()
    finally:
        sc.close(); vc.close()

def _init_persistent_db():
    vol = (os.environ.get("VLPH_DB_PATH") or "").strip()
    seed = os.path.join(BASE_DIR, "vlph.db")
    if not vol:
        return
    try:
        os.makedirs(os.path.dirname(vol) or ".", exist_ok=True)
        reseed = os.environ.get("VLPH_DB_RESEED") == "1"
        seed_is_real = os.path.exists(seed) and not os.path.islink(seed)
        # Seed the volume from the committed DB on first run (or forced reseed).
        if seed_is_real and (not os.path.exists(vol) or reseed):
            shutil.copy2(seed, vol)
        # Refresh the static reference tables from the committed seed (must run
        # while `seed` is still the real committed file, before the symlink).
        if seed_is_real and os.path.exists(vol) and os.path.realpath(seed) != os.path.realpath(vol):
            try:
                _refresh_seed_tables(seed, vol, _SEED_REFRESH_TABLES)
            except Exception as _re:
                print(f"[db] static-table refresh skipped: {_re}")
        # Point the in-repo path at the volume so all DB access is persistent.
        if os.path.realpath(seed) != os.path.realpath(vol):
            if os.path.islink(seed) or os.path.exists(seed):
                os.remove(seed)
            os.symlink(vol, seed)
        print(f"[db] using persistent volume: {vol} (refreshed {_SEED_REFRESH_TABLES})")
    except Exception as e:
        # Never leave the app without a database: if the symlink couldn't be
        # created (e.g. symlinks unsupported), restore a working copy at the
        # in-repo path. Edits then won't persist across redeploys, but the app
        # keeps running rather than failing to open a missing vlph.db.
        if not os.path.exists(seed) and os.path.exists(vol):
            try:
                shutil.copy2(vol, seed)
            except Exception:
                pass
        print(f"[db] persistent volume init failed ({e}); using in-repo vlph.db")

_init_persistent_db()
DB_PATH = os.path.join(BASE_DIR, "vlph.db")  # symlink → volume when VLPH_DB_PATH set

# Documentation site (Antora, pre-built into ./docs_site) served at /docs.
# Guarded so a missing build never breaks app start-up.
from fastapi.staticfiles import StaticFiles
_DOCS_DIR = os.path.join(BASE_DIR, "docs_site")
if os.path.isdir(_DOCS_DIR):
    app.mount("/docs", StaticFiles(directory=_DOCS_DIR, html=True), name="docs")
UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
QUOTES_FOLDER = os.path.join(BASE_DIR, "quotes")

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(QUOTES_FOLDER, exist_ok=True)

# ─────────────────────────────────────────────────────────────────────────────
# Authentication — login gate for the whole portal (two roles: admin / user).
# ─────────────────────────────────────────────────────────────────────────────
from engine.auth import (
    SESSION_COOKIE, SESSION_MAX_AGE,
    verify_credentials, make_token, verify_token,
)

# Reachable without logging in.
_PUBLIC_PATHS = {"/login", "/logout", "/health", "/favicon.ico"}


def _auth_is_admin_only(method: str, path: str) -> bool:
    """Endpoints that edit pricing/data or expose the raw DB are admin-only."""
    if path == "/viewer":                       # raw DB editor page
        return True
    if path == "/api/pricelist-summary":        # vendor catalog prices (burner/blower/etc.)
        return True
    if method in ("PUT", "DELETE"):             # all rate/data edits use PUT/DELETE
        return True
    if method == "POST" and (
        path.startswith("/api/stock/") or path.startswith("/upload-excel")
    ):
        return True
    return False


@app.middleware("http")
async def auth_gate(request: Request, call_next):
    path = request.url.path
    # CORS preflight, public pages and the docs/login assets pass straight through.
    if request.method == "OPTIONS" or path in _PUBLIC_PATHS:
        return await call_next(request)

    payload = verify_token(request.cookies.get(SESSION_COOKIE))
    if not payload:
        accept = request.headers.get("accept", "")
        # Browser page loads → redirect to the login screen; API calls → 401.
        if request.method == "GET" and "text/html" in accept:
            return RedirectResponse(url=f"/login?next={path}", status_code=303)
        return JSONResponse({"detail": "Not authenticated"}, status_code=401)

    role = payload.get("r")
    if _auth_is_admin_only(request.method, path) and role != "admin":
        return JSONResponse(
            {"detail": "Admin access required for this action."}, status_code=403
        )

    request.state.user = payload.get("u")
    request.state.role = role
    return await call_next(request)


@app.get("/health")
def health():
    return {"status": "ok"}


@app.get("/login", response_class=HTMLResponse)
def login_page():
    with open(os.path.join(BASE_DIR, "login.html"), "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())


@app.post("/login")
async def login_submit(request: Request):
    form = await request.form()
    username = (form.get("username") or "").strip()
    password = form.get("password") or ""
    next_url = form.get("next") or "/"
    if not next_url.startswith("/"):
        next_url = "/"
    role = verify_credentials(username, password)
    if not role:
        sep = "&" if next_url != "/" else ""
        nxt = f"&next={next_url}" if next_url != "/" else ""
        return RedirectResponse(url=f"/login?error=1{nxt}", status_code=303)
    resp = RedirectResponse(url=next_url, status_code=303)
    resp.set_cookie(
        SESSION_COOKIE, make_token(username, role),
        max_age=SESSION_MAX_AGE, httponly=True, samesite="lax",
    )
    return resp


@app.get("/logout")
def logout():
    resp = RedirectResponse(url="/login", status_code=303)
    resp.delete_cookie(SESSION_COOKIE)
    return resp


@app.get("/api/me")
def whoami(request: Request):
    """Current session's user + role (for the UI to show/hide admin controls)."""
    return {
        "user": getattr(request.state, "user", None),
        "role": getattr(request.state, "role", None),
    }

VALID_TABLES = None

COUNTER_FILE = os.path.join(BASE_DIR, "quote_counter.txt")

def next_quote_seq():
    if os.path.exists(COUNTER_FILE):
        with open(COUNTER_FILE) as f:
            n = int(f.read().strip()) + 1
    else:
        n = 1
    with open(COUNTER_FILE, "w") as f:
        f.write(str(n))
    return str(n).zfill(3)


def peek_quote_seq() -> str:
    """Return the sequence the next generate-quote call WILL receive,
    without consuming it. Used by the form to preview the auto ref."""
    if os.path.exists(COUNTER_FILE):
        with open(COUNTER_FILE) as f:
            return str(int(f.read().strip()) + 1).zfill(3)
    return "001"


def _with_salutation(salutation: str, name: str) -> str:
    """Return 'Mr. Ashish Gupta' if both supplied; else whichever is present."""
    s = (salutation or "").strip()
    n = (name or "").strip()
    if s and n:
        return f"{s} {n}"
    return n or s


def _greeting(salutation: str) -> str:
    """Cover-letter greeting word: 'Sir' for Mr., 'Ma'am' otherwise."""
    return "Sir" if (salutation or "").strip().rstrip(".").lower() == "mr" else "Ma'am"


def _person_initials(name: str) -> str:
    """First letter of each whitespace-separated word, uppercased.
    'Jyotirmoy Rabha' -> 'JR'. Returns '' for empty input."""
    parts = [p for p in (name or "").strip().split() if p]
    return "".join(p[0].upper() for p in parts)


_LOCATION_CODES = {        # ENCON branch codes used in the enquiry ref
    "goa":       "GOA",
    "vadodara":  "VAD",
    "faridabad": "FBD",
}


def _location_code(location: str) -> str:
    return _LOCATION_CODES.get((location or "").strip().lower(), "")


def build_enquiry_ref(seq: str, technical_person: str,
                      location: str = "",
                      year: Optional[int] = None) -> str:
    """ENCON.04026.{seq}/{LOC}/{initials} DT.{DD/MM/YYYY}.

    The '04026' segment is ENCON's fixed code (applies to ladle and
    tundish offers alike). Location code (GOA/VDD/FBD) and the
    technical-person initials get omitted only when missing.
    """
    from datetime import datetime as _dt
    today = _dt.now()
    ini = _person_initials(technical_person)
    # ENCON operates from Faridabad only — the location code is always FBD,
    # regardless of any location passed in.
    loc = "FBD"
    # Assemble the slash segment: include only the parts we have.
    parts = [seq]
    if loc: parts.append(loc)
    if ini: parts.append(ini)
    body = "/".join(parts)
    date_str = today.strftime("%d/%m/%Y")
    return f"ENCON.04026.{body} DT.{date_str}"

def ensure_log_table():
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""CREATE TABLE IF NOT EXISTS table_update_log (
        table_name TEXT PRIMARY KEY, updated_at TEXT, uploaded_by TEXT)""")
    # Every generated offer is recorded here — the source for the dashboard.
    conn.execute("""CREATE TABLE IF NOT EXISTS quotes_log (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        quote_no TEXT, ref_no TEXT, company_name TEXT, poc_name TEXT,
        email TEXT, mobile_no TEXT, project_name TEXT, equipment_type TEXT,
        location TEXT, ladle_tons REAL, grand_total REAL,
        marketing_person TEXT, technical_person TEXT, file_path TEXT,
        margin_pct REAL, created_at TEXT)""")
    # Self-heal an older quotes_log that predates these columns.
    for col, decl in (("location", "TEXT"), ("margin_pct", "REAL")):
        try:
            conn.execute(f"ALTER TABLE quotes_log ADD COLUMN {col} {decl}")
        except sqlite3.OperationalError:
            pass
    # ── CRM (dashboard Phase 2) ──────────────────────────────────────────
    conn.execute("""CREATE TABLE IF NOT EXISTS enquiries (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        enquiry_no TEXT, company_name TEXT, contact_name TEXT, email TEXT,
        phone TEXT, location TEXT, product TEXT, source TEXT,
        stage TEXT DEFAULT 'new', value_est REAL DEFAULT 0, owner TEXT,
        notes TEXT, created_at TEXT, updated_at TEXT)""")
    conn.execute("""CREATE TABLE IF NOT EXISTS projects (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        project_no TEXT, company_name TEXT, title TEXT, product TEXT,
        value REAL DEFAULT 0, status TEXT DEFAULT 'active',
        progress INTEGER DEFAULT 0, owner TEXT, start_date TEXT,
        target_date TEXT, enquiry_id INTEGER, notes TEXT,
        created_at TEXT, updated_at TEXT)""")
    conn.execute("""CREATE TABLE IF NOT EXISTS activity_log (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        ts TEXT, kind TEXT, title TEXT, detail TEXT)""")
    conn.commit()
    conn.close()

ensure_log_table()


def ensure_oil_burner_master():
    """Self-heal oil_burner_master if it's mis-parsed: an earlier one-off import
    grabbed the left blocks of the ' Oil Burner' sheet, ran past each block's
    TOTAL into the junk below (S.G.-Assembly sub-table, stray H.V. rows), and
    never reached the 7A block which sits in the offset columns AB–AI. Rebuild
    cleanly from the four side-by-side blocks (start cols A/J/S/AB, rows 4 down
    to each block's TOTAL). Idempotent — only runs when the table looks wrong."""
    try:
        conn = sqlite3.connect(DB_PATH)
        try:
            need = conn.execute(
                "SELECT (SELECT COUNT(*) FROM oil_burner_master WHERE burner_type='7A')=0 "
                "OR EXISTS(SELECT 1 FROM oil_burner_master "
                "WHERE particular='Burner Size' OR particular LIKE 'H.V.%')").fetchone()[0]
        except sqlite3.OperationalError:
            conn.close(); return            # table doesn't exist — nothing to heal
        if not need:
            conn.close(); return
        src = os.path.join(BASE_DIR, "uploads", "Pricelist WorkBook 28-08-2025.xlsx")
        if not os.path.exists(src):
            conn.close(); return
        import openpyxl
        ws = openpyxl.load_workbook(src, data_only=True)[" Oil Burner"]

        def cv(r, c):
            v = ws.cell(r, c).value
            return v.strip() if isinstance(v, str) else v

        rows = []
        for bt, sc in (("2A/3A", 1), ("4A", 10), ("5A/6A", 19), ("7A", 28)):
            total_row = None
            for r in range(3, 60):
                if any(isinstance(ws.cell(r, c).value, str) and
                       ws.cell(r, c).value.strip().upper() == "TOTAL"
                       for c in range(sc, sc + 8)):
                    total_row = r
                    break
            for r in range(4, total_row or 60):
                s_no, particular = cv(r, sc), cv(r, sc + 1)
                if not particular and s_no in (None, ""):
                    continue
                rows.append((s_no, particular, cv(r, sc + 2), cv(r, sc + 3),
                             cv(r, sc + 4), cv(r, sc + 5), cv(r, sc + 6), cv(r, sc + 7), bt))
        conn.execute("DELETE FROM oil_burner_master")
        conn.executemany(
            "INSERT INTO oil_burner_master (s_no,particular,qty,unit,rate,amount,"
            "mc_cost,total_amount,burner_type) VALUES (?,?,?,?,?,?,?,?,?)", rows)
        conn.commit()
        conn.close()
        print(f"ensure_oil_burner_master: rebuilt {len(rows)} clean rows (incl. 7A)")
    except Exception as e:
        print(f"WARN: ensure_oil_burner_master failed: {e}")


ensure_oil_burner_master()


def ensure_hv_oil_burner_master():
    """Build hv_oil_burner_master from the 'HV  Oil Burner' sheet — same 4-block
    layout as the oil burner (2A/3A · 4A · 5A/6A · 7A at cols A/J/S/AB; header row
    2, parts from row 3 down to each block's TOTAL). Built once (skips if non-empty)."""
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.execute("CREATE TABLE IF NOT EXISTS hv_oil_burner_master "
                     "(s_no TEXT, particular TEXT, qty TEXT, unit TEXT, rate TEXT, "
                     "amount TEXT, mc_cost TEXT, total_amount TEXT, burner_type TEXT)")
        if conn.execute("SELECT COUNT(*) FROM hv_oil_burner_master").fetchone()[0] > 0:
            conn.close(); return
        src = os.path.join(BASE_DIR, "uploads", "Pricelist WorkBook 28-08-2025.xlsx")
        if not os.path.exists(src):
            conn.close(); return
        import openpyxl
        ws = openpyxl.load_workbook(src, data_only=True)["HV  Oil Burner"]

        def cv(r, c):
            v = ws.cell(r, c).value
            return v.strip() if isinstance(v, str) else v

        rows = []
        for bt, sc in (("2A/3A", 1), ("4A", 10), ("5A/6A", 19), ("7A", 28)):
            total_row = None
            for r in range(3, 40):
                if any(isinstance(ws.cell(r, c).value, str) and
                       ws.cell(r, c).value.strip().upper() == "TOTAL"
                       for c in range(sc, sc + 8)):
                    total_row = r
                    break
            for r in range(3, total_row or 40):
                s_no, particular = cv(r, sc), cv(r, sc + 1)
                if not particular and s_no in (None, ""):
                    continue
                rows.append((s_no, particular, cv(r, sc + 2), cv(r, sc + 3),
                             cv(r, sc + 4), cv(r, sc + 5), cv(r, sc + 6), cv(r, sc + 7), bt))
        conn.executemany(
            "INSERT INTO hv_oil_burner_master (s_no,particular,qty,unit,rate,amount,"
            "mc_cost,total_amount,burner_type) VALUES (?,?,?,?,?,?,?,?,?)", rows)
        conn.commit()
        conn.close()
        print(f"ensure_hv_oil_burner_master: built {len(rows)} rows")
    except Exception as e:
        print(f"WARN: ensure_hv_oil_burner_master failed: {e}")


ensure_hv_oil_burner_master()


# --- Oil-burner rate master (the ' Oil Burner' Rate column pulls from Rates!) ---
# cell -> (display name, unit, category). Mirrors the Rates! sheet items the
# parts reference: the "Bought-out / Casting" K-column, plus the M.S. raw
# materials (the steel the fabricated parts are made from).
_BO = "Bought-out / Casting"
_RM = "Raw Material"
RATE_LABELS = {
    "K6":  ("SS Assembly 2A/3A", "", _BO), "K7": ("SS Assembly 4A", "", _BO),
    "K8":  ("SS Assembly 5A/6A", "", _BO), "K9": ("SS Assembly 7A", "", _BO),
    "K10": ("Micro Valve 2A/3A", "", _BO), "K11": ("Micro Valve 4A", "", _BO),
    "K12": ("Micro Valve 5A/6A", "", _BO), "K13": ("Micro Valve 7A", "", _BO),
    "K14": ("Flexible Hose 15NB x750 (Oil)", "", _BO), "K15": ("Flexible Hose 15NB x1000 (Oil)", "", _BO),
    "K16": ("Flexible Hose 20NB x1000", "", _BO), "K17": ("Flexible Hose 25NB x1000 (Air)", "", _BO),
    "K21": ('Butterfly Valve 2.5"', "", _BO), "K22": ('Butterfly Valve 4"', "", _BO),
    "K23": ('Butterfly Valve 6"', "", _BO), "K24": ("Y-Strainer 20NB", "", _BO),
    "K25": ("Whyteheat K (Burner Block)", "Per Kg", _BO),
    "YSTR25": ("Y-Strainer 25NB", "", _BO),
    "ADP_OIL": ("Adopter 15x20 NB (Oil)", "", _BO),
    "ADP_AIR": ("Adopter 15x15 NB (Air)", "", _BO),
    "K5":  ("Casting Burner Parts (MS)", "Per Kg", _RM),
    "RM_PLATE": ("M.S. Plate / Round", "Per Kg", _RM),
    "RM_FAB":   ("M.S. Sheet (fab. tube/pipe)", "Per Kg", _RM),
    "SPACER7A": ("Burner Spacer", "Per Kg", _RM),
    "TIKKY7A":  ("M.S. Tikky", "Per Kg", _RM),
    "MIXER7A":  ("Air Mixer", "Per Kg", _RM),
    "FLANGE7A": ("M.S. Flange", "Per Kg", _RM),
    "MSPIPE7A": ("M.S. Pipe", "Per Kg", _RM),
    "C22":      ('M.S. Tube "B" Class 1.5in', "Per Kg", _RM),
}
# rate source per part, by position within each group (None = fixed/independent
# rate, e.g. bought-out adopters). 7A's MS-fab parts link to the M.S. materials.
PART_RATE_REFS = {
    "2A/3A": ["K5", "K5", "K5", "K5", "K6", "K21", "K24", "K14", "K14", "K10", "ADP_OIL", "ADP_AIR", "K5", "K25", None, None],
    "4A":    ["K5", "K5", "K5", "K5", "K7", "K22", "K24", "K14", "K14", "K11", "ADP_OIL", "ADP_AIR", "K5", "K25", None, None],
    "5A/6A": ["K5", "K5", "K5", "K5", "K8", "K22", "K24", "K15", "K16", "K12", "ADP_OIL", "ADP_AIR", "K5", "K25", None, None],
    "7A":    ["K5", "FLANGE7A", "MSPIPE7A", "MIXER7A", "TIKKY7A", "K5", "TIKKY7A", "TIKKY7A", "SPACER7A",
              "K9", "K13", "K16", "K17", "K23", "YSTR25", "K25", "K5", "RM_FAB", None, None],
}
# HV oil-burner part -> rate cell, by position within each block. None = lump/HV-
# specific (H.V. Burner, H.V. Burner block, labour, paint) — no rate link.
PART_RATE_REFS_HV = {
    "2A/3A": [None, None, "K21", "K24", "K14", "K14", "K10", "ADP_OIL", "ADP_AIR", None, None],
    "4A":    [None, None, "K22", "K24", "K14", "K14", "K11", "ADP_OIL", "ADP_AIR", None, None],
    "5A/6A": [None, None, "K22", "K24", "K15", "K16", "K14", "K12", "ADP_OIL", "ADP_AIR", None, None],
    "7A":    ["K5", "FLANGE7A", "MSPIPE7A", "MIXER7A", "TIKKY7A", "K5", "TIKKY7A", "TIKKY7A", "SPACER7A",
              "K9", "K13", "K16", "K17", "K23", "YSTR25", None, "RM_FAB", None, None],
}
# rate cell -> component_price_master item: these rates are SOURCED LIVE from the
# Pricelist (single source of truth). The rest (butterfly K21-23, hoses K14-17,
# M.S. plate/sheet RM_*) stay burner-local until the Pricelist carries them 1:1.
RATE_CPM = {
    "K5":  "CASTING BURNER PARTS",
    "K6":  "SS ASSLY 2A/3A", "K7": "SS ASSLY 4A", "K8": "SS ASSLY 5A/6A", "K9": "SS ASSLY 7A",
    "K10": "MICRO VALVE 2A/3A", "K11": "MICRO VALVE 4A", "K12": "MICRO VALVE 5A/6A", "K13": "MICRO VALVE 7A",
    "K24": "ENCON-Y-STRAINER 20 NB", "YSTR25": "ENCON-Y-STRAINER 25 NB", "K25": "WHYTEHEAT K",
    "SPACER7A": "Burner Spacer", "TIKKY7A": "M.S. Tikky", "MIXER7A": "Air Mixer",
    "FLANGE7A": "M.S. Flange", "MSPIPE7A": "M.S. Pipe",
    "ADP_OIL": "ADOPTER 15 NB*20 NB (Oil)", "ADP_AIR": "ADOPTER 15 NB*15 NB (Air)",
    "C22": "M.S. Tube B Class 1.5 in",
    # butterfly: inch × 25.4 → nearest NB (2.5"→65, 4"→100, 6"→150)
    "K21": "BUTTERFLY VALVE 65 NB", "K22": "BUTTERFLY VALVE 100 NB", "K23": "BUTTERFLY VALVE 150 NB",
    # flexible hoses — by NB size to the matching Pricelist hose
    "K14": "FLEXIBLE HOSE 15 NB 1500mm", "K15": "FLEXIBLE HOSE 15 NB 1500mm",
    "K16": "FLEXIBLE HOSE-20NB*1000MM (OIL)", "K17": "FLEXIBLE HOSE-25NB*1000MM (AIR)",
}


def ensure_rate_master():
    """Link the oil_burner_master Rate column to a rate master, so editing a rate
    (e.g. MS/kg) reprices every part that uses it. Adds oil_burner_master.rate_ref
    (assigned by position within each size), and seeds the rate_master table from
    the parts' current rates. Idempotent; preserves user-edited master values."""
    try:
        conn = sqlite3.connect(DB_PATH)
        cols = [r[1] for r in conn.execute("PRAGMA table_info(oil_burner_master)").fetchall()]
        if "rate_ref" not in cols:
            conn.execute("ALTER TABLE oil_burner_master ADD COLUMN rate_ref TEXT")
        # assign rate_ref by row position within each group (PART_RATE_REFS is
        # authoritative, so re-mappings — e.g. 7A Y-strainer K24->YSTR25 — apply)
        for group, refs in PART_RATE_REFS.items():
            rids = [r[0] for r in conn.execute(
                "SELECT rowid FROM oil_burner_master WHERE burner_type=? ORDER BY rowid", (group,)).fetchall()]
            for i, rid in enumerate(rids):
                ref = refs[i] if i < len(refs) else None
                conn.execute("UPDATE oil_burner_master SET rate_ref=? WHERE rowid=?", (ref, rid))
        conn.execute("CREATE TABLE IF NOT EXISTS rate_master "
                     "(cell TEXT PRIMARY KEY, name TEXT, value REAL, unit TEXT, sort INTEGER)")
        rmcols = [r[1] for r in conn.execute("PRAGMA table_info(rate_master)").fetchall()]
        if "category" not in rmcols:
            conn.execute("ALTER TABLE rate_master ADD COLUMN category TEXT")
        for i, cell in enumerate(RATE_LABELS):
            name, unit, cat = RATE_LABELS[cell]
            if conn.execute("SELECT 1 FROM rate_master WHERE cell=?", (cell,)).fetchone():
                conn.execute("UPDATE rate_master SET name=?, unit=?, category=?, sort=? WHERE cell=?",
                             (name, unit, cat, i, cell))   # refresh labels, keep edited value
                continue
            r = conn.execute("SELECT rate FROM oil_burner_master WHERE rate_ref=? "
                             "AND rate IS NOT NULL LIMIT 1", (cell,)).fetchone()
            try:
                val = float(r[0]) if r and r[0] not in (None, "") else None
            except (TypeError, ValueError):
                val = None
            conn.execute("INSERT INTO rate_master (cell, name, value, unit, sort, category) "
                         "VALUES (?,?,?,?,?,?)", (cell, name, val, unit, i, cat))
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"WARN: ensure_rate_master failed: {e}")


ensure_rate_master()


def ensure_hv_rate_refs():
    """Link the HV oil-burner parts to the same rate cells as the oil burner
    (add hv_oil_burner_master.rate_ref, assigned by position). Authoritative."""
    try:
        conn = sqlite3.connect(DB_PATH)
        cols = [r[1] for r in conn.execute("PRAGMA table_info(hv_oil_burner_master)").fetchall()]
        if not cols:
            conn.close(); return
        if "rate_ref" not in cols:
            conn.execute("ALTER TABLE hv_oil_burner_master ADD COLUMN rate_ref TEXT")
        for group, refs in PART_RATE_REFS_HV.items():
            rids = [r[0] for r in conn.execute(
                "SELECT rowid FROM hv_oil_burner_master WHERE burner_type=? ORDER BY rowid", (group,)).fetchall()]
            for i, rid in enumerate(rids):
                ref = refs[i] if i < len(refs) else None
                conn.execute("UPDATE hv_oil_burner_master SET rate_ref=? WHERE rowid=?", (ref, rid))
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"WARN: ensure_hv_rate_refs failed: {e}")


ensure_hv_rate_refs()


def cleanup_ciplate_pricelist():
    """C.I. Plate was briefly (and wrongly) seeded into the Pricelist as a Raw
    Material; it is a computed burner price (burner-plate × 1.8), not a rate.
    Remove those rows on startup so the live (volume) DB is cleaned — deleting
    the seeding code alone can't remove rows already written to the volume.
    Idempotent no-op once gone."""
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.execute("DELETE FROM component_price_master WHERE item LIKE 'C.I. PLATE %'")
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"WARN: cleanup_ciplate_pricelist failed: {e}")


cleanup_ciplate_pricelist()


def cleanup_agr_fixed_ratio():
    """Standardise AGR on 1:1-to-1:10 (both threaded & flanged): drop the fixed
    1:1 rows from the Pricelist (component_price_master) and the BOM source
    (agr_master). Runs on startup so the live (volume) DB is cleaned. The gas
    BOM selector now requests ratio='1:1 to 1:10'."""
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.execute("DELETE FROM component_price_master "
                     "WHERE category='Air Gas Regulator' AND specification='1:1'")
        try:
            conn.execute("DELETE FROM agr_master WHERE ratio='1:1'")
        except sqlite3.OperationalError:
            pass   # agr_master may not exist in some DBs
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"WARN: cleanup_agr_fixed_ratio failed: {e}")


cleanup_agr_fixed_ratio()


# Pricelist (Bought Out) items the burner needs that the base Pricelist lacks —
# seeded if missing, user edits preserved.
#  (item, category, price, specification)  — spec = A-frame size for per-size rows
BURNER_PRICELIST_SEED = [
    ("ENCON-Y-STRAINER 25 NB", "Bought Out", 420, ""),   # 7A uses a 25 NB Y-strainer
    ("Burner Spacer", "Raw Material", 72, ""),           # 7A spacer
    ("M.S. Tikky", "Raw Material", 72, ""),              # 7A tikkies (×3)
    ("Air Mixer", "Raw Material", 80, ""),               # 7A air mixer
    ("M.S. Flange", "Raw Material", 72, ""),             # 7A flange
    ("M.S. Pipe", "Raw Material", 135, ""),              # 7A M.S. pipe (100NB)
    # Spare for ENCON Film Burner — S.G. Assembly + Air Resistor, per size
    # (shown under its own heading inside the Raw Material tab; size in Size col)
    ("S.G. Assembly 2A", "Spare for ENCON Film Burner", 7500, "2A"),
    ("S.G. Assembly 3A", "Spare for ENCON Film Burner", 7750, "3A"),
    ("S.G. Assembly 4A", "Spare for ENCON Film Burner", 10218, "4A"),
    ("S.G. Assembly 5A", "Spare for ENCON Film Burner", 17000, "5A"),
    ("S.G. Assembly 6A", "Spare for ENCON Film Burner", 17200, "6A"),
    ("S.G. Assembly 7A", "Spare for ENCON Film Burner", 52000, "7A"),
    ("Air Resistor 2A", "Spare for ENCON Film Burner", 2900, "2A"),
    ("Air Resistor 3A", "Spare for ENCON Film Burner", 2900, "3A"),
    ("Air Resistor 4A", "Spare for ENCON Film Burner", 3600, "4A"),
    ("Air Resistor 5A", "Spare for ENCON Film Burner", 6600, "5A"),
    ("Air Resistor 6A", "Spare for ENCON Film Burner", 6600, "6A"),
    ("Air Resistor 7A", "Spare for ENCON Film Burner", 8100, "7A"),
    # Burner Block for High Velocity Burners (per size) — placeholder ×1.2.
    ("HV Burner Block 3A", "Burner Block for High Velocity Burners", 52800, "3A"),
    ("HV Burner Block 4A", "Burner Block for High Velocity Burners", 70800, "4A"),
    ("HV Burner Block 5A", "Burner Block for High Velocity Burners", 85800, "5A"),
    ("HV Burner Block 6A", "Burner Block for High Velocity Burners", 193200, "6A"),
    ("HV Burner Block 7A", "Burner Block for High Velocity Burners", 241560, "7A"),
]
# one-time correction: HV burner blocks were seeded at the JUL'22 placeholder;
# bump any still at the old value to the final ×1.2 (item -> old value).
_HV_BLOCK_OLD = {"HV Burner Block 3A": 44000, "HV Burner Block 4A": 59000,
                 "HV Burner Block 5A": 71500, "HV Burner Block 6A": 161000,
                 "HV Burner Block 7A": 201300}


def ensure_burner_pricelist_seed():
    try:
        conn = sqlite3.connect(DB_PATH)
        for item, cat, price, spec in BURNER_PRICELIST_SEED:
            if not conn.execute("SELECT 1 FROM component_price_master WHERE item=?", (item,)).fetchone():
                conn.execute("INSERT INTO component_price_master (item, category, price, previous_price, "
                             "specification) VALUES (?,?,?,?,?)", (item, cat, price, price, spec or None))
            else:
                if spec:   # backfill the size on rows seeded before spec was added
                    conn.execute("UPDATE component_price_master SET specification=? WHERE item=? "
                                 "AND (specification IS NULL OR specification='')", (spec, item))
                if item in _HV_BLOCK_OLD:   # one-time placeholder -> ×1.2 correction
                    conn.execute("UPDATE component_price_master SET price=?, previous_price=? "
                                 "WHERE item=? AND price=?", (price, price, item, _HV_BLOCK_OLD[item]))
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"WARN: ensure_burner_pricelist_seed failed: {e}")


ensure_burner_pricelist_seed()


def ensure_casting_category():
    """Casting Burner Parts is a per-kg MS rate — it belongs in Raw Material, not
    Bought Out. Move it on startup so the live (volume) Pricelist reflects it."""
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.execute("UPDATE component_price_master SET category='Raw Material' "
                     "WHERE item='CASTING BURNER PARTS' AND category!='Raw Material'")
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"WARN: ensure_casting_category failed: {e}")


ensure_casting_category()


def _log_quote(*, quote_no="", ref_no="", company_name="", poc_name="",
               email="", mobile_no="", project_name="", equipment_type="",
               location="", ladle_tons=0.0, grand_total=0.0,
               marketing_person="", technical_person="", file_path="",
               margin_pct=None):
    """Record a generated offer in quotes_log (dashboard source). Never raises
    — a logging failure must never break offer generation."""
    try:
        from datetime import datetime as _dt
        conn = sqlite3.connect(DB_PATH)
        conn.execute("""
            INSERT INTO quotes_log (
                quote_no, ref_no, company_name, poc_name, email, mobile_no,
                project_name, equipment_type, location, ladle_tons, grand_total,
                marketing_person, technical_person, file_path, margin_pct, created_at
            ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
        """, (
            quote_no or "", ref_no or "", company_name or "", poc_name or "",
            email or "", mobile_no or "", project_name or "", equipment_type or "",
            location or "", float(ladle_tons or 0), float(grand_total or 0),
            marketing_person or "", technical_person or "", file_path or "",
            (float(margin_pct) if margin_pct is not None else None),
            _dt.now().isoformat(timespec="seconds"),
        ))
        conn.commit()
        conn.close()
        _log_activity("quote", f"Quote · {company_name or 'Client'}",
                      f"{equipment_type or 'Offer'} — ₹{float(grand_total or 0):,.0f}")
    except Exception as log_err:
        print(f"WARN: quotes_log insert failed: {log_err}")


def _log_equipment_quote(cust, equipment_type: str, grand_total, file_path: str):
    """quotes_log entry for an equipment offer (HPU/PU/Blower/Burner) whose
    customer is an HpuCustomer. Wraps the field mapping in one place."""
    _log_quote(
        quote_no=getattr(cust, "ref_no", "") or "",
        ref_no=getattr(cust, "ref_no", ""),
        company_name=getattr(cust, "company", ""),
        poc_name=_with_salutation(getattr(cust, "salutation", ""), getattr(cust, "name", "")),
        email=getattr(cust, "email", ""), mobile_no=getattr(cust, "phone", ""),
        project_name=getattr(cust, "subject", ""), equipment_type=equipment_type,
        location=getattr(cust, "location", ""), grand_total=grand_total,
        marketing_person=_with_salutation(getattr(cust, "marketing_salutation", ""), getattr(cust, "marketing", "")),
        technical_person=_with_salutation(getattr(cust, "technical_salutation", ""), getattr(cust, "technical", "")),
        file_path=file_path,
    )


def _log_activity(kind: str, title: str, detail: str = ""):
    """Append to activity_log (dashboard feed). Never raises."""
    try:
        from datetime import datetime as _dt
        conn = sqlite3.connect(DB_PATH)
        conn.execute("INSERT INTO activity_log (ts, kind, title, detail) VALUES (?,?,?,?)",
                     (_dt.now().isoformat(timespec="seconds"), kind, title, detail))
        conn.commit()
        conn.close()
    except Exception as act_err:
        print(f"WARN: activity_log insert failed: {act_err}")


def _compute_notifications(conn):
    """Derived alerts for the dashboard — overdue projects, stale enquiries,
    recent wins. `conn` is an open connection with row_factory=sqlite3.Row."""
    from datetime import datetime as _dt, timedelta as _td
    now = _dt.now()
    out = []
    try:
        for r in conn.execute("SELECT project_no, company_name, target_date FROM projects "
                              "WHERE status='active' AND IFNULL(target_date,'')<>''"):
            try:
                if _dt.fromisoformat(r["target_date"]) < now:
                    out.append({"kind": "warn", "title": f"Project overdue · {r['company_name']}",
                                "detail": f"{r['project_no']} · target {r['target_date']}"})
            except Exception:
                pass
        cutoff = (now - _td(days=14)).isoformat()
        for r in conn.execute("SELECT company_name FROM enquiries WHERE stage IN ('new','qualified') "
                              "AND IFNULL(updated_at,'')<>'' AND updated_at<? ORDER BY updated_at LIMIT 8", (cutoff,)):
            out.append({"kind": "info", "title": f"Follow up · {r['company_name']}",
                        "detail": "No movement in 14+ days"})
        recent = (now - _td(days=7)).isoformat()
        for r in conn.execute("SELECT company_name FROM enquiries WHERE stage='won' "
                              "AND updated_at>? ORDER BY updated_at DESC LIMIT 6", (recent,)):
            out.append({"kind": "good", "title": f"Won · {r['company_name']}",
                        "detail": "Ready to convert into a project"})
    except Exception as e:
        print(f"WARN: notifications failed: {e}")
    return out[:12]


def ensure_valve_sizes():
    """Ensure rotary joint table has entries up to 600 NB.
    (butterfly_valve_master was dropped — selector now reads from
    component_price_master and lt_butterfly_valve_master.)"""
    conn = sqlite3.connect(DB_PATH)
    rj_defaults = [(400, 65000), (450, 75000), (500, 90000), (600, 110000)]
    try:
        for nb, price in rj_defaults:
            conn.execute("INSERT OR IGNORE INTO rotary_joint_master (nb, price) VALUES (?,?)", (nb, price))
        conn.commit()
    except sqlite3.OperationalError:
        pass  # rotary_joint_master may not exist yet on a fresh DB
    conn.close()

ensure_valve_sizes()


def ensure_extra_columns():
    """Add updated_at and company columns to component_price_master if not present."""
    conn = sqlite3.connect(DB_PATH)
    cols = [r[1] for r in conn.execute("PRAGMA table_info(component_price_master)").fetchall()]
    if "updated_at" not in cols:
        conn.execute("ALTER TABLE component_price_master ADD COLUMN updated_at TEXT")
    if "company" not in cols:
        conn.execute("ALTER TABLE component_price_master ADD COLUMN company TEXT")
    conn.commit()
    conn.close()

ensure_extra_columns()

# Duplicate name pairs in component_price_master — canonical -> [aliases]
_RATE_DUPLICATES = [
    ("FLEXIBLE HOSE-15NB*1000MM (OIL)", ["FLEXIBLE HOSE-15NB*1000MM (OIL )"]),
    ("FLEXIBLE HOSE-15NB*750MM (OIL)",  ["FLEXIBLE HOSE-15NB*750MM (OIL )"]),
    ("FLEXIBLE HOSE-20NB*1000MM (OIL)", ["FLEXIBLE HOSE-20NB*1000MM (OIL )"]),
    ("FLEXIBLE HOSE-25NB*1000MM (AIR)", ["FLEXIBLE HOSE-25NB*1000MM (AIR )"]),
    ("M.S. Sheet 2mm",       ["M.S. Sheet  2mm"]),
    ("M.S. Sheet 3mm",       ["M.S. Sheet  3mm"]),
    ("M.S. Plate 16mm*5mm",  ["M.S. Plate 16mm* 5mm"]),
    ("M.S. Tube B Class 1.5 in", ['M.S. Tube "B" Class 1.5 in']),
    ("M.S. Tube C Class 1.5 in", ['M.S. Tube "C" Class 1.5 in']),
    ("M.S. Channel",         ["M.S. Chanel", "M.S.Chanel"]),
    ("Plumber block with Bearing", ["Plumber Block with Bearing"]),
    ("Pulley with V belt",   ["Pulley with V Belt"]),
    ("SS Pipe 304 60x3mm (per mtr)",  ["SS Pipe 304 60x3mm",  "SS Pipe 304 60 X 3mm"]),
    ("SS Pipe 304 76x3mm (per mtr)",  ["SS Pipe 304 76x3mm",  "SS Pipe 304 76 X 3mm"]),
    ("SS Pipe 304 100x3mm (per mtr)", ["SS Pipe 304 100x3mm", "SS Pipe 304 100 X 3mm"]),
    ("ID FAN (ARE 35)",      ["ID FAN  (ARE 35)"]),
    ("SEQUENCE CONTROLLER",  ["SEQUENCE"]),
]

def clean_duplicate_rates(conn):
    """Remove known duplicate/alias rows from component_price_master."""
    for canonical, aliases in _RATE_DUPLICATES:
        for alias in aliases:
            exists = conn.execute(
                "SELECT 1 FROM component_price_master WHERE item=?", (alias,)
            ).fetchone()
            if not exists:
                continue
            canonical_exists = conn.execute(
                "SELECT 1 FROM component_price_master WHERE item=?", (canonical,)
            ).fetchone()
            if canonical_exists:
                conn.execute("DELETE FROM component_price_master WHERE item=?", (alias,))
            else:
                conn.execute(
                    "UPDATE component_price_master SET item=? WHERE item=?", (canonical, alias)
                )
    conn.commit()

import glob
import tempfile

ALLOWED_EDIT_TABLES = {
    'hpu_master', 'oil_burner_parts_master', 'hv_oil_burner_parts_master',
    'gas_burner_parts_master', 'horizontal_master', 'vertical_master',
    'recuperator_master', 'blower_pricelist_master',
    'rad_heat_master', 'rad_heat_tata_master', 'gail_gas_burner_master',
    'rotary_joint_master',
}

def _find_latest_pricebook():
    """Find the most recently uploaded full pricebook Excel file."""
    candidates = []
    search_dirs = [UPLOAD_FOLDER, BASE_DIR]
    for d in search_dirs:
        for f in glob.glob(os.path.join(d, "*.xlsx")):
            # Skip stock files and other non-pricebook files
            bn = os.path.basename(f).lower()
            if "stock" in bn or "costing" in bn or "regen" in bn or "sample" in bn:
                continue
            try:
                xl = pd.ExcelFile(f)
                if len(xl.sheet_names) >= 8:
                    candidates.append((os.path.getmtime(f), f))
            except Exception:
                pass
    return sorted(candidates, reverse=True)[0][1] if candidates else None

def _find_regen_file():
    """Find the Regen Standard Costing workbook in BASE_DIR. Skips Excel
    lock/temp stubs (~$…) so a stray open-file lock can't be picked and crash
    the parser."""
    for f in glob.glob(os.path.join(BASE_DIR, "*.xlsx")):
        bn = os.path.basename(f).lower()
        if bn.startswith("~$"):
            continue
        if "regen" in bn and "costing" in bn:
            return f
    return None

def _ensure_regen_costing():
    """Parse the regen costing file if found and table not yet populated."""
    try:
        regen_file = _find_regen_file()
        if not regen_file:
            return
        conn = sqlite3.connect(DB_PATH)
        try:
            count = conn.execute("SELECT COUNT(*) FROM regen_costing_items").fetchone()[0]
            if count > 0:
                conn.close()
                return
        except Exception:
            pass
        from bom.regen_parser import parse_regen_costing
        parse_regen_costing(regen_file, conn)
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"Regen costing parse error: {e}")

_ensure_regen_costing()

def _cascade_recalculate(xl_path: str, conn):
    """Patch the Rates sheet with current DB rates, re-run all parsers."""
    import openpyxl
    # Get current rates with cell positions
    rate_cells = {}
    for row in conn.execute(
        "SELECT price, excel_row, excel_col FROM component_price_master WHERE excel_row IS NOT NULL AND excel_col IS NOT NULL"
    ):
        price, er, ec = row
        if er and ec:
            rate_cells[(int(er), int(ec))] = price

    wb = openpyxl.load_workbook(xl_path, data_only=False)
    rates_sn = next((s for s in wb.sheetnames if s.strip().lower() == 'rates'), None)
    if rates_sn and rate_cells:
        ws = wb[rates_sn]
        for (er, ec), price in rate_cells.items():
            ws.cell(er, ec).value = price

    tmp_fd, tmp_path = tempfile.mkstemp(suffix='.xlsx')
    os.close(tmp_fd)
    wb.save(tmp_path)
    wb.close()

    results = _parse_pricelist_all(tmp_path, conn)
    try:
        os.unlink(tmp_path)
    except Exception:
        pass
    return results

def ensure_rate_columns():
    conn = sqlite3.connect(DB_PATH)
    for col in ['excel_row', 'excel_col']:
        try:
            conn.execute(f"ALTER TABLE component_price_master ADD COLUMN {col} INTEGER")
            conn.commit()
        except Exception:
            pass
    conn.close()

ensure_rate_columns()

@app.get("/", response_class=HTMLResponse)
def root():
    html_path = os.path.join(BASE_DIR, "dashboard.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())


@app.get("/internal-costing", response_class=HTMLResponse)
def internal_costing_page():
    with open(os.path.join(BASE_DIR, "internal_costing.html"), "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())


@app.get("/api/internal-costing/oil-burner-prices")
def api_ic_oil_burner_prices():
    """Oil-burner price tables from burner_pricelist_master: the Film Burner &
    Accessories table and the Spares table, pivoted (sizes x components)."""
    LABELS = {"BURNER ALONE": "Burner Alone", "MICRO VALVE": "Micro Valve",
              "C.I.BURNER PLATE": "C.I. Plate", "HIGH AL. WHYTEHEAT K BURNER BLOCK": "Burner Block",
              "FLEXIBLE HOSES SET": "Flex Hoses", "Y TYPE STRAINER": "Y-Strainer",
              "BUTTERFLY VALVE": "Butterfly", "BURNER SET": "Burner Set",
              "BALL VALVE": "Ball Valve",
              "S.G. ASSEMBLY": "S.G. Assembly", "AIR RESISTOR": "Air Resistor"}

    import re as _re

    def _kw_map(conn):
        """Firing-capacity range (kW) per burner size token (2A/3A/…), from
        burner_selection_master. kcal/hr ÷ 860 = kW; envelope across both
        pressures. The largest model has an open-ended max (sentinel) → 'NNN+'."""
        m = {}
        for model, mn, mx in conn.execute(
                "SELECT model, MIN(min_firing_kcal_hr), MAX(max_firing_kcal_hr) "
                "FROM burner_selection_master GROUP BY model"):
            t = _re.search(r'(\d+\s*A)\b', str(model), _re.I)
            if not t:
                continue
            try:
                lo, hi = round(float(mn) / 860), round(float(mx) / 860)
            except (TypeError, ValueError):
                continue
            m[t.group(1).replace(' ', '').upper()] = (f"{lo}+" if hi > 20000 else f"{lo}–{hi}")
        return m

    def _kw_for(size, km):
        t = _re.search(r'(\d+\s*A)\b', str(size), _re.I)
        return km.get(t.group(1).replace(' ', '').upper()) if t else None

    def table(section, km):
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        rows = conn.execute("SELECT burner_size, component, price FROM burner_pricelist_master "
                            "WHERE section=? ORDER BY rowid", (section,)).fetchall()
        conn.close()
        comps, data = [], {}
        for r in rows:
            if r["component"] not in comps:
                comps.append(r["component"])
            data.setdefault(r["burner_size"], {})[r["component"]] = r["price"]
        return {"section": section,
                "columns": [{"label": LABELS.get(c, c.title()), "comp": c} for c in comps],
                "rows": [{"size": sz, "kw": _kw_for(sz, km), "cells": [vals.get(c) for c in comps]}
                         for sz, vals in data.items()]}
    try:
        _c0 = sqlite3.connect(DB_PATH)
        km = _kw_map(_c0)
        _c0.close()
        return {"film":   table("PRICE FOR VARIOUS SIZES OF ENCON 'FILM' BURNER & ACCESSORIES", km),
                "dual":   table("PRICE FOR VARIOUS SIZES OF ENCON DUAL FUEL BURNER & ACCESSORIES", km),
                "gas":    table("PRICE FOR VARIOUS SIZES OF ENCON 'GAS' BURNER & ACCESSORIES", km),
                "hv_oil": table("PRICE LIST FOR HIGH VELOCITY OIL BURNERS", km),
                "hv_gas": table("PRICE LIST FOR HIGH VELOCITY GAS BURNERS", km),
                "spares": table("PRICE LIST FOR SPARES OF IIP ENCON OIL FILM BURNERS", km)}
    except Exception as e:
        return {"film": {"columns": [], "rows": []}, "dual": {"columns": [], "rows": []},
                "gas": {"columns": [], "rows": []}, "hv_oil": {"columns": [], "rows": []},
                "hv_gas": {"columns": [], "rows": []},
                "spares": {"columns": [], "rows": []}, "error": str(e)}


# Cost -> price markups baked into the burner tables, now editable (Markup Master).
# key -> (label, default)
BURNER_MARKUPS = [
    ("ba",             "Burner Alone ×",             2.5),
    ("micro",          "Micro Valve ×",              2.0),
    ("ciplate",        "C.I. Plate ×",               1.8),
    ("block",          "Burner Block ×",             2.0),
    ("block_5a6a",     "Burner Block × (5A/6A)",     2.2),
    ("hoses",          "Flexible Hoses ×",           2.0),
    ("ystrainer",      "Y-Strainer ×",               2.5),
    ("butterfly",      "Butterfly ×",                2.0),
    ("ballvalve",      "Ball Valve ×",               2.0),
    ("ballvalve_disc", "Ball Valve discount ×",      0.78),
    ("dual_hoses",     "Dual Fuel Hoses ×",          1.5),
    ("dual_hoses_off", "Dual Fuel Hoses offset (+)", -5.0),
    ("hv_hoses",       "HV Oil Hoses ×",             2.0),
]


def ensure_burner_markups():
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.execute("CREATE TABLE IF NOT EXISTS burner_markup "
                     "(key TEXT PRIMARY KEY, label TEXT, value REAL, sort INTEGER)")
        for i, (k, label, default) in enumerate(BURNER_MARKUPS):
            if conn.execute("SELECT 1 FROM burner_markup WHERE key=?", (k,)).fetchone():
                conn.execute("UPDATE burner_markup SET label=?, sort=? WHERE key=?", (label, i, k))
            else:
                conn.execute("INSERT INTO burner_markup (key, label, value, sort) VALUES (?,?,?,?)",
                             (k, label, default, i))
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"WARN: ensure_burner_markups failed: {e}")


ensure_burner_markups()


def _load_markups(conn):
    """Current markup factors as a dict (falls back to defaults)."""
    m = {k: d for k, _, d in BURNER_MARKUPS}
    try:
        for k, v in conn.execute("SELECT key, value FROM burner_markup").fetchall():
            if v is not None:
                m[k] = v
    except sqlite3.OperationalError:
        pass
    return m


# HV burner block price cascades from the Whyteheat K refractory rate:
#   block price = Whyteheat ₹/kg × block weight (kg)
# so a 10% rise in the Whyteheat rate lifts every HV block 10%. The weights
# are the kg implied by the JUL'25 block prices at the ₹60/kg Whyteheat rate.
HV_BLOCK_WHYTE_ITEM = "WHYTEHEAT K"
HV_BLOCK_WHYTE_KG = {
    "HV Burner Block 3A": 880,
    "HV Burner Block 4A": 1180,
    "HV Burner Block 5A": 1430,
    "HV Burner Block 6A": 3220,
    "HV Burner Block 7A": 4026,
}


def recompute_hv_blocks(conn):
    """HV burner block Pricelist prices = Whyteheat K ₹/kg × per-size weight."""
    try:
        row = conn.execute(
            "SELECT price FROM component_price_master WHERE item=?",
            (HV_BLOCK_WHYTE_ITEM,)).fetchone()
        if not row or row[0] is None:
            return
        rate = float(row[0])
        for item, kg in HV_BLOCK_WHYTE_KG.items():
            conn.execute(
                "UPDATE component_price_master SET price=? WHERE item=?",
                (round(rate * kg), item))
    except Exception as e:
        print(f"WARN: recompute_hv_blocks failed: {e}")


def recompute_burner_prices(conn):
    """Recompute the derived BURNER prices from the oil_burner_master part totals,
    exactly as the workbook formulas do (part totals x markup). Independent cells
    — Y-Strainer, Air Resistor, and the 3A hand-set Burner Alone & C.I. Plate —
    are read but never overwritten. Mirrors the ' Oil Burner'!->BURNER! formulas.

      Burner Alone = Σ(first part totals) × 2.5      Butterfly   = part × 2
      C.I. Plate   = burner-plate total × 1.8        Micro Valve = part × 2
      Burner Block = block total × 2 (5A/6A ×2.2)    Flex Hoses  = Σ(2 hoses) × 2
      Burner Set   = Σ(the 7 film components)        S.G. Assy   = S.S.Assy × n
    """
    FILM = "PRICE FOR VARIOUS SIZES OF ENCON 'FILM' BURNER & ACCESSORIES"
    SPARE = "PRICE LIST FOR SPARES OF IIP ENCON OIL FILM BURNERS"

    # Refresh the HV block Pricelist prices from the live Whyteheat rate first,
    # so the HV oil/gas sections below read the cascaded values.
    recompute_hv_blocks(conn)

    def f(x):
        try:
            return float(x)
        except (TypeError, ValueError):
            return 0.0

    def part_totals(group):
        rows = conn.execute("SELECT amount, mc_cost, total_amount FROM oil_burner_master "
                            "WHERE burner_type=? ORDER BY rowid", (group,)).fetchall()
        return [(f(t) if t not in (None, "") else f(a) + f(m)) for a, m, t in rows]

    def stored(section, size, comp):
        r = conn.execute("SELECT price FROM burner_pricelist_master "
                         "WHERE section=? AND burner_size=? AND component=?",
                         (section, size, comp)).fetchone()
        return f(r[0]) if r else 0.0

    def put(section, size, comp, val):
        conn.execute("UPDATE burner_pricelist_master SET price=? "
                     "WHERE section=? AND burner_size=? AND component=?",
                     (round(val, 2), section, size, comp))

    # part index (1-based, within the group's ordered rows) for each role.
    # Layout A = the 16-row sizes (2A/3A, 4A, 5A/6A); B = the 20-row 7A.
    LAYOUTS = {
        "A": {"ba": [1, 2, 3, 4, 5], "butterfly": 6, "ystrainer": 7, "hoses": [8, 9],
              "micro": 10, "ciplate": 13, "block": 14, "ssassy": 5},
        "B": {"ba": [1, 2, 3, 4, 5, 6, 7, 8, 9, 10], "micro": 11, "hoses": [12, 13],
              "butterfly": 14, "ystrainer": 15, "block": 16, "ciplate": 17, "ssassy": 10},
    }
    # (price size, parts group, layout, block markup, S.G.-Assembly rule)
    SIZES = [("ENCON 2A", "2A/3A", "A", 2.0, ("mult", 3.0)),
             ("ENCON 3A", "2A/3A", "A", 2.0, ("plus", "ENCON 2A", 250)),
             ("ENCON 4A", "4A",    "A", 2.0, ("mult", 3.93)),
             ("ENCON 5A", "5A/6A", "A", 2.2, ("mult", 3.4)),
             ("ENCON 6A", "5A/6A", "A", 2.2, ("plus", "ENCON 5A", 200)),
             ("ENCON 7A", "7A",    "B", 2.0, ("mult", 2.0))]
    # 2A and 3A share the 2A/3A part group, so every component (Burner Alone,
    # C.I. Plate, …) computes identically — no 3A hand-set overrides remain.
    INDEP = set()
    FILMCOMPS = ["BURNER ALONE", "MICRO VALVE", "C.I.BURNER PLATE",
                 "HIGH AL. WHYTEHEAT K BURNER BLOCK", "FLEXIBLE HOSES SET",
                 "Y TYPE STRAINER", "BUTTERFLY VALVE"]
    MK = _load_markups(conn)
    sg_seen = {}
    for size, group, layout, bmult, sgcfg in SIZES:
        L = LAYOUTS[layout]
        T = part_totals(group)

        def g(i, _T=T):
            return _T[i - 1] if 0 < i <= len(_T) else 0.0

        c = {}
        c["BURNER ALONE"] = (stored(FILM, size, "BURNER ALONE") if (size, "BURNER ALONE") in INDEP
                             else sum(g(i) for i in L["ba"]) * MK["ba"])
        c["MICRO VALVE"] = g(L["micro"]) * MK["micro"]
        c["C.I.BURNER PLATE"] = (stored(FILM, size, "C.I.BURNER PLATE") if (size, "C.I.BURNER PLATE") in INDEP
                                 else g(L["ciplate"]) * MK["ciplate"])
        c["HIGH AL. WHYTEHEAT K BURNER BLOCK"] = g(L["block"]) * (MK["block_5a6a"] if group == "5A/6A" else MK["block"])
        c["FLEXIBLE HOSES SET"] = sum(g(i) for i in L["hoses"]) * MK["hoses"]
        c["Y TYPE STRAINER"] = g(L["ystrainer"]) * MK["ystrainer"]
        c["BUTTERFLY VALVE"] = g(L["butterfly"]) * MK["butterfly"]
        for k in ("BURNER ALONE", "MICRO VALVE", "C.I.BURNER PLATE", "HIGH AL. WHYTEHEAT K BURNER BLOCK",
                  "FLEXIBLE HOSES SET", "Y TYPE STRAINER", "BUTTERFLY VALVE"):
            if (size, k) not in INDEP:
                put(FILM, size, k, c[k])
        put(FILM, size, "BURNER SET", sum(c[k] for k in FILMCOMPS))

        # S.G. Assembly + Air Resistor are sourced from the Pricelist ("Spare for
        # ENCON Film Burner") so editing those rows cascades to the Spares table
        # AND Dual Fuel (which adds S.G. Assembly). Fall back to the computed
        # S.S.Assy × markup only if the Pricelist row is missing.
        _short = size.replace("ENCON ", "")
        _sgr = conn.execute("SELECT price FROM component_price_master WHERE item=?",
                            ("S.G. Assembly " + _short,)).fetchone()
        if _sgr and _sgr[0] is not None:
            sgval = f(_sgr[0])
        else:
            ssassy = g(L["ssassy"])
            sgval = ssassy * sgcfg[1] if sgcfg[0] == "mult" else sg_seen[sgcfg[1]] + sgcfg[2]
        sg_seen[size] = sgval
        put(SPARE, size.replace("ENCON ", "ENCON-"), "S.G. ASSEMBLY", sgval)
        _arr = conn.execute("SELECT price FROM component_price_master WHERE item=?",
                            ("Air Resistor " + _short,)).fetchone()
        if _arr and _arr[0] is not None:
            put(SPARE, size.replace("ENCON ", "ENCON-"), "AIR RESISTOR", f(_arr[0]))

    recompute_dualfuel_prices(conn)
    recompute_gas_prices(conn)
    recompute_hv_oil_prices(conn)
    recompute_hv_gas_prices(conn)


# ball-valve item per size (2A/3A→20NB, 4A→25NB, 5A→32NB, 6A/7A→40NB) — shared by
# the Dual Fuel and Gas builders; price = 2 × Pricelist NB ball-valve × 0.78.
BURNER_BALL_ITEM = {"ENCON 2A": "BALL VALVE 20 NB #01 L3RBTC/L3RSWC",
                    "ENCON 3A": "BALL VALVE 20 NB #01 L3RBTC/L3RSWC",
                    "ENCON 4A": "BALL VALVE 25 NB #01 L3RBTC/L3RSWC",
                    "ENCON 5A": "BALL VALVE 32 NB #01 L3RBTC/L3RSWC",
                    "ENCON 6A": "BALL VALVE 40 NB #01 L3RBTC/L3RSWC",
                    "ENCON 7A": "BALL VALVE 40 NB #01 L3RBTC/L3RSWC"}


def recompute_gas_prices(conn):
    """ENCON Gas burner = Film components + a Ball Valve (no micro valve / Y-
    strainer), per the workbook. Derived from the live Film section.
      Burner Alone / C.I. Plate / High Al / Flex Hoses / Butterfly = Film (same
      size) ;  Ball Valve = 2 × Pricelist NB ball-valve × 0.78 ;  Burner Set = Σ"""
    FILM = "PRICE FOR VARIOUS SIZES OF ENCON 'FILM' BURNER & ACCESSORIES"
    GAS = "PRICE FOR VARIOUS SIZES OF ENCON 'GAS' BURNER & ACCESSORIES"
    BLOCK = "HIGH AL. WHYTEHEAT K BURNER BLOCK"

    def f(x):
        try:
            return float(x)
        except (TypeError, ValueError):
            return 0.0

    def film(size, comp):
        r = conn.execute("SELECT price FROM burner_pricelist_master WHERE section=? "
                         "AND burner_size=? AND component=?", (FILM, size, comp)).fetchone()
        return f(r[0]) if r else 0.0

    def cpm(item):
        r = conn.execute("SELECT price FROM component_price_master WHERE item=?", (item,)).fetchone()
        return f(r[0]) if r else 0.0

    def put(size, comp, val):
        conn.execute("UPDATE burner_pricelist_master SET price=? WHERE section=? "
                     "AND burner_size=? AND component=?", (round(val, 2), GAS, size, comp))

    MK = _load_markups(conn)
    COMPS = ["BURNER ALONE", "BALL VALVE", "C.I.BURNER PLATE", BLOCK,
             "FLEXIBLE HOSES SET", "BUTTERFLY VALVE"]
    for size in ("ENCON 2A", "ENCON 3A", "ENCON 4A", "ENCON 5A", "ENCON 6A", "ENCON 7A"):
        c = {
            "BURNER ALONE": film(size, "BURNER ALONE"),
            "BALL VALVE": MK["ballvalve"] * cpm(BURNER_BALL_ITEM[size]) * MK["ballvalve_disc"],
            "C.I.BURNER PLATE": film(size, "C.I.BURNER PLATE"),
            BLOCK: film(size, BLOCK),
            "FLEXIBLE HOSES SET": film(size, "FLEXIBLE HOSES SET"),
            "BUTTERFLY VALVE": film(size, "BUTTERFLY VALVE"),
        }
        for comp in COMPS:
            put(size, comp, c[comp])
        put(size, "BURNER SET", sum(c[comp] for comp in COMPS))


def recompute_dualfuel_prices(conn):
    """ENCON Dual Fuel burner = Film burner + S.G. Assembly + a Ball Valve, per
    the workbook. Derived from the live Film section so it tracks rate edits.
      Burner Alone = Film Burner Alone + S.G. Assembly
      Micro/C.I.Plate/Y-type/Butterfly = Film ;  High Al = Film block (5A/6A 11850)
      Flex Hoses = Film hoses × 1.5 (−5 for 2A/3A/4A, off the 2A hose)
      Ball Valve = 2 × Pricelist NB ball-valve × 0.78 ;  Burner Set = Σ
    """
    FILM = "PRICE FOR VARIOUS SIZES OF ENCON 'FILM' BURNER & ACCESSORIES"
    SPARE = "PRICE LIST FOR SPARES OF IIP ENCON OIL FILM BURNERS"
    DUAL = "PRICE FOR VARIOUS SIZES OF ENCON DUAL FUEL BURNER & ACCESSORIES"
    BLOCK = "HIGH AL. WHYTEHEAT K BURNER BLOCK"

    def f(x):
        try:
            return float(x)
        except (TypeError, ValueError):
            return 0.0

    def film(size, comp):
        r = conn.execute("SELECT price FROM burner_pricelist_master WHERE section=? "
                         "AND burner_size=? AND component=?", (FILM, size, comp)).fetchone()
        return f(r[0]) if r else 0.0

    def sg(size):
        r = conn.execute("SELECT price FROM burner_pricelist_master WHERE section=? "
                         "AND burner_size=? AND component='S.G. ASSEMBLY'",
                         (SPARE, size.replace("ENCON ", "ENCON-"))).fetchone()
        return f(r[0]) if r else 0.0

    def cpm(item):
        r = conn.execute("SELECT price FROM component_price_master WHERE item=?", (item,)).fetchone()
        return f(r[0]) if r else 0.0

    def put(size, comp, val):
        conn.execute("UPDATE burner_pricelist_master SET price=? WHERE section=? "
                     "AND burner_size=? AND component=?", (round(val, 2), DUAL, size, comp))

    HIGH_AL_FIXED = {"ENCON 5A": 11850.0, "ENCON 6A": 11850.0}
    HOSE_SRC = {"ENCON 2A": ("ENCON 2A", -5), "ENCON 3A": ("ENCON 2A", -5), "ENCON 4A": ("ENCON 2A", -5),
                "ENCON 5A": ("ENCON 5A", 0), "ENCON 6A": ("ENCON 5A", 0), "ENCON 7A": ("ENCON 7A", 0)}
    COMPS = ["BURNER ALONE", "MICRO VALVE", "C.I.BURNER PLATE", BLOCK, "FLEXIBLE HOSES SET",
             "BALL VALVE", "Y TYPE STRAINER", "BUTTERFLY VALVE"]

    MK = _load_markups(conn)
    for size in ("ENCON 2A", "ENCON 3A", "ENCON 4A", "ENCON 5A", "ENCON 6A", "ENCON 7A"):
        hs_size, hs_off = HOSE_SRC[size]
        c = {
            "BURNER ALONE": film(size, "BURNER ALONE") + sg(size),
            "MICRO VALVE": film(size, "MICRO VALVE"),
            "C.I.BURNER PLATE": film(size, "C.I.BURNER PLATE"),
            BLOCK: HIGH_AL_FIXED.get(size, film(size, BLOCK)),
            "FLEXIBLE HOSES SET": film(hs_size, "FLEXIBLE HOSES SET") * MK["dual_hoses"] + (MK["dual_hoses_off"] if hs_off else 0),
            "BALL VALVE": MK["ballvalve"] * cpm(BURNER_BALL_ITEM[size]) * MK["ballvalve_disc"],
            "Y TYPE STRAINER": film(size, "Y TYPE STRAINER"),
            "BUTTERFLY VALVE": film(size, "BUTTERFLY VALVE"),
        }
        for comp in COMPS:
            put(size, comp, c[comp])
        put(size, "BURNER SET", sum(c[comp] for comp in COMPS))


def recompute_hv_oil_prices(conn):
    """HV Oil burner (HV-3A..7A): Burner Alone/Y-Strainer/Butterfly = Film (same
    size), Micro Valve = Film (offset −1 size), Burner Block = Pricelist HV block,
    Flex Hoses = Σ(HV-oil hose parts in the block) × 2. Burner Set = Σ."""
    FILM = "PRICE FOR VARIOUS SIZES OF ENCON 'FILM' BURNER & ACCESSORIES"
    HVOIL = "PRICE LIST FOR HIGH VELOCITY OIL BURNERS"

    def f(x):
        try:
            return float(x)
        except (TypeError, ValueError):
            return 0.0

    def film(size, comp):
        r = conn.execute("SELECT price FROM burner_pricelist_master WHERE section=? "
                         "AND burner_size=? AND component=?", (FILM, size, comp)).fetchone()
        return f(r[0]) if r else 0.0

    def cpm(item):
        r = conn.execute("SELECT price FROM component_price_master WHERE item=?", (item,)).fetchone()
        return f(r[0]) if r else 0.0

    MK = _load_markups(conn)

    def hose2x(block):
        rows = conn.execute("SELECT amount, mc_cost, total_amount FROM hv_oil_burner_master "
                            "WHERE burner_type=? AND particular LIKE '%FLEXIBLE HOSE%'", (block,)).fetchall()
        return sum((f(t) if t not in (None, "") else f(a) + f(m)) for a, m, t in rows) * MK["hv_hoses"]

    def put(size, comp, val):
        conn.execute("UPDATE burner_pricelist_master SET price=? WHERE section=? "
                     "AND burner_size=? AND component=?", (round(val, 2), HVOIL, size, comp))

    # hv size, Film size (BA/Ystr/Butterfly), Film size for Micro (offset), HV
    # parts block for hoses, Pricelist HV-block item
    HV_OIL = [("ENCON HV-3A", "ENCON 3A", "ENCON 2A", "2A/3A", "HV Burner Block 3A"),
              ("ENCON HV-4A", "ENCON 4A", "ENCON 3A", "4A",    "HV Burner Block 4A"),
              ("ENCON HV-5A", "ENCON 5A", "ENCON 4A", "4A",    "HV Burner Block 5A"),
              ("ENCON HV-6A", "ENCON 6A", "ENCON 5A", "5A/6A", "HV Burner Block 6A"),
              ("ENCON HV-7A", "ENCON 7A", "ENCON 7A", "7A",    "HV Burner Block 7A")]
    for hv, fs, ms, hb, bi in HV_OIL:
        c = {"BURNER ALONE": film(fs, "BURNER ALONE"),
             "BURNER BLOCK": cpm(bi),
             "MICRO VALVE": film(ms, "MICRO VALVE"),
             "FLEXIBLE HOSES SET": hose2x(hb),
             "Y TYPE STRAINER": film(fs, "Y TYPE STRAINER"),
             "BUTTERFLY VALVE": film(fs, "BUTTERFLY VALVE")}
        for comp, val in c.items():
            put(hv, comp, val)
        put(hv, "BURNER SET", sum(c.values()))


def recompute_hv_gas_prices(conn):
    """HV Gas burner (HV-3A..7A): Burner Alone/Hoses/Butterfly = Film (same size),
    Ball Valve = Gas (same size), Burner Block = Pricelist HV block. Burner Set = Σ."""
    FILM = "PRICE FOR VARIOUS SIZES OF ENCON 'FILM' BURNER & ACCESSORIES"
    GAS = "PRICE FOR VARIOUS SIZES OF ENCON 'GAS' BURNER & ACCESSORIES"
    HVGAS = "PRICE LIST FOR HIGH VELOCITY GAS BURNERS"

    def f(x):
        try:
            return float(x)
        except (TypeError, ValueError):
            return 0.0

    def price(section, size, comp):
        r = conn.execute("SELECT price FROM burner_pricelist_master WHERE section=? "
                         "AND burner_size=? AND component=?", (section, size, comp)).fetchone()
        return f(r[0]) if r else 0.0

    def cpm(item):
        r = conn.execute("SELECT price FROM component_price_master WHERE item=?", (item,)).fetchone()
        return f(r[0]) if r else 0.0

    def put(size, comp, val):
        conn.execute("UPDATE burner_pricelist_master SET price=? WHERE section=? "
                     "AND burner_size=? AND component=?", (round(val, 2), HVGAS, size, comp))

    HV_GAS = [("ENCON HV-3A", "ENCON 3A", "HV Burner Block 3A"),
              ("ENCON HV-4A", "ENCON 4A", "HV Burner Block 4A"),
              ("ENCON HV-5A", "ENCON 5A", "HV Burner Block 5A"),
              ("ENCON HV-6A", "ENCON 6A", "HV Burner Block 6A"),
              ("ENCON HV-7A", "ENCON 7A", "HV Burner Block 7A")]
    for hv, fs, bi in HV_GAS:
        c = {"BURNER ALONE": price(FILM, fs, "BURNER ALONE"),
             "BURNER BLOCK": cpm(bi),
             "BALL VALVE": price(GAS, fs, "BALL VALVE"),
             "FLEXIBLE HOSES SET": price(FILM, fs, "FLEXIBLE HOSES SET"),
             "BUTTERFLY VALVE": price(FILM, fs, "BUTTERFLY VALVE")}
        for comp, val in c.items():
            put(hv, comp, val)
        put(hv, "BURNER SET", sum(c.values()))


def sync_cpm_rates(conn):
    """Pull the live Pricelist (component_price_master) into the burner for the
    cleanly-mapped rates (RATE_CPM): refresh the rate-master value, repush into
    every linked part (rate, Amount, Total), then recompute the burner prices.
    Makes the Pricelist the single source of truth for those rates."""
    def _f(x):
        try:
            return float(x)
        except (TypeError, ValueError):
            return 0.0
    for cell, item in RATE_CPM.items():
        r = conn.execute("SELECT price FROM component_price_master WHERE item=?", (item,)).fetchone()
        if not r or r[0] is None:
            continue
        price = _f(r[0])
        conn.execute("UPDATE rate_master SET value=? WHERE cell=?", (price, cell))
        _push_rate_to_parts(conn, cell, price, _f)
    recompute_burner_prices(conn)


def _push_rate_to_parts(conn, cell, price, _f):
    """Set rate/Amount/Total on every oil AND HV burner part linked to `cell`."""
    for table in ("oil_burner_master", "hv_oil_burner_master"):
        try:
            parts = conn.execute(f"SELECT rowid, qty, mc_cost FROM {table} WHERE rate_ref=?",
                                 (cell,)).fetchall()
        except sqlite3.OperationalError:
            continue   # table (or rate_ref column) not present
        for rid, qty, mc in parts:
            amount = _f(qty) * price
            conn.execute(f"UPDATE {table} SET rate=?, amount=?, total_amount=? WHERE rowid=?",
                         (price, amount, amount + _f(mc), rid))


def _startup_cpm_sync():
    try:
        conn = sqlite3.connect(DB_PATH)
        sync_cpm_rates(conn)
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"WARN: startup sync_cpm_rates failed: {e}")


_startup_cpm_sync()


def _startup_seed_hpu_catalog():
    """Insert the HPU raw-material / component catalogue into component_price_
    master if missing. Idempotent and non-destructive (existing rows and live
    Price-Master edits are preserved), so it can run on every deploy — this is
    how the catalogue reaches the persistent Railway volume, which is never
    seed-refreshed."""
    try:
        from bom.hpu_pricelist import (seed_hpu_catalog, consolidate_raw_materials,
                                        apply_hpu_renames, reclassify_bought_sources)
        conn = sqlite3.connect(DB_PATH)
        n = seed_hpu_catalog(conn)
        # One-time: de-duplicate HPU raw materials against the generic rows and
        # fix the "M.S. Chanel" -> "M.S. Channel" spelling (higher rate wins).
        merged = consolidate_raw_materials(conn)
        apply_hpu_renames(conn)          # e.g. "TEMP. GAUGE 0 -150 *C" -> "TEMPERATURE GAUGE"
        reclassify_bought_sources(conn)  # HPU pressure gauge -> small HGURU (Instrumentation)
        conn.close()
        if n:
            print(f"[db] seeded {n} HPU catalogue rows into component_price_master")
        if merged:
            print("[db] consolidated HPU raw materials onto generic rows")
    except Exception as e:
        print(f"WARN: startup seed_hpu_catalog failed: {e}")


_startup_seed_hpu_catalog()


def _startup_seed_blower_alone():
    """Create the 'Blower Alone' pricelist rows (PERKINS) so the blower-alone
    price is editable in the Rates tab and fetched by the Blower tab/offer.
    Idempotent; reaches the persistent volume on deploy."""
    try:
        from bom.blower_pricelist import seed_blower_alone, seed_blower_motor
        conn = sqlite3.connect(DB_PATH)
        n = seed_blower_alone(conn)
        m = seed_blower_motor(conn)
        conn.close()
        if n:
            print(f"[db] seeded {n} 'Blower Alone' pricelist rows")
        if m:
            print(f"[db] seeded {m} 'Blower Motor' pricelist rows")
    except Exception as e:
        print(f"WARN: startup seed_blower_alone failed: {e}")


_startup_seed_blower_alone()


def _startup_purge_regen_pricelist():
    """Regen prices live in code (bom/regen_builder), NOT the Pricelist. Remove
    any REGEN_* rows from component_price_master on startup so they never show
    up in Price-Master — this also clears them from a persistent volume that may
    still carry rows seeded by an earlier build."""
    try:
        conn = sqlite3.connect(DB_PATH)
        n = conn.execute("DELETE FROM component_price_master "
                         "WHERE category LIKE 'REGEN %'").rowcount
        conn.commit()
        conn.close()
        if n:
            print(f"[db] removed {n} REGEN pricelist rows from component_price_master")
    except Exception as e:
        print(f"WARN: purge regen pricelist failed: {e}")


_startup_purge_regen_pricelist()


def _startup_discount_madas_solenoid_valves():
    """MADAS 'Solenoid Valve - Automatic Reset' pricelist rows carry a 45%
    discount off list (price = list × 0.55, previous_price = list). Apply it to
    every row idempotently — rows already ~45%-discounted are skipped, so manual
    Price-Master edits and the persistent volume aren't disturbed."""
    try:
        conn = sqlite3.connect(DB_PATH)
        rows = conn.execute(
            "SELECT rowid, price, previous_price FROM component_price_master "
            "WHERE category='Solenoid Valve - Automatic Reset'").fetchall()
        n = 0
        for rid, price, prev in rows:
            if not price:
                continue
            list_price = prev if prev else price      # pre-discount list price
            already = prev and abs(price - list_price * 0.55) < 1.0
            if already:
                continue
            conn.execute(
                "UPDATE component_price_master SET previous_price=?, price=? WHERE rowid=?",
                (list_price, round(list_price * 0.55, 2), rid))
            n += 1
        conn.commit()
        conn.close()
        if n:
            print(f"[db] applied 45% discount to {n} MADAS solenoid valve rows")
    except Exception as e:
        print(f"WARN: solenoid valve discount migration failed: {e}")


_startup_discount_madas_solenoid_valves()


def _startup_ensure_thermocouple_small():
    """Add the 'Thermocouple Small' pricelist row (₹5,000) if missing — a small
    bought-out TC, alongside the full THERMOCOUPLE (₹36,000). Idempotent; reaches
    the persistent volume. (The regen 'Thermocouple with TT' price maps here.)"""
    try:
        conn = sqlite3.connect(DB_PATH)
        exists = conn.execute("SELECT 1 FROM component_price_master WHERE item=? LIMIT 1",
                              ("Thermocouple Small",)).fetchone()
        if not exists:
            conn.execute(
                "INSERT INTO component_price_master (item, category, unit, price, "
                "previous_price, company) VALUES (?,?,?,?,?,?)",
                ("Thermocouple Small", "Bought Out", "nos", 5000, 5000, "TEMPSENS"))
            conn.commit()
            print("[db] added 'Thermocouple Small' (₹5,000) to component_price_master")
        conn.close()
    except Exception as e:
        print(f"WARN: ensure Thermocouple Small failed: {e}")


_startup_ensure_thermocouple_small()


def _startup_ensure_regen_pricelist_extras():
    """Add the Pricelist rows regen needs that weren't there: DPT flow meters by
    size (general, 'Flow Meter (DPT)' — a DPT + orifice + flanges assembly) and
    the PLC-with-HMI tiers by pair-count. Idempotent; reaches the volume."""
    FLOW = {32:48000, 40:49000, 50:49700, 65:50000, 80:51000, 100:52000,
            125:54000, 150:54000, 200:57000, 250:58000, 300:60000, 350:64000, 400:70500}
    PLC = [("1-2 Pair", 300000), ("3 Pair", 600000), ("4 Pair", 750000),
           ("5 Pair", 800000), ("6 Pair", 900000)]
    # Regen control panel, per burner KW.
    PANEL = {500:300000, 1000:300000, 1500:450000, 2000:450000,
             2500:600000, 3000:600000, 4500:600000, 6000:700000}
    # Burner + regenerator, per KW (computed material x rates, now a Pricelist row).
    BURNER = {500:124429.13, 1000:162998.63, 1500:196797.43, 2000:346253.53,
              2500:356349.12, 3000:474790.79, 4500:663539.31, 6000:868568.06}
    # Pneumatic (flue-gas) damper, by NB — distinct from the manual damper.
    DAMPER = {200:80000, 250:125000, 300:148000, 350:177000, 400:350000, 450:350000,
              500:350000, 550:350000, 600:350000, 650:350000, 700:350000, 900:350000}
    # Oil line (Regen_Oil_Testing.xlsx) — (item, category, price, company).
    OIL = [
        ("Solenoid Valve (Oil Line) 25 NB",   "Oil Line", 14000, "MADAS"),
        ("Gate Valve (Oil Line) 25 NB",       "Oil Line",  5000, "L&T"),
        ("Flexible Hose Pipe (Oil Line) 25 NB","Oil Line",  1750, "BENGAL INDUSTRIES"),
        ("Globe Type Oil Control Valve 25 NB", "Oil Line", 80000, "DEMBLA"),
        ("Oil Flow Meter",                     "Oil Line", 90000, "HONEYWELL"),
        ("TT in Oil Line",                     "Oil Line",  5000, "HONEYWELL"),
        ("PT in Oil Line",                     "Oil Line", 12000, "HONEYWELL"),
        # Oil auxiliaries (Regen_Oil_Testing.xlsx). HPU is computed by the HPU
        # calculator, not a row; these two are single Pricelist lines.
        ("Paperless Recorder",                 "Oil Line",160000, "EUROTHERM"),
        ("ID Fan 15 HP",                       "Oil Line",200000, "ENCON"),
        # Oil BURNER line (NB20) — Jefferson solenoids, L&T/INTERVALVE ball
        # valves + flameless-mode variants, hose, pressure gauge.
        ("Solenoid Valve (Oil Line) 20 NB",           "Oil Line", 11813, "JEFFERSON"),
        ("Solenoid Valve Flameless (Oil Line) 20 NB", "Oil Line", 11813, "JEFFERSON"),
        ("Ball Valve (Oil Line) 20 NB",               "Oil Line",  1900, "L&T / INTERVALVE"),
        ("Ball Valve Flameless (Oil Line) 20 NB",     "Oil Line",  1900, "L&T / INTERVALVE"),
        ("Flexible Hose Pipe (Oil Line) 20 NB",       "Oil Line",  1750, "BENGAL INDUSTRIES"),
        ("Pressure Gauge 0-500 (Oil Line)",           "Oil Line",  4000, "H GURU / BAUMER"),
        ("Oil Control Valve (DN125)",                 "Oil Line", 111000, "DEMBLA"),
        # Packaged gas train for the NG/LPG pilot burner (oil regen).
        ("Packaged Gas Train for NG/LPG Pilot Burner","Gas Train", 46333, "ENCON"),
    ]
    # HLPH trolley geared motor mechanism (10-30T = 1 HP, >30T = 3 HP). The HP
    # is in the item name so the pricelist Size column shows it.
    GEARED = [
        ("GEARED MOTOR MECHANISM 1 HP", "HLPH Trolley", 100000, "POWERTEK"),
        ("GEARED MOTOR MECHANISM 3 HP", "HLPH Trolley", 210000, "POWERTEK"),
    ]
    # Gas burner-line flexible hoses regen needs but the base Pricelist lacks
    # (65/80/350 NB). BENGAL INDUSTRIES, category 'Flexible Hose', so valve_price
    # resolves them by size instead of snapping down to the 50 NB row.
    FLEXHOSE = [
        ("Flexible Hose 65 NB",  "Flexible Hose",  4200, "BENGAL INDUSTRIES"),
        ("Flexible Hose 80 NB",  "Flexible Hose",  6900, "BENGAL INDUSTRIES"),
        ("Flexible Hose 350 NB", "Flexible Hose", 90000, "BENGAL INDUSTRIES"),
    ]
    try:
        conn = sqlite3.connect(DB_PATH)
        def ins(item, cat, price, company):
            if conn.execute("SELECT 1 FROM component_price_master WHERE item=? AND category=? LIMIT 1",
                            (item, cat)).fetchone():
                return 0
            conn.execute("INSERT INTO component_price_master (item, category, unit, price, "
                         "previous_price, company) VALUES (?,?,?,?,?,?)",
                         (item, cat, "nos", price, price, company))
            return 1
        n = 0
        for nb, price in FLOW.items():
            n += ins(f"DPT Flow Meter {nb} NB", "Flow Meter (DPT)", price, "HONEYWELL")
        # (item, company) is UNIQUE, so the pair size goes in the item name too.
        for label, price in PLC:
            n += ins(f"PLC with HMI ({label})", "PLC with HMI", price, "SIEMENS")
        for kw, price in PANEL.items():
            n += ins(f"Control Panel {kw} KW", "Control Panel", price, "ENCON")
        for kw, price in BURNER.items():
            n += ins(f"Burner with Regenerator {kw} KW", "Burner with Regenerator", price, "ENCON")
        for nb, price in DAMPER.items():
            n += ins(f"Pneumatic Damper {nb} NB", "Pneumatic Damper", price, "ENCON")
        for item, cat, price, company in OIL:
            n += ins(item, cat, price, company)
        # rename legacy 'GEARED MOTOR MECHANISM' (no HP) -> '... 3 HP' so the
        # pricelist Size column shows the HP (idempotent).
        if not conn.execute("SELECT 1 FROM component_price_master WHERE "
                            "item='GEARED MOTOR MECHANISM 3 HP'").fetchone():
            conn.execute("UPDATE component_price_master SET item='GEARED MOTOR MECHANISM 3 HP' "
                         "WHERE item='GEARED MOTOR MECHANISM'")
        for item, cat, price, company in GEARED:
            n += ins(item, cat, price, company)
        for item, cat, price, company in FLEXHOSE:
            n += ins(item, cat, price, company)
        # Normalise 'BENGAL' → 'BENGAL INDUSTRIES' so sized-valve lookups (which
        # filter by company) resolve the flexible-hose rows (idempotent).
        conn.execute("UPDATE component_price_master SET company='BENGAL INDUSTRIES' "
                     "WHERE company='BENGAL'")
        conn.commit()
        conn.close()
        if n:
            print(f"[db] added {n} regen-support Pricelist rows (DPT flow meters / PLC)")
    except Exception as e:
        print(f"WARN: ensure regen pricelist extras failed: {e}")


_startup_ensure_regen_pricelist_extras()


def _startup_seed_markups():
    """Editable cost→price markups for HPU and Blower (like the burner Markup
    Master). Seeded once; the tabs AND the offers read from here."""
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.execute("CREATE TABLE IF NOT EXISTS product_markup ("
                     "product TEXT, key TEXT, label TEXT, value REAL, sort INTEGER, "
                     "PRIMARY KEY(product, key))")
        defaults = [
            ("hpu",    "markup",        "Selling markup — Material cost ×", 1.8, 1),
            ("blower", "without_motor", "Price w/o motor — Blower Alone ×", 1.8, 1),
            ("blower", "motor",         "Motor add-on — Motor ×",           1.5, 2),
        ]
        for product, key, label, value, sort in defaults:
            conn.execute("INSERT OR IGNORE INTO product_markup "
                         "(product, key, label, value, sort) VALUES (?,?,?,?,?)",
                         (product, key, label, value, sort))
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"WARN: startup seed_markups failed: {e}")


_startup_seed_markups()


def _product_markup(conn, product: str, key: str, default: float) -> float:
    """Read one product markup value, falling back to `default`."""
    try:
        r = conn.execute("SELECT value FROM product_markup WHERE product=? AND key=?",
                         (product, key)).fetchone()
        return float(r[0]) if r and r[0] is not None else default
    except Exception:
        return default


# One-off pricelist adjustments, applied exactly once per database (tracked in
# _price_ops) so they reach the persistent Railway volume on the next deploy
# without re-applying on every boot.
_PRICE_OPS = [
    ("dembla_psov_discount_45",
     "45% discount on DEMBLA Pneumatic Shut Off Valve prices",
     "UPDATE component_price_master SET previous_price=price, "
     "price=ROUND(price*0.55, 2) "
     "WHERE category='Pneumatic Shut Off Valve' AND UPPER(company)='DEMBLA'"),
    # Blower Alone rows now hold the AMOUNT (not the ×1.8 price); price is
    # derived as Blower Alone × 1.8. Convert any existing ×1.8-price rows.
    ("blower_alone_store_amount",
     "Store Amount (not the x1.8 price) in Blower Alone rates rows",
     "UPDATE component_price_master SET previous_price=price, price=ROUND("
     "(SELECT b.per_kg_amount FROM blower_pricelist_master b "
     " WHERE b.model=component_price_master.item "
     " AND b.section IN ('MEDIUM PRESSURE','HIGH PRESSURE') LIMIT 1)) "
     "WHERE category IN ('Blower Alone (28 inch)','Blower Alone (40 inch)')"),
]


def _startup_price_ops():
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.execute("CREATE TABLE IF NOT EXISTS _price_ops "
                     "(op_id TEXT PRIMARY KEY, applied_at TEXT)")
        for op_id, desc, sql in _PRICE_OPS:
            if conn.execute("SELECT 1 FROM _price_ops WHERE op_id=?",
                            (op_id,)).fetchone():
                continue
            n = conn.execute(sql).rowcount
            conn.execute("INSERT INTO _price_ops(op_id, applied_at) "
                         "VALUES (?, datetime('now'))", (op_id,))
            conn.commit()
            print(f"[db] price op '{op_id}' applied to {n} rows ({desc})")
        conn.close()
    except Exception as e:
        print(f"WARN: startup price ops failed: {e}")


_startup_price_ops()


class _PartEdit(BaseModel):
    rid: int
    field: str
    value: Optional[str] = ""


@app.post("/api/internal-costing/update-part")
def api_ic_update_part(u: _PartEdit):
    """Edit one cell of oil_burner_master, then cascade: recompute this row's
    Amount (=Qty×Rate) and Total (=Amount+M/C), then the derived burner prices."""
    allowed = {"s_no", "particular", "qty", "unit", "rate", "amount", "mc_cost", "total_amount"}
    numeric = {"qty", "rate", "amount", "mc_cost", "total_amount"}
    if u.field not in allowed:
        return {"success": False, "error": "field not editable"}
    try:
        val = (u.value or "").strip()
        if u.field in numeric:
            val = val.replace(",", "")
            try:
                val = float(val)
            except ValueError:
                val = None if val == "" else val
        else:
            val = val or None
        conn = sqlite3.connect(DB_PATH)
        conn.execute(f"UPDATE oil_burner_master SET {u.field}=? WHERE rowid=?", (val, u.rid))

        def _f(x):
            try:
                return float(x)
            except (TypeError, ValueError):
                return 0.0
        # rows that carry a Qty AND Rate are formula rows: Amount=Qty×Rate,
        # Total=Amount+M/C. Lump rows (labour, paint) have neither — leave their
        # hand-entered Total alone.
        row = conn.execute("SELECT qty, rate, mc_cost FROM oil_burner_master WHERE rowid=?",
                           (u.rid,)).fetchone()
        if row and row[0] not in (None, "") and row[1] not in (None, ""):
            amount = _f(row[0]) * _f(row[1])
            conn.execute("UPDATE oil_burner_master SET amount=?, total_amount=? WHERE rowid=?",
                         (amount, amount + _f(row[2]), u.rid))
        recompute_burner_prices(conn)
        conn.commit()
        conn.close()
        return {"success": True}
    except Exception as e:
        return {"success": False, "error": str(e)}


class _PriceEdit(BaseModel):
    section: str
    burner_size: str
    component: str
    value: Optional[str] = ""


@app.post("/api/internal-costing/update-price")
def api_ic_update_price(u: _PriceEdit):
    """Edit one price cell of burner_pricelist_master (the film/spares tables)."""
    try:
        try:
            price = float((u.value or "").replace(",", "").strip())
        except ValueError:
            price = None
        conn = sqlite3.connect(DB_PATH)
        conn.execute("UPDATE burner_pricelist_master SET price=? "
                     "WHERE section=? AND burner_size=? AND component=?",
                     (price, u.section, u.burner_size, u.component))
        # an independent cell (Y-Strainer / Air Resistor / 3A inputs) feeds the
        # Burner Set total, so re-roll the derived prices.
        recompute_burner_prices(conn)
        conn.commit()
        conn.close()
        return {"success": True}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/api/internal-costing/markups")
def api_ic_markups():
    """The editable cost→price markup factors (Markup Master)."""
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        rows = [dict(r) for r in conn.execute(
            "SELECT key, label, value FROM burner_markup ORDER BY sort")]
        conn.close()
        return {"markups": rows}
    except Exception as e:
        return {"markups": [], "error": str(e)}


class _MarkupEdit(BaseModel):
    key: str
    value: Optional[str] = ""


@app.post("/api/internal-costing/update-markup")
def api_ic_update_markup(u: _MarkupEdit):
    """Edit a markup factor, then recompute every burner price."""
    try:
        try:
            val = float((u.value or "").replace(",", "").strip())
        except ValueError:
            val = None
        conn = sqlite3.connect(DB_PATH)
        conn.execute("UPDATE burner_markup SET value=? WHERE key=?", (val, u.key))
        recompute_burner_prices(conn)
        conn.commit()
        conn.close()
        return {"success": True}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/api/internal-costing/product-markups")
def api_ic_product_markups(product: str):
    """Editable markups for a product (hpu / blower) — its Markup Master."""
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        rows = [dict(r) for r in conn.execute(
            "SELECT key, label, value FROM product_markup WHERE product=? ORDER BY sort",
            (product,))]
        conn.close()
        return {"markups": rows}
    except Exception as e:
        return {"markups": [], "error": str(e)}


class _ProductMarkupEdit(BaseModel):
    product: str
    key: str
    value: float


@app.post("/api/internal-costing/product-markup")
def api_ic_update_product_markup(u: _ProductMarkupEdit):
    """Edit an HPU/Blower markup factor. Prices recompute on next read (tab +
    offer both read from product_markup)."""
    try:
        conn = sqlite3.connect(DB_PATH)
        n = conn.execute("UPDATE product_markup SET value=? WHERE product=? AND key=?",
                         (u.value, u.product, u.key)).rowcount
        conn.commit()
        conn.close()
        return {"success": n > 0}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/api/internal-costing/rates")
def api_ic_rates():
    """The oil-burner rate master (Rates! items the part rates pull from)."""
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        rows = [dict(r) for r in conn.execute(
            "SELECT cell, name, value, unit, category FROM rate_master ORDER BY sort")]
        conn.close()
        for r in rows:                       # mark Pricelist-sourced rates
            r["linked"] = r["cell"] in RATE_CPM
        return {"rates": rows}
    except Exception as e:
        return {"rates": [], "error": str(e)}


class _RateEdit(BaseModel):
    cell: str
    value: Optional[str] = ""


@app.post("/api/internal-costing/update-rate")
def api_ic_update_rate(u: _RateEdit):
    """Edit a rate-master value, then cascade: push it into every part that
    references this rate (rate, Amount=Qty×rate, Total=Amount+M/C), then re-roll
    all derived burner prices."""
    if u.cell in RATE_CPM:
        return {"success": False, "error": "This rate is sourced from the Pricelist — edit it there."}
    try:
        try:
            newrate = float((u.value or "").replace(",", "").strip())
        except ValueError:
            newrate = None
        conn = sqlite3.connect(DB_PATH)
        conn.execute("UPDATE rate_master SET value=? WHERE cell=?", (newrate, u.cell))

        def _f(x):
            try:
                return float(x)
            except (TypeError, ValueError):
                return 0.0
        _push_rate_to_parts(conn, u.cell, newrate or 0.0, _f)
        recompute_burner_prices(conn)
        conn.commit()
        conn.close()
        return {"success": True}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/api/internal-costing/oil-burner")
def api_ic_oil_burner():
    """Oil-burner internal parts costing (oil_burner_master), grouped by size."""
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        rows = [dict(r) for r in conn.execute(
            "SELECT rowid AS _rid, s_no, particular, qty, unit, rate, amount, "
            "mc_cost, total_amount, burner_type, rate_ref "
            "FROM oil_burner_master ORDER BY burner_type, rowid")]
        conn.close()
        def _f(v):
            try:
                return float(v)
            except (TypeError, ValueError):
                return 0.0
        groups = {}
        for r in rows:
            groups.setdefault(r["burner_type"] or "—", []).append(r)
        out = []
        for bt, items in groups.items():
            for it in items:
                rt = _f(it["total_amount"])
                if rt == 0.0:   # total not pre-summed — derive from amount + M/C cost
                    rt = _f(it["amount"]) + _f(it["mc_cost"])
                it["row_total"] = round(rt)
            total = sum(i["row_total"] for i in items)
            out.append({"burner_type": bt, "items": items, "total": round(total)})
        return {"groups": out}
    except Exception as e:
        return {"groups": [], "error": str(e)}


@app.get("/api/internal-costing/hv-oil-burner")
def api_ic_hv_oil_burner():
    """HV oil-burner parts costing (hv_oil_burner_master), grouped by size."""
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        rows = [dict(r) for r in conn.execute(
            "SELECT rowid AS _rid, s_no, particular, qty, unit, rate, amount, "
            "mc_cost, total_amount, burner_type, rate_ref FROM hv_oil_burner_master ORDER BY burner_type, rowid")]
        conn.close()

        def _f(v):
            try:
                return float(v)
            except (TypeError, ValueError):
                return 0.0
        groups = {}
        for r in rows:
            groups.setdefault(r["burner_type"] or "—", []).append(r)
        out = []
        for bt, items in groups.items():
            for it in items:
                rt = _f(it["total_amount"])
                if rt == 0.0:
                    rt = _f(it["amount"]) + _f(it["mc_cost"])
                it["row_total"] = round(rt)
            out.append({"burner_type": bt, "items": items,
                        "total": round(sum(i["row_total"] for i in items))})
        return {"groups": out}
    except Exception as e:
        return {"groups": [], "error": str(e)}


@app.get("/api/internal-costing/hpu")
def api_ic_hpu():
    """HPU (Heating & Pumping Unit) internal parts costing (hpu_master).
    Grouped by KW rating, then by variant (Simplex / Duplex 1 / Duplex 2),
    each variant a full BOM breakdown. Every line's RATE is pulled live from
    the pricelist (component_price_master, HPU_* categories) via the resolver;
    qty stays per-line, amount = qty × rate. Selling price = cost × 1.8 (see
    bom/hpu_calculator.HPU_MARKUP). LABOUR is a fixed line (stored amount)."""
    from bom import hpu_pricelist as _hp
    HPU_MARKUP = 1.8
    VARIANT_ORDER = ["Simplex", "Duplex 1", "Duplex 2"]
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        rows = [dict(r) for r in conn.execute(
            "SELECT rowid AS _rid, unit_kw, variant, item, qty, unit, rate, amount "
            "FROM hpu_master ORDER BY unit_kw, rowid")]
        rates = _hp.load_rates(conn)
        HPU_MARKUP = _product_markup(conn, "hpu", "markup", 1.8)   # editable
        conn.close()

        def _f(v):
            try:
                return float(v)
            except (TypeError, ValueError):
                return 0.0

        # data[kw][variant] = {items, total, selling}
        data = {}
        for r in rows:
            kw = r["unit_kw"]
            var = r["variant"] or "—"
            r["labour"] = _hp.is_labour(r["item"])
            if r["labour"]:
                # Labour is not a material — keep its stored amount, no link.
                amount = _f(r["amount"])
                r["rate"] = None
                r["rate_ref"] = None
            else:
                rate, src = _hp.resolve_rate(r["item"], r["unit"], rates)
                r["rate"] = round(rate, 2)
                r["rate_ref"] = src          # pricelist SKU the rate came from
                amount = _f(r["qty"]) * rate
            r["amount"] = round(amount, 2)
            bucket = data.setdefault(kw, {})
            grp = bucket.setdefault(var, {"items": [], "total": 0.0})
            grp["items"].append(r)
            grp["total"] += amount

        ratings = sorted(data.keys())
        # Number each variant's rows and finalise totals/selling in fixed order.
        out = {}
        for kw in ratings:
            out[str(kw)] = {}
            for var in VARIANT_ORDER:
                grp = data[kw].get(var)
                if not grp:
                    continue
                for i, it in enumerate(grp["items"], 1):
                    it["s_no"] = i
                total = round(grp["total"])
                out[str(kw)][var] = {
                    "items":   grp["items"],
                    "total":   total,
                    "selling": round(total * HPU_MARKUP),
                }
        return {"ratings": ratings, "variants": VARIANT_ORDER,
                "markup": HPU_MARKUP, "data": out}
    except Exception as e:
        return {"ratings": [], "variants": [], "markup": HPU_MARKUP,
                "data": {}, "error": str(e)}


class _HpuEdit(BaseModel):
    rid: int
    field: str          # 'qty' | 'amount'
    value: float


@app.post("/api/internal-costing/hpu-update")
def api_ic_hpu_update(req: _HpuEdit):
    """Edit a per-line HPU BOM value that isn't from the pricelist — the line
    Qty, or the LABOUR amount. Rates stay pricelist-linked; amount recomputes
    from qty × rate on next read."""
    if req.field not in ("qty", "amount"):
        return {"success": False, "error": f"bad field {req.field!r}"}
    try:
        conn = sqlite3.connect(DB_PATH)
        n = conn.execute(
            f"UPDATE hpu_master SET {req.field}=? WHERE rowid=?",
            (req.value, req.rid)).rowcount
        conn.commit()
        conn.close()
        return {"success": n > 0}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/api/internal-costing/blower")
def api_ic_blower():
    """Blower pricing on the PERKIN Rates basis (blower_pricelist_master,
    MEDIUM / HIGH PRESSURE). Price without motor = Amount × 1.8; with motor
    = + Motor × 1.5. This is the basis for the blower equipment offer too."""
    from bom import blower_pricelist as _bp
    try:
        conn = sqlite3.connect(DB_PATH)
        data = _bp.blower_models(conn)
        legacy = _bp.legacy_models(conn)
        wo_mk, motor_mk = _bp.blower_markups(conn)   # editable Markup Master
        conn.close()
        sections = [s for s in _bp.SECTIONS if s in data] + \
                   [s for s in data if s not in _bp.SECTIONS]
        legacy_sections = [s for s in _bp.LEGACY_SECTIONS if s in legacy] + \
                          [s for s in legacy if s not in _bp.LEGACY_SECTIONS]
        return {"sections": sections, "data": data,
                "without_markup": wo_mk, "motor_markup": motor_mk,
                "legacy_sections": legacy_sections, "legacy": legacy,
                "legacy_overhead": _bp.LEGACY_OVERHEAD, "legacy_markup": _bp.LEGACY_MARKUP}
    except Exception as e:
        return {"sections": [], "data": {}, "without_markup": 1.8,
                "motor_markup": 1.5, "legacy_sections": [], "legacy": {},
                "legacy_overhead": 1.3, "legacy_markup": 1.8, "error": str(e)}


class _BlowerEdit(BaseModel):
    model: str
    field: str          # 'weight' | 'amount' | 'motor'
    value: float


@app.post("/api/internal-costing/blower-update")
def api_ic_blower_update(req: _BlowerEdit):
    """Edit a PERKIN blower field (Weight / Amount / Motor) in
    blower_pricelist_master. Price w/o & w/ motor recompute on next read."""
    COLS = {"weight": "blower_weight", "amount": "per_kg_amount",
            "motor": "motor_price_abb"}
    col = COLS.get(req.field)
    if not col:
        return {"success": False, "error": f"bad field {req.field!r}"}
    try:
        conn = sqlite3.connect(DB_PATH)
        n = conn.execute(
            f"UPDATE blower_pricelist_master SET {col}=? "
            "WHERE model=? AND section IN ('MEDIUM PRESSURE','HIGH PRESSURE')",
            (req.value, req.model)).rowcount
        conn.commit()
        conn.close()
        return {"success": n > 0}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/api/internal-costing/regen")
def api_ic_regen():
    """Regenerative-burner sizing & costing breakdown — the 'Burner Sizing and
    costing' sheet: per-KW material weights × material rates → burner cost.
    Weights are from regen_builder._BURNER_WEIGHTS; the material rates are the
    editable regen_material_rates (MS/SS/Refractory/Ceramic, +labour, ×wastage).
    The resulting per-model total is the burner_cost that seeds the REGEN
    pricelist (Regen: Burner with Regenerator …)."""
    from bom.regen_builder import _BURNER_WEIGHTS, MODEL_KWS, REGEN_MODELS
    # component -> (display label, weight key, material used for the rate)
    COMPONENTS = [
        ("Burner Body — MS",           "burner_ms",     "MS"),
        ("Burner Body — Refractory",   "burner_refrac", "Refractory"),
        ("Regenerator — MS",           "regen_ms",      "MS"),
        ("Regenerator — SS",           "regen_ss",      "SS"),
        ("Regenerator — Refractory",   "regen_refrac",  "Refractory"),
        ("Regenerator — Ceramic Balls","regen_ceramic", "Ceramic Balls"),
        ("Burner Block — Refractory",  "block_refrac",  "Refractory"),
    ]
    try:
        conn = sqlite3.connect(DB_PATH)
        rates, rate_rows = {}, []
        for material, wastage, mc, lc in conn.execute(
                "SELECT material, wastage, material_cost, labor_cost "
                "FROM regen_material_rates"):
            wa = wastage or 0; mc = mc or 0; lc = lc or 0
            eff = round((mc + lc) * (1 + wa), 2)
            rates[material] = eff
            rate_rows.append({"material": material, "wastage": wa,
                              "material_cost": mc, "labor_cost": lc, "rate": eff})
        conn.close()
        models = []
        for kw in MODEL_KWS:
            w = _BURNER_WEIGHTS[kw]
            comps, total = [], 0.0
            for label, wk, mat in COMPONENTS:
                wt = w[wk]; rate = rates.get(mat, 0.0); cost = round(wt * rate, 2)
                total += cost
                comps.append({"label": label, "material": mat,
                              "weight": wt, "rate": rate, "cost": cost})
            models.append({"kw": kw, "components": comps,
                           "total_cost": round(total, 2),
                           "pricelist_cost": REGEN_MODELS[kw]["burner_cost"]})
        return {"component_labels": [c[0] for c in COMPONENTS],
                "material_rates": rate_rows, "models": models}
    except Exception as e:
        return {"component_labels": [], "material_rates": [], "models": [],
                "error": str(e)}


@app.get("/enquiries", response_class=HTMLResponse)
def enquiries_page():
    with open(os.path.join(BASE_DIR, "enquiries.html"), "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())


@app.get("/projects", response_class=HTMLResponse)
def projects_page():
    with open(os.path.join(BASE_DIR, "projects.html"), "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())


@app.get("/api/dashboard")
def api_dashboard():
    """Live dashboard metrics, all derived from quotes_log (every generated
    offer is recorded there). Never raises — returns empty buckets on error."""
    from datetime import datetime as _dt
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        now = _dt.now()
        ym = now.strftime("%Y-%m")

        def scalar(sql, args=()):
            r = conn.execute(sql, args).fetchone()
            return (r[0] if r and r[0] is not None else 0)

        total_quotes = scalar("SELECT COUNT(*) FROM quotes_log")
        total_value  = scalar("SELECT SUM(grand_total) FROM quotes_log")
        month_quotes = scalar("SELECT COUNT(*) FROM quotes_log WHERE substr(created_at,1,7)=?", (ym,))
        month_value  = scalar("SELECT SUM(grand_total) FROM quotes_log WHERE substr(created_at,1,7)=?", (ym,))
        avg_value    = scalar("SELECT AVG(grand_total) FROM quotes_log WHERE grand_total>0")
        companies    = scalar("SELECT COUNT(DISTINCT company_name) FROM quotes_log WHERE company_name<>''")

        by_product = [
            {"type": (r["equipment_type"] or "Other"), "count": r["c"], "value": (r["v"] or 0)}
            for r in conn.execute(
                "SELECT equipment_type, COUNT(*) c, SUM(grand_total) v FROM quotes_log "
                "GROUP BY equipment_type ORDER BY c DESC, v DESC")
        ]

        # Last 12 calendar months (oldest -> newest), zero-filled.
        seq = []
        for i in range(11, -1, -1):
            mm, yy = now.month - i, now.year
            while mm <= 0:
                mm += 12; yy -= 1
            seq.append(f"{yy:04d}-{mm:02d}")
        mrows = {r["ym"]: (r["c"], r["v"] or 0) for r in conn.execute(
            "SELECT substr(created_at,1,7) ym, COUNT(*) c, SUM(grand_total) v "
            "FROM quotes_log GROUP BY ym")}
        by_month = [{"month": k, "count": mrows.get(k, (0, 0))[0],
                     "value": mrows.get(k, (0, 0))[1]} for k in seq]

        recent = [
            {"created_at": r["created_at"], "company": r["company_name"],
             "product": r["equipment_type"], "ref": (r["ref_no"] or r["quote_no"] or ""),
             "value": (r["grand_total"] or 0), "location": (r["location"] or ""),
             "poc": (r["poc_name"] or ""), "file": os.path.basename(r["file_path"] or "")}
            for r in conn.execute("SELECT * FROM quotes_log ORDER BY id DESC LIMIT 12")
        ]

        # ── CRM metrics (Phase 2) ──
        open_enquiries  = scalar("SELECT COUNT(*) FROM enquiries WHERE stage NOT IN ('won','lost')")
        pending_rfq     = scalar("SELECT COUNT(*) FROM enquiries WHERE stage IN ('new','qualified')")
        active_projects = scalar("SELECT COUNT(*) FROM projects WHERE status='active'")
        stage_rows = {r["stage"]: (r["c"], r["v"] or 0) for r in conn.execute(
            "SELECT stage, COUNT(*) c, SUM(value_est) v FROM enquiries GROUP BY stage")}
        pipeline = [{"stage": st, "count": stage_rows.get(st, (0, 0))[0],
                     "value": stage_rows.get(st, (0, 0))[1]} for st in ENQUIRY_STAGES]
        activity = [dict(r) for r in conn.execute(
            "SELECT ts, kind, title, detail FROM activity_log ORDER BY id DESC LIMIT 10")]
        notifications = _compute_notifications(conn)
        conn.close()
        return {
            "kpis": {
                "total_quotes": total_quotes, "month_quotes": month_quotes,
                "total_value": total_value, "month_value": month_value,
                "avg_value": round(avg_value or 0), "companies": companies,
                "open_enquiries": open_enquiries, "active_projects": active_projects,
                "pending_rfq": pending_rfq,
            },
            "by_product": by_product, "by_month": by_month, "recent": recent,
            "pipeline": pipeline, "activity": activity, "notifications": notifications,
            "month_label": now.strftime("%b %Y"),
        }
    except Exception as e:
        import traceback
        return {"error": str(e), "trace": traceback.format_exc(),
                "kpis": {}, "by_product": [], "by_month": [], "recent": [],
                "pipeline": [], "activity": [], "notifications": []}


# ══════════════════════════════════════════════════════════════════════════
#  CRM — Enquiries & Projects (dashboard Phase 2)
# ══════════════════════════════════════════════════════════════════════════
ENQUIRY_STAGES = ["new", "qualified", "quoted", "won", "lost"]
PROJECT_STATUSES = ["active", "on_hold", "completed", "cancelled"]


class EnquiryIn(BaseModel):
    id: Optional[int] = None
    company_name: str = ""
    contact_name: str = ""
    email: str = ""
    phone: str = ""
    location: str = ""
    product: str = ""
    source: str = ""
    stage: str = "new"
    value_est: float = 0
    owner: str = ""
    notes: str = ""


class ProjectIn(BaseModel):
    id: Optional[int] = None
    company_name: str = ""
    title: str = ""
    product: str = ""
    value: float = 0
    status: str = "active"
    progress: int = 0
    owner: str = ""
    start_date: str = ""
    target_date: str = ""
    enquiry_id: Optional[int] = None
    notes: str = ""


class IdIn(BaseModel):
    id: int


class StageIn(BaseModel):
    id: int
    stage: str


class ProjStatusIn(BaseModel):
    id: int
    status: Optional[str] = None
    progress: Optional[int] = None


@app.post("/api/enquiry")
def api_enquiry_save(e: EnquiryIn):
    """Create (no id) or update (id) an enquiry."""
    from datetime import datetime as _dt
    now = _dt.now().isoformat(timespec="seconds")
    try:
        conn = sqlite3.connect(DB_PATH)
        if e.id:
            conn.execute("""UPDATE enquiries SET company_name=?,contact_name=?,email=?,phone=?,
                location=?,product=?,source=?,stage=?,value_est=?,owner=?,notes=?,updated_at=?
                WHERE id=?""", (e.company_name, e.contact_name, e.email, e.phone, e.location,
                e.product, e.source, e.stage, e.value_est, e.owner, e.notes, now, e.id))
            eid = e.id
            created = False
        else:
            n = conn.execute("SELECT COUNT(*) FROM enquiries").fetchone()[0] + 1
            cur = conn.execute("""INSERT INTO enquiries (enquiry_no,company_name,contact_name,email,
                phone,location,product,source,stage,value_est,owner,notes,created_at,updated_at)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)""", (f"ENQ-{n:04d}", e.company_name,
                e.contact_name, e.email, e.phone, e.location, e.product, e.source,
                e.stage or "new", e.value_est, e.owner, e.notes, now, now))
            eid = cur.lastrowid
            created = True
        conn.commit()
        conn.close()
        if created:   # log AFTER the write commits/closes (avoids a SQLite write-lock)
            _log_activity("enquiry", f"Enquiry · {e.company_name or 'Client'}",
                          f"{e.product or 'New enquiry'}{(' · ' + e.location) if e.location else ''}")
        return {"success": True, "id": eid}
    except Exception as ex:
        return {"success": False, "error": str(ex)}


@app.get("/api/enquiries")
def api_enquiries():
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        rows = [dict(r) for r in conn.execute("SELECT * FROM enquiries ORDER BY id DESC")]
        conn.close()
        return {"enquiries": rows, "stages": ENQUIRY_STAGES}
    except Exception as ex:
        return {"enquiries": [], "stages": ENQUIRY_STAGES, "error": str(ex)}


@app.post("/api/enquiry/stage")
def api_enquiry_stage(s: StageIn):
    from datetime import datetime as _dt
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        r = conn.execute("SELECT company_name FROM enquiries WHERE id=?", (s.id,)).fetchone()
        conn.execute("UPDATE enquiries SET stage=?, updated_at=? WHERE id=?",
                     (s.stage, _dt.now().isoformat(timespec="seconds"), s.id))
        conn.commit()
        conn.close()
        _log_activity("enquiry", f"Enquiry → {s.stage}", (r["company_name"] if r else ""))
        return {"success": True}
    except Exception as ex:
        return {"success": False, "error": str(ex)}


@app.post("/api/enquiry/delete")
def api_enquiry_delete(d: IdIn):
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.execute("DELETE FROM enquiries WHERE id=?", (d.id,))
        conn.commit()
        conn.close()
        return {"success": True}
    except Exception as ex:
        return {"success": False, "error": str(ex)}


@app.post("/api/enquiry/convert")
def api_enquiry_convert(d: IdIn):
    """Create a project from an enquiry and mark the enquiry 'won'."""
    from datetime import datetime as _dt
    now = _dt.now().isoformat(timespec="seconds")
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        e = conn.execute("SELECT * FROM enquiries WHERE id=?", (d.id,)).fetchone()
        if not e:
            conn.close()
            return {"success": False, "error": "enquiry not found"}
        n = conn.execute("SELECT COUNT(*) FROM projects").fetchone()[0] + 1
        pno = f"PRJ-{n:04d}"
        cur = conn.execute("""INSERT INTO projects (project_no,company_name,title,product,value,
            status,progress,owner,start_date,target_date,enquiry_id,notes,created_at,updated_at)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)""", (pno, e["company_name"],
            e["product"] or "Project", e["product"], e["value_est"] or 0, "active", 0,
            e["owner"], now[:10], "", e["id"], e["notes"], now, now))
        conn.execute("UPDATE enquiries SET stage='won', updated_at=? WHERE id=?", (now, d.id))
        pid = cur.lastrowid
        conn.commit()
        conn.close()
        _log_activity("project", f"Won → Project · {e['company_name']}", pno)
        return {"success": True, "project_id": pid, "project_no": pno}
    except Exception as ex:
        return {"success": False, "error": str(ex)}


@app.post("/api/project")
def api_project_save(p: ProjectIn):
    from datetime import datetime as _dt
    now = _dt.now().isoformat(timespec="seconds")
    try:
        conn = sqlite3.connect(DB_PATH)
        if p.id:
            conn.execute("""UPDATE projects SET company_name=?,title=?,product=?,value=?,status=?,
                progress=?,owner=?,start_date=?,target_date=?,notes=?,updated_at=? WHERE id=?""",
                (p.company_name, p.title, p.product, p.value, p.status, p.progress, p.owner,
                 p.start_date, p.target_date, p.notes, now, p.id))
            pid = p.id
            created = False
        else:
            n = conn.execute("SELECT COUNT(*) FROM projects").fetchone()[0] + 1
            cur = conn.execute("""INSERT INTO projects (project_no,company_name,title,product,value,
                status,progress,owner,start_date,target_date,enquiry_id,notes,created_at,updated_at)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)""", (f"PRJ-{n:04d}", p.company_name, p.title,
                p.product, p.value, p.status or "active", p.progress, p.owner, p.start_date,
                p.target_date, p.enquiry_id, p.notes, now, now))
            pid = cur.lastrowid
            created = True
        conn.commit()
        conn.close()
        if created:   # log AFTER the write commits/closes (avoids a SQLite write-lock)
            _log_activity("project", f"Project · {p.company_name or 'Client'}",
                          f"{p.title or p.product or 'New project'}")
        return {"success": True, "id": pid}
    except Exception as ex:
        return {"success": False, "error": str(ex)}


@app.get("/api/projects")
def api_projects():
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        rows = [dict(r) for r in conn.execute("SELECT * FROM projects ORDER BY id DESC")]
        conn.close()
        return {"projects": rows, "statuses": PROJECT_STATUSES}
    except Exception as ex:
        return {"projects": [], "statuses": PROJECT_STATUSES, "error": str(ex)}


@app.post("/api/project/status")
def api_project_status(s: ProjStatusIn):
    from datetime import datetime as _dt
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        r = conn.execute("SELECT company_name FROM projects WHERE id=?", (s.id,)).fetchone()
        sets, args = [], []
        if s.status is not None:
            sets.append("status=?"); args.append(s.status)
        if s.progress is not None:
            sets.append("progress=?"); args.append(s.progress)
        sets.append("updated_at=?"); args.append(_dt.now().isoformat(timespec="seconds"))
        args.append(s.id)
        conn.execute(f"UPDATE projects SET {','.join(sets)} WHERE id=?", args)
        conn.commit()
        conn.close()
        _log_activity("project", f"Project {s.status or 'updated'}", (r["company_name"] if r else ""))
        return {"success": True}
    except Exception as ex:
        return {"success": False, "error": str(ex)}


@app.post("/api/project/delete")
def api_project_delete(d: IdIn):
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.execute("DELETE FROM projects WHERE id=?", (d.id,))
        conn.commit()
        conn.close()
        return {"success": True}
    except Exception as ex:
        return {"success": False, "error": str(ex)}


@app.get("/api/clients")
def api_clients():
    """Company-level rollup across quotes, enquiries and projects."""
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        q = {r["company_name"]: dict(r) for r in conn.execute(
            "SELECT company_name, COUNT(*) quotes, SUM(grand_total) value, "
            "MAX(created_at) last_quote, MAX(location) location, MAX(poc_name) poc "
            "FROM quotes_log WHERE company_name<>'' GROUP BY company_name")}
        enq = {r["company_name"]: dict(r) for r in conn.execute(
            "SELECT company_name, COUNT(*) enquiries, "
            "SUM(CASE WHEN stage NOT IN ('won','lost') THEN 1 ELSE 0 END) open_enq "
            "FROM enquiries WHERE company_name<>'' GROUP BY company_name")}
        prj = {r["company_name"]: dict(r) for r in conn.execute(
            "SELECT company_name, COUNT(*) projects, "
            "SUM(CASE WHEN status='active' THEN 1 ELSE 0 END) active_proj, "
            "SUM(value) proj_value FROM projects WHERE company_name<>'' GROUP BY company_name")}
        conn.close()
        clients = []
        for nm in (set(q) | set(enq) | set(prj)):
            qd, ed, pd = q.get(nm, {}), enq.get(nm, {}), prj.get(nm, {})
            clients.append({
                "company": nm, "location": qd.get("location") or "",
                "poc": qd.get("poc") or "",
                "quotes": qd.get("quotes", 0) or 0, "quoted_value": qd.get("value", 0) or 0,
                "enquiries": ed.get("enquiries", 0) or 0, "open_enq": ed.get("open_enq", 0) or 0,
                "projects": pd.get("projects", 0) or 0, "active_proj": pd.get("active_proj", 0) or 0,
                "proj_value": pd.get("proj_value", 0) or 0, "last_quote": qd.get("last_quote") or "",
            })
        clients.sort(key=lambda c: (c["quoted_value"], c["quotes"]), reverse=True)
        return {"clients": clients}
    except Exception as ex:
        return {"clients": [], "error": str(ex)}


@app.get("/api/reports")
def api_reports():
    """Analytics aggregates for the Reports page."""
    from datetime import datetime as _dt
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        now = _dt.now()

        def scal(sql, a=()):
            r = conn.execute(sql, a).fetchone()
            return (r[0] if r and r[0] is not None else 0)

        total_quotes = scal("SELECT COUNT(*) FROM quotes_log")
        total_value = scal("SELECT SUM(grand_total) FROM quotes_log")
        won = scal("SELECT COUNT(*) FROM enquiries WHERE stage='won'")
        lost = scal("SELECT COUNT(*) FROM enquiries WHERE stage='lost'")
        total_enq = scal("SELECT COUNT(*) FROM enquiries")
        proj_value = scal("SELECT SUM(value) FROM projects WHERE status<>'cancelled'")
        win_rate = round(won / (won + lost) * 100) if (won + lost) else 0
        conversion = round(won / total_enq * 100) if total_enq else 0

        def group(sql):
            return [{"label": (r["k"] or "—"), "count": r["c"], "value": (r["v"] or 0)}
                    for r in conn.execute(sql)]
        by_product = group("SELECT equipment_type k, COUNT(*) c, SUM(grand_total) v "
                           "FROM quotes_log GROUP BY equipment_type ORDER BY v DESC")
        by_location = group("SELECT location k, COUNT(*) c, SUM(grand_total) v "
                            "FROM quotes_log GROUP BY location ORDER BY c DESC")
        by_owner = group("SELECT COALESCE(NULLIF(marketing_person,''),NULLIF(technical_person,'')) k, "
                         "COUNT(*) c, SUM(grand_total) v FROM quotes_log GROUP BY k ORDER BY v DESC")
        top_clients = group("SELECT company_name k, COUNT(*) c, SUM(grand_total) v FROM quotes_log "
                            "WHERE company_name<>'' GROUP BY company_name ORDER BY v DESC LIMIT 10")

        seq = []
        for i in range(11, -1, -1):
            mm, yy = now.month - i, now.year
            while mm <= 0:
                mm += 12; yy -= 1
            seq.append(f"{yy:04d}-{mm:02d}")
        mrows = {r["ym"]: (r["c"], r["v"] or 0) for r in conn.execute(
            "SELECT substr(created_at,1,7) ym, COUNT(*) c, SUM(grand_total) v FROM quotes_log GROUP BY ym")}
        by_month = [{"label": k, "count": mrows.get(k, (0, 0))[0],
                     "value": mrows.get(k, (0, 0))[1]} for k in seq]
        pipe = {r["stage"]: r["c"] for r in conn.execute(
            "SELECT stage, COUNT(*) c FROM enquiries GROUP BY stage")}
        pipeline = [{"label": st, "count": pipe.get(st, 0)} for st in ENQUIRY_STAGES]
        conn.close()
        return {
            "kpis": {"total_quotes": total_quotes, "total_value": total_value,
                     "win_rate": win_rate, "conversion": conversion,
                     "total_enquiries": total_enq, "order_book": proj_value,
                     "won": won, "lost": lost},
            "by_product": by_product, "by_location": by_location, "by_owner": by_owner,
            "by_month": by_month, "top_clients": top_clients, "pipeline": pipeline,
        }
    except Exception as ex:
        import traceback
        return {"error": str(ex), "trace": traceback.format_exc(), "kpis": {}}


@app.get("/api/export/{kind}.csv")
def api_export_csv(kind: str):
    """Download quotes / enquiries / projects as CSV."""
    import csv, io
    from fastapi.responses import Response
    tables = {"quotes": "quotes_log", "enquiries": "enquiries", "projects": "projects"}
    tbl = tables.get(kind)
    if not tbl:
        return {"error": "unknown export"}
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    rows = [dict(r) for r in conn.execute(f"SELECT * FROM {tbl} ORDER BY id DESC")]
    conn.close()
    buf = io.StringIO()
    if rows:
        w = csv.DictWriter(buf, fieldnames=list(rows[0].keys()))
        w.writeheader()
        w.writerows(rows)
    return Response(content=buf.getvalue(), media_type="text/csv",
                    headers={"Content-Disposition": f'attachment; filename="encon_{kind}.csv"'})


@app.get("/clients", response_class=HTMLResponse)
def clients_page():
    with open(os.path.join(BASE_DIR, "clients.html"), "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())


@app.get("/reports", response_class=HTMLResponse)
def reports_page():
    with open(os.path.join(BASE_DIR, "reports.html"), "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())


# ── Google Drive OAuth ────────────────────────────────────────────────────
# One-time browser sign-in by an encon.in user (process@encon.in is the
# intended account). The captured refresh token is stored in vlph.db and
# used by drive_uploader for every subsequent offer upload.

@app.get("/auth/drive/login")
def drive_oauth_login():
    """Redirect the user to Google's OAuth consent screen."""
    from fastapi.responses import RedirectResponse
    client_id = os.environ.get("GOOGLE_OAUTH_CLIENT_ID", "").strip()
    redirect_uri = os.environ.get("GOOGLE_OAUTH_REDIRECT_URI", "").strip()
    if not client_id or not redirect_uri:
        return HTMLResponse(
            "<h3>Drive OAuth not configured</h3>"
            "<p>Set GOOGLE_OAUTH_CLIENT_ID and GOOGLE_OAUTH_REDIRECT_URI in Railway env vars.</p>",
            status_code=500,
        )
    from urllib.parse import urlencode
    params = {
        "client_id":     client_id,
        "redirect_uri":  redirect_uri,
        "response_type": "code",
        "scope":         "https://www.googleapis.com/auth/drive.file",
        "access_type":   "offline",   # required to get a refresh token
        "prompt":        "consent",   # forces refresh-token reissue every time
        "include_granted_scopes": "true",
    }
    return RedirectResponse(
        f"https://accounts.google.com/o/oauth2/v2/auth?{urlencode(params)}"
    )


@app.get("/auth/drive/callback")
def drive_oauth_callback(code: str = "", error: str = ""):
    """Handle Google's redirect: exchange the authorisation code for a
    refresh token and persist it. Then show a small confirmation page."""
    if error:
        return HTMLResponse(
            f"<h3>Drive auth failed</h3><p>{error}</p>", status_code=400
        )
    if not code:
        return HTMLResponse(
            "<h3>Drive auth failed</h3><p>No code returned by Google.</p>",
            status_code=400,
        )
    client_id     = os.environ.get("GOOGLE_OAUTH_CLIENT_ID", "").strip()
    client_secret = os.environ.get("GOOGLE_OAUTH_CLIENT_SECRET", "").strip()
    redirect_uri  = os.environ.get("GOOGLE_OAUTH_REDIRECT_URI", "").strip()
    if not (client_id and client_secret and redirect_uri):
        return HTMLResponse(
            "<h3>Drive OAuth not configured</h3>"
            "<p>Missing CLIENT_ID / CLIENT_SECRET / REDIRECT_URI env vars.</p>",
            status_code=500,
        )
    import urllib.request, urllib.parse, json as _json
    body = urllib.parse.urlencode({
        "code":          code,
        "client_id":     client_id,
        "client_secret": client_secret,
        "redirect_uri":  redirect_uri,
        "grant_type":    "authorization_code",
    }).encode()
    try:
        with urllib.request.urlopen(
            "https://oauth2.googleapis.com/token", data=body, timeout=15
        ) as resp:
            token_data = _json.loads(resp.read().decode())
    except Exception as e:
        return HTMLResponse(
            f"<h3>Token exchange failed</h3><pre>{e}</pre>", status_code=500
        )
    refresh_token = token_data.get("refresh_token")
    if not refresh_token:
        return HTMLResponse(
            "<h3>No refresh token returned</h3>"
            "<p>Try /auth/drive/login again. Make sure 'prompt=consent' is set "
            "(it should be — but if Google has already issued a refresh token "
            "to this client and account, it skips it). You may need to revoke "
            "access at https://myaccount.google.com/permissions and retry.</p>"
            f"<pre>{_json.dumps(token_data, indent=2)}</pre>",
            status_code=500,
        )
    from engine.drive_uploader import save_refresh_token
    save_refresh_token(refresh_token)
    # Show the refresh token so the user can persist it as a Railway env
    # var (vlph.db on Railway gets wiped on every redeploy — env vars
    # survive). Once GOOGLE_DRIVE_REFRESH_TOKEN is set the uploader reads
    # from the env var first and never has to re-auth.
    return HTMLResponse(
        "<div style='font-family:Calibri,Arial,sans-serif;max-width:780px;margin:24px auto;padding:24px;'>"
        "<h3 style='color:green'>Google Drive connected ✓</h3>"
        "<p>The offer generator will now upload every new offer (docx + pdf): "
        "VLPH/HLPH go to the Ladle folder, Tundish to the Tundish folder.</p>"
        "<div style='background:#fff7ed;border:1px solid #fdba74;border-radius:8px;padding:14px;margin:18px 0;'>"
        "<b>Important — make it permanent:</b><br/>"
        "On Railway redeploys the local token gets wiped. To keep auth alive "
        "across deploys, copy the refresh token below and set it as a Railway "
        "env var named <code>GOOGLE_DRIVE_REFRESH_TOKEN</code>.<br/>"
        f"<textarea readonly style='width:100%;height:60px;margin-top:10px;font-family:monospace;font-size:12px;'>{refresh_token}</textarea>"
        "<p style='font-size:0.85rem;color:#7c2d12;margin-top:8px;'>Treat this like a password. Don't paste it in chat or commit it to git.</p>"
        "</div>"
        "<p>You can close this tab once the env var is saved.</p>"
        "</div>"
    )


@app.get("/auth/drive/status")
def drive_oauth_status():
    """JSON endpoint the dashboard can poll to show 'Connected' vs not."""
    from engine.drive_uploader import is_authorized
    return {"authorized": is_authorized()}


@app.get("/api/drive/ensure-combined")
def api_drive_ensure_combined():
    """Diagnostic / manual trigger: create (or find) the 'Combined Offers'
    folder now and report exactly what happened, so we can see why it isn't
    auto-creating (usually: Drive not connected on this deploy)."""
    from engine.drive_uploader import is_authorized, _get_service, _ensure_combined_folder
    if not is_authorized():
        return {"authorized": False,
                "error": "Drive not connected. Open /auth/drive/login first, "
                         "then set GOOGLE_DRIVE_REFRESH_TOKEN in Railway."}
    svc = _get_service()
    if svc is None:
        return {"authorized": True,
                "error": "Drive service unavailable — missing GOOGLE_OAUTH_CLIENT_ID/SECRET?"}
    try:
        fid = _ensure_combined_folder(svc)
        if fid:
            return {"authorized": True, "folder_id": fid,
                    "link": f"https://drive.google.com/drive/folders/{fid}"}
        return {"authorized": True,
                "error": "Folder could not be created (check Railway logs for the WARN line)."}
    except Exception as e:
        import traceback
        return {"authorized": True, "error": str(e), "trace": traceback.format_exc()}


@app.get("/api/last-pricebook-update")
def last_pricebook_update():
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("SELECT updated_at FROM table_update_log ORDER BY updated_at DESC LIMIT 1")
        row = c.fetchone()
        conn.close()
        if row:
            dt = datetime.fromisoformat(row[0])
            return {"date": dt.strftime("%d %b %Y")}
        return {"date": None}
    except Exception:
        return {"date": None}


@app.get("/api/system-health")
def system_health():
    """
    Returns health status for each key master table plus recent activity log.
    Used by the dashboard to show a product manager what data is loaded.
    """
    KEY_TABLES = [
        ("component_price_master", "Rates / Component Prices",      "Rates sheet"),
        ("hpu_master",             "HPU Components",                 "HPU sheet"),
        ("burner_pricelist_master","Burner Selling Prices",          "BURNER sheet"),
        ("blower_master",          "Blower Models & Prices",         "Blower sheet"),
        ("vertical_master",        "VLPH Structure Costs",           "Vertical sheet"),
        ("horizontal_master",      "HLPH Structure Costs",           "Horizontal sheet"),
        ("gas_burner_parts_master","Gas Burner Fabrication Parts",   "Gas Burner sheet"),
        ("gail_gas_burner_master", "GAIL Burner Prices",             "GAIL GAS Burner sheet"),
        ("ng_gas_train_master",    "Gas Train Models",               "manual / init_db"),
        ("agr_master",             "AGR Models",                     "manual / init_db"),
        ("rotary_joint_master",    "Rotary Joint Models",            "manual / init_db"),
        ("blower_pricelist_master","Blower Pricelist (raw)",         "Blower sheet"),
        ("recuperator_master",     "Recuperator Models",             "Recuperator sheet"),
    ]

    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()

        c.execute("SELECT table_name, updated_at FROM table_update_log")
        log = {r[0]: r[1] for r in c.fetchall()}

        tables = []
        for table, label, source in KEY_TABLES:
            try:
                c.execute(f"SELECT COUNT(*) FROM [{table}]")
                count = c.fetchone()[0]
            except Exception:
                count = 0
            updated_at = log.get(table)
            if updated_at:
                dt = datetime.fromisoformat(updated_at)
                updated_str = dt.strftime("%d %b %Y, %H:%M")
            else:
                updated_str = None
            tables.append({
                "table":      table,
                "label":      label,
                "source":     source,
                "rows":       count,
                "ok":         count > 0,
                "updated_at": updated_str,
            })

        # Recent activity: last 8 table updates
        c.execute("SELECT table_name, updated_at FROM table_update_log ORDER BY updated_at DESC LIMIT 8")
        activity = []
        for tname, upd in c.fetchall():
            try:
                dt = datetime.fromisoformat(upd)
                activity.append({"table": tname, "at": dt.strftime("%d %b %Y, %H:%M")})
            except Exception:
                pass

        conn.close()
        return {"tables": tables, "activity": activity}
    except Exception as e:
        return {"error": str(e), "tables": [], "activity": []}

@app.get("/quote", response_class=HTMLResponse)
def quote_form():
    html_path = os.path.join(BASE_DIR, "quote_form.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())

@app.get("/viewer", response_class=HTMLResponse)
def db_viewer():
    html_path = os.path.join(BASE_DIR, "db_viewer.html")
    with open(html_path, "r", encoding="utf-8") as f:
        content = f.read()
    content = content.replace('const API = "http://127.0.0.1:8000"', 'const API = ""')
    return HTMLResponse(content=content)


@app.get("/api/catalog")
def get_catalog():
    try:
        conn = sqlite3.connect(DB_PATH)
        def q(sql):
            return pd.read_sql(sql, conn).to_dict(orient="records")
        # Price master items (if table exists)
        pm_raw, pm_bo, pm_ep = [], [], []
        try:
            pm_raw = q("SELECT item as id, item as label, price as base_price, unit FROM component_price_master WHERE category='Raw Material' ORDER BY item")
            pm_bo  = q("SELECT item as id, item as label, price as base_price, unit FROM component_price_master WHERE category='Bought Out' ORDER BY item")
            pm_ep  = q("SELECT item as id, item as label, price as base_price, unit FROM component_price_master WHERE category='ENCON Purchase' ORDER BY item")
        except Exception:
            pass

        catalog = {
            "Horizontal Ladle Preheater": q("SELECT DISTINCT model as id, model as label FROM horizontal_master"),
            "Vertical Ladle Preheater": q("SELECT DISTINCT model as id, model as label FROM vertical_master"),
            "Blower": q("SELECT DISTINCT model as id, model as label, section, price_without_motor, price_with_motor, hp FROM blower_pricelist_master WHERE price_without_motor IS NOT NULL ORDER BY section, hp"),
            "HPU": q("SELECT DISTINCT unit_kw as id, CAST(unit_kw AS TEXT) || ' KW - ' || variant as label, unit_kw, variant FROM hpu_master ORDER BY unit_kw"),
            "Burner (Film)": q("SELECT DISTINCT burner_size as id, burner_size as label, price as base_price FROM burner_pricelist_master WHERE component='BURNER ALONE' AND section LIKE '%FILM%' GROUP BY burner_size"),
            "Burner (Dual Fuel)": q("SELECT DISTINCT burner_size as id, burner_size as label, price as base_price FROM burner_pricelist_master WHERE component='BURNER ALONE' AND section LIKE '%DUAL%' GROUP BY burner_size"),
            "Recuperator": q("SELECT DISTINCT model as id, type || ' - ' || model as label FROM recuperator_master"),
            "Rad Heat": q("SELECT item as id, item || ' (' || output_kw || ')' as label, price_with_ss_tubing as base_price FROM rad_heat_master WHERE section='MODEL'"),
            "GAIL Gas Burner": q("SELECT burner_size as id, burner_size as label, burner_set as base_price FROM gail_gas_burner_master WHERE burner_set IS NOT NULL"),
            "Raw Material": pm_raw,
            "Bought Out Item": pm_bo,
            "ENCON Component": pm_ep,
        }
        conn.close()
        return catalog
    except Exception as e:
        return {"error": str(e)}

@app.get("/api/price")
def get_price(product_type: str, model: str, qty: int = 1,
              with_motor: bool = False, variant: str = "Duplex 1"):
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        price = 0
        breakdown = []

        if product_type == "Horizontal Ladle Preheater":
            cursor.execute("SELECT particular, amount FROM horizontal_master WHERE model=? AND amount IS NOT NULL", (model,))
            for particular, amount in cursor.fetchall():
                if particular not in ("COMBUSTION EQUIPMENT:", "S.NO.") and amount:
                    breakdown.append({"item": particular, "amount": float(amount)})
                    price += float(amount)

        elif product_type == "Vertical Ladle Preheater":
            cursor.execute("SELECT particular, amount FROM vertical_master WHERE model=? AND amount IS NOT NULL", (model,))
            for particular, amount in cursor.fetchall():
                if particular not in ("COMBUSTION EQUIPMENT:", "S.NO.") and amount:
                    breakdown.append({"item": particular, "amount": float(amount)})
                    price += float(amount)

        elif product_type == "Blower":
            from bom.blower_pricelist import blower_price
            price = blower_price(conn, model, with_motor=with_motor)   # PERKIN basis
            if price:
                breakdown = [{"item": f"{model} ({'with' if with_motor else 'without'} motor)", "amount": price}]

        elif product_type == "HPU":
            from bom.hpu_pricelist import hpu_material_cost
            price = hpu_material_cost(conn, model, variant)   # pricelist-linked
            if price:
                breakdown = [{"item": f"HPU {model} KW ({variant})", "amount": price}]

        elif "Burner" in product_type:
            section_filter = "FILM" if "Film" in product_type else "DUAL"
            cursor.execute("SELECT price FROM burner_pricelist_master WHERE burner_size=? AND component='BURNER ALONE' AND section LIKE ? LIMIT 1", (model, f"%{section_filter}%"))
            row = cursor.fetchone()
            if row and row[0]:
                price = float(row[0])
                breakdown = [{"item": f"{model} Burner", "amount": price}]

        elif product_type == "Rad Heat":
            cursor.execute("SELECT price_with_ss_tubing FROM rad_heat_master WHERE item=? AND section='MODEL' LIMIT 1", (model,))
            row = cursor.fetchone()
            if row and row[0]:
                price = float(row[0])
                breakdown = [{"item": f"Rad Heat {model}", "amount": price}]

        elif product_type == "GAIL Gas Burner":
            cursor.execute("SELECT burner_set FROM gail_gas_burner_master WHERE burner_size=? LIMIT 1", (model,))
            row = cursor.fetchone()
            if row and row[0]:
                price = float(row[0])
                breakdown = [{"item": f"GAIL Gas Burner {model} Set", "amount": price}]

        elif product_type in ("Raw Material", "Bought Out Item", "ENCON Component"):
            cursor.execute("SELECT price, unit FROM component_price_master WHERE item=? LIMIT 1", (model,))
            row = cursor.fetchone()
            if row and row[0]:
                price = float(row[0])
                unit = row[1] or "nos"
                breakdown = [{"item": f"{model} ({unit})", "amount": price}]

        conn.close()
        return {"unit_price": price, "qty": qty, "total": price * qty, "breakdown": breakdown}
    except Exception as e:
        return {"error": str(e)}


class RateUpdateRequest(BaseModel):
    item: str
    price: float

class CompanyUpdateRequest(BaseModel):
    item: str
    company: str

class ItemUpdateRequest(BaseModel):
    table: str
    rowid: int
    qty: Optional[float] = None
    rate: Optional[float] = None
    price: Optional[float] = None

class VLPHCalcRequest(BaseModel):
    mode: str = "calc"                          # "calc" or "direct"
    Ti: float = 650.0
    Tf: float = 1200.0
    refractory_weight: float = 21500.0
    fuel_cv: float = 8500.0
    time_taken_hr: float = 2.0
    refractory_heat_factor: float = 0.25
    efficiency: float = 0.52
    ladle_tons: float = 10.0
    fuel1_type: str = "ng"
    fuel1_cv: float = 8500.0
    fuel2_type: str = "none"
    fuel2_cv: float = 0.0
    direct_burner_capacity: float = 0.0         # kW (direct mode — per burner for tundish)
    direct_heat_input_kcal_hr: float = 0.0      # kcal/hr (direct mode — by heat; takes precedence over capacity)
    blower_pressure: str = "28"                  # "28" or "40" (WG inches)
    control_mode: str = "automatic"              # "manual" or "automatic"
    auto_control_type: str = "agr"               # "plc", "plc_agr", "pid"
    control_valve_vendor: str = "dembla"         # "dembla" or "cair"
    shutoff_valve_vendor: str = "aira"           # "dembla", "aira", or "cair" — shut off valve vendor
    butterfly_valve_vendor: str = "lt_lever"     # "lt_lever" or "lt_gear" (L&T butterfly variants)
    pressure_gauge_vendor: str = "baumer"        # "baumer" or "hguru"
    hpu_variant: str = "Duplex 1"                # "Simplex" | "Duplex 1" | "Duplex 2" — for oil fuels
    # burner_pressure_wg is derived from blower_pressure (28→24, 40→36)
    pilot_burner: str = "auto"                   # "auto" | "lpg_10" | "nglpg_100" | "cog_100"
    pipeline_weight_kg: float = 1000.0           # Air-gas pipeline weight (700–2000 kg, step 100)
    purging_line: str = "no"                     # "yes" | "no" — nitrogen purging line for MG/COG
    manual_pilot_burner: str = "yes"             # "yes" | "no" — include pilot burner in manual BOM
    pilot_line_fuel: str = "lpg"                 # "lpg" | "ng" — pilot line fuel type (manual mode)
    num_burners: int = 1                         # Number of burners (tundish: splits firing rate, multiplies burner-line items)
    ms_structure_kg_override: float = 0.0        # Tundish: override MS structure weight used for fabrication cost
    ceramic_rolls_override: int = 0              # Tundish: override ceramic-fibre roll count
    hood_type: str = "up_down"                   # "up_down" | "swivel_manual" | "swivel_geared"
    special_auto_ignition: bool = False            # Special Requirements: auto-ignition requested
    special_auto_controls: bool = False            # Special Requirements: auto-controls requested


class QuoteItem(BaseModel):
    product_type: str
    model: str
    description: Optional[str] = None
    qty: int = 1
    with_motor: bool = False
    variant: str = "Duplex 1"
    unit_price: Optional[float] = None   # pre-set price (e.g. from VLPH costing)
    total: Optional[float] = None

class QuoteRequest(BaseModel):
    # Company / customer
    company_name: str
    company_address: Optional[str] = ""
    company_city: Optional[str] = ""
    company_state: Optional[str] = ""
    company_pin: Optional[str] = ""
    company_gstin: Optional[str] = ""
    # Point of contact
    poc_name: Optional[str] = ""
    poc_designation: Optional[str] = ""
    mobile_no: Optional[str] = ""
    email: Optional[str] = ""
    # Enquiry
    project_name: Optional[str] = ""
    subject: Optional[str] = ""
    ref_no: Optional[str] = ""
    your_ref: Optional[str] = ""
    enquiry_ref: Optional[str] = ""
    currency: Optional[str] = "INR"     # "USD" → offer priced in USD
    fx_rate: Optional[float] = 0        # INR → USD, used only when currency == "USD"
    extra_context: dict = {}            # product-specific template vars (e.g. regen: fuel_word, kw, price_in_words)
    marketing_person: Optional[str] = ""
    marketing_phone: Optional[str] = ""
    marketing_email: Optional[str] = ""
    technical_person: Optional[str] = ""
    technical_phone: Optional[str] = ""
    technical_email: Optional[str] = ""
    # Technical data (passed from lastCalc to populate template tech table)
    ladle_tons: Optional[float] = 0
    ladle_dim: Optional[str] = ""
    ladle_drawing_no: Optional[str] = ""
    refractory_weight_kg: Optional[str] = ""
    heating_schedule: Optional[str] = ""
    fuel_cv: Optional[str] = ""
    fuel_consumption: Optional[str] = ""
    burner_model: Optional[str] = ""
    blower_model: Optional[str] = ""
    blower_size: Optional[str] = ""
    blower_capacity: Optional[str] = ""
    hydraulic_motor_hp: Optional[str] = ""
    max_electrical_load: Optional[str] = ""
    total_in_words: Optional[str] = ""
    heating_time: Optional[str] = ""
    fuel_name: Optional[str] = ""
    burner_capacity_range: Optional[str] = ""
    pumping_unit: Optional[str] = ""
    hood_movement: Optional[str] = ""
    hood_type: Optional[str] = "up_down"   # 'swivel' or 'up_down'
    pilot_gas_type: Optional[str] = "LPG"  # 'LPG' or 'NG'
    ignition_method: Optional[str] = ""
    # Tundish dual-fuel fields
    num_burners: Optional[str] = ""
    fuel2_cv: Optional[str] = ""
    fuel2_consumption: Optional[str] = ""
    max_fuel_consumption1: Optional[str] = ""
    max_fuel_consumption2: Optional[str] = ""
    is_oil: Optional[bool] = False   # True for oil fuels — drives scope-of-supply rendering
    is_dual: Optional[bool] = False  # True for dual-fuel (gas + oil)
    control_mode: Optional[str] = "automatic"  # "manual" or "automatic" — drives scope text
    auto_control_type: Optional[str] = "plc"   # "plc" | "plc_agr" | "pid" — drives PDF scope wording
    control_valve_type: Optional[str] = "pneumatic"  # "pneumatic" | "motorised" — wording in I&C / Temp Control
    special_auto_ignition: Optional[bool] = False  # Special Requirements: drives pilot-burner sections
    special_auto_controls:  Optional[bool] = False  # Special Requirements
    vertical_qty:   Optional[int] = 1   # Annexure I scope-of-supply header — Vertical units count
    horizontal_qty: Optional[int] = 1   # Annexure I scope-of-supply header — Horizontal units count
    transport_amt:  float = 0           # Transport (flat Rs.) — shown as its own price-schedule line
    purging_line:    Optional[str] = "no"   # "yes" | "no" — drives the Nitrogen Purging block in the offer
    hpu_variant:     Optional[str] = "Duplex 1"   # "Simplex" | "Duplex 1" | "Duplex 2" — Pumping Unit type
    burner_kw_value: Optional[str] = ""   # Pre-formatted total kW for the ENCON burner body line
    location:        Optional[str] = ""   # "Goa" | "Vadodara" | "Faridabad" — appended to enquiry ref
    poc_salutation:        Optional[str] = ""   # "Mr." | "Mrs." | "Miss" | "Dr." — prefixed to poc_name in offer
    marketing_salutation:  Optional[str] = ""
    technical_salutation:  Optional[str] = ""
    # Editable Annexure IV — Terms & Conditions
    tnc_prices:               Optional[str] = "EX Bhagola (Ex-Works)"
    tnc_delivery:             Optional[str] = "10 – 12 weeks from our ENCON works, from the date of receipt of advance or drawing approval, whichever is later."
    tnc_gst:                  Optional[str] = "18% extra."
    tnc_hsn_code:             Optional[str] = "84541000"
    tnc_pan_gst:              Optional[str] = "PAN: AAACE0327M  |  GST: 06AAACE0327M1ZV"
    tnc_payment_terms:        Optional[str] = "30% advance with the purchase order\n70% against proforma invoice prior to dispatch"
    tnc_packing_forwarding:   Optional[str] = "4% and 2% respectively."
    tnc_freight:              Optional[str] = "In client's scope."
    tnc_transit_insurance:    Optional[str] = "To be arranged by the client."
    tnc_validity:             Optional[str] = "45 days from the date of our offer."
    tnc_inspection:           Optional[str] = "If required, materials can be inspected at our works before dispatch, at the client's cost and with prior intimation."
    tnc_guarantee:            Optional[str] = "Materials are guaranteed for 18 months from the date of dispatch, or 12 months from commissioning, whichever is earlier — against any manufacturing defects or poor workmanship."
    bom_items: Optional[List[dict]] = []   # [{item, make, media, ref}, ...] for offer scope + MAKE LIST
    # Items & commercial
    items: List[QuoteItem]
    gst_percent: float = 18
    freight: float = 0
    valid_days: int = 30


@app.get("/costing", response_class=HTMLResponse)
def costing_form():
    html_path = os.path.join(BASE_DIR, "vlph_costing.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())

@app.get("/hlph", response_class=HTMLResponse)
def hlph_costing_form():
    html_path = os.path.join(BASE_DIR, "hlph_costing.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())

@app.get("/tundish", response_class=HTMLResponse)
def tundish_costing_form():
    html_path = os.path.join(BASE_DIR, "tundish_costing.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())

@app.get("/tundish-dryer", response_class=HTMLResponse)
def tundish_dryer_costing_form():
    html_path = os.path.join(BASE_DIR, "tundish_dryer_costing.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())

@app.get("/tundish-cooling", response_class=HTMLResponse)
def tundish_cooling_costing_form():
    html_path = os.path.join(BASE_DIR, "tundish_cooling_costing.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())

@app.get("/api/tundish-cooling/options")
def tundish_cooling_options():
    """Cooling-fan choices + rates for the Tundish Cooling module (editable in
    the pricelist under category 'Tundish Cooling')."""
    conn = sqlite3.connect(DB_PATH)
    fans = [{"item": it, "price": float(pr)} for it, pr in conn.execute(
        "SELECT item, price FROM component_price_master "
        "WHERE category='Tundish Cooling' AND item LIKE 'COOLING FAN%' ORDER BY price")]
    def _one(name, default):
        r = conn.execute("SELECT price FROM component_price_master "
                         "WHERE category='Tundish Cooling' AND item=? LIMIT 1", (name,)).fetchone()
        return float(r[0]) if r and r[0] is not None else default
    fr = conn.execute("SELECT price FROM component_price_master WHERE item='FABRICATION RATE' LIMIT 1").fetchone()
    out = {"fans": fans, "fan_qty": 2,
           "ms_rate": float(fr[0]) if fr and fr[0] is not None else 110.0,  # single shared fabrication rate
           "damper_price": _one("DAMPER MANUAL", 50000.0),
           "markup": 1.8}
    conn.close()
    return out

@app.get("/sen-stove", response_class=HTMLResponse)
def sen_stove_costing_form():
    html_path = os.path.join(BASE_DIR, "sen_stove_costing.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())

@app.get("/api/sen-stove/bom")
def sen_stove_bom():
    """Fixed BOM for the SEN Preheating Stove. Each line's unit cost is pulled
    LIVE from the pricelist (component_price_master, category 'SEN Preheating
    Stove') via its price_key, so pricelist edits cascade; falls back to the
    stored basic. Bought-out lines are marked up; ENCON lines added at face."""
    conn = sqlite3.connect(DB_PATH)

    def _resolve(pk, fallback):
        # Every line's cost comes from the main pricelist (no SEN-only category):
        #  '@BALLVALVE:<nb>' -> cheapest Ball Valve for that NB
        #  '@PILOT:<kw>'     -> cheapest <kw> Pilot Burner (COG/LPG/NG)
        #  '@FABRICATION'    -> the single shared FABRICATION RATE
        #  anything else     -> that exact item name (cheapest if it recurs)
        if pk and pk.startswith("@BALLVALVE:"):
            nb = pk.split(":", 1)[1]
            r = conn.execute(
                "SELECT price FROM component_price_master WHERE category='Ball Valve' "
                "AND item LIKE ? ORDER BY price ASC LIMIT 1", (f"BALL VALVE {nb} NB%",)).fetchone()
        elif pk and pk.startswith("@PILOT:"):
            kw = pk.split(":", 1)[1]
            r = conn.execute(
                "SELECT price FROM component_price_master WHERE category='Pilot Burner' "
                "AND item LIKE ? ORDER BY price ASC LIMIT 1", (f"%{kw} KW%",)).fetchone()
        elif pk == "@FABRICATION":
            r = conn.execute("SELECT price FROM component_price_master WHERE item='FABRICATION RATE' LIMIT 1").fetchone()
        elif pk:
            r = conn.execute("SELECT price FROM component_price_master WHERE item=? ORDER BY price ASC LIMIT 1", (pk,)).fetchone()
        else:
            r = None
        return float(r[0]) if r and r[0] is not None else float(fallback or 0)

    rows = []
    for s, m, it, rf, q, u, mk, b, pk in conn.execute(
            "SELECT section, media, item, ref, qty, unit, make, basic, price_key "
            "FROM sen_stove_bom ORDER BY sno"):
        basic = _resolve(pk, b)      # live pricelist cost
        rows.append({"section": s, "media": m, "item": it, "ref": rf,
                     "qty": q, "unit": u, "make": mk, "basic": basic,
                     "total": float(q or 0) * basic})
    conn.close()
    return {"rows": rows, "markup": 1.8}

@app.get("/equipment-offer", response_class=HTMLResponse)
def equipment_offer_hub():
    """Hub page that lets the user pick which stand-alone equipment
    offer to generate (HPU, Blower or Burner)."""
    html_path = os.path.join(BASE_DIR, "equipment_offer.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())


@app.get("/hpu", response_class=HTMLResponse)
def hpu_picker():
    """Sub-hub under Equipment Offer that lets the user pick between
    HPU (Heating + Pumping Unit) and PU (Pumping Unit only)."""
    html_path = os.path.join(BASE_DIR, "hpu_picker.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())


@app.get("/hpu/heating-pumping", response_class=HTMLResponse)
def hpu_costing_form():
    """HPU offer form — pumping skid with in-built electric heater
    (model codes HPS / HPD / HPDD)."""
    html_path = os.path.join(BASE_DIR, "hpu_costing.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())


@app.get("/hpu/pumping-only", response_class=HTMLResponse)
def pu_costing_form():
    """PU offer form — pumping skid without heater, for pre-heated
    oils such as LDO / LSHS (model codes PUS / PUD / PUDD)."""
    html_path = os.path.join(BASE_DIR, "pu_costing.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())

@app.get("/regen", response_class=HTMLResponse)
def regen_costing_form():
    html_path = os.path.join(BASE_DIR, "regen_costing.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())

@app.get("/price-master", response_class=HTMLResponse)
def price_master_page():
    html_path = os.path.join(BASE_DIR, "price_master.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())


@app.get("/pricelist", response_class=HTMLResponse)
def pricelist_viewer_page():
    html_path = os.path.join(BASE_DIR, "pricelist_viewer.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())


@app.get("/api/pricelist-summary")
def pricelist_summary():
    """Return all live-computed prices from every master table."""
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()

        def q(sql, *args):
            return c.execute(sql, args).fetchall()

        def _parts_sections(table, markup=1.25):
            """Convert section-header / item rows into [{title, items, fabrication_cost, selling_price}]."""
            sections, cur_title, cur_items, cur_total = [], None, [], 0.0
            for rowid, _, part, qty, unit, rate, amt in conn.execute(
                f"SELECT rowid, section, particular, qty, unit, rate, amount FROM {table} ORDER BY rowid"
            ):
                if amt is None:
                    if cur_title is not None:
                        sections.append({"title": cur_title, "items": cur_items,
                                         "fabrication_cost": round(cur_total, 2),
                                         "selling_price": round(cur_total * markup, 2)})
                    cur_title, cur_items, cur_total = part, [], 0.0
                else:
                    cur_items.append({"rowid": rowid, "particular": part, "qty": qty, "unit": unit,
                                      "rate": rate, "amount": amt})
                    cur_total += amt or 0
            if cur_title is not None:
                sections.append({"title": cur_title, "items": cur_items,
                                 "fabrication_cost": round(cur_total, 2),
                                 "selling_price": round(cur_total * markup, 2)})
            return sections

        # ── HPU ───────────────────────────────────────────────────────────
        _variant_order = "CASE variant WHEN 'Duplex 1' THEN 1 WHEN 'Duplex 2' THEN 2 WHEN 'Simplex' THEN 3 ELSE 4 END"
        hpu_kws = [r[0] for r in q("SELECT DISTINCT unit_kw FROM hpu_master ORDER BY unit_kw")]
        hpu = []
        for kw in hpu_kws:
            variants_data = []
            for (variant,) in q(f"SELECT DISTINCT variant FROM hpu_master WHERE unit_kw=? ORDER BY {_variant_order}", kw):
                rows = q("SELECT rowid, item, qty, unit, rate, amount FROM hpu_master WHERE unit_kw=? AND variant=? ORDER BY rowid", kw, variant)
                items = [{"rowid": r[0], "item": r[1], "qty": r[2], "unit": r[3], "rate": r[4], "amount": r[5]} for r in rows]
                mat = sum((r[5] or 0) for r in rows)
                variants_data.append({"name": variant, "material_cost": round(mat, 2),
                                      "selling_price": round(mat * 1.8, 2), "items": items})
            hpu.append({"kw": kw, "variants": variants_data})

        # ── Burners ───────────────────────────────────────────────────────
        burner_rows = q("SELECT section, burner_size, component, price FROM burner_pricelist_master ORDER BY section, burner_size")
        burners = {}
        for sec, size, comp, price in burner_rows:
            burners.setdefault(sec, {}).setdefault(size, {})[comp] = price

        # ── Blowers ───────────────────────────────────────────────────────
        blower_cols = [d[0] for d in c.execute("SELECT * FROM blower_pricelist_master LIMIT 0").description]
        blowers = [dict(zip(blower_cols, r)) for r in q("SELECT * FROM blower_pricelist_master ORDER BY section, CAST(hp AS REAL), model")]
        try:
            dm_cols = [d[0] for d in c.execute("SELECT * FROM blower_dm_idm_master LIMIT 0").description]
            blower_dm_idm = [dict(zip(dm_cols, r)) for r in q("SELECT * FROM blower_dm_idm_master ORDER BY section, CAST(SUBSTR(model,1,INSTR(model,'/')-1) AS REAL)")]
        except Exception:
            blower_dm_idm = []

        # ── Recuperator ───────────────────────────────────────────────────
        try:
            # Use cursor description (avoids variable-name clash with outer cursor 'c')
            c.execute("SELECT * FROM recuperator_master ORDER BY type, CAST(model AS REAL)")
            rcols = [d[0] for d in c.description]
            recup_rows = [dict(zip(rcols, r)) for r in c.fetchall()]
            # Fetch component rates used in recuperator calculations
            recup_rate_items = {
                'tube_c':     "M.S. Tube \"C\" Class",
                'tube_b':     "M.S. Tube \"B\" Class",
                'ci_gills':   "C.I. Gills",
                'ang_6550':   "M.S. Angle 65,50",
                'ang_100100': "M.S. Angle 100,100",
                'plate_16_5': "M.S. Plate 16mm*",
                'bolt':       "Hardware Bolt",
                'channel':    "M.S. Channel",
                'plate_16_10':"M.S. Plate 16mm*10mm",
                'plate_5':    "M.S. Plate 5mm",
                'ang_50':     "M.S. Angle 50*6",
                'ss_sheet':   "S.S. Sheet 3mm",
            }
            recup_rates = {}
            for key, prefix in recup_rate_items.items():
                row = conn.execute(
                    "SELECT price FROM component_price_master WHERE item LIKE ? LIMIT 1",
                    (prefix + '%',)
                ).fetchone()
                recup_rates[key] = float(row[0]) if row else None
            recup = {"rows": recup_rows, "rates": recup_rates}
        except Exception:
            recup = {"rows": [], "rates": {}}

        # ── Rad Heat ──────────────────────────────────────────────────────
        def _rad_rows(table):
            models, spares = [], []
            for r in q(f"SELECT * FROM {table} ORDER BY section, item"):
                d = dict(zip([x[0] for x in conn.execute(f"SELECT * FROM {table} LIMIT 0").description], r))
                if d["section"] == "MODEL":
                    models.append(d)
                else:
                    spares.append({"item": d["item"], "price": d["price_with_ss_tubing"]})
            return {"models": models, "spares": spares}
        rad      = _rad_rows("rad_heat_master")
        rad_tata = _rad_rows("rad_heat_tata_master")

        # ── GAIL Gas Burner ───────────────────────────────────────────────
        gail_cols = [d[0] for d in c.execute("SELECT * FROM gail_gas_burner_master LIMIT 0").description]
        gail = [dict(zip(gail_cols, r)) for r in q("SELECT * FROM gail_gas_burner_master ORDER BY section, burner_size")]

        # ── Rotary Joint ──────────────────────────────────────────────────
        rj_cols = ["rowid"] + [d[0] for d in c.execute("SELECT * FROM rotary_joint_master LIMIT 0").description]
        rotary_joint = [
            dict(zip(rj_cols, r))
            for r in q("SELECT rowid, * FROM rotary_joint_master ORDER BY nb")
        ]

        # ── Burner Parts (Oil / HV Oil / Gas) ─────────────────────────────
        oil_parts = _parts_sections("oil_burner_parts_master",    markup=1.25)
        hv_parts  = _parts_sections("hv_oil_burner_parts_master", markup=1.25)
        gas_parts = _parts_sections("gas_burner_parts_master",    markup=1.25)

        # ── Horizontal LPH ────────────────────────────────────────────────
        hlph = []
        for (model,) in q("SELECT DISTINCT model FROM horizontal_master ORDER BY model"):
            rows = q("SELECT particular, qty, amount FROM horizontal_master WHERE model=?", model)
            items = [{"particular": r[0], "qty": r[1], "amount": r[2]} for r in rows if r[2]]
            if items:
                hlph.append({"model": model, "total": round(sum(i["amount"] for i in items), 2), "items": items})

        # ── Vertical LPH ─────────────────────────────────────────────────
        vlph = []
        for (model,) in q("SELECT DISTINCT model FROM vertical_master ORDER BY model"):
            rows = q("SELECT particular, qty, amount FROM vertical_master WHERE model=?", model)
            items = [{"particular": r[0], "qty": r[1], "amount": r[2]} for r in rows if r[2]]
            if items:
                vlph.append({"model": model, "total": round(sum(i["amount"] for i in items), 2), "items": items})

        # ── Regen Costing ──────────────────────────────────────────────────
        try:
            rci_cols = [d[0] for d in conn.execute("SELECT * FROM regen_costing_items LIMIT 0").description]
            rci_rows = [dict(zip(rci_cols, r)) for r in conn.execute(
                "SELECT * FROM regen_costing_items ORDER BY kw, row_order"
            ).fetchall()]
            rsz_cols = [d[0] for d in conn.execute("SELECT * FROM regen_sizing LIMIT 0").description]
            rsz_rows = [dict(zip(rsz_cols, r)) for r in conn.execute(
                "SELECT * FROM regen_sizing ORDER BY kw"
            ).fetchall()]
            rpl_cols = [d[0] for d in conn.execute("SELECT * FROM regen_pricelist LIMIT 0").description]
            rpl_rows = [dict(zip(rpl_cols, r)) for r in conn.execute(
                "SELECT * FROM regen_pricelist ORDER BY kw"
            ).fetchall()]
            rmr_cols = [d[0] for d in conn.execute("SELECT * FROM regen_material_rates LIMIT 0").description]
            rmr_rows = [dict(zip(rmr_cols, r)) for r in conn.execute(
                "SELECT * FROM regen_material_rates"
            ).fetchall()]
            rnz_cols = [d[0] for d in conn.execute("SELECT * FROM regen_nozzle_sizing LIMIT 0").description]
            rnz_rows = [dict(zip(rnz_cols, r)) for r in conn.execute(
                "SELECT * FROM regen_nozzle_sizing ORDER BY power_kw"
            ).fetchall()]
            rps_cols = [d[0] for d in conn.execute("SELECT * FROM regen_pipe_sizes LIMIT 0").description]
            rps_rows = [dict(zip(rps_cols, r)) for r in conn.execute(
                "SELECT * FROM regen_pipe_sizes ORDER BY gas_type, burner_size_kw"
            ).fetchall()]
            regen_costing = {
                "items": rci_rows, "sizing": rsz_rows,
                "pricelist": rpl_rows, "material_rates": rmr_rows,
                "nozzle_sizing": rnz_rows, "pipe_sizes": rps_rows,
            }
        except Exception:
            regen_costing = {"items": [], "sizing": [], "pricelist": [], "material_rates": [], "nozzle_sizing": [], "pipe_sizes": []}

        conn.close()
        return {
            "hpu": hpu,
            "burners": burners,
            "blowers": blowers,
            "blower_dm_idm": blower_dm_idm,
            "recuperator": recup,
            "rad_heat": rad,
            "rad_heat_tata": rad_tata,
            "gail_gas_burner": gail,
            "oil_burner_parts": oil_parts,
            "hv_oil_burner_parts": hv_parts,
            "gas_burner_parts": gas_parts,
            "horizontal_lph": hlph,
            "vertical_lph": vlph,
            "rotary_joint": rotary_joint,
            "regen_costing": regen_costing,
        }
    except Exception as e:
        import traceback
        return {"error": str(e), "detail": traceback.format_exc()}


@app.get("/api/pricelist/rates")
def get_pricelist_rates():
    """Return component_price_master for editing."""
    try:
        conn = sqlite3.connect(DB_PATH)
        # FUEL DENSITY rows live in this table only because burner sizing
        # uses _get_fuel_density() to look them up by name. They are not
        # editable rates, so exclude them from the Rates view.
        rows = conn.execute(
            "SELECT rowid, item, category, price, previous_price, updated_at, company, specification "
            "FROM component_price_master "
            "WHERE item NOT LIKE 'FUEL DENSITY%' "
            "  AND (category IS NULL OR category != 'Fuel Density') "
            "ORDER BY category, price, item"
        ).fetchall()
        conn.close()
        return [{"rowid": r[0], "item": r[1], "category": r[2],
                 "price": r[3], "previous_price": r[4], "updated_at": r[5],
                 "company": r[6], "specification": r[7]} for r in rows]
    except Exception as e:
        return {"error": str(e)}


@app.get("/api/fabrication-mapping")
def fabrication_mapping():
    """Ladle fabrication / pipeline / ceramic weights per capacity & mechanism,
    for the pricelist 'Fabrication' tab. Ceramic rolls = ceil(ceramic_kg / 14)."""
    import math
    try:
        conn = sqlite3.connect(DB_PATH)
        rows = conn.execute(
            "SELECT ladle_capacity_ton, preheater_type, hood_type, "
            "fabrication_kg, pipeline_kg, ceramic_kg "
            "FROM fabrication_ladle_mapping ORDER BY ladle_capacity_ton"
        ).fetchall()
        conn.close()
        return [{
            "ton": t, "preheater": pt, "hood": ht or "",
            "fabrication_kg": fab, "pipeline_kg": pipe, "ceramic_kg": cer,
            "ceramic_rolls": int(math.ceil(cer / 14)) if cer else 0,
        } for t, pt, ht, fab, pipe, cer in rows]
    except Exception as e:
        return {"error": str(e)}


@app.put("/api/pricelist/rate")
def update_pricelist_rate(req: RateUpdateRequest):
    """Update a rate in component_price_master and cascade-recalculate all tables."""
    import re as _re

    def _norm(s):
        s = str(s or "").upper()
        s = _re.sub(r"[\s\-_\(\)\*\.\,\"\']+", " ", s).strip()
        s = _re.sub(r"\s+", " ", s)  # collapse multiple spaces
        # Also create a no-space version for comparison
        return s

    def _norm_compact(s):
        """Extra compact normalization — removes ALL spaces and dots."""
        s = str(s or "").upper()
        s = _re.sub(r"[^A-Z0-9]", "", s)
        return s

    try:
        conn = sqlite3.connect(DB_PATH)

        # Get old price before update
        old_row = conn.execute(
            "SELECT price FROM component_price_master WHERE item=?", (req.item,)
        ).fetchone()
        old_price = old_row[0] if old_row else None

        # Update master table
        conn.execute(
            "UPDATE component_price_master SET previous_price=price, price=?, updated_at=? WHERE item=?",
            (req.price, datetime.now().strftime("%Y-%m-%d %H:%M"), req.item)
        )
        conn.commit()

        # ── Direct cascade to all parts tables ──────────────────────────────
        # Update rows where name fuzzy-matches AND rate = old_price
        PARTS_TABLES = [
            ("oil_burner_parts_master",    "particular"),
            ("hv_oil_burner_parts_master", "particular"),
            ("gas_burner_parts_master",    "particular"),
            ("hpu_master",                 "item"),
        ]
        norm_item = _norm(req.item)
        compact_item = _norm_compact(req.item)
        cascade_counts = {}

        # ── Method 1: Use formula mapping table (exact legacy links) ──
        mapped_targets = conn.execute(
            "SELECT target_table, target_item FROM rate_cascade_map WHERE rates_item = ?",
            (req.item,)
        ).fetchall()

        for target_table, target_item in mapped_targets:
            name_col = "particular" if "parts" in target_table else "item"
            rows = conn.execute(
                f"SELECT rowid, {name_col}, qty FROM {target_table} WHERE {name_col} = ?",
                (target_item,)
            ).fetchall()
            updated = 0
            for rowid, part_name, qty in rows:
                new_amount = round(float(qty or 0) * req.price, 2) if qty is not None else None
                conn.execute(
                    f"UPDATE {target_table} SET rate=?, amount=? WHERE rowid=?",
                    (req.price, new_amount, rowid)
                )
                updated += 1
            if updated:
                cascade_counts[target_table] = cascade_counts.get(target_table, 0) + updated

        # ── Method 2: Fuzzy name match (for items not in mapping table) ──
        for table, name_col in PARTS_TABLES:
            rows = conn.execute(
                f"SELECT rowid, {name_col}, qty, rate FROM {table} WHERE rate IS NOT NULL"
            ).fetchall()
            updated = 0
            for rowid, part_name, qty, rate in rows:
                norm_part = _norm(part_name or "")
                compact_part = _norm_compact(part_name or "")
                if norm_part != norm_item and compact_part != compact_item:
                    continue
                new_amount = round(float(qty or 0) * req.price, 2) if qty is not None else None
                conn.execute(
                    f"UPDATE {table} SET rate=?, amount=? WHERE rowid=?",
                    (req.price, new_amount, rowid)
                )
                updated += 1
            if updated:
                cascade_counts[table] = cascade_counts.get(table, 0) + updated

        conn.commit()

        # also cascade into the internal-costing oil_burner_master (rate-master
        # linked to the Pricelist) so the burner tab reflects the new rate.
        try:
            sync_cpm_rates(conn)
            conn.commit()
        except Exception as _e:
            print(f"WARN: burner sync after pricelist edit failed: {_e}")

        conn.close()
        return {
            "success": True,
            "cascaded": True,
            "direct_cascade": cascade_counts,
        }
    except Exception as e:
        import traceback
        return {"success": False, "error": str(e), "detail": traceback.format_exc()}


@app.put("/api/pricelist/company")
def update_pricelist_company(req: CompanyUpdateRequest):
    """Update the company/supplier name for an item in component_price_master."""
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.execute(
            "UPDATE component_price_master SET company=? WHERE item=?",
            (req.company.strip() or None, req.item)
        )
        conn.commit()
        conn.close()
        return {"success": True}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.put("/api/pricelist/item")
def update_pricelist_item(req: ItemUpdateRequest):
    """Update qty and/or rate for a row in any master table; recalculate amount."""
    if req.table not in ALLOWED_EDIT_TABLES:
        return {"success": False, "error": f"Table '{req.table}' not editable"}
    try:
        conn = sqlite3.connect(DB_PATH)
        cols = {r[1] for r in conn.execute(f"PRAGMA table_info({req.table})").fetchall()}
        if req.qty is not None and 'qty' in cols:
            conn.execute(f"UPDATE {req.table} SET qty=? WHERE rowid=?", (req.qty, req.rowid))
        if req.rate is not None and 'rate' in cols:
            conn.execute(f"UPDATE {req.table} SET rate=? WHERE rowid=?", (req.rate, req.rowid))
        if req.price is not None and 'price' in cols:
            conn.execute(f"UPDATE {req.table} SET price=? WHERE rowid=?", (req.price, req.rowid))
        # Recalculate amount if qty/rate/amount all exist
        if {'qty', 'rate', 'amount'} <= cols:
            row = conn.execute(f"SELECT qty, rate FROM {req.table} WHERE rowid=?", (req.rowid,)).fetchone()
            if row and row[0] is not None and row[1] is not None:
                amount = round(float(row[0]) * float(row[1]), 2)
                conn.execute(f"UPDATE {req.table} SET amount=? WHERE rowid=?", (amount, req.rowid))
        conn.commit()
        select_cols = [c for c in ('qty', 'rate', 'amount', 'price') if c in cols]
        row2 = conn.execute(
            f"SELECT {','.join(select_cols)} FROM {req.table} WHERE rowid=?", (req.rowid,)
        ).fetchone() if select_cols else None
        conn.close()
        result = {"success": True}
        if row2 and select_cols:
            result.update(dict(zip(select_cols, row2)))
        return result
    except Exception as e:
        return {"success": False, "error": str(e)}


# ── Vendor / component price tables + Recup rate constants ──────────────
# Surfaces vendor pricelists (Cair / Dembla / L&T / regulator / orifice /
# hose / gas-train / etc.) and the standalone recup_rates table so they can
# be inspected and edited from /pricelist without raw SQL.

VENDOR_TABLES: dict[str, dict] = {
    # name -> { label, key, numeric_cols, order_by }
    # `order_by` is a SQL fragment (no semicolons) used to sort the rows
    # in ascending order. Falls back to the natural key when not provided.
    "aira_valve_master":            {"label": "AIRA",             "key": ["id"], "numeric_cols": ["list_price", "discount_pct", "net_price"], "order_by": "valve_type, nb"},
    "cair_motorized_valve_master":  {"label": "Cair MOV",         "key": ["id"], "numeric_cols": ["list_price", "discount_pct", "net_price"], "order_by": "valve_type, nb"},
    "dembla_valve_master":          {"label": "Dembla CV",        "key": ["id"], "numeric_cols": ["list_price", "discount_pct", "net_price"], "order_by": "valve_type, nb"},
    "lt_ball_valve_master":         {"label": "L&T Ball Valve",   "key": ["cat_no"], "numeric_cols": ["nb_15","nb_20","nb_25","nb_32","nb_40","nb_50","nb_65","nb_80","nb_100","nb_125","nb_150","nb_200","nb_250","nb_300","nb_350","nb_400","nb_450","nb_500","nb_600"], "order_by": "cat_no"},
    "lt_butterfly_valve_master":    {"label": "L&T Butterfly",    "key": ["id"], "numeric_cols": ["price"], "order_by": "operation, nb"},
    "gas_regulator_master":         {"label": "Gas Regulator",    "key": ["id"], "numeric_cols": ["list_price"], "order_by": "category, nb"},
    "orifice_plate_master":         {"label": "Orifice Plate",    "key": ["nb"], "numeric_cols": ["flanges_price", "plate_price", "fasteners_price", "total_price"], "order_by": "nb"},
    "compensator_master":           {"label": "Compensator",      "key": ["nb"], "numeric_cols": ["price"], "order_by": "nb"},
    "flexible_hose_master":         {"label": "Flexible Hose",    "key": ["id"], "numeric_cols": ["price"], "order_by": "dn, length_mm"},
    "gas_train_master":             {"label": "Gas Train",        "key": ["sr_no"], "numeric_cols": ["price_inr"], "order_by": "sr_no"},
    "motorized_valve_master":       {"label": "Motorized Valve",  "key": ["nb"], "numeric_cols": ["price"], "order_by": "nb"},
    "solenoidvalve_component_master": {"label": "Solenoid Valve", "key": ["id"], "numeric_cols": ["list_price"], "order_by": "section, size, description"},
}


class VendorCellUpdate(BaseModel):
    table: str
    key_values: dict           # { key_col: value, ... }  identifies the row
    column: str                # the numeric column being edited
    new_value: float


class RecupRateUpdate(BaseModel):
    key: str
    value: float


@app.get("/api/pricelist/vendor-tables")
def get_vendor_tables():
    """Return all vendor / component price tables as
    { name: { label, columns, key, numeric_cols, rows } }."""
    try:
        conn = sqlite3.connect(DB_PATH)
        out = {}
        for table, meta in VENDOR_TABLES.items():
            try:
                cols = [r[1] for r in conn.execute(f'PRAGMA table_info("{table}")').fetchall()]
                if not cols:
                    out[table] = {"label": meta["label"], "error": "table missing"}
                    continue
                order_clause = f' ORDER BY {meta["order_by"]}' if meta.get("order_by") else ''
                rows = conn.execute(f'SELECT * FROM "{table}"{order_clause}').fetchall()
                out[table] = {
                    "label":        meta["label"],
                    "columns":      cols,
                    "key":          meta["key"],
                    "numeric_cols": meta["numeric_cols"],
                    "rows":         [dict(zip(cols, r)) for r in rows],
                }
            except Exception as inner:
                out[table] = {"label": meta["label"], "error": str(inner)}
        conn.close()
        return out
    except Exception as e:
        return {"error": str(e)}


@app.put("/api/pricelist/vendor-cell")
def update_vendor_cell(req: VendorCellUpdate):
    """Update one numeric cell in a whitelisted vendor table."""
    if req.table not in VENDOR_TABLES:
        return {"success": False, "error": f"Table '{req.table}' not editable here"}
    meta = VENDOR_TABLES[req.table]
    if req.column not in meta["numeric_cols"]:
        return {"success": False, "error": f"Column '{req.column}' not editable"}
    key_cols = meta["key"]
    if set(req.key_values.keys()) != set(key_cols):
        return {"success": False, "error": f"Expected key cols {key_cols}, got {list(req.key_values.keys())}"}
    try:
        conn = sqlite3.connect(DB_PATH)
        where = " AND ".join(f'"{k}"=?' for k in key_cols)
        params = [req.new_value] + [req.key_values[k] for k in key_cols]
        cur = conn.execute(
            f'UPDATE "{req.table}" SET "{req.column}"=? WHERE {where}', params
        )
        # For Cair / Dembla: when list_price or discount_pct changes, recompute net_price.
        if req.column in ("list_price", "discount_pct") and req.table in ("cair_motorized_valve_master", "dembla_valve_master"):
            row = conn.execute(
                f'SELECT list_price, discount_pct FROM "{req.table}" WHERE {where}',
                [req.key_values[k] for k in key_cols]
            ).fetchone()
            if row and row[0] is not None and row[1] is not None:
                net = round(float(row[0]) * (1 - float(row[1]) / 100.0), 2)
                conn.execute(
                    f'UPDATE "{req.table}" SET net_price=? WHERE {where}',
                    [net] + [req.key_values[k] for k in key_cols]
                )
        conn.commit()
        # Read back the affected row so the client can refresh derived cells.
        all_cols = [r[1] for r in conn.execute(f'PRAGMA table_info("{req.table}")').fetchall()]
        row = conn.execute(
            f'SELECT * FROM "{req.table}" WHERE {where}',
            [req.key_values[k] for k in key_cols]
        ).fetchone()
        conn.close()
        return {
            "success":      True,
            "rows_changed": cur.rowcount,
            "row":          dict(zip(all_cols, row)) if row else None,
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/api/pricelist/recup-rates")
def get_recup_rates():
    """Return every key/value in recup_rates (with unit + notes)."""
    try:
        conn = sqlite3.connect(DB_PATH)
        cols = [r[1] for r in conn.execute('PRAGMA table_info("recup_rates")').fetchall()]
        rows = conn.execute('SELECT * FROM recup_rates ORDER BY key').fetchall()
        conn.close()
        return [dict(zip(cols, r)) for r in rows]
    except Exception as e:
        return {"error": str(e)}


@app.put("/api/pricelist/recup-rate")
def update_recup_rate(req: RecupRateUpdate):
    """Update a single key in recup_rates."""
    try:
        conn = sqlite3.connect(DB_PATH)
        cur = conn.execute(
            "UPDATE recup_rates SET value=? WHERE key=?", (req.value, req.key)
        )
        if cur.rowcount == 0:
            conn.close()
            return {"success": False, "error": f"Key '{req.key}' not found"}
        conn.commit()
        conn.close()
        return {"success": True, "key": req.key, "value": req.value}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/api/ladle-mapping")
def ladle_mapping(tons: float, type: str = "vertical", hood_type: str = ""):
    """Return auto-filled fabrication/pipeline/ceramic values for a ladle
    capacity. The UI calls this to populate read-only form fields.
    hood_type is used only for vertical ('swivel' or 'up_down'); defaults
    to 'swivel' in the lookup if empty."""
    from bom.vlph_builder import lookup_ladle_fab_pipeline
    return lookup_ladle_fab_pipeline(tons, type, hood_type or None) or {}


@app.post("/api/vlph-calculate")
def vlph_calculate(req: VLPHCalcRequest):
    try:
        from calculations.burner import BurnerInputs, calculate_burner
        from calculations.pipes import PipeInputs, calculate_pipe_sizes, select_oil_pipe_nb
        from bom.selectors.selection_engine import select_equipment
        from bom.vlph_builder import build_vlph_120t_df, build_vlph_manual_df

        FUEL_NAMES = {
            "ng": "Natural Gas", "lpg": "LPG", "cog": "COG", "bg": "BFG", "rlng": "RLNG", "mg": "Mixed Gas",
            "hsd": "Diesel (HSD)", "ldo": "LDO", "hdo": "HDO", "fo": "Furnace Oil",
            "sko": "Kerosene (SKO)", "cfo": "CFO", "lshs": "LSHS",
        }
        OIL_FUELS = {"hsd", "ldo", "hdo", "fo", "sko", "cfo", "lshs"}

        f1_cv = req.fuel1_cv if req.fuel1_cv > 0 else req.fuel_cv

        # Circulating Ladle Preheater: a hot ladle (initial temperature above the
        # threshold) is already in circulation, so it only needs topping up — the
        # effective refractory mass to heat is HALF the input weight. Detected in
        # calc mode only (direct mode enters burner capacity, not weight).
        CIRCULATING_TI_THRESHOLD = 600
        is_circulating = (req.mode == "calc") and (req.Ti > CIRCULATING_TI_THRESHOLD)
        effective_refractory_weight = (
            req.refractory_weight / 2 if is_circulating else req.refractory_weight
        )

        if req.mode == "direct":
            # --- Direct mode: user enters burner capacity (kW) or heat input (kcal/hr) ---
            # For tundish: capacity is PER BURNER; total = capacity × num_burners.
            # Heat input takes precedence (already total).
            is_dual = req.fuel2_type != "none" and req.fuel2_cv > 0
            n_burners = max(1, int(req.num_burners or 1))
            if req.direct_heat_input_kcal_hr and req.direct_heat_input_kcal_hr > 0:
                heat_kcal_hr = req.direct_heat_input_kcal_hr
            else:
                # kW → kcal/hr (× 860), then × num_burners for total
                heat_kcal_hr = req.direct_burner_capacity * 860 * n_burners
            # Each fuel's flow = heat / its CV
            ng_flow = heat_kcal_hr / f1_cv
            air_flow = heat_kcal_hr * 118 / 100000
            br1 = None
        else:
            # --- Calc mode: calculate from process params ---
            br1 = calculate_burner(BurnerInputs(
                Ti=req.Ti, Tf=req.Tf,
                refractory_weight=effective_refractory_weight,
                fuel_cv=f1_cv,
                time_taken_hr=req.time_taken_hr,
                refractory_heat_factor=req.refractory_heat_factor,
                efficiency=req.efficiency,
            ))
            ng_flow = br1.extra_firing_rate_nm3hr
            air_flow = br1.air_qty_nm3hr

        # --- Fuel 2 calculation (if dual fuel) ---
        is_dual = req.fuel2_type != "none" and req.fuel2_cv > 0

        # Burner pressure derived from blower pressure: 28" -> 24" w.g., 40" -> 36" w.g.
        burner_pressure_wg = 36 if req.blower_pressure == "40" else 24

        # For tundish with multiple burners: blower/gas-train/main-line are sized for
        # TOTAL flow, burner is sized for PER-BURNER flow.
        n_burners = max(1, int(req.num_burners or 1))
        per_burner_ng  = ng_flow / n_burners
        per_burner_air = air_flow / n_burners

        # For dual fuel: pre-compute fuel 2 air flow so blower is sized for
        # whichever fuel needs MORE air (lower CV = higher flow).
        blower_air_flow = air_flow
        if is_dual:
            if req.mode == "direct":
                air_flow2_pre = air_flow  # same heat → same air
            else:
                br2_pre = calculate_burner(BurnerInputs(
                    Ti=req.Ti, Tf=req.Tf,
                    refractory_weight=effective_refractory_weight,
                    fuel_cv=req.fuel2_cv,
                    time_taken_hr=req.time_taken_hr,
                    refractory_heat_factor=req.refractory_heat_factor,
                    efficiency=req.efficiency,
                ))
                air_flow2_pre = br2_pre.air_qty_nm3hr
            blower_air_flow = max(air_flow, air_flow2_pre)

        # Pre-compute fuel2 oil LPH for blower CFM sizing in dual-fuel
        f2_oil_lph_for_blower = 0
        if is_dual and req.fuel2_type in OIL_FUELS:
            # heat_kcal_hr exists in direct mode; in calc mode derive from fuel1 flow × CV
            heat_for_f2 = heat_kcal_hr if req.mode == "direct" else (ng_flow * f1_cv)
            f2_flow_kghr = heat_for_f2 / req.fuel2_cv if req.fuel2_cv > 0 else 0
            from bom.selectors.encon_burner import _get_fuel_density
            f2_density = _get_fuel_density(req.fuel2_type)
            f2_oil_lph_for_blower = f2_flow_kghr / f2_density if f2_density else 0

        equip1 = select_equipment(
            ng_flow_nm3hr=ng_flow,
            air_flow_nm3hr=blower_air_flow,
            is_dual_fuel=is_dual,
            fuel_cv=f1_cv,
            blower_pressure=req.blower_pressure,
            fuel_type=req.fuel1_type,
            hpu_variant=req.hpu_variant,
            burner_pressure_wg=burner_pressure_wg,
            butterfly_valve_vendor=req.butterfly_valve_vendor,
            shutoff_valve_vendor=req.shutoff_valve_vendor,
            control_mode=req.control_mode,
            auto_control_type=req.auto_control_type,
            fuel2_lph=f2_oil_lph_for_blower,
        )

        # Size pipes AFTER blower selection. Air pipe sized from the airflow the
        # blower actually moves — cfm × 1.7 (max of gas combustion air and oil
        # atomisation air) — matching the air-line NB from select_equipment.
        pipes1 = calculate_pipe_sizes(PipeInputs(
            ng_flow_nm3hr=ng_flow,
            air_flow_nm3hr=equip1.get("air_line_flow") or air_flow,
        ))

        # Tundish multi-burner: re-select burner for per-burner flow
        if n_burners > 1:
            equip_pb = select_equipment(
                ng_flow_nm3hr=per_burner_ng,
                air_flow_nm3hr=per_burner_air,
                is_dual_fuel=is_dual,
                fuel_cv=f1_cv,
                blower_pressure=req.blower_pressure,
                fuel_type=req.fuel1_type,
                hpu_variant=req.hpu_variant,
                burner_pressure_wg=burner_pressure_wg,
                butterfly_valve_vendor=req.butterfly_valve_vendor,
                shutoff_valve_vendor=req.shutoff_valve_vendor,
                control_mode=req.control_mode,
                auto_control_type=req.auto_control_type,
            )
            equip1["burner"] = equip_pb["burner"]

        f1_is_oil = req.fuel1_type in OIL_FUELS
        f1_oil_lph = equip1["burner"].get("equivalent_lph", 0) if f1_is_oil else 0
        f1_oil_nb = select_oil_pipe_nb(f1_oil_lph) if f1_is_oil else 0

        br2 = None
        pipes2 = None
        equip2 = None
        ng_flow2 = 0
        air_flow2 = 0
        if is_dual:
            if req.mode == "direct":
                # Same heat, different CV → different flow
                ng_flow2 = heat_kcal_hr / req.fuel2_cv
                air_flow2 = air_flow  # air is CV-independent (same heat)
            else:
                br2 = calculate_burner(BurnerInputs(
                    Ti=req.Ti, Tf=req.Tf,
                    refractory_weight=effective_refractory_weight,
                    fuel_cv=req.fuel2_cv,
                    time_taken_hr=req.time_taken_hr,
                    refractory_heat_factor=req.refractory_heat_factor,
                    efficiency=req.efficiency,
                ))
                ng_flow2 = br2.extra_firing_rate_nm3hr
                air_flow2 = br2.air_qty_nm3hr
            pipes2 = calculate_pipe_sizes(PipeInputs(
                ng_flow_nm3hr=ng_flow2,
                air_flow_nm3hr=air_flow2,
            ))
            equip2 = select_equipment(
                ng_flow_nm3hr=ng_flow2,
                air_flow_nm3hr=air_flow2,
                is_dual_fuel=is_dual,
                fuel_cv=req.fuel2_cv,
                blower_pressure=req.blower_pressure,
                fuel_type=req.fuel2_type,
                hpu_variant=req.hpu_variant,
                burner_pressure_wg=burner_pressure_wg,
                butterfly_valve_vendor=req.butterfly_valve_vendor,
                shutoff_valve_vendor=req.shutoff_valve_vendor,
                control_mode=req.control_mode,
                auto_control_type=req.auto_control_type,
            )

        # Auto-fill fabrication (ms_structure) AND ceramic-fibre rolls from
        # the fabrication_ladle_mapping table (nearest ladle capacity for the
        # selected hood type). Pipeline weight is always user-supplied — the
        # file leaves pipeline kg blank on most rows so an auto-fill would
        # silently zero the user's input.
        from bom.vlph_builder import lookup_ladle_fab_pipeline
        _mapped = lookup_ladle_fab_pipeline(req.ladle_tons, "vertical", req.hood_type)
        _pipeline_kg  = req.pipeline_weight_kg
        # User-supplied override wins; fall back to DB-mapped value only when override is 0.
        _ms_override  = req.ms_structure_kg_override or (_mapped.get("fabrication_kg") if _mapped else 0)
        _ceramic_rolls = req.ceramic_rolls_override or (_mapped.get("ceramic_rolls") if _mapped else 0)

        # Air is CV-independent, so use fuel1 for air sizing
        if req.control_mode == "manual":
            bom_df = build_vlph_manual_df(
                equipment=equip1,
                ladle_tons=req.ladle_tons,
                fuel1_type=req.fuel1_type,
                fuel2_type=req.fuel2_type,
                equipment2=equip2,
                purging_line=req.purging_line,
                pressure_gauge_vendor=req.pressure_gauge_vendor,
                pilot_burner=req.pilot_burner,
                pipeline_weight_kg=_pipeline_kg,
                include_pilot=req.manual_pilot_burner == "yes",
                pilot_line_fuel=req.pilot_line_fuel,
                hood_type=req.hood_type,
                ceramic_rolls_override=_ceramic_rolls,
            )
        else:
            bom_df = build_vlph_120t_df(
                equipment=equip1,
                ladle_tons=req.ladle_tons,
                fuel1_type=req.fuel1_type,
                fuel2_type=req.fuel2_type,
                equipment2=equip2,
                control_mode=req.control_mode,
                auto_control_type=req.auto_control_type,
                control_valve_vendor=req.control_valve_vendor,
                butterfly_valve_vendor=req.butterfly_valve_vendor,
                shutoff_valve_vendor=req.shutoff_valve_vendor,
                pressure_gauge_vendor=req.pressure_gauge_vendor,
                pilot_burner=req.pilot_burner,
                pilot_line_fuel=req.pilot_line_fuel,
                pipeline_weight_kg=_pipeline_kg,
                purging_line=req.purging_line,
                num_burners=n_burners,
                ms_structure_kg_override=_ms_override,
                ceramic_rolls_override=_ceramic_rolls,
                hood_type=req.hood_type,
                special_auto_ignition=req.special_auto_ignition,
            )

        # Split summary rows from detail rows
        detail = bom_df[bom_df["MEDIA"] != ""].copy()
        bought_out_total = float(bom_df.loc[bom_df["ITEM NAME"] == "BOUGHT OUT ITEMS",   "TOTAL"].values[0]) if "BOUGHT OUT ITEMS" in bom_df["ITEM NAME"].values else 0
        encon_total      = float(bom_df.loc[bom_df["ITEM NAME"] == "ENCON ITEMS",        "TOTAL"].values[0]) if "ENCON ITEMS" in bom_df["ITEM NAME"].values else 0
        grand_total      = float(bom_df.loc[bom_df["ITEM NAME"] == "GRAND TOTAL",        "TOTAL"].values[0]) if "GRAND TOTAL" in bom_df["ITEM NAME"].values else 0

        # Build response — blower HP at user-selected pressure
        # Match the CFM logic in selection_engine: max(gas_cfm, oil_cfm)
        gas_cfm = blower_air_flow / 1.7
        oil_cfm = 0
        if f1_is_oil and equip1["burner"].get("equivalent_lph"):
            oil_cfm = equip1["burner"]["equivalent_lph"] * 10
        if f2_oil_lph_for_blower > 0:
            oil_cfm = max(oil_cfm, f2_oil_lph_for_blower * 10)
        cfm = max(gas_cfm, oil_cfm) if oil_cfm else gas_cfm
        blower_hp_calc = cfm * int(req.blower_pressure) / 3200
        resp = {
            "calculations": {
                "mode": req.mode,
                "Ti": req.Ti,
                "Tf": req.Tf,
                "refractory_weight": req.refractory_weight,
                "is_circulating": is_circulating,
                "preheater_type": "Circulating Ladle Preheater" if is_circulating else "Vertical Ladle Preheater",
                "effective_refractory_weight": effective_refractory_weight,
                "fuel_cv": f1_cv,
                "fuel1_type": req.fuel1_type,
                "fuel1_name": FUEL_NAMES.get(req.fuel1_type, req.fuel1_type),
                "fuel1_cv": f1_cv,
                "control_mode": req.control_mode,
                "auto_control_type": req.auto_control_type if req.control_mode == "automatic" else None,
                "time_taken_hr": req.time_taken_hr,
                "num_burners":                    n_burners,
                "avg_temp_rise":                  round(br1.avg_temp_rise, 2) if br1 else 0,
                "firing_rate_kcal":               round(br1.firing_rate_kcal, 2) if br1 else 0,
                "heat_load_kcal":                 round(br1.heat_load_kcal, 2) if br1 else 0,
                "fuel_consumption_nm3":           round(br1.fuel_consumption_nm3, 2) if br1 else 0,
                "calculated_firing_rate_nm3hr":   round(br1.calculated_firing_rate_nm3hr, 2) if br1 else round(ng_flow / 1.1, 2),
                "extra_firing_rate_nm3hr":        round(ng_flow, 2),
                "equivalent_lph":                 round(equip1["burner"].get("equivalent_lph", 0), 2),
                "fuel_density":                   equip1["burner"].get("fuel_density", 0),
                "fuel_density_unit":              equip1["burner"].get("fuel_density_unit", "kg/ltr"),
                "final_firing_rate_mw":           round(br1.final_firing_rate_mw, 2) if br1 else round(ng_flow * f1_cv / (860 * 1000), 2),
                "air_qty_nm3hr":                  round(air_flow, 2),
                "cfm":                            round(cfm, 2),
                "blower_hp_calc":                 round(blower_hp_calc, 2),
            },
            "pipes": {
                "fuel1_label": FUEL_NAMES.get(req.fuel1_type, "Fuel 1"),
                "fuel1_is_oil": f1_is_oil,
                "fuel1_oil_lph": round(f1_oil_lph, 2) if f1_is_oil else None,
                "ng_flow":      round(ng_flow, 2),
                "ng_velocity":  12.7 if not f1_is_oil else 0,
                "ng_dia_mm":    round(pipes1.ng_pipe_inner_dia_mm, 2) if not f1_is_oil else f1_oil_nb,
                "ng_nb":        pipes1.ng_pipe_nb if not f1_is_oil else f1_oil_nb,
                "gas_train_flow": round(equip1["ng_gas_train"]["max_flow"], 0) if not f1_is_oil else None,
                "gas_train_model": f'{equip1["ng_gas_train"]["inlet_nb"]} x {equip1["ng_gas_train"]["outlet_nb"]}' if not f1_is_oil else None,
                "air_flow":     round(air_flow, 2),
                "air_velocity": 15.0,
                "air_dia_mm":   round(pipes1.air_pipe_inner_dia_mm, 2),
                "air_nb":       pipes1.air_pipe_nb,
                # Actual combustion-air line NB the BOM components are sized to
                # (hydraulic pipe size, floored to 125 & air duct) — see selection_engine.
                "air_line_nb":  equip1.get("air_line_nb") or pipes1.air_pipe_nb,
                "air_line_flow": equip1.get("air_line_flow") or round(air_flow),
            },
            "equipment": {
                "burner_model":   equip1["burner"]["model"],
                "burner_max_kcal_hr": equip1["burner"].get("max_firing_kcal_hr", 0),
                "burner_max_lph":     equip1["burner"].get("max_firing_lph", 0),
                "blower_model":   equip1["blower"]["model"],
                "blower_hp":      equip1["blower"]["hp"],
                "blower_airflow": equip1["blower"]["airflow_nm3hr"],
                "blower_cfm":     equip1["blower"].get("cfm", 0),
                "ng_gas_train":   f'{equip1["ng_gas_train"]["inlet_nb"]} x {equip1["ng_gas_train"]["outlet_nb"]} NB',
                # In dual-fuel offers the HPU may live on equip2 instead of
                # equip1 (e.g. fuel 1 = NG gas, fuel 2 = LDO oil). Pick whichever
                # side actually carries it so the Step-4 'Pumping Unit' row is
                # populated regardless of which fuel slot the oil is in.
                "hpu": (
                    (lambda h: f'{h["model"]} — {h["unit_kw"]} KW {h["variant"]}')(
                        (equip1 or {}).get("hpu") or (equip2 or {}).get("hpu")
                    )
                    if ((equip1 or {}).get("hpu") or (equip2 or {}).get("hpu"))
                    else None
                ),
            },
            "bom": detail[["MEDIA","ITEM NAME","REFERENCE","QTY","MAKE","UNIT PRICE","TOTAL"]].to_dict(orient="records"),
            "cost_summary": {
                "bought_out_total": round(bought_out_total, 2),
                "encon_total":      round(encon_total, 2),
                "grand_total":      round(grand_total, 2),
            },
        }

        # Add fuel2 data if dual fuel
        if is_dual and pipes2:
            resp["calculations"]["is_dual"] = True
            resp["calculations"]["fuel2_type"] = req.fuel2_type
            resp["calculations"]["fuel2_name"] = FUEL_NAMES.get(req.fuel2_type, req.fuel2_type)
            resp["calculations"]["fuel2_cv"] = req.fuel2_cv
            resp["calculations"]["fuel2_consumption_nm3"] = round(br2.fuel_consumption_nm3, 2) if br2 else 0
            resp["calculations"]["fuel2_firing_rate_nm3hr"] = round(br2.calculated_firing_rate_nm3hr, 2) if br2 else round(ng_flow2 / 1.1, 2)
            resp["calculations"]["fuel2_extra_firing_rate_nm3hr"] = round(ng_flow2, 2)
            resp["calculations"]["fuel2_equivalent_lph"] = round(equip2["burner"].get("equivalent_lph", 0), 2) if equip2 else 0
            resp["calculations"]["fuel2_fuel_density"] = equip2["burner"].get("fuel_density", 0) if equip2 else 0
            resp["calculations"]["fuel2_fuel_density_unit"] = equip2["burner"].get("fuel_density_unit", "kg/ltr") if equip2 else "kg/ltr"
            resp["pipes"]["fuel2_label"] = FUEL_NAMES.get(req.fuel2_type, "Fuel 2")
            resp["pipes"]["fuel2_flow"] = round(ng_flow2, 2)
            f2_is_oil = req.fuel2_type in OIL_FUELS
            f2_oil_lph = (equip2["burner"].get("equivalent_lph", 0) if (f2_is_oil and equip2) else 0)
            f2_oil_nb = select_oil_pipe_nb(f2_oil_lph) if f2_is_oil else 0
            resp["pipes"]["fuel2_is_oil"] = f2_is_oil
            resp["pipes"]["fuel2_oil_lph"] = round(f2_oil_lph, 2) if f2_is_oil else None
            resp["pipes"]["fuel2_dia_mm"] = round(pipes2.ng_pipe_inner_dia_mm, 2) if not f2_is_oil else f2_oil_nb
            resp["pipes"]["fuel2_nb"] = pipes2.ng_pipe_nb if not f2_is_oil else f2_oil_nb
            if not f2_is_oil and equip2:
                resp["pipes"]["fuel2_gas_train_flow"] = round(equip2["ng_gas_train"]["max_flow"], 0)
                resp["pipes"]["fuel2_gas_train_model"] = f'{equip2["ng_gas_train"]["inlet_nb"]} x {equip2["ng_gas_train"]["outlet_nb"]}'

        return resp
    except Exception as e:
        import traceback
        return {"error": str(e), "trace": traceback.format_exc()}


# ──────────────────────────────────────────────────────────────────────────
# Cost Variations — runs the VLPH calc once per single-axis swap, applies
# the user's markup factor, and also surfaces markup-only variations
# (1.75 / 1.70 / 1.65 / 1.60) so the user can see total cost across every
# realistic lever for negotiating a deal closer to a target price.
# ──────────────────────────────────────────────────────────────────────────
class CostVariationsRequest(VLPHCalcRequest):
    markup: float = 1.80    # Frontend's bought-out markup multiplier.


@app.post("/api/cost-variations")
def cost_variations(req: CostVariationsRequest):
    """Iterate over the variation axes (vendor swaps, hpu_variant, hood_type,
    auto_control_type, control_mode, special_auto_ignition) plus markup
    levels and return each alternative's GRAND TOTAL (= bought_out * markup
    + encon_total).

    Response shape:
      {
        "current_total":   <float>,
        "current_markup":  <float>,
        "variations":      [{axis, value, total, saves}, ...],
        "best_combined":   {total, saves, swaps:[{axis,value}, ...]} | null
      }
    Variations that don't save money vs the baseline are dropped.
    """
    OIL_FUELS = {"hsd", "ldo", "hdo", "fo", "sko", "cfo", "lshs"}
    is_oil_offer = (
        req.fuel1_type in OIL_FUELS
        or (req.fuel2_type and req.fuel2_type in OIL_FUELS)
    )

    LABEL_MAP = {
        "control_valve_vendor": {
            "_axis": "Control Valve Vendor",
            "dembla": "DEMBLA", "aira": "AIRA", "cair": "CAIR",
        },
        "shutoff_valve_vendor": {
            "_axis": "Shutoff Valve Vendor",
            "dembla": "DEMBLA", "aira": "AIRA", "cair": "CAIR",
        },
        "butterfly_valve_vendor": {
            "_axis": "Butterfly Valve Vendor",
            "lt_lever": "L&T Lever", "lt_gear": "L&T Gear",
        },
        "pressure_gauge_vendor": {
            "_axis": "Pressure Gauge Vendor",
            "baumer": "BAUMER", "hguru": "HGURU",
        },
        "hpu_variant": {
            "_axis": "HPU Variant",
            "Simplex": "Simplex", "Duplex 1": "Duplex-I", "Duplex 2": "Duplex-II",
        },
        "hood_type": {
            "_axis": "Hood Movement",
            "up_down": "Up/Down (hydraulic)",
            "swivel_manual": "Manual Swivelling",
            "swivel_geared": "Geared Swivelling",
        },
        "auto_control_type": {
            "_axis": "Auto Control Type",
            "plc": "PLC", "plc_agr": "PLC + AGR", "pid": "PID",
        },
        "control_mode": {
            "_axis": "Control Mode",
            "automatic": "Automatic", "manual": "Manual",
        },
        "special_auto_ignition": {
            "_axis": "Auto Ignition",
            True: "Auto Ignition ON", False: "Auto Ignition OFF",
        },
    }

    axes = [
        ("control_valve_vendor",   ["dembla", "aira", "cair"]),
        ("shutoff_valve_vendor",   ["dembla", "aira", "cair"]),
        ("butterfly_valve_vendor", ["lt_lever", "lt_gear"]),
        ("pressure_gauge_vendor",  ["baumer", "hguru"]),
        ("hood_type",              ["up_down", "swivel_manual", "swivel_geared"]),
        ("special_auto_ignition",  [True, False]),
    ]
    if is_oil_offer:
        axes.append(("hpu_variant", ["Simplex", "Duplex 1", "Duplex 2"]))
    if req.control_mode == "automatic":
        axes.append(("auto_control_type", ["plc", "plc_agr", "pid"]))
    other_mode = "manual" if req.control_mode == "automatic" else "automatic"
    axes.append(("control_mode", [other_mode]))

    # Strip the markup field before passing to vlph_calculate so it round-trips
    # cleanly through the existing VLPHCalcRequest schema.
    def _to_calc_req(modified):
        d = modified.dict()
        d.pop("markup", None)
        return VLPHCalcRequest(**d)

    def _calc_costs(modified):
        """Returns (bought_out, encon) for the given request, or None."""
        try:
            resp = vlph_calculate(_to_calc_req(modified))
            cs = resp.get("cost_summary", {})
            return (
                float(cs.get("bought_out_total") or 0),
                float(cs.get("encon_total") or 0),
            )
        except Exception:
            return None

    def _grand(bought, encon, markup):
        return bought * markup + encon

    base = _calc_costs(req)
    if base is None:
        return {"error": "Could not compute baseline cost"}
    base_bought, base_encon = base
    base_total = _grand(base_bought, base_encon, req.markup)
    if base_total <= 0:
        return {"error": "Baseline cost is zero — check inputs"}

    variations = []
    best_per_axis = {}  # field -> (alt, total) of cheapest swap on this axis

    for field, alts in axes:
        cur_value = getattr(req, field)
        for alt in alts:
            if alt == cur_value:
                continue
            modified = req.copy(update={field: alt})
            costs = _calc_costs(modified)
            if costs is None:
                continue
            total = _grand(costs[0], costs[1], req.markup)
            # Drop variations where the calc returned 0 (silent failure) or
            # didn't actually save money against the baseline.
            if total <= 0:
                continue
            saves = base_total - total
            if saves <= 0:
                continue
            label = LABEL_MAP.get(field, {})
            variations.append({
                "axis":  label.get("_axis", field),
                "value": label.get(alt, str(alt)),
                "field": field,
                "alt":   alt,
                "total": round(total, 2),
                "saves": round(saves, 2),
            })
            cur_best = best_per_axis.get(field)
            if cur_best is None or total < cur_best[1]:
                best_per_axis[field] = (alt, total)

    # Markup-only variations — same BOM, different markup multiplier.
    for mk in [1.75, 1.70, 1.65, 1.60]:
        if abs(mk - req.markup) < 1e-6:
            continue
        total = _grand(base_bought, base_encon, mk)
        saves = base_total - total
        if saves <= 0:
            continue
        variations.append({
            "axis":  "Markup",
            "value": f"{mk:.2f}×",
            "field": "markup",
            "alt":   mk,
            "total": round(total, 2),
            "saves": round(saves, 2),
        })

    variations.sort(key=lambda r: r["saves"], reverse=True)

    # Combined-best: cheapest alt for each axis applied at once + lowest markup.
    best_combined = None
    update = {field: best for field, (best, _) in best_per_axis.items()}
    combined_req = req.copy(update=update) if update else req
    costs = _calc_costs(combined_req)
    if costs is not None:
        lowest_markup = 1.60 if req.markup > 1.60 else req.markup
        combined_total = _grand(costs[0], costs[1], lowest_markup)
        if combined_total > 0 and combined_total < base_total:
            swaps = []
            for field, (alt, _) in best_per_axis.items():
                label = LABEL_MAP.get(field, {})
                swaps.append({
                    "axis":  label.get("_axis", field),
                    "value": label.get(alt, str(alt)),
                })
            if abs(lowest_markup - req.markup) > 1e-6:
                swaps.append({"axis": "Markup", "value": f"{lowest_markup:.2f}×"})
            best_combined = {
                "total": round(combined_total, 2),
                "saves": round(base_total - combined_total, 2),
                "swaps": swaps,
            }

    return {
        "current_total":  round(base_total, 2),
        "current_markup": round(req.markup, 2),
        "variations":     variations,
        "best_combined":  best_combined,
    }


class RegenCalcRequest(BaseModel):
    # Direct-selection mode (Size / Fuel / Pairs) — the primary path.
    model_kw: Optional[int] = None     # BURNER size 500..6000; if set, skips the heat-load calc
    regen_kw: Optional[int] = None     # REGENERATOR size 500..6000; defaults to model_kw
    fuel: Optional[str] = None         # Natural Gas / Blast Furnace Gas / Coke Oven Gas / Producer Gas / Oil
    num_pairs: Optional[int] = None    # number of burner pairs
    markup: float = 1.80
    # Legacy heat-load sizing inputs — used only when model_kw is not given.
    material_weight_kg: float = 0.0
    Ti: float = 0.0
    Tf: float = 0.0
    Cp: float = 0.48
    cycle_time_hr: float = 2.0
    efficiency: float = 0.65
    num_pairs_override: int = 0


@app.post("/api/regen-calculate")
def regen_calculate(req: RegenCalcRequest):
    try:
        from bom.regen_builder import build_regen_df, select_model, get_supplementary_data

        result = None
        if req.model_kw:
            # ── Direct-selection mode: user picks Size + Fuel + Pairs ──────────
            model_kw  = select_model(int(req.model_kw))   # snap to a valid burner model
            regen_kw  = select_model(int(req.regen_kw)) if req.regen_kw else model_kw
            num_pairs = max(1, int(req.num_pairs or 1))
        else:
            # ── Legacy heat-load sizing (charge weight → kW) ──────────────────
            from calculations.regen import RegenInputs, calculate_regen
            result = calculate_regen(RegenInputs(
                material_weight_kg=req.material_weight_kg,
                Ti=req.Ti,
                Tf=req.Tf,
                Cp=req.Cp,
                cycle_time_hr=req.cycle_time_hr,
                efficiency=req.efficiency,
                num_pairs_override=req.num_pairs_override,
            ))
            # Select the appropriate KW model (smallest >= required_kw per pair)
            kw_per_pair = result.required_kw / max(1, result.num_pairs)
            model_kw    = select_model(kw_per_pair)
            regen_kw    = model_kw
            num_pairs   = result.num_pairs

        model_markup = req.markup if req.markup != 1.80 else None  # None → use model default

        bom_df = build_regen_df(model_kw, model_markup, num_pairs=num_pairs,
                                db_path=DB_PATH, fuel=req.fuel or "Natural Gas",
                                regen_kw=regen_kw)
        supplementary = get_supplementary_data(model_kw)

        # Augment supplementary with full sizing + nozzle + legacy rates from DB
        try:
            with sqlite3.connect(DB_PATH) as _c:
                # ALL sizing rows (all KW models) — for full Excel-like tables
                sz_cols = [d[0] for d in _c.execute("SELECT * FROM regen_sizing LIMIT 0").description]
                sz_all  = [dict(zip(sz_cols, r)) for r in _c.execute("SELECT * FROM regen_sizing ORDER BY kw").fetchall()]
                supplementary['burner_sizing']['all_sizing'] = sz_all
                # Single selected-KW row (for inline detail)
                sz_row = next((r for r in sz_all if r['kw'] == model_kw), None)
                if sz_row:
                    supplementary['burner_sizing']['dimensions'] = {
                        k: sz_row.get(k) for k in [
                            'shell_thick','retainer_thick','refractory_thick',
                            'dim_L','dim_H','dim_W','bottom_h',
                            'vol_total','vol_effective','vol_refractory',
                            'density_castable','wt_refractory_insulation',
                            'loose_density_balls','vol_available_balls','balls_filling_pct',
                        ]
                    }
                    supplementary['burner_sizing']['weight_detail'] = {
                        k: sz_row.get(k) for k in [
                            'bb_dia_inner','bb_dia_outer','bb_depth','wt_burner_block',
                            'burner_length','burner_dia',
                            'wt_burner_shell','wt_burner_refrac_detail','wt_burner_total',
                            'wt_shell','wt_ss_plate','wt_ceramic_balls_burner',
                            'wt_regen_total','wt_grand_total','bloom_approx_wt',
                        ]
                    }
                # Legacy material rates (from Excel "Burner Sizing and costing" sheet)
                mr_rows = _c.execute("SELECT material, wastage, material_cost, labor_cost FROM regen_material_rates").fetchall()
                legacy_rates = {r[0]: {'material': r[0], 'wastage': r[1], 'mat_cost': r[2], 'labour_cost': r[3] or 0} for r in mr_rows}
                if legacy_rates:
                    supplementary['burner_sizing']['legacy_material_rates'] = list(legacy_rates.values())
                    # Rebuild cost_detail using legacy rates so formula matches Excel exactly
                    from bom.regen_builder import _BURNER_WEIGHTS
                    w = _BURNER_WEIGHTS[model_kw]
                    def _legacy_rate(mat_key):
                        r = legacy_rates.get(mat_key, {})
                        mc, lc, wa = r.get('mat_cost', 0) or 0, r.get('labour_cost', 0) or 0, r.get('wastage', 0) or 0
                        return round((mc + lc) * (1 + wa), 4), mc, lc, wa
                    ms_rate,   ms_m,  ms_l,  ms_w  = _legacy_rate('MS')
                    ss_rate,   ss_m,  ss_l,  ss_w  = _legacy_rate('SS')
                    rf_rate,   rf_m,  rf_l,  rf_w  = _legacy_rate('Refractory')
                    cb_rate,   cb_m,  cb_l,  cb_w  = _legacy_rate('Ceramic Balls')
                    supplementary['burner_sizing']['cost_detail'] = [
                        dict(component='Burner Body',  material='MS',           weight_kg=w['burner_ms'],    mat_cost=ms_m, labour_cost=ms_l, wastage=ms_w, rate=ms_rate, cost=round(w['burner_ms']    * ms_rate, 2)),
                        dict(component='Burner Body',  material='Refractory',   weight_kg=w['burner_refrac'],mat_cost=rf_m, labour_cost=rf_l, wastage=rf_w, rate=rf_rate, cost=round(w['burner_refrac'] * rf_rate, 2)),
                        dict(component='Regenerator',  material='MS',           weight_kg=w['regen_ms'],     mat_cost=ms_m, labour_cost=ms_l, wastage=ms_w, rate=ms_rate, cost=round(w['regen_ms']     * ms_rate, 2)),
                        dict(component='Regenerator',  material='SS',           weight_kg=w['regen_ss'],     mat_cost=ss_m, labour_cost=ss_l, wastage=ss_w, rate=ss_rate, cost=round(w['regen_ss']     * ss_rate, 2)),
                        dict(component='Regenerator',  material='Refractory',   weight_kg=w['regen_refrac'], mat_cost=rf_m, labour_cost=rf_l, wastage=rf_w, rate=rf_rate, cost=round(w['regen_refrac'] * rf_rate, 2)),
                        dict(component='Regenerator',  material='Ceramic Balls',weight_kg=w['regen_ceramic'],mat_cost=cb_m, labour_cost=cb_l, wastage=cb_w, rate=cb_rate, cost=round(w['regen_ceramic'] * cb_rate, 2)),
                        dict(component='Burner Block', material='Refractory',   weight_kg=w['block_refrac'], mat_cost=rf_m, labour_cost=rf_l, wastage=rf_w, rate=rf_rate, cost=round(w['block_refrac'] * rf_rate, 2)),
                    ]
                    supplementary['burner_sizing']['total_unit_cost'] = round(sum(d['cost'] for d in supplementary['burner_sizing']['cost_detail']), 2)
                    supplementary['burner_sizing']['total_pair_cost'] = round(supplementary['burner_sizing']['total_unit_cost'] * 2, 2)
                # Nozzle sizing (all burners)
                nz_cols = [d[0] for d in _c.execute("SELECT * FROM regen_nozzle_sizing LIMIT 0").description]
                nz_rows = _c.execute("SELECT * FROM regen_nozzle_sizing ORDER BY power_kw").fetchall()
                supplementary['nozzle_sizing'] = [dict(zip(nz_cols, r)) for r in nz_rows]
                # All pipe sizes (all KW + all gas types) for full Excel-like table
                ps_cols = [d[0] for d in _c.execute("SELECT * FROM regen_pipe_sizes LIMIT 0").description]
                ps_rows = _c.execute("SELECT * FROM regen_pipe_sizes ORDER BY gas_type, burner_size_kw").fetchall()
                supplementary['all_pipe_sizes'] = [dict(zip(ps_cols, r)) for r in ps_rows]
        except Exception:
            pass  # DB not ready yet — supplementary still has the hardcoded data

        # Blower + ID-fan sizing worked example (shown under the costing sheet).
        try:
            from bom.regen_builder import compute_fan_flows
            with sqlite3.connect(DB_PATH) as _fc:
                supplementary['fan_sizing'] = compute_fan_flows(
                    model_kw, num_pairs, req.fuel or "Natural Gas", _fc)
        except Exception:
            pass

        total_cost    = float(bom_df["TOTAL COST"].sum())
        total_selling = float(bom_df["TOTAL SELLING"].sum())

        calculations = {
            "mode": "direct" if req.model_kw else "heat",
            "fuel": req.fuel or "Natural Gas",
            "num_pairs": num_pairs,
            "model_kw": model_kw,          # burner size
            "regen_kw": regen_kw,          # regenerator size
            "total_kw": model_kw * num_pairs,
        }
        if result is not None:
            # Legacy heat-load mode — include the derivation for the calc sheet.
            calculations.update({
                "material_weight_kg": req.material_weight_kg,
                "Ti": req.Ti,
                "Tf": req.Tf,
                "Cp": req.Cp,
                "delta_T": round(result.delta_T, 2),
                "heat_required_kj": round(result.heat_required_kj, 2),
                "heat_required_kcal": round(result.heat_required_kcal, 2),
                "cycle_time_hr": req.cycle_time_hr,
                "efficiency": req.efficiency,
                "required_kw": round(result.required_kw, 2),
            })
        return {
            "calculations": calculations,
            "bom": bom_df.to_dict(orient="records"),
            "cost_summary": {
                "total_cost": round(total_cost, 2),
                "total_selling": round(total_selling, 2),
                "markup": req.markup,
            },
            "supplementary": supplementary,
        }
    except Exception as e:
        import traceback
        return {"error": str(e), "trace": traceback.format_exc()}


@app.post("/api/hlph-calculate")
def hlph_calculate(req: VLPHCalcRequest):
    try:
        from calculations.burner import BurnerInputs, calculate_burner
        from calculations.pipes import PipeInputs, calculate_pipe_sizes, select_oil_pipe_nb
        from bom.selectors.selection_engine import select_equipment
        from bom.hlph_builder import build_hlph_df, build_hlph_manual_df

        FUEL_NAMES = {
            "ng": "Natural Gas", "lpg": "LPG", "cog": "COG", "bg": "BFG", "rlng": "RLNG", "mg": "Mixed Gas",
            "hsd": "Diesel (HSD)", "ldo": "LDO", "hdo": "HDO", "fo": "Furnace Oil",
            "sko": "Kerosene (SKO)", "cfo": "CFO", "lshs": "LSHS",
        }
        OIL_FUELS = {"hsd", "ldo", "hdo", "fo", "sko", "cfo", "lshs"}

        f1_cv = req.fuel1_cv if req.fuel1_cv > 0 else req.fuel_cv
        is_dual = req.fuel2_type != "none" and req.fuel2_cv > 0

        # Circulating Ladle Preheater: a hot ladle (initial temperature above the
        # threshold) is already in circulation, so it only needs topping up — the
        # effective refractory mass to heat is HALF the input weight. Detected in
        # calc mode only (direct mode enters burner capacity, not weight).
        CIRCULATING_TI_THRESHOLD = 600
        is_circulating = (req.mode == "calc") and (req.Ti > CIRCULATING_TI_THRESHOLD)
        effective_refractory_weight = (
            req.refractory_weight / 2 if is_circulating else req.refractory_weight
        )

        if req.mode == "direct":
            # Direct mode: burner capacity in kW (× 860 → kcal/hr)
            if req.direct_heat_input_kcal_hr and req.direct_heat_input_kcal_hr > 0:
                heat_kcal_hr = req.direct_heat_input_kcal_hr
            else:
                heat_kcal_hr = req.direct_burner_capacity * 860
            ng_flow  = heat_kcal_hr / f1_cv
            air_flow = heat_kcal_hr * 118 / 100000
            br = None
        else:
            burner_inputs = BurnerInputs(
                Ti=req.Ti, Tf=req.Tf,
                refractory_weight=effective_refractory_weight,
                fuel_cv=f1_cv,
                time_taken_hr=req.time_taken_hr,
                refractory_heat_factor=req.refractory_heat_factor,
                efficiency=req.efficiency,
            )
            br = calculate_burner(burner_inputs)
            ng_flow = br.extra_firing_rate_nm3hr
            air_flow = br.air_qty_nm3hr

        burner_pressure_wg = 36 if req.blower_pressure == "40" else 24

        # Dual fuel: blower sized for max(air1, air2)
        blower_air_flow = air_flow
        if is_dual:
            if req.mode == "direct":
                air_flow2_pre = air_flow
            else:
                br2_pre = calculate_burner(BurnerInputs(
                    Ti=req.Ti, Tf=req.Tf,
                    refractory_weight=effective_refractory_weight,
                    fuel_cv=req.fuel2_cv,
                    time_taken_hr=req.time_taken_hr,
                    refractory_heat_factor=req.refractory_heat_factor,
                    efficiency=req.efficiency,
                ))
                air_flow2_pre = br2_pre.air_qty_nm3hr
            blower_air_flow = max(air_flow, air_flow2_pre)

        # Pre-compute fuel2 oil LPH for dual-fuel blower CFM
        f2_oil_lph_for_blower = 0
        if is_dual and req.fuel2_type in OIL_FUELS:
            heat_for_f2 = ng_flow * f1_cv if req.mode != "direct" else heat_kcal_hr
            f2_flow_kghr = heat_for_f2 / req.fuel2_cv if req.fuel2_cv > 0 else 0
            from bom.selectors.encon_burner import _get_fuel_density
            f2_density = _get_fuel_density(req.fuel2_type)
            f2_oil_lph_for_blower = f2_flow_kghr / f2_density if f2_density else 0

        equip1 = select_equipment(
            ng_flow_nm3hr=ng_flow, air_flow_nm3hr=blower_air_flow,
            is_dual_fuel=is_dual, fuel_cv=f1_cv,
            blower_pressure=req.blower_pressure, fuel_type=req.fuel1_type,
            hpu_variant=req.hpu_variant, burner_pressure_wg=burner_pressure_wg,
            butterfly_valve_vendor=req.butterfly_valve_vendor,
            shutoff_valve_vendor=req.shutoff_valve_vendor,
            control_mode=req.control_mode, auto_control_type=req.auto_control_type,
            fuel2_lph=f2_oil_lph_for_blower,
        )

        # Size pipes AFTER blower selection — gas pipe for fuel 1's flow,
        # air pipe for the combustion-air flow (the burner's actual air demand).
        pipes1 = calculate_pipe_sizes(PipeInputs(ng_flow_nm3hr=ng_flow, air_flow_nm3hr=equip1.get("air_line_flow") or air_flow))

        f1_is_oil = req.fuel1_type in OIL_FUELS
        f1_oil_lph = equip1["burner"].get("equivalent_lph", 0) if f1_is_oil else 0

        # Fuel 2
        br2, equip2, ng_flow2, air_flow2 = None, None, 0, 0
        if is_dual:
            if req.mode == "direct":
                # Same heat output, fuel2 flow at fuel2 CV; air is CV-independent
                ng_flow2  = heat_kcal_hr / req.fuel2_cv
                air_flow2 = air_flow
            else:
                br2 = calculate_burner(BurnerInputs(
                    Ti=req.Ti, Tf=req.Tf,
                    refractory_weight=effective_refractory_weight,
                    fuel_cv=req.fuel2_cv,
                    time_taken_hr=req.time_taken_hr,
                    refractory_heat_factor=req.refractory_heat_factor,
                    efficiency=req.efficiency,
                ))
                ng_flow2  = br2.extra_firing_rate_nm3hr
                air_flow2 = br2.air_qty_nm3hr
            equip2 = select_equipment(
                ng_flow_nm3hr=ng_flow2, air_flow_nm3hr=air_flow2,
                is_dual_fuel=is_dual, fuel_cv=req.fuel2_cv,
                blower_pressure=req.blower_pressure, fuel_type=req.fuel2_type,
                hpu_variant=req.hpu_variant, burner_pressure_wg=burner_pressure_wg,
                butterfly_valve_vendor=req.butterfly_valve_vendor,
                shutoff_valve_vendor=req.shutoff_valve_vendor,
                control_mode=req.control_mode, auto_control_type=req.auto_control_type,
            )
            # Pipe sizing for fuel-2 — needed by the front-end calc-result
            # panel ('Fuel-2 line size: gas flow / pipe NB'). Without this
            # the JS reads undefined and renders 'NaN Nm3/hr / undefined NB'.
            pipes2 = calculate_pipe_sizes(PipeInputs(
                ng_flow_nm3hr=ng_flow2, air_flow_nm3hr=air_flow2,
            ))
        else:
            pipes2 = None

        # Auto-fill fabrication + ceramic-fibre rolls from
        # fabrication_ladle_mapping (nearest horizontal row). Pipeline weight
        # is always user-supplied (file carries no horizontal pipeline data).
        from bom.vlph_builder import lookup_ladle_fab_pipeline
        _mapped_h = lookup_ladle_fab_pipeline(req.ladle_tons, "horizontal")
        _pipeline_kg_h  = req.pipeline_weight_kg
        # User-supplied override wins; fall back to DB-mapped value only when override is 0.
        _ms_override_h  = req.ms_structure_kg_override or (_mapped_h.get("fabrication_kg") if _mapped_h else 0)
        _ceramic_rolls_h = req.ceramic_rolls_override or (_mapped_h.get("ceramic_rolls") if _mapped_h else 0)

        if req.control_mode == "manual":
            bom_df = build_hlph_manual_df(
                equipment=equip1, ladle_tons=req.ladle_tons,
                fuel1_type=req.fuel1_type,
                purging_line=req.purging_line,
                pressure_gauge_vendor=req.pressure_gauge_vendor,
                pilot_burner=req.pilot_burner,
                pipeline_weight_kg=_pipeline_kg_h,
                include_pilot=req.manual_pilot_burner == "yes",
                pilot_line_fuel=req.pilot_line_fuel,
                ceramic_rolls_override=_ceramic_rolls_h,
            )
        else:
            bom_df = build_hlph_df(
                equipment=equip1, ladle_tons=req.ladle_tons,
                fuel1_type=req.fuel1_type, fuel2_type=req.fuel2_type,
                equipment2=equip2,
                control_mode=req.control_mode, auto_control_type=req.auto_control_type,
                control_valve_vendor=req.control_valve_vendor,
                butterfly_valve_vendor=req.butterfly_valve_vendor,
                shutoff_valve_vendor=req.shutoff_valve_vendor,
                pressure_gauge_vendor=req.pressure_gauge_vendor,
                pilot_burner=req.pilot_burner, pilot_line_fuel=req.pilot_line_fuel,
                pipeline_weight_kg=_pipeline_kg_h,
                purging_line=req.purging_line,
                ms_structure_kg_override=_ms_override_h,
                ceramic_rolls_override=_ceramic_rolls_h,
                special_auto_ignition=req.special_auto_ignition,
            )

        detail = bom_df[bom_df["MEDIA"] != ""].copy()
        bought_out_total = float(bom_df.loc[bom_df["ITEM NAME"] == "BOUGHT OUT ITEMS", "TOTAL"].values[0]) if "BOUGHT OUT ITEMS" in bom_df["ITEM NAME"].values else 0
        encon_total = float(bom_df.loc[bom_df["ITEM NAME"] == "ENCON ITEMS", "TOTAL"].values[0]) if "ENCON ITEMS" in bom_df["ITEM NAME"].values else 0
        grand_total = float(bom_df.loc[bom_df["ITEM NAME"] == "GRAND TOTAL", "TOTAL"].values[0]) if "GRAND TOTAL" in bom_df["ITEM NAME"].values else 0

        # Match CFM logic in selection_engine: max(gas_cfm, oil_cfm)
        gas_cfm = blower_air_flow / 1.7
        oil_cfm = 0
        if f1_is_oil and equip1["burner"].get("equivalent_lph"):
            oil_cfm = equip1["burner"]["equivalent_lph"] * 10
        if f2_oil_lph_for_blower > 0:
            oil_cfm = max(oil_cfm, f2_oil_lph_for_blower * 10)
        cfm = max(gas_cfm, oil_cfm) if oil_cfm else gas_cfm
        blower_hp_calc = cfm * int(req.blower_pressure) / 3200

        resp = {
            "calculations": {
                "Ti": req.Ti, "Tf": req.Tf,
                "refractory_weight": req.refractory_weight,
                "is_circulating": is_circulating,
                "preheater_type": "Circulating Ladle Preheater (Horizontal)" if is_circulating else "Horizontal Ladle Preheater",
                "effective_refractory_weight": effective_refractory_weight,
                "fuel_cv": f1_cv,
                "fuel1_type": req.fuel1_type,
                "fuel1_name": FUEL_NAMES.get(req.fuel1_type, req.fuel1_type),
                "fuel1_cv": f1_cv,
                "is_dual": is_dual,
                "time_taken_hr": req.time_taken_hr,
                "mode": req.mode,
                "avg_temp_rise":              round(br.avg_temp_rise, 2)              if br else 0,
                "firing_rate_kcal":           round(br.firing_rate_kcal, 2)           if br else 0,
                "heat_load_kcal":             round(br.heat_load_kcal, 2)             if br else 0,
                "fuel_consumption_nm3":       round(br.fuel_consumption_nm3, 2)       if br else 0,
                "calculated_firing_rate_nm3hr": round(br.calculated_firing_rate_nm3hr, 2) if br else round(ng_flow / 1.1, 2),
                "extra_firing_rate_nm3hr":    round(ng_flow, 2),
                "equivalent_lph":             round(equip1["burner"].get("equivalent_lph", 0), 2),
                "fuel_density":               equip1["burner"].get("fuel_density", 0),
                "fuel_density_unit":          equip1["burner"].get("fuel_density_unit", "kg/ltr"),
                "final_firing_rate_mw":       round(br.final_firing_rate_mw, 2)       if br else round(ng_flow * f1_cv / (860 * 1000), 2),
                "air_qty_nm3hr":              round(air_flow, 2),
                "cfm": round(cfm, 2),
                "blower_hp_calc": round(blower_hp_calc, 2),
            },
            "pipes": {
                "fuel1_label": FUEL_NAMES.get(req.fuel1_type, "Fuel 1"),
                "fuel1_is_oil": f1_is_oil,
                "fuel1_oil_lph": round(f1_oil_lph, 2) if f1_is_oil else None,
                "ng_flow": round(ng_flow, 2),
                "ng_nb": pipes1.ng_pipe_nb,
                "air_flow": round(air_flow, 2),
                "air_nb": pipes1.air_pipe_nb,
                "air_line_nb": equip1.get("air_line_nb") or pipes1.air_pipe_nb,
                "air_line_flow": equip1.get("air_line_flow") or round(air_flow),
                "gas_train_flow": round(equip1["ng_gas_train"]["max_flow"], 0) if equip1.get("ng_gas_train") else 0,
                "gas_train_model": f'{equip1["ng_gas_train"]["inlet_nb"]} x {equip1["ng_gas_train"]["outlet_nb"]}' if equip1.get("ng_gas_train") else "",
            },
            "equipment": {
                "burner_model": equip1["burner"]["model"],
                "burner_max_kcal_hr": equip1["burner"].get("max_firing_kcal_hr", 0),
                "burner_max_lph":     equip1["burner"].get("max_firing_lph", 0),
                "blower_model": equip1["blower"]["model"],
                "blower_hp": equip1["blower"]["hp"],
                "blower_airflow": equip1["blower"]["airflow_nm3hr"],
                "blower_cfm":   equip1["blower"].get("cfm", 0),
                "ng_gas_train": f'{equip1["ng_gas_train"]["inlet_nb"]} x {equip1["ng_gas_train"]["outlet_nb"]} NB' if equip1.get("ng_gas_train") else "",
                # In dual-fuel offers the HPU may live on equip2 (e.g. fuel 1 NG
                # + fuel 2 LDO). Pick whichever side actually carries it.
                "hpu": (
                    (lambda h: f'{h["model"]} — {h["unit_kw"]} KW {h.get("variant", "")}'.strip())(
                        (equip1 or {}).get("hpu") or (equip2 or {}).get("hpu")
                    )
                    if ((equip1 or {}).get("hpu") or (equip2 or {}).get("hpu"))
                    else None
                ),
            },
            "bom": detail.to_dict(orient="records"),
            "cost_summary": {
                "bought_out_total": round(bought_out_total, 2),
                "encon_total": round(encon_total, 2),
                "grand_total": round(grand_total, 2),
            },
        }

        # Surface fuel-2 fields the front-end + offer pipeline rely on (mirrors
        # the VLPH endpoint).
        if is_dual and equip2 is not None:
            resp["calculations"]["fuel2_type"] = req.fuel2_type
            resp["calculations"]["fuel2_name"] = FUEL_NAMES.get(req.fuel2_type, req.fuel2_type)
            resp["calculations"]["fuel2_cv"]   = req.fuel2_cv
            resp["calculations"]["fuel2_extra_firing_rate_nm3hr"] = round(ng_flow2, 2)
            # Calc-mode fields the front-end needs to render Fuel Consumption
            # and Calc. Firing Rate in the result table. Direct mode falls
            # back to deriving from ng_flow2 (the +10% extra rate).
            resp["calculations"]["fuel2_consumption_nm3"] = (
                round(br2.fuel_consumption_nm3, 2) if br2 else 0
            )
            resp["calculations"]["fuel2_firing_rate_nm3hr"] = (
                round(br2.calculated_firing_rate_nm3hr, 2) if br2 else round(ng_flow2 / 1.1, 2)
            )
            resp["calculations"]["fuel2_equivalent_lph"] = round(
                equip2["burner"].get("equivalent_lph", 0), 2
            )
            resp["calculations"]["fuel2_fuel_density"] = equip2["burner"].get("fuel_density", 0)
            # Pipe-line size info for the Client Data Summary card. For oil
            # fuel-2 we use the oil pipe NB (sized from LPH); for gas we use
            # the standard pipe NB from calculate_pipe_sizes.
            f2_is_oil = req.fuel2_type in OIL_FUELS
            f2_oil_lph = (equip2["burner"].get("equivalent_lph", 0) if f2_is_oil else 0)
            f2_oil_nb = select_oil_pipe_nb(f2_oil_lph) if f2_is_oil else 0
            resp["pipes"]["fuel2_label"]  = FUEL_NAMES.get(req.fuel2_type, "Fuel 2")
            resp["pipes"]["fuel2_flow"]   = round(ng_flow2, 2)
            resp["pipes"]["fuel2_is_oil"] = f2_is_oil
            resp["pipes"]["fuel2_oil_lph"] = round(f2_oil_lph, 2) if f2_is_oil else None
            if pipes2 is not None:
                resp["pipes"]["fuel2_dia_mm"] = (
                    round(pipes2.ng_pipe_inner_dia_mm, 2) if not f2_is_oil else f2_oil_nb
                )
                resp["pipes"]["fuel2_nb"] = (
                    pipes2.ng_pipe_nb if not f2_is_oil else f2_oil_nb
                )
            if not f2_is_oil and equip2.get("ng_gas_train"):
                resp["pipes"]["fuel2_gas_train_flow"] = round(
                    equip2["ng_gas_train"]["max_flow"], 0
                )
                resp["pipes"]["fuel2_gas_train_model"] = (
                    f'{equip2["ng_gas_train"]["inlet_nb"]} x '
                    f'{equip2["ng_gas_train"]["outlet_nb"]}'
                )

        return resp
    except Exception as e:
        import traceback
        return {"error": str(e), "trace": traceback.format_exc()}


# ──────────────────────────────────────────────────────────────────────────
# HLPH Cost Variations — same idea as /api/cost-variations but runs the
# horizontal-ladle-preheater calc instead.
# ──────────────────────────────────────────────────────────────────────────
@app.post("/api/hlph-cost-variations")
def hlph_cost_variations(req: CostVariationsRequest):
    """Run hlph_calculate for every single-axis swap + markup level and
    return each alternative's grand total. Same response shape as
    /api/cost-variations."""
    OIL_FUELS = {"hsd", "ldo", "hdo", "fo", "sko", "cfo", "lshs"}
    is_oil_offer = (
        req.fuel1_type in OIL_FUELS
        or (req.fuel2_type and req.fuel2_type in OIL_FUELS)
    )

    LABEL_MAP = {
        "control_valve_vendor": {
            "_axis": "Control Valve Vendor",
            "dembla": "DEMBLA", "aira": "AIRA", "cair": "CAIR",
        },
        "shutoff_valve_vendor": {
            "_axis": "Shutoff Valve Vendor",
            "dembla": "DEMBLA", "aira": "AIRA", "cair": "CAIR",
        },
        "butterfly_valve_vendor": {
            "_axis": "Butterfly Valve Vendor",
            "lt_lever": "L&T Lever", "lt_gear": "L&T Gear",
        },
        "pressure_gauge_vendor": {
            "_axis": "Pressure Gauge Vendor",
            "baumer": "BAUMER", "hguru": "HGURU",
        },
        "hpu_variant": {
            "_axis": "HPU Variant",
            "Simplex": "Simplex", "Duplex 1": "Duplex-I", "Duplex 2": "Duplex-II",
        },
        "auto_control_type": {
            "_axis": "Auto Control Type",
            "plc": "PLC", "plc_agr": "PLC + AGR", "pid": "PID",
        },
        "control_mode": {
            "_axis": "Control Mode",
            "automatic": "Automatic", "manual": "Manual",
        },
        "special_auto_ignition": {
            "_axis": "Auto Ignition",
            True: "Auto Ignition ON", False: "Auto Ignition OFF",
        },
    }

    # HLPH does NOT vary by hood_type (horizontal uses a trolley drive,
    # not a swivel/up-down hood mechanism).
    axes = [
        ("control_valve_vendor",   ["dembla", "aira", "cair"]),
        ("shutoff_valve_vendor",   ["dembla", "aira", "cair"]),
        ("butterfly_valve_vendor", ["lt_lever", "lt_gear"]),
        ("pressure_gauge_vendor",  ["baumer", "hguru"]),
        ("special_auto_ignition",  [True, False]),
    ]
    if is_oil_offer:
        axes.append(("hpu_variant", ["Simplex", "Duplex 1", "Duplex 2"]))
    if req.control_mode == "automatic":
        axes.append(("auto_control_type", ["plc", "plc_agr", "pid"]))
    other_mode = "manual" if req.control_mode == "automatic" else "automatic"
    axes.append(("control_mode", [other_mode]))

    def _to_calc_req(modified):
        d = modified.dict()
        d.pop("markup", None)
        return VLPHCalcRequest(**d)

    def _calc_costs(modified):
        try:
            resp = hlph_calculate(_to_calc_req(modified))
            cs = resp.get("cost_summary", {})
            return (
                float(cs.get("bought_out_total") or 0),
                float(cs.get("encon_total") or 0),
            )
        except Exception:
            return None

    def _grand(bought, encon, markup):
        return bought * markup + encon

    base = _calc_costs(req)
    if base is None:
        return {"error": "Could not compute baseline cost"}
    base_bought, base_encon = base
    base_total = _grand(base_bought, base_encon, req.markup)
    if base_total <= 0:
        return {"error": "Baseline cost is zero — check inputs"}

    variations = []
    best_per_axis = {}

    for field, alts in axes:
        cur_value = getattr(req, field)
        for alt in alts:
            if alt == cur_value:
                continue
            modified = req.copy(update={field: alt})
            costs = _calc_costs(modified)
            if costs is None:
                continue
            total = _grand(costs[0], costs[1], req.markup)
            if total <= 0:
                continue
            saves = base_total - total
            if saves <= 0:
                continue
            label = LABEL_MAP.get(field, {})
            variations.append({
                "axis":  label.get("_axis", field),
                "value": label.get(alt, str(alt)),
                "field": field,
                "alt":   alt,
                "total": round(total, 2),
                "saves": round(saves, 2),
            })
            cur_best = best_per_axis.get(field)
            if cur_best is None or total < cur_best[1]:
                best_per_axis[field] = (alt, total)

    for mk in [1.75, 1.70, 1.65, 1.60]:
        if abs(mk - req.markup) < 1e-6:
            continue
        total = _grand(base_bought, base_encon, mk)
        saves = base_total - total
        if saves <= 0:
            continue
        variations.append({
            "axis":  "Markup",
            "value": f"{mk:.2f}×",
            "field": "markup",
            "alt":   mk,
            "total": round(total, 2),
            "saves": round(saves, 2),
        })

    variations.sort(key=lambda r: r["saves"], reverse=True)

    best_combined = None
    update = {field: best for field, (best, _) in best_per_axis.items()}
    combined_req = req.copy(update=update) if update else req
    costs = _calc_costs(combined_req)
    if costs is not None:
        lowest_markup = 1.60 if req.markup > 1.60 else req.markup
        combined_total = _grand(costs[0], costs[1], lowest_markup)
        if combined_total > 0 and combined_total < base_total:
            swaps = []
            for field, (alt, _) in best_per_axis.items():
                label = LABEL_MAP.get(field, {})
                swaps.append({
                    "axis":  label.get("_axis", field),
                    "value": label.get(alt, str(alt)),
                })
            if abs(lowest_markup - req.markup) > 1e-6:
                swaps.append({"axis": "Markup", "value": f"{lowest_markup:.2f}×"})
            best_combined = {
                "total": round(combined_total, 2),
                "saves": round(base_total - combined_total, 2),
                "swaps": swaps,
            }

    return {
        "current_total":  round(base_total, 2),
        "current_markup": round(req.markup, 2),
        "variations":     variations,
        "best_combined":  best_combined,
    }


# ── Box Type Furnace ────────────────────────────────────────────────────────

@app.get("/btf", response_class=HTMLResponse)
def btf_costing_form():
    html_path = os.path.join(BASE_DIR, "btf_costing.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return f.read()


class BTFCalcRequest(BaseModel):
    combustion_mode: str = "onoff"
    markup: float = 1.8


@app.post("/api/btf-calculate")
def btf_calculate(req: BTFCalcRequest):
    try:
        from bom.btf_builder import build_btf_df, get_supplementary
        import json
        df, summary = build_btf_df(combustion_mode=req.combustion_mode, markup=req.markup)
        bom = json.loads(df.to_json(orient="records"))
        return {
            "bom": bom,
            "cost_summary": summary,
            "supplementary": get_supplementary(),
        }
    except Exception as e:
        import traceback
        return {"error": str(e), "detail": traceback.format_exc()}


# ── SNSF BRF (Billet Reheating Furnace) ─────────────────────────────────────

@app.get("/snsf-brf", response_class=HTMLResponse)
def snsf_brf_costing_form():
    html_path = os.path.join(BASE_DIR, "snsf_brf_costing.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return f.read()


class SNSFBRFCalcRequest(BaseModel):
    include_ng_optional: bool = False
    include_client_scope: bool = False


@app.post("/api/snsf-brf-calculate")
def snsf_brf_calculate(req: SNSFBRFCalcRequest):
    try:
        from bom.snsf_brf_builder import build_snsf_brf_df, get_supplementary
        import json
        df, summary = build_snsf_brf_df(
            include_ng_optional=req.include_ng_optional,
            include_client_scope=req.include_client_scope,
        )
        bom = json.loads(df.to_json(orient="records"))
        return {
            "bom": bom,
            "cost_summary": summary,
            "supplementary": get_supplementary(),
        }
    except Exception as e:
        import traceback
        return {"error": str(e), "detail": traceback.format_exc()}


# ── Recuperator (USK Exports model) ─────────────────────────────────────────

@app.get("/recup", response_class=HTMLResponse)
def recup_costing_form():
    html_path = os.path.join(BASE_DIR, "recup_costing.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())


class RecupCalcRequest(BaseModel):
    # Flue gas (only flow + inlet temp are inputs; mass and final temp are derived)
    flue_flow_nm3hr:    float = 1900.0
    flue_temp_in_C:     float = 800.0
    cp_flue_kcal_kgC:   float = 0.23
    # Combustion air
    air_volume_nm3hr:   float = 1750.0
    air_temp_in_C:      float = 35.0
    air_temp_out_C:     float = 400.0
    cp_air_kcal_kgC:    float = 0.247
    # Geometry / overrides
    heat_transfer_coef: float = 30.0
    pipe_dia_mm:        float = 48.3
    pipe_thick_mm:      float = 2.77
    pipe_kg_per_m:      float = 3.16
    pipe_length_m_per_bank: float = 0.63
    bank_gap_mm:        float = 150.0
    pipes_total_override: int = 0    # 0 = auto-derive from surface area
    hot_bank_material:    str = "SS" # "SS" (Rs 250/kg) or "MS" (Rs 70/kg)
    cold_bank_material:   str = "SS" # can differ from hot bank
    side_hood_kg:         float = 1500.0   # MS side hood weight
    cai_rate_override:    float = 0.0      # CAI Assembly Rs/kg override; 0 = use DB default


@app.post("/api/recup-calculate")
def recup_calculate(req: RecupCalcRequest):
    try:
        from calculations.recup import RecupInputs, calculate_recup
        from bom.recup_builder import build_recup_df, recup_summary, _load_rates

        results = calculate_recup(RecupInputs(
            flue_flow_nm3hr=req.flue_flow_nm3hr,
            flue_temp_in_C=req.flue_temp_in_C,
            cp_flue_kcal_kgC=req.cp_flue_kcal_kgC,
            air_volume_nm3hr=req.air_volume_nm3hr,
            air_temp_in_C=req.air_temp_in_C,
            air_temp_out_C=req.air_temp_out_C,
            cp_air_kcal_kgC=req.cp_air_kcal_kgC,
            heat_transfer_coef=req.heat_transfer_coef,
            pipe_dia_mm=req.pipe_dia_mm,
            pipe_thick_mm=req.pipe_thick_mm,
            pipe_kg_per_m=req.pipe_kg_per_m,
            pipe_length_m_per_bank=req.pipe_length_m_per_bank,
            bank_gap_mm=req.bank_gap_mm,
            pipes_total_override=req.pipes_total_override,
            hot_bank_material=req.hot_bank_material,
            cold_bank_material=req.cold_bank_material,
            side_hood_kg=req.side_hood_kg,
            cai_rate_override=req.cai_rate_override,
        ))

        rates = _load_rates()
        df = build_recup_df(results, rates)
        summary = recup_summary(results, rates)

        detail = df[df["MEDIA"] != ""].copy()

        return {
            "calculations": {
                "flue_mass_kghr":       results.flue_mass_kghr,
                "heat_required_kcal":   results.heat_required_kcal,
                "flue_temp_out_C":      results.flue_temp_out_C,
                "lmtd_C":               results.lmtd_C,
                "surface_area_m2":      results.surface_area_m2,
                "pipes_total_raw":      results.pipes_total_raw,
                "pipes_total":          results.pipes_total,
                "pipes_in_row":         results.pipes_in_row,
                "pipes_in_column":      results.pipes_in_column,
                "pipes_per_bank":       results.pipes_per_bank,
                "bank_length_mm":       results.bank_length_mm,
                "bank_width_mm":        results.bank_width_mm,
                "weight_per_pipe_kg":   results.weight_per_pipe_kg,
                "weight_hot_bank_kg":   results.weight_hot_bank_kg,
                "weight_cold_bank_kg":  results.weight_cold_bank_kg,
                "weight_total_pipes_kg": results.weight_total_pipes_kg,
                "ms_outer_shell_kg":    results.ms_outer_shell_kg,
                "ms_air_inlet_duct_kg": results.ms_air_inlet_duct_kg,
                "ms_hot_outlet_duct_kg": results.ms_hot_outlet_duct_kg,
                "ms_pipe_holding_kg":   results.ms_pipe_holding_kg,
                "ms_bottom_box_kg":     results.ms_bottom_box_kg,
                "hot_bank_material":    results.hot_bank_material,
                "cold_bank_material":   results.cold_bank_material,
                "cai_rate_override":    results.cai_rate_override,
                # Echo of inputs — the offer template needs these for the
                # Designing Parameters table (flue/air conditions are
                # process inputs, not derived values).
                "flue_flow_nm3hr":      req.flue_flow_nm3hr,
                "flue_temp_in_C":       req.flue_temp_in_C,
                "air_volume_nm3hr":     req.air_volume_nm3hr,
                "air_temp_in_C":        req.air_temp_in_C,
                "air_temp_out_C":       req.air_temp_out_C,
                "pipe_dia_mm":          req.pipe_dia_mm,
                "pipe_thick_mm":        req.pipe_thick_mm,
                "pipe_length_m_per_bank": req.pipe_length_m_per_bank,
            },
            "bom": detail[["MEDIA","ITEM NAME","REFERENCE","QTY","MAKE","UNIT PRICE","TOTAL"]].to_dict(orient="records"),
            "cost_summary": summary,
        }
    except Exception as e:
        import traceback
        return {"error": str(e), "trace": traceback.format_exc()}


class RecupQuoteRequest(BaseModel):
    # Customer / project
    salutation:      str = ""
    poc_name:        str = ""
    poc_designation: str = ""
    company_name:    str = ""
    company_address: str = ""
    email:           str = ""
    mobile_no:       str = ""
    project_name:    str = ""
    subject:         str = ""
    application:     str = ""
    client_enq_ref:  str = ""
    # Marketing / ref
    marketing_person: str = ""
    marketing_email:  str = ""
    marketing_phone:  str = ""
    technical_person: str = ""
    technical_phone:  str = ""
    technical_email:  str = ""
    location:         str = "FBD"
    # Supervision charges (VLPH-style supervision table in Annexure III)
    supervision_mech: str = ""
    supervision_plc:  str = ""
    # T&C
    tnc_prices:             str = ""
    tnc_delivery:           str = ""
    tnc_gst:                str = ""
    tnc_hsn_code:           str = ""
    tnc_pan_gst:            str = ""
    tnc_payment_terms:      str = ""
    tnc_packing_forwarding: str = ""
    tnc_freight:            str = ""
    tnc_transit_insurance:  str = ""
    tnc_validity:           str = ""
    tnc_inspection:         str = ""
    tnc_guarantee:          str = ""
    # Recup calc payload
    calculations: dict = {}
    bom:          list = []        # the full BOM detail rows for the price schedule
    final_total:  float = 0.0
    grand_total:  float = 0.0
    transport_amt: float = 0          # Transport (flat Rs.) — own price-schedule line
    qty:          int = 1
    # Price schedule style — "single" (VLPH-like, one row + amount in words)
    # or "full" (Viraj-like, every BOM line + supervision + totals)
    price_schedule_style: str = "single"
    # Supervision charges row (optional — only meaningful in 'full' style)
    supervision_include: bool = False
    supervision_rate:    str = "Rs. 10,500 / Man / Day"
    supervision_note:    str = "Plus: To and Fro fare from Delhi to Site, Boarding and Lodging, Local Conveyance, Medical assistance"


def _material_of_construction(hot_mat: str, cold_mat: str, pipe_dia_mm: float = 48.3) -> dict:
    """Translate Step-2's Hot/Cold tube-material picks into the four
    placeholders the Material of Construction table in Annexure I needs.
    SS -> 'SS-316' (typical hot-bank grade), MS -> 'Mild Steel (CS)'.
    The Tube line carries the bend/straight + size + schedule spec."""
    def block(mat: str | None, bend: bool):
        m = (mat or "SS").upper()
        if m == "MS":
            plate = "Mild Steel (CS)"
            tube  = f"CS, ERW {'Bended' if bend else 'Straight'} pipes, OD: {pipe_dia_mm}, Schedule-40"
        else:
            plate = "SS-316"
            tube  = f"SS-316 {'Bended' if bend else 'Straight'}, ERW pipes, OD: {pipe_dia_mm}, Schedule-40"
        return plate, tube
    hot_plate, hot_tube   = block(hot_mat,  bend=True)   # hot bank is bent (for expansion)
    cold_plate, cold_tube = block(cold_mat, bend=False)  # cold bank is straight
    return {
        "hot_tube_plate_material":  hot_plate,
        "hot_tube_material":        hot_tube,
        "cold_tube_plate_material": cold_plate,
        "cold_tube_material":       cold_tube,
    }


def _bom_rows_for_offer(bom: list, *, supervision_include: bool,
                        supervision_rate: str, supervision_note: str,
                        single_line_total: float = 0.0,
                        single_line_mode: bool = False) -> dict:
    """Translate the calculator's BOM detail (12 line items + 3 summary
    rows) into the placeholders the offer's Price Schedule needs:
      - bom_rows               iterable list of {sno,item,qty,unit_price,total}
      - bought_out_total       formatted ₹
      - encon_total            formatted ₹
      - grand_total            formatted ₹
      - grand_total_in_words   Indian-English words
      - supervision_*          row content (only rendered when flag is on)
    Summary rows from the BOM dataframe (MEDIA == '') are skipped here;
    we recompute totals from the line items so user-edited unit prices
    on Step 3 propagate to the offer.
    """
    from engine.quote_writer import amount_in_words_indian as _words, _format_inr

    rows = []
    bought_out = 0.0
    encon      = 0.0
    for r in bom or []:
        media = (r.get("MEDIA") or "").strip()
        if not media:
            continue  # skip the BOUGHT OUT / ENCON / GRAND TOTAL summary rows
        total = float(r.get("TOTAL") or 0)
        unit  = float(r.get("UNIT PRICE") or 0)
        qty   = r.get("QTY") or 1
        rows.append({
            "sno":        len(rows) + 1,
            "item":       r.get("ITEM NAME", ""),
            "qty":        qty,
            "unit_price": _format_inr(unit),
            "total":      _format_inr(total),
        })
        if media == "ENCON ITEMS":
            encon += total
        else:
            bought_out += total
    grand = bought_out + encon
    # In single-line mode the footer should show the FINAL price (the one
    # rendered in the lone Price Schedule row), not the raw BOM sum.
    footer_total = single_line_total if single_line_mode else grand
    return {
        "bom_rows":            rows,
        "bought_out_total":    _format_inr(bought_out),
        "encon_total":         _format_inr(encon),
        "grand_total":         _format_inr(footer_total),
        "grand_total_in_words": f"INR. {_words(footer_total)} ONLY.",
        "supervision_include": bool(supervision_include),
        "supervision_rate":    supervision_rate or "",
        "supervision_note":    supervision_note or "",
    }


@app.post("/api/generate-recup-quote")
def generate_recup_quote(req: RecupQuoteRequest):
    """Render Recup_Offer_Template.docx with payload data, save to
    quotes/, and optionally convert to PDF.

    Template note: the recup template is cloned from the VLPH offer
    template (Offer_Template.docx) via build_recup_template_from_vlph.py
    — see that script for the full layout (cover page, Annexures, MoC
    sub-table + 3D image block + Designing Parameters table + scope)."""
    try:
        from datetime import datetime as _dt
        from docxtpl import DocxTemplate
        from engine.quote_writer import amount_in_words_indian, _format_inr

        seq = next_quote_seq()
        full_ref = build_enquiry_ref(seq, req.technical_person or "", req.location or "")
        # Split full ref (e.g. "ENCON.04026.050/FBD/JS DT.16/05/2026")
        # into the short part and the date.
        short_ref = full_ref.split(" DT.")[0]
        date_str = full_ref.split(" DT.")[-1] if " DT." in full_ref else _dt.now().strftime("%d/%m/%Y")

        c = req.calculations or {}
        unit_price = float(req.final_total or 0)
        qty = max(1, int(req.qty or 1))
        total_price = unit_price * qty

        ctx = {
            "project_name":     req.project_name or (f"Recuperator for {req.application}" if req.application else "Recuperator"),
            "subject":          req.subject or (f"Offer for Recuperator — {req.application}" if req.application else "Offer for Recuperator"),
            "application":      req.application or "Furnace",
            # VLPH-style top-of-document equipment name (renders on cover
            # page + 'About the Equipment' section + Annexure I banner).
            "equipment_name":   f"Recuperator for {req.application}" if req.application else "Recuperator",
            "company_name":     req.company_name,
            "company_address":  req.company_address,
            "email":            req.email,
            "mobile_no":        req.mobile_no,
            "poc_name":         _with_salutation(req.salutation, req.poc_name),
            "poc_greeting":     _greeting(req.salutation),
            "poc_designation":  req.poc_designation or "",
            "client_enq_ref":   req.client_enq_ref,
            "enquiry_ref":      full_ref,
            "enquiry_ref_short": short_ref,
            "enquiry_date_str":  date_str,
            "marketing_person": req.marketing_person,
            "marketing_email":  req.marketing_email,
            "marketing_phone":  req.marketing_phone,
            # Technical person (cover-letter signature in VLPH layout)
            "technical_person": req.technical_person or req.marketing_person,
            "technical_phone":  req.technical_phone  or req.marketing_phone,
            "technical_email":  req.technical_email  or req.marketing_email,
            # Supervision charges - VLPH renders these in a small sub-table
            # right after the Price Schedule. For recup we expose mech +
            # plc lines (PLC line may stay blank if not in scope).
            "supervision_mech": req.supervision_mech or "",
            "supervision_plc":  req.supervision_plc  or "",
            # Designing parameters table
            "flue_flow_nm3hr":  f"{c.get('flue_flow_nm3hr', 0):,.0f}" if c.get('flue_flow_nm3hr') else "",
            "flue_temp_in_C":   f"{int(round(c.get('flue_temp_in_C', 0)))}" if c.get('flue_temp_in_C') else "",
            "flue_temp_out_C":  f"{int(round(c.get('flue_temp_out_C', 0)))}" if c.get('flue_temp_out_C') else "",
            "air_volume_nm3hr": f"{c.get('air_volume_nm3hr', 0):,.0f}" if c.get('air_volume_nm3hr') else "",
            "air_temp_in_C":    f"{int(round(c.get('air_temp_in_C', 0)))}" if c.get('air_temp_in_C') else "",
            "air_temp_out_C":   f"{int(round(c.get('air_temp_out_C', 0)))}" if c.get('air_temp_out_C') else "",
            # Marketing-friendly rounded-up version of the air outlet
            # temperature: ceil to the next multiple of 50 °C. Used in
            # the Annexure I description paragraph (e.g. an exact
            # value of 412 °C reads as 'about 450 °C').
            "air_temp_out_C_rounded": (
                f"{int(((c.get('air_temp_out_C', 0) + 49.999) // 50) * 50)}"
                if c.get('air_temp_out_C') else ""
            ),
            "surface_area_m2":  f"{c.get('surface_area_m2', 0):.2f}",
            "pipe_dia_mm":      f"{c.get('pipe_dia_mm', 48.3):.1f}" if c.get('pipe_dia_mm') else "48.3",
            "pipe_length_m":    f"{c.get('pipe_length_m_per_bank', 0.63):.2f}" if c.get('pipe_length_m_per_bank') else "0.63",
            "pipe_thick_mm":    f"{c.get('pipe_thick_mm', 2.77):.2f}" if c.get('pipe_thick_mm') else "2.77",
            # Material of Construction (Annexure I — auto-derived from
            # the engineer's Hot/Cold material picks on Step 2).
            **_material_of_construction(c.get('hot_bank_material'),
                                        c.get('cold_bank_material'),
                                        c.get('pipe_dia_mm', 48.3)),
            # Price schedule (legacy single-line placeholders — kept for
            # backwards compatibility in case any text still references them)
            "recup_qty":           f"{qty:02d} No.",
            "recup_unit_price":    _format_inr(unit_price),
            "recup_total_price":   _format_inr(total_price),
            "recup_total_in_words": f"INR. {amount_in_words_indian(total_price)} ONLY.",
            # Price-schedule mode flag (template branches on this)
            "price_schedule_style": (req.price_schedule_style or "single").lower(),
            # Full-BOM iterable price schedule (Annexure III, 'full' mode)
            **_bom_rows_for_offer(req.bom or [],
                                  supervision_include=req.supervision_include,
                                  supervision_rate=req.supervision_rate,
                                  supervision_note=req.supervision_note,
                                  single_line_total=total_price,
                                  single_line_mode=(req.price_schedule_style or 'single').lower() == 'single'),
            # T&C
            "tnc_prices":             req.tnc_prices,
            "tnc_delivery":           req.tnc_delivery,
            "tnc_gst":                req.tnc_gst,
            "tnc_hsn_code":           req.tnc_hsn_code,
            "tnc_pan_gst":            req.tnc_pan_gst,
            "tnc_payment_terms":      req.tnc_payment_terms,
            "tnc_packing_forwarding": req.tnc_packing_forwarding,
            "tnc_freight":            req.tnc_freight,
            "tnc_transit_insurance":  req.tnc_transit_insurance,
            "tnc_validity":           req.tnc_validity,
            "tnc_inspection":         req.tnc_inspection,
            "tnc_guarantee":          req.tnc_guarantee,
        }

        tpl_path = os.path.join(BASE_DIR, "Recup_Offer_Template.docx")
        tpl = DocxTemplate(tpl_path)

        # 3D recuperator images (Annexure I) — embed only if the JPEGs are
        # present in static/recup/. Missing files render as blank paragraphs
        # rather than crashing the whole offer.
        from docxtpl import InlineImage
        from docx.shared import Mm
        _img_dir = os.path.join(BASE_DIR, "static", "recup")
        for key, fname, width_mm in (
            ("image_recup_side",  "recup_3d_side.jpeg",  120),
            ("image_recup_front", "recup_3d_front.jpeg", 120),
            ("image_recup_top",   "recup_3d_top.jpeg",   120),
        ):
            fpath = os.path.join(_img_dir, fname)
            ctx[key] = InlineImage(tpl, fpath, width=Mm(width_mm)) if os.path.exists(fpath) else ""

        tpl.render(ctx)

        # Save with a safe, sequenced filename
        safe_company = "".join(ch for ch in (req.company_name or "Client") if ch.isalnum() or ch in " _-").strip().replace(" ", "_") or "Client"
        docx_name = f"Recup_Offer_{safe_company}_{seq}.docx"
        docx_path = os.path.join(QUOTES_FOLDER, docx_name)
        tpl.save(docx_path)

        # PDF (best effort — LibreOffice may not be present locally)
        pdf_name = docx_name.replace(".docx", ".pdf")
        pdf_path = os.path.join(QUOTES_FOLDER, pdf_name)
        pdf_ok = _docx_to_pdf(docx_path, pdf_path)

        # Mirror docx + pdf to Google Drive (recup-specific folder) in a
        # background thread. product_type='recuperator' triggers the
        # recup routing in drive_uploader._folder_id_for_product.
        # Failures are logged but never break offer generation.
        try:
            from engine.drive_uploader import upload_offer_async
            upload_offer_async(docx_path, docx_name, "recuperator")
            if pdf_ok:
                upload_offer_async(pdf_path, pdf_name, "recuperator")
        except Exception as _drv_err:
            print(f"WARN: drive upload kickoff failed: {_drv_err}")

        _log_quote(
            quote_no=full_ref, ref_no=req.client_enq_ref, company_name=req.company_name,
            poc_name=_with_salutation(req.salutation, req.poc_name), email=req.email,
            mobile_no=req.mobile_no, project_name=req.project_name,
            equipment_type="Recuperator", location=req.location, grand_total=total_price,
            marketing_person=req.marketing_person, technical_person=req.technical_person,
            file_path=docx_path,
        )

        return {
            "filename":     docx_name,
            "pdf_filename": pdf_name if pdf_ok else None,
            "download_url": f"/api/download-quote/{docx_name}",
            "pdf_url":      f"/api/pdf-quote/{pdf_name}" if pdf_ok else None,
            # Preview endpoint converts DOCX -> HTML via mammoth, so it
            # always uses the docx filename (works whether PDF rendered
            # or not).
            "preview_url":  f"/api/preview-quote/{docx_name}",
            "quote_no":     full_ref,
            "enquiry_ref":  full_ref,
            "final_total":  unit_price,
            "total_price":  total_price,
        }
    except Exception as e:
        import traceback
        return {"error": str(e), "trace": traceback.format_exc()}


# ── HPU stand-alone offer ────────────────────────────────────────────────────
class HpuCustomer(BaseModel):
    salutation: Optional[str] = ""
    name: Optional[str] = ""
    email: Optional[str] = ""
    phone: Optional[str] = ""
    company: Optional[str] = ""
    designation: Optional[str] = ""
    address: Optional[str] = ""
    city: Optional[str] = ""
    state: Optional[str] = ""
    pin: Optional[str] = ""
    gstin: Optional[str] = ""
    ref_no: Optional[str] = ""
    location: Optional[str] = ""
    subject: Optional[str] = ""
    marketing_salutation: Optional[str] = ""
    marketing: Optional[str] = ""
    marketing_email: Optional[str] = ""
    marketing_phone: Optional[str] = ""
    technical_salutation: Optional[str] = ""
    technical: Optional[str] = ""
    technical_email: Optional[str] = ""
    technical_phone: Optional[str] = ""
    # Annexure IV — Terms & Conditions (HPU Step 3, mirrored from VLPH).
    tnc_prices: Optional[str] = ""
    tnc_delivery: Optional[str] = ""
    tnc_gst: Optional[str] = ""
    tnc_hsn_code: Optional[str] = ""
    tnc_pan_gst: Optional[str] = ""
    tnc_payment_terms: Optional[str] = ""
    tnc_packing_forwarding: Optional[str] = ""
    tnc_freight: Optional[str] = ""
    tnc_transit_insurance: Optional[str] = ""
    tnc_validity: Optional[str] = ""
    tnc_inspection: Optional[str] = ""
    tnc_guarantee: Optional[str] = ""


class HpuQuoteRequest(BaseModel):
    customer: HpuCustomer
    hpu_variant: str
    hpu_kw: float
    qty: int = 1
    # Oil flow rate (LPH) is catalog-driven — the backend looks it up
    # from pumping_unit_price.flow_lph using (kw, variant). The client
    # may post it as a hint but it is overridden by the DB value.
    fuel_lph: Optional[float] = None
    # Legacy: kept optional so older clients posting fuel_type don't 400.
    # No longer surfaced in the offer doc.
    fuel_type: Optional[str] = ""
    # Commercial adjustments (Packaging & Forwarding, Designing, Negotiation,
    # Transport). final_total = the form's sell price with all of these applied;
    # transport_amt is broken onto its own price-schedule line. Default 0 so the
    # offer falls back to the catalog price when the form doesn't send them.
    final_total:   float = 0.0
    grand_total:   float = 0.0
    transport_amt: float = 0
    pf_pct:        float = 0    # Packaging & Forwarding (% of catalog price)
    design_pct:    float = 0    # Designing (%)
    neg_pct:       float = 0    # Negotiation (%, added)


@app.get("/api/hpu/flow-lph")
def hpu_flow_lph(kw: float, variant: str):
    """Look up the catalog oil flow rate (LPH) for a given (kW, variant)
    from pumping_unit_price.flow_lph. Used by the HPU form to auto-fill
    the Oil Flow Rate field so the user can't type a wrong number."""
    try:
        conn = sqlite3.connect(DB_PATH)
        row = conn.execute(
            "SELECT flow_lph, model_code, sell_price FROM pumping_unit_price "
            "WHERE unit_kw = ? AND variant = ? LIMIT 1",
            (int(round(kw)), variant),
        ).fetchone()
        conn.close()
        if not row:
            return {"error": f"no row for {variant} @ {kw} kW", "flow_lph": None}
        return {
            "flow_lph":   float(row[0]) if row[0] is not None else None,
            "model_code": row[1] or "",
            "sell_price": float(row[2]) if row[2] is not None else None,
        }
    except Exception as e:
        return {"error": str(e), "flow_lph": None}


@app.get("/api/hpu/price")
def hpu_price(kw: float, variant: str, mode: str = "hpu"):
    """Catalog sell price for an HPU/PU (kW, variant) so the form can show a live
    total before generating. HPU = SUM(hpu_master.amount) × 1.8;
    PU = pumping_unit_price.sell_price."""
    try:
        kw_int = int(round(kw))
        conn = sqlite3.connect(DB_PATH)
        HPU_MARKUP = _product_markup(conn, "hpu", "markup", 1.8)   # editable
        if (mode or "hpu").lower() == "pu":
            row = conn.execute(
                "SELECT sell_price FROM pumping_unit_price WHERE unit_kw = ? AND variant = ? LIMIT 1",
                (kw_int, variant)).fetchone()
            conn.close()
            if not row or row[0] is None:
                return {"error": "no price", "unit_price": None}
            return {"unit_price": float(row[0])}
        from bom.hpu_pricelist import hpu_material_cost
        material = hpu_material_cost(conn, kw_int, variant)   # pricelist-linked
        conn.close()
        if not material:
            return {"error": "no price", "unit_price": None}
        return {"unit_price": round(material * HPU_MARKUP, 2)}
    except Exception as e:
        return {"error": str(e), "unit_price": None}


def _generate_pumping_unit_offer(req: "HpuQuoteRequest", *, mode: str) -> dict:
    """Shared body for the HPU and PU stand-alone offer endpoints.

    The two products share the form, the template shell and the
    render path; only price source, model-code prefix, equipment
    label, filename and template file differ. `mode` is "hpu" or
    "pu" and picks all five.
    """
    from engine.quote_engine import calculate_quote
    from engine.quote_writer import generate_quote_docx
    from bom.selectors.hpu_selector import VARIANT_PREFIX, PU_VARIANT_PREFIX

    cust = req.customer
    kw_int = int(round(req.hpu_kw))

    # ── 1. Unit price + catalog flow rate ──────────────────────────────
    # Flow rate (LPH) is identical for HPU and PU since the pump skid
    # is the same — read from pumping_unit_price.flow_lph either way.
    conn = sqlite3.connect(DB_PATH)
    HPU_MARKUP = _product_markup(conn, "hpu", "markup", 1.8)   # editable
    flow_row = conn.execute(
        "SELECT flow_lph FROM pumping_unit_price "
        "WHERE unit_kw = ? AND variant = ? LIMIT 1",
        (kw_int, req.hpu_variant),
    ).fetchone()

    if mode == "hpu":
        # HPU sell price = pricelist-linked material cost × HPU_MARKUP.
        # Material cost is Σ(qty × live pricelist rate) + labour, shared with
        # the Internal-Costing HPU tab (bom/hpu_pricelist.hpu_material_cost).
        from bom.hpu_pricelist import hpu_material_cost
        material = hpu_material_cost(conn, kw_int, req.hpu_variant)
        conn.close()
        if not material:
            return {"error": f"No HPU rows in hpu_master for "
                             f"{req.hpu_variant} @ {req.hpu_kw} kW."}
        unit_price = round(material * HPU_MARKUP, 2)
        model_code = f"{VARIANT_PREFIX[req.hpu_variant]}-{kw_int}"
        equipment_label   = "Heating and Pumping Unit"
        item_product_type = "Hydraulic Pumping Unit"   # gates preheater scope
        filename_infix    = "HPU"
        template_name     = "HPU_Offer_Template.docx"
    else:  # mode == "pu"
        # PU sell price is already markup-included in pumping_unit_price
        # (rebuilt from hpu_master minus heater/thermostat × 1.8).
        price_row = conn.execute(
            "SELECT sell_price FROM pumping_unit_price "
            "WHERE unit_kw = ? AND variant = ? LIMIT 1",
            (kw_int, req.hpu_variant),
        ).fetchone()
        conn.close()
        if not price_row or price_row[0] is None:
            return {"error": f"No PU price in pumping_unit_price for "
                             f"{req.hpu_variant} @ {req.hpu_kw} kW."}
        unit_price = float(price_row[0])
        model_code = f"{PU_VARIANT_PREFIX[req.hpu_variant]}-{kw_int}"
        equipment_label   = "Pumping Unit"
        item_product_type = "Pumping Unit"
        filename_infix    = "PU"
        template_name     = "PU_Offer_Template.docx"

    catalog_lph = float(flow_row[0]) if flow_row and flow_row[0] is not None \
                  else (req.fuel_lph or 0.0)
    qty = max(1, int(req.qty or 1))
    # Commercial adjustments applied to the catalog price: Packaging & Forwarding
    # and Designing add, Negotiation all add; Transport is a flat amount shown
    # on its own price-schedule line. The per-unit offer price folds Transport in
    # (so _break_out_transport can split it back out below). Final total rounded
    # to the nearest Rs.1000, matching the other equipment offers.
    _pf  = float(getattr(req, "pf_pct", 0) or 0)
    _des = float(getattr(req, "design_pct", 0) or 0)
    _neg = float(getattr(req, "neg_pct", 0) or 0)
    _trn = float(getattr(req, "transport_amt", 0) or 0)
    _adj_total  = unit_price * qty * (1 + (_pf + _des + _neg) / 100) + _trn
    _final_incl = (round(_adj_total / 1000) * 1000) if (_pf or _des or _neg or _trn) else _adj_total
    offer_unit  = _final_incl / qty

    # ── 2. Enquiry ref (canonical ENCON pattern) ──────────────────────
    seq = next_quote_seq()
    auto_ref = build_enquiry_ref(seq, cust.technical or "", cust.location or "")
    # OUR REF from the user's input (Enquiry / Ref No.) + today's date; the
    # quote_writer splits "<ref> DT.<date>" into enquiry_ref_short / date.
    _ref_in = (cust.ref_no or "").split(" DT.")[0].strip()
    auto_ref = (f"{_ref_in} DT.{datetime.now().strftime('%d/%m/%Y')}"
                if _ref_in else auto_ref)

    # ── 3. Build form_data mirroring VLPH/Tundish shape ──────────────
    equipment_name = (
        f"{equipment_label} – {req.hpu_variant}, "
        f"{req.hpu_kw:g} kW @ {catalog_lph:g} LPH"
    )
    form_data = {
        "quote_seq": seq,
        "customer": {
            "company_name":     cust.company or "",
            "company_city":     cust.city or "",
            "company_state":    cust.state or "",
            "address":          ", ".join(filter(None, [
                                    cust.address or "",
                                    cust.city or "",
                                    cust.state or "",
                                    cust.pin or "",
                                ])),
            "poc_name":         _with_salutation(cust.salutation, cust.name),
            "poc_greeting":     _greeting(cust.salutation),
            "poc_designation":  cust.designation or "",
            "mobile_no":        cust.phone or "",
            "email":            cust.email or "",
            "project_name":     cust.subject or equipment_name,
            "subject":          cust.subject or equipment_name,
            "ref_no":           auto_ref,
            "your_ref":         cust.ref_no or auto_ref,
            "enquiry_ref":      auto_ref,
            "marketing_person": _with_salutation(cust.marketing_salutation, cust.marketing),
            "marketing_phone":  cust.marketing_phone or "",
            "marketing_email":  cust.marketing_email or "",
            "technical_person": _with_salutation(cust.technical_salutation, cust.technical),
            "technical_phone":  cust.technical_phone or "",
            "technical_email":  cust.technical_email or "",
            "gstin":            cust.gstin or "",
            # Picked up by _build_equipment_name + by the HPU/PU template
            # via {{ hpu_variant }} / {{ hpu_kw }} / {{ hpu_lph }} /
            # {{ hpu_qty }} (the PU template reuses the hpu_* keys so a
            # single context dict serves both).
            "equipment_name_override": equipment_name,
            "hpu_variant":      req.hpu_variant,
            "hpu_kw":           f"{req.hpu_kw:g}",
            "hpu_lph":          f"{catalog_lph:g}",
            "hpu_model":        model_code,
            "hpu_qty":          str(qty),
            # Preheater-specific tech-data: blank so _strip_empty_tech_rows
            # drops the rows entirely.
            "ladle_tons":          "", "ladle_dim":           "",
            "ladle_drawing_no":    "", "refractory_weight_kg": "",
            "heating_schedule":    "", "heating_time":        "",
            "fuel_cv":             "", "fuel_consumption":    "",
            "fuel2_cv":            "", "fuel2_consumption":   "",
            "burner_model":        "", "blower_model":        "",
            "blower_size":         "", "blower_capacity":     "",
            "hydraulic_motor_hp":  "", "max_electrical_load": "",
            "fuel_name":           req.fuel_type or "",
            "burner_capacity_range": "",
            "pumping_unit":        f"{req.hpu_variant}, {req.hpu_kw:g} kW",
            "hood_movement":       "", "hood_type":           "",
            "pilot_gas_type":      "", "ignition_method":     "",
            "num_burners":         "",
            "max_fuel_consumption1": "", "max_fuel_consumption2": "",
            "is_oil":              True,
            "is_dual":             False,
            # PU offer = pumping only (no heater); HPU offer = heating + pumping.
            "force_pumping_only":  (filename_infix == "PU"),
            # Stand-alone equipment offer → use the detailed ENCON pumping-unit
            # scope wording + rich price description (quote_writer gates on this).
            "standalone_pumping":  True,
            "control_mode":        "manual",
            "auto_control_type":   "",
            "control_valve_type":  "",
            "special_auto_ignition": False,
            "special_auto_controls": False,
            "vertical_qty":        qty,
            "horizontal_qty":      0,
            "nitrogen_purging":    False,
            "burner_kw_value":     "",
            "bom_items":           [],
            # Annexure IV — T&Cs sourced from Step 3 (mirrors VLPH).
            "tnc_prices":             cust.tnc_prices or "",
            "tnc_delivery":           cust.tnc_delivery or "",
            "tnc_gst":                cust.tnc_gst or "",
            "tnc_hsn_code":           cust.tnc_hsn_code or "",
            "tnc_pan_gst":            cust.tnc_pan_gst or "",
            "tnc_payment_terms":      cust.tnc_payment_terms or "",
            "tnc_packing_forwarding": cust.tnc_packing_forwarding or "",
            "tnc_freight":            cust.tnc_freight or "",
            "tnc_transit_insurance":  cust.tnc_transit_insurance or "",
            "tnc_validity":           cust.tnc_validity or "",
            "tnc_inspection":         cust.tnc_inspection or "",
            "tnc_guarantee":          cust.tnc_guarantee or "",
        },
        # Single-line price schedule. The product_type isn't 'Vertical
        # Ladle Preheater' / 'Horizontal Ladle Preheater' / 'Tundish'
        # so quote_writer's is_vertical / is_horizontal / is_tundish
        # flags stay False and the preheater scope sections are
        # suppressed; the total still pours into the vertical price
        # slot via the 'else' branch at quote_writer:692.
        "items": [{
            "product_type": item_product_type,
            "model":        f"{req.hpu_variant} {req.hpu_kw:g} kW",
            "description":  equipment_name,
            "qty":          qty,
            # Sell price with Packaging & Forwarding, Designing, Negotiation and
            # Transport applied (Transport is broken onto its own line below).
            "unit_price":   offer_unit,
        }],
        "valid_days": 30,
    }

    quote_data = calculate_quote(form_data)
    total_price = float(quote_data.get("grand_total") or offer_unit * qty)

    # ── 4. Filename: {YYYY-MM-DD}_{Customer}_{HPU|PU}-{kW}kW-{variant}.docx ─
    _safe_company = "".join(ch for ch in (cust.company or "Client")
                            if ch.isalnum() or ch in " _-").strip().replace(" ", "_") or "Client"
    _safe_variant = req.hpu_variant.replace(" ", "")
    _date = datetime.now().strftime("%Y-%m-%d")
    filename = f"{_date}_{_safe_company}_{filename_infix}-{kw_int}kW-{_safe_variant}.docx"
    output_path = os.path.join(QUOTES_FOLDER, filename)

    template_path = os.path.join(BASE_DIR, template_name)
    generate_quote_docx(quote_data, output_path, template_path=template_path)
    _drop_marketing_if_empty(output_path)
    # Pull Transport onto its own price-schedule line (no-op if 0).
    try:
        _break_out_transport(output_path, req.transport_amt)
    except Exception as _trn_err:
        print(f"WARN: recup transport line break-out failed: {_trn_err}")

    _log_equipment_quote(cust, mode.upper(), total_price, output_path)  # HPU / PU

    return {
        "success":      True,
        "filename":     filename,
        "download_url": f"/api/download-quote/{filename}",
        "preview_url":  f"/api/preview-quote/{filename}",
        "unit_price":   offer_unit,
        "total_price":  total_price,
        "variant":      req.hpu_variant,
        "kw":           req.hpu_kw,
        "lph":          catalog_lph,
        "model_code":   model_code,
        "qty":          qty,
        "mode":         mode,
    }


@app.post("/api/generate-hpu-quote")
def generate_hpu_quote(req: HpuQuoteRequest):
    """Stand-alone HPU offer — pumping skid WITH in-built electric heater
    (model codes HPS / HPD / HPDD). Price = SUM(hpu_master) × 1.8."""
    try:
        return _generate_pumping_unit_offer(req, mode="hpu")
    except Exception as e:
        import traceback
        return {"error": str(e), "trace": traceback.format_exc()}


@app.post("/api/generate-pu-quote")
def generate_pu_quote(req: HpuQuoteRequest):
    """Stand-alone PU offer — pumping skid WITHOUT heater, for pre-heated
    oils (LDO / LSHS). Model codes PUS / PUD / PUDD. Price comes from
    pumping_unit_price.sell_price (already includes the markup)."""
    try:
        return _generate_pumping_unit_offer(req, mode="pu")
    except Exception as e:
        import traceback
        return {"error": str(e), "trace": traceback.format_exc()}


# ══════════════════════════════════════════════════════════════════════════
#  STAND-ALONE BLOWER & BURNER OFFERS
#  Catalog-pick offers modelled on the HPU/PU + Recuperator pattern:
#  template (Blower_/Burner_Offer_Template.docx, built by
#  build_blower_burner_templates.py) rendered self-contained via docxtpl.
# ══════════════════════════════════════════════════════════════════════════

@app.get("/blower", response_class=HTMLResponse)
def blower_costing_form():
    """Stand-alone Blower offer form (catalog pick from blower_master)."""
    html_path = os.path.join(BASE_DIR, "blower_costing.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())


@app.get("/burner", response_class=HTMLResponse)
def burner_costing_form():
    """Stand-alone Burner offer form (catalog pick — ENCON oil/gas/dual + GAIL)."""
    html_path = os.path.join(BASE_DIR, "burner_costing.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())


def _finite_price(v):
    """Coerce to a finite float > 0, or None. Guards against NaN/inf values
    (e.g. incomplete blower_master rows) that aren't JSON-serialisable and
    would 500 the response."""
    try:
        f = float(v)
    except (TypeError, ValueError):
        return None
    if f != f or f in (float("inf"), float("-inf")) or f <= 0:
        return None
    return f


def _fmt_num(v) -> str:
    """'10.0' -> '10', '7.5' -> '7.5', '2040.0' -> '2,040'."""
    try:
        f = float(v)
    except (TypeError, ValueError):
        return str(v or "")
    if f != f or f in (float("inf"), float("-inf")):  # NaN / inf -> blank
        return ""
    return f"{int(f):,}" if f == int(f) else f"{f:g}"


@app.get("/api/blower/catalog")
def blower_catalog():
    """ENCON blower models for the offer dropdown. Prices on the PERKIN basis
    (Blower Alone × 1.8; with-motor + Motor × 1.5) — both are returned so the
    client can pick with or without motor."""
    from bom.blower_pricelist import blower_price
    conn = sqlite3.connect(DB_PATH)
    rows = conn.execute(
        "SELECT model, hp, airflow, pressure "
        "FROM blower_master ORDER BY pressure, CAST(hp AS REAL)"
    ).fetchall()
    items = []
    for r in rows:
        pw = blower_price(conn, r[0], with_motor=True)
        po = blower_price(conn, r[0], with_motor=False)
        if not pw and not po:
            continue  # no price → unsellable
        items.append({
            "model":               r[0],
            "hp":                  _fmt_num(r[1]),
            "airflow":             _fmt_num(r[2]),
            "pressure":            (r[3] or "").strip(),
            "price":               pw,      # default (with motor), back-compat
            "price_with_motor":    pw,
            "price_without_motor": po,
        })
    conn.close()
    return {"items": items}


# Burner catalog config: group key -> (label, fuel text, pricelist section).
_BURNER_SECTIONS = {
    "oil":  ("ENCON Oil (Film) Burner", "Oil (LDO / HSD / SKO)",
             "PRICE FOR VARIOUS SIZES OF ENCON 'FILM' BURNER & ACCESSORIES"),
    "gas":  ("ENCON Gas Burner",        "Gas (NG / LPG / COG)",
             "PRICE FOR VARIOUS SIZES OF ENCON 'GAS' BURNER & ACCESSORIES"),
    "dual": ("ENCON Dual Fuel Burner",  "Dual Fuel (Gas + Oil)",
             "PRICE FOR VARIOUS SIZES OF ENCON DUAL FUEL BURNER & ACCESSORIES"),
    "hv":   ("ENCON High Velocity Burner", "Oil (LDO / HSD)",
             "PRICE LIST FOR HIGH VELOCITY OIL BURNERS"),
}


def _burner_firing_capacity(conn, model: str) -> str:
    """Max firing rate (LPH at 24\" W.G.) for an ENCON burner size, if known."""
    row = conn.execute(
        "SELECT max_firing_lph FROM burner_selection_master "
        "WHERE model = ? AND pressure_wg = 24 LIMIT 1", (model,)
    ).fetchone()
    if row and row[0]:
        return f"up to {_fmt_num(row[0])} LPH"
    return ""


def _burner_capacity_kw(conn, model: str) -> str:
    """Max firing capacity in kW (= max kcal/hr / 860) for an ENCON burner size."""
    row = conn.execute(
        "SELECT max_firing_kcal_hr FROM burner_selection_master "
        "WHERE model = ? AND pressure_wg = 24 LIMIT 1", (model,)
    ).fetchone()
    if row and row[0]:
        return f"up to {_fmt_num(round(row[0] / 860 / 10) * 10)} kW"
    return ""


def _burner_lookup(group: str, model: str):
    """Return (price, fuel_text, capacity, label) for a burner pick, or None."""
    conn = sqlite3.connect(DB_PATH)
    try:
        if group == "gail":
            row = conn.execute(
                "SELECT burner_set FROM gail_gas_burner_master "
                "WHERE burner_size = ? AND burner_set IS NOT NULL LIMIT 1", (model,)
            ).fetchone()
            price = _finite_price(row[0]) if row else None
            if price is None:
                return None
            # Capacity is inherent in the GAIL model name (e.g. 'ENCON-500 KW').
            import re as _re
            m = _re.search(r"(\d[\d,]*)\s*KW", model, _re.I)
            cap = f"{m.group(1)} kW" if m else ""
            return price, "Natural Gas", cap, "GAIL Gas Burner"
        cfg = _BURNER_SECTIONS.get(group)
        if not cfg:
            return None
        label, fuel, section = cfg
        row = conn.execute(
            "SELECT price FROM burner_pricelist_master "
            "WHERE section = ? AND burner_size = ? AND component = 'BURNER SET' LIMIT 1",
            (section, model),
        ).fetchone()
        price = _finite_price(row[0]) if row else None
        if price is None:
            return None
        return price, fuel, _burner_firing_capacity(conn, model), label
    finally:
        conn.close()


@app.get("/api/burner/catalog")
def burner_catalog():
    """Burner models grouped by type for the offer dropdown."""
    conn = sqlite3.connect(DB_PATH)
    groups = []
    for key, (label, fuel, section) in _BURNER_SECTIONS.items():
        rows = conn.execute(
            "SELECT burner_size, price FROM burner_pricelist_master "
            "WHERE section = ? AND component = 'BURNER SET' ORDER BY burner_size",
            (section,),
        ).fetchall()
        items = []
        for r in rows:
            price = _finite_price(r[1])
            if price is None:
                continue
            items.append({"model": r[0], "price": price,
                          "capacity": _burner_firing_capacity(conn, r[0]),
                          "capacity_kw": _burner_capacity_kw(conn, r[0])})
        groups.append({"key": key, "label": label, "fuel": fuel, "items": items})
    conn.close()
    return {"groups": groups}


class BlowerQuoteRequest(BaseModel):
    customer: HpuCustomer
    blower_model: str
    with_motor: bool = True     # supply blower with (True) or without (False) motor
    qty: int = 1
    pf_pct:        float = 0    # Packaging & Forwarding (%)
    design_pct:    float = 0    # Designing (%)
    neg_pct:       float = 0    # Negotiation (%, added)
    transport_amt: float = 0    # Transport (flat Rs., own price line)


class BurnerQuoteRequest(BaseModel):
    customer: HpuCustomer
    burner_group: str = "gas"   # oil | gas | dual | gail
    burner_model: str
    qty: int = 1
    pf_pct:        float = 0    # Packaging & Forwarding (%)
    design_pct:    float = 0    # Designing (%)
    neg_pct:       float = 0    # Negotiation (%, added)
    transport_amt: float = 0    # Transport (flat Rs., own price line)


def _fill_blower_specs(docx_path: str, spec_rows: list, note: str = ""):
    """Populate the 'Blower Specifications' table with the detailed spec rows
    (label / value), reusing the existing row's formatting. If `note` is given
    (e.g. the recommended motor when the blower is quoted without motor), add it
    as an italic paragraph right after the table. Runs after render."""
    import copy
    from docx import Document as _Doc
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn
    d = _Doc(docx_path)
    tbl = None
    for t in d.tables:
        if t.rows and t.rows[0].cells and \
           t.rows[0].cells[0].text.strip().lower().startswith("blower specification"):
            tbl = t
            break
    if tbl is None or len(tbl.rows) < 2:
        return

    def _set_cell(cell, text):
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = text
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.add_run(text)

    template_tr = copy.deepcopy(tbl.rows[1]._tr)      # a data row → keeps borders/fonts
    for row in list(tbl.rows[1:]):                    # clear everything but the header
        row._tr.getparent().remove(row._tr)
    for s in spec_rows:
        tbl._tbl.append(copy.deepcopy(template_tr))
        row = tbl.rows[-1]
        _set_cell(row.cells[0], str(s.get("label", "")))
        _set_cell(row.cells[1], str(s.get("value", "")))

    if note:
        p_el = OxmlElement("w:p")
        r_el = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        rpr.append(OxmlElement("w:b"))
        rpr.append(OxmlElement("w:i"))
        r_el.append(rpr)
        t_el = OxmlElement("w:t")
        t_el.text = note
        t_el.set(_qn("xml:space"), "preserve")
        r_el.append(t_el)
        p_el.append(r_el)
        tbl._tbl.addnext(p_el)

    d.save(docx_path)


def _drop_marketing_if_empty(docx_path: str):
    """Remove the 'MARKETING PERSON DETAILS' heading + table when its value
    column is blank (equipment offers usually have no marketing contact)."""
    try:
        from docx import Document as _Doc
        from docx.oxml.ns import qn as _qn
        from docx.table import Table as _Tbl
        d = _Doc(docx_path)
        body = d.element.body
        kids = list(body.iterchildren())
        for i, ch in enumerate(kids):
            if ch.tag != _qn("w:p"):
                continue
            txt = "".join(x.text or "" for x in ch.findall(".//" + _qn("w:t"))).strip()
            if txt != "MARKETING PERSON DETAILS":
                continue
            for j in range(i + 1, len(kids)):
                if kids[j].tag == _qn("w:tbl"):
                    tbl = _Tbl(kids[j], d)
                    if all(len(r.cells) < 2 or not r.cells[1].text.strip() for r in tbl.rows):
                        for k in kids[i:j + 1]:
                            body.remove(k)
                        d.save(docx_path)
                    return
                if kids[j].tag == _qn("w:p") and "".join(
                        x.text or "" for x in kids[j].findall(".//" + _qn("w:t"))).strip():
                    return
            return
    except Exception as _e:
        print(f"WARN: marketing-table drop failed: {_e}")


def _generate_equipment_offer(cust: HpuCustomer, *, equipment_name: str,
                              specs: dict, unit_price: float, qty: int,
                              template_name: str, filename_infix: str,
                              drive_product: str,
                              pf_pct: float = 0.0, design_pct: float = 0.0,
                              neg_pct: float = 0.0, transport_amt: float = 0.0) -> dict:
    """Shared minimal-offer generator for stand-alone equipment (blower /
    burner). Builds the docxtpl context from the customer block + the
    equipment specs, renders template_name, saves to quotes/, best-effort
    PDF + Drive upload. Returns the API response dict."""
    from datetime import datetime as _dt
    from docxtpl import DocxTemplate
    from engine.quote_writer import amount_in_words_indian, _format_inr
    from equipment_advantages import tnc_value as _tnc

    def _fmt0(x):  # whole-rupee Indian format (no paise)
        s = _format_inr(round(float(x or 0)))
        return s[:-3] if s.endswith(".00") else s

    qty = max(1, int(qty or 1))
    # Commercial adjustments: P&F + Designing add, Negotiation all add; Transport
    # is a flat amount broken onto its own price-schedule line below. Final total
    # rounded to the nearest Rs.1000 when any adjustment is applied.
    _pf, _des, _neg, _trn = (pf_pct or 0), (design_pct or 0), (neg_pct or 0), (transport_amt or 0)
    _adj_total  = float(unit_price) * qty * (1 + (_pf + _des + _neg) / 100) + _trn
    _final_incl = (round(_adj_total / 1000) * 1000) if (_pf or _des or _neg or _trn) else _adj_total
    offer_unit  = _final_incl / qty
    total_price = _final_incl

    seq = next_quote_seq()
    full_ref = build_enquiry_ref(seq, cust.technical or "", cust.location or "")
    # OUR REF = the backend ENCON ref (real sequence + branch + initials), the
    # SAME value shown on the cover page. The client-side ref_no is only a form
    # preview (literal "XXX" sequence), so we never use it here — that mismatch
    # is exactly the bug we're avoiding. DATE is always today.
    short_ref = full_ref.split(" DT.")[0]
    date_str = _dt.now().strftime("%d/%m/%Y")

    company_address = ", ".join(filter(None, [
        (cust.address or "").strip(), (cust.city or "").strip(),
        (cust.state or "").strip(), (cust.pin or "").strip()]))

    ctx = {
        "project_name":      cust.subject or equipment_name,
        "subject":           cust.subject or f"Offer for {equipment_name}",
        "application":       equipment_name,
        "equipment_name":    equipment_name,
        "company_name":      cust.company or "",
        "company_address":   company_address,
        "email":             cust.email or "",
        "mobile_no":         cust.phone or "",
        "poc_name":          _with_salutation(cust.salutation, cust.name),
        "poc_greeting":      _greeting(cust.salutation),
        "poc_designation":   cust.designation or "",
        "client_enq_ref":    "",
        "enquiry_ref":       full_ref,
        "enquiry_ref_short": short_ref,
        "enquiry_date_str":  date_str,
        "marketing_person":  _with_salutation(cust.marketing_salutation, cust.marketing),
        "marketing_email":   cust.marketing_email or "",
        "marketing_phone":   cust.marketing_phone or "",
        "technical_person":  _with_salutation(cust.technical_salutation, cust.technical) or _with_salutation(cust.marketing_salutation, cust.marketing),
        "technical_phone":   cust.technical_phone or cust.marketing_phone or "",
        "technical_email":   cust.technical_email or cust.marketing_email or "",
        # Price schedule (single line)
        "item_qty":          f"{qty:02d} No.",
        "unit_price":        _fmt0(offer_unit),
        "total_price":       _fmt0(total_price),
        "grand_total":       _fmt0(total_price),
        "grand_total_in_words": f"INR. {amount_in_words_indian(total_price)} ONLY.",
        # T&C (Annexure renumbered to III) — standard ENCON defaults when blank
        "tnc_prices":             _tnc("tnc_prices", cust.tnc_prices),
        "tnc_delivery":           _tnc("tnc_delivery", cust.tnc_delivery),
        "tnc_gst":                _tnc("tnc_gst", cust.tnc_gst),
        "tnc_hsn_code":           _tnc("tnc_hsn_code", cust.tnc_hsn_code),
        "tnc_pan_gst":            _tnc("tnc_pan_gst", cust.tnc_pan_gst),
        "tnc_payment_terms":      _tnc("tnc_payment_terms", cust.tnc_payment_terms),
        "tnc_packing_forwarding": _tnc("tnc_packing_forwarding", cust.tnc_packing_forwarding),
        "tnc_freight":            _tnc("tnc_freight", cust.tnc_freight),
        "tnc_transit_insurance":  _tnc("tnc_transit_insurance", cust.tnc_transit_insurance),
        "tnc_validity":           _tnc("tnc_validity", cust.tnc_validity),
        "tnc_inspection":         _tnc("tnc_inspection", cust.tnc_inspection),
        "tnc_guarantee":          _tnc("tnc_guarantee", cust.tnc_guarantee),
    }
    ctx.update(specs)

    # Price-table item description: bold heading + spec sentence + component
    # bullets (real bulleted paragraphs) + notes, from specs['price_desc'].
    _pd = specs.get("price_desc") or {}
    _ph = _pd.get("heading") or ctx.get("equipment_name", "")
    ctx["price_heading"] = _ph or ""   # price schedule shows the heading only
    ctx["price_body"]    = ""
    ctx["price_bullets"] = [{"item": b} for b in _pd.get("bullets", [])]
    ctx["price_notes"]   = [{"item": n} for n in _pd.get("notes", [])]

    # Dynamic per-equipment Advantages section (per fuel type for burners).
    from equipment_advantages import build_advantages_ctx
    ctx.update(build_advantages_ctx(specs.get("advantages_kind") or drive_product))

    tpl_path = os.path.join(BASE_DIR, template_name)
    tpl = DocxTemplate(tpl_path)
    tpl.render(ctx, autoescape=True)   # preserve '&' in values (e.g. "Steel & Alloys")

    safe_company = "".join(ch for ch in (cust.company or "Client")
                           if ch.isalnum() or ch in " _-").strip().replace(" ", "_") or "Client"
    docx_name = f"{filename_infix}_Offer_{safe_company}_{seq}.docx"
    docx_path = os.path.join(QUOTES_FOLDER, docx_name)
    tpl.save(docx_path)
    if specs.get("blower_specs"):
        try:
            _fill_blower_specs(docx_path, specs["blower_specs"],
                               specs.get("blower_specs_note") or "")
        except Exception as _bs_err:
            print(f"WARN: blower spec fill failed: {_bs_err}")
    _drop_marketing_if_empty(docx_path)
    # Transport onto its own price-schedule line (no-op if 0).
    try:
        _break_out_transport(docx_path, _trn)
    except Exception as _trn_err:
        print(f"WARN: {filename_infix} transport break-out failed: {_trn_err}")

    pdf_name = docx_name.replace(".docx", ".pdf")
    pdf_path = os.path.join(QUOTES_FOLDER, pdf_name)
    pdf_ok = _docx_to_pdf(docx_path, pdf_path)

    try:
        from engine.drive_uploader import upload_offer_async
        upload_offer_async(docx_path, docx_name, drive_product)
        if pdf_ok:
            upload_offer_async(pdf_path, pdf_name, drive_product)
    except Exception as _drv_err:
        print(f"WARN: drive upload kickoff failed: {_drv_err}")

    _PROD_LABEL = {"blower": "Blower", "burner": "Burner"}
    _log_equipment_quote(
        cust, _PROD_LABEL.get((drive_product or "").lower(), (drive_product or "Equipment").title()),
        total_price, docx_path)

    return {
        "success":      True,
        "filename":     docx_name,
        "pdf_filename": pdf_name if pdf_ok else None,
        "download_url": f"/api/download-quote/{docx_name}",
        "pdf_url":      f"/api/pdf-quote/{pdf_name}" if pdf_ok else None,
        "preview_url":  f"/api/preview-quote/{docx_name}",
        "quote_no":     full_ref,
        "enquiry_ref":  full_ref,
        "unit_price":   offer_unit,
        "total_price":  total_price,
        "qty":          qty,
    }


@app.post("/api/generate-blower-quote")
def generate_blower_quote(req: BlowerQuoteRequest):
    """Stand-alone Blower offer — catalog pick from blower_master
    (price = price_premium, i.e. with motor)."""
    try:
        from engine.quote_writer import _format_inr as _finr
        conn = sqlite3.connect(DB_PATH)
        row = conn.execute(
            "SELECT model, hp, airflow, cfm, pressure, price_basic, price_premium "
            "FROM blower_master WHERE model = ? LIMIT 1", (req.blower_model,)
        ).fetchone()
        conn.close()
        if not row:
            return {"error": f"unknown blower model: {req.blower_model}"}
        model, hp, airflow, cfm, pressure, price_basic, price = row
        from bom.blower_pricelist import blower_price, blower_spec_rows
        conn2 = sqlite3.connect(DB_PATH)
        unit_price = blower_price(conn2, model, with_motor=req.with_motor)  # PERKIN basis
        conn2.close()
        if not unit_price:
            return {"error": f"blower '{model}' has no valid price in the catalog"}
        import re as _re
        _pm = _re.search(r"\d+", pressure or "")
        ptype = "High" if (_pm and int(_pm.group()) >= 40) else "Medium"
        _qty = max(1, int(req.qty or 1))
        _model_short = (model or "").replace("ENCON", "").strip()
        _press = (pressure or "").strip()
        _wm = bool(req.with_motor)
        _motor_clause = (f"fitted with a {_fmt_num(hp)} HP, 2900 rpm motor of reputed "
                         f"make such as ABB, Crompton, etc."
                         if _wm else
                         f"suitable for a {_fmt_num(hp)} HP motor (motor NOT included — "
                         f"in customer's scope).")
        _motor_tag = "with Motor" if _wm else "without Motor"
        equipment_name = f"Centrifugal Blower – {model} ({_motor_tag})"
        scope_intro = (f"Supply ex-works of {_qty} No. ENCON {ptype} Pressure Blower, "
                       f"model {model}, having a capacity of {_fmt_num(cfm)} CFM at {_press} "
                       f"pressure, {_motor_clause}")
        scope_items = []   # blower scope is the intro sentence only — no bullets
        _spec_rows = blower_spec_rows(model, hp, cfm, pressure, with_motor=_wm)
        _spec_rows.append({"label": "Quantity", "value": f"{_qty:02d} No."})
        _specs_note = ("" if _wm else
                       f"Recommended motor to use: {_fmt_num(hp)} HP, 2-Pole "
                       f"(≈ 2880 rpm), 3-Phase, 415 V AC — to be supplied by the customer.")
        specs = {
            "blower_model":    model,
            "blower_hp":       _fmt_num(hp),
            "blower_airflow":  _fmt_num(airflow),
            "blower_pressure": _press,
            "blower_specs":    _spec_rows,
            "blower_specs_note": _specs_note,
            "scope_intro":     scope_intro,
            "scope_items":     scope_items,
            "price_desc": {
                "heading": f"{ptype.upper()} PRESSURE BLOWER, MODEL {_model_short} ({_motor_tag.upper()})",
                "body":    scope_intro,
                "bullets": [],
                "notes":   [],
            },
        }
        result = _generate_equipment_offer(
            req.customer, equipment_name=equipment_name, specs=specs,
            unit_price=unit_price, qty=req.qty,
            template_name="Blower_Offer_Template.docx",
            filename_infix="Blower", drive_product="blower",
            pf_pct=req.pf_pct, design_pct=req.design_pct,
            neg_pct=req.neg_pct, transport_amt=req.transport_amt)
        result.update({"model": model, "config": f"{_fmt_num(hp)} HP • {(pressure or '').strip()}"})
        return result
    except Exception as e:
        import traceback
        return {"error": str(e), "trace": traceback.format_exc()}


@app.post("/api/generate-burner-quote")
def generate_burner_quote(req: BurnerQuoteRequest):
    """Stand-alone Burner offer — ENCON oil/gas/dual (burner_pricelist_master,
    BURNER SET) or GAIL pre-assembled set (gail_gas_burner_master)."""
    try:
        looked = _burner_lookup(req.burner_group, req.burner_model)
        if not looked:
            return {"error": f"unknown burner: {req.burner_group}/{req.burner_model}"}
        unit_price, fuel, capacity, label = looked
        equipment_name = f"{label} – {req.burner_model}"
        _bqty = max(1, int(req.qty or 1))
        _bg = (req.burner_group or "oil").lower()
        # Per fuel-type: display name, operation phrase, and component list
        # (mirrors the standard ENCON burner scope wording).
        _BNAME = {"dual": "ENCON Dual Fuel Burner", "oil": "IIP-ENCON Film Burner",
                  "gas": "ENCON Gas Burner", "hv": "ENCON High Velocity Burner"}
        _BOPER = {"dual": "dual operation using liquid fuel (LDO) and gaseous fuel (such as Natural Gas / LPG)",
                  "oil": "operation using liquid fuel (LDO)",
                  "gas": "operation using gaseous fuel (such as Natural Gas / LPG)",
                  "hv": "high-velocity operation using liquid fuel (LDO / HSD)"}
        _BCOMP = {"dual": ["Burner Alone", "Micro Valve", "C.I Burner Plate", "Burner Block",
                           "Flexible Hoses set", "Ball Valve", "\"Y\" type Strainer", "Butterfly Valve"],
                  "oil": ["Burner Alone", "Micro Valve", "C.I Burner Plate", "Burner Block",
                          "Flexible Hoses set", "\"Y\" type Strainer", "Butterfly Valve"],
                  "gas": ["Burner Alone", "Ball Valve", "C.I Burner Plate", "Burner Block",
                          "Flexible Hoses set", "Butterfly Valve"],
                  "hv": ["Burner Alone", "Burner Block", "Micro Valve",
                         "Flexible Hoses set", "\"Y\" type Strainer", "Butterfly Valve"]}
        # Advantages copy per fuel type.
        _BADV = {"dual": "burner_dual", "gas": "burner_gas",
                 "hv": "burner_hv", "oil": "burner_film"}
        _bname = _BNAME.get(_bg, label)
        _boper = _BOPER.get(_bg, "operation using the selected fuel")
        _bcomp = _BCOMP.get(_bg, ["Burner Alone", "Ball Valve", "C.I Burner Plate",
                                  "Burner Block", "Flexible Hoses set", "Butterfly Valve"])
        _rate = f", having a firing rate of {capacity}" if capacity else ""
        _bsentence = (f"Supply ex-works of {_bqty} no. {_bname}, Model {req.burner_model}, "
                      f"suitable for {_boper}{_rate}, fitted with necessary accessories of "
                      f"reputed make.")
        scope_intro = _bsentence + " Equipment with Burner:"
        scope_items = [{"item": x} for x in _bcomp]
        specs = {
            "burner_model":    req.burner_model,
            "burner_fuel":     fuel,
            "burner_capacity": capacity,
            "scope_intro":     scope_intro,
            "scope_items":     scope_items,
            "advantages_kind": _BADV.get(_bg, "burner_film"),
            "price_desc": {       # price cell: short sentence only, no bullets
                "heading": _bname.upper(),
                "body":    _bsentence,
                "bullets": [],
                "notes":   [],
            },
        }
        result = _generate_equipment_offer(
            req.customer, equipment_name=equipment_name, specs=specs,
            unit_price=unit_price, qty=req.qty,
            template_name="Burner_Offer_Template.docx",
            filename_infix="Burner", drive_product="burner",
            pf_pct=req.pf_pct, design_pct=req.design_pct,
            neg_pct=req.neg_pct, transport_amt=req.transport_amt)
        result.update({"model": req.burner_model,
                       "config": f"{fuel}{(' • ' + capacity) if capacity else ''}"})
        return result
    except Exception as e:
        import traceback
        return {"error": str(e), "trace": traceback.format_exc()}


# ══════════════════════════════════════════════════════════════════════════
#  COMBINED OFFER — merge several individually-generated offers into one .docx
# ══════════════════════════════════════════════════════════════════════════

@app.get("/combined", response_class=HTMLResponse)
@app.get("/combined-builder", response_class=HTMLResponse)
def combined_builder_page():
    """The Combined Offer builder: customer + T&C once, an input section per
    equipment, then one unified offer + a costing Excel. ('/combined-builder'
    is kept as an alias for older links/bookmarks.)"""
    html_path = os.path.join(BASE_DIR, "combined_builder.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())


# ── Combined costing workbook: one sheet per equipment + a summary ──────────
class CombinedEquipment(BaseModel):
    name: str                          # e.g. "Vertical Ladle Preheater (10 T)"
    bom: List[dict] = []               # rows: {media?, item, ref?, qty, unit_price, total}
    total: Optional[float] = None      # equipment grand total (falls back to sum of rows)
    qty: int = 1                       # number of this equipment (for per-unit breakdown)
    unit_price: Optional[float] = None # sell price per unit (falls back to total / qty)


class CombinedCostingRequest(BaseModel):
    project_name: Optional[str] = ""
    company_name: Optional[str] = ""
    equipments: List[CombinedEquipment] = []
    # Commercial adjustments (applied once to the combined grand total) so the
    # costing Excel shows the same P&F / Designing / Negotiation / Transport /
    # Final Total box the builder displays on screen.
    pf_pct: float = 0
    design_pct: float = 0
    neg_pct: float = 0
    transport_amt: float = 0


@app.get("/api/download-xlsx/{filename}")
def download_xlsx(filename: str):
    """Serve a generated .xlsx from the quotes folder with the right type."""
    safe = os.path.basename(filename)
    file_path = os.path.join(QUOTES_FOLDER, safe)
    if not safe.lower().endswith(".xlsx") or not os.path.exists(file_path):
        return {"error": "File not found"}
    return FileResponse(
        path=file_path, filename=safe,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


@app.post("/api/combined-costing-excel")
def combined_costing_excel(req: CombinedCostingRequest):
    """Build one .xlsx with a sheet per equipment (full itemised BOM) plus a
    Summary sheet totalling all equipment. Returns a download link."""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

        if not req.equipments:
            return {"error": "No equipment supplied."}

        navy = "1A3A5C"
        hdr_fill = PatternFill("solid", fgColor=navy)
        hdr_font = Font(bold=True, color="FFFFFF")
        bold = Font(bold=True)
        thin = Side(style="thin", color="D0D7DE")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)
        money = '#,##0.00'

        wb = openpyxl.Workbook()

        # ── Summary sheet (first) ────────────────────────────────────────
        ws = wb.active
        ws.title = "Summary"
        ws["A1"] = "COMBINED COSTING SUMMARY"; ws["A1"].font = Font(bold=True, size=14, color=navy)
        ws["A2"] = f"Project: {req.project_name or '—'}"
        ws["A3"] = f"Client: {req.company_name or '—'}"
        r = 5
        for c, h in enumerate(["S. No.", "Equipment", "Total (Rs.)"], start=1):
            cell = ws.cell(r, c, h); cell.fill = hdr_fill; cell.font = hdr_font; cell.border = border
        grand = 0.0
        eq_totals = []
        for i, eq in enumerate(req.equipments, start=1):
            t = eq.total if eq.total is not None else sum(float(x.get("total") or 0) for x in eq.bom)
            t = float(t or 0); grand += t; eq_totals.append(t)
            ws.cell(r + i, 1, i).border = border
            ws.cell(r + i, 2, eq.name).border = border
            tc = ws.cell(r + i, 3, t); tc.number_format = money; tc.border = border
        gr = r + len(req.equipments) + 1
        ws.cell(gr, 2, "GRAND TOTAL").font = bold
        gc = ws.cell(gr, 3, grand); gc.font = bold; gc.number_format = money
        # (Commercial adjustments / final total now live on the Cost Breakdown sheet.)

        ws.column_dimensions["A"].width = 8; ws.column_dimensions["B"].width = 46; ws.column_dimensions["C"].width = 18

        # ── Cost Breakdown sheet (full build-up + additional-cost distribution) ──
        import math as _math
        wsb = wb.create_sheet("Cost Breakdown")
        lite    = PatternFill("solid", fgColor="EFF6FF")   # subtle navy-tint for totals
        right   = Alignment(horizontal="right")
        ctr     = Alignment(horizontal="center", wrap_text=True)
        SUBTOTAL_NAMES = {"BOUGHT OUT ITEMS", "ENCON ITEMS", "GRAND TOTAL", "BOM TOTAL"}
        _negp = float(req.neg_pct or 0); _pfp = float(req.pf_pct or 0)
        _desp = float(req.design_pct or 0); _frt = float(req.transport_amt or 0)

        def _yhead(r_, c_, v_):
            cell = wsb.cell(r_, c_, v_); cell.fill = hdr_fill
            cell.font = hdr_font; cell.border = border; cell.alignment = ctr
            return cell

        # Per-equipment build-up: bought-out vs in-house split from the BOM media
        # tags (preheaters/recup have them; catalogue picks fall to in-house).
        rows_data = []
        for eq in req.equipments:
            q = max(1, int(eq.qty or 1))
            bought = inhouse = 0.0; has_media = False
            for b in (eq.bom or []):
                if str(b.get("item", "")).strip().upper() in SUBTOTAL_NAMES:
                    continue
                t = float(b.get("total") or (float(b.get("qty") or 0) * float(b.get("unit_price") or 0)))
                media = str(b.get("media", "")).strip().upper()
                if media:
                    has_media = True
                if media == "ENCON ITEMS":
                    inhouse += t
                elif media:
                    bought += t
                else:
                    inhouse += t          # catalogue pick — no split
            unit = float(eq.unit_price if eq.unit_price is not None else ((eq.total or 0) / q))
            if has_media and bought > 0:
                sell_bought = unit - inhouse          # marked-up bought-out (reconciles to unit)
                cost = bought + inhouse
            else:
                bought = sell_bought = 0.0; inhouse = cost = unit
            neg_amt = unit * _negp / 100
            sell_rounded = round((unit + neg_amt) / 10000) * 10000   # nearest Rs.10,000
            rows_data.append(dict(name=eq.name, bought=bought, sell_bought=sell_bought,
                inhouse=inhouse, cost=cost, unit=unit, neg=neg_amt,
                sell=sell_rounded, qty=q, total=sell_rounded * q))

        wsb["A1"] = "COMBINED COSTING — BREAKDOWN"; wsb["A1"].font = Font(bold=True, size=13, color=navy)
        hr = 3
        heads = ["S. No.", "Item Description", "Bought Out", "Sell Price", "In-house",
                 "Cost Price", "Unit Price", f"Negotiation ({_negp:g}%)", "Sell Price", "Qty", "Total Price"]
        for c, h in enumerate(heads, start=1):
            _yhead(hr, c, h)
        rr = hr + 1
        basic = 0.0
        for i, d in enumerate(rows_data, start=1):
            basic += d["total"]
            vals = [i, d["name"], d["bought"], d["sell_bought"], d["inhouse"], d["cost"],
                    d["unit"], d["neg"], d["sell"], d["qty"], d["total"]]
            for c, v in enumerate(vals, start=1):
                cell = wsb.cell(rr, c, v); cell.border = border
                if c in (3, 4, 5, 6, 7, 8, 9, 11):
                    cell.number_format = money; cell.alignment = right
                elif c in (1, 10):
                    cell.alignment = Alignment(horizontal="center")
            rr += 1
        wsb.cell(rr, 2, "BASIC SELL PRICE").font = Font(bold=True)
        bc = wsb.cell(rr, 11, basic); bc.font = Font(bold=True); bc.number_format = money; bc.fill = lite
        rr += 2

        # Additional-cost summary
        pf_amt  = basic * _pfp / 100
        des_amt = basic * _desp / 100
        total_add = pf_amt + des_amt + _frt
        incr_ratio = (total_add / basic) if basic else 0.0
        summ = [("Basic Sell Price", basic), (f"Additional Cost (P&F) ({_pfp:g}%)", pf_amt)]
        if _desp:
            summ.append((f"Designing ({_desp:g}%)", des_amt))
        summ += [("Freight Charges", _frt), ("Total Additional Cost", total_add)]
        for label, amt in summ:
            lc = wsb.cell(rr, 2, label); lc.font = Font(bold=True); lc.border = border
            vc = wsb.cell(rr, 3, amt); vc.number_format = money; vc.border = border; vc.alignment = right
            rr += 1
        ic = wsb.cell(rr, 2, "Increase Amount"); ic.font = Font(bold=True); ic.border = border
        pc = wsb.cell(rr, 3, incr_ratio); pc.number_format = '0.000000%'; pc.border = border; pc.alignment = right
        rr += 2

        # Distribution of the additional cost back into per-equipment per-unit prices
        for c, h in enumerate(["ITEMS", f"Increase ({incr_ratio * 100:.3f}%)",
                               "P&F + Freight included", "Qty", "Per Unit", "Round Up"], start=1):
            _yhead(rr, c, h)
        rr += 1
        grand_final = 0.0
        for d in rows_data:
            increase = d["total"] * incr_ratio
            include = d["total"] + increase
            grand_final += include
            per_unit = include / d["qty"]
            round_up = _math.ceil(per_unit / 1000) * 1000
            vals = [d["name"], increase, include, d["qty"], per_unit, round_up]
            for c, v in enumerate(vals, start=1):
                cell = wsb.cell(rr, c, v); cell.border = border
                if c in (2, 3, 5):
                    cell.number_format = money; cell.alignment = right
                elif c == 4:
                    cell.alignment = Alignment(horizontal="center")
                elif c == 6:
                    cell.number_format = money; cell.alignment = right
                    cell.fill = lite; cell.font = Font(bold=True, color=navy)
            rr += 1
        wsb.cell(rr, 2, "GRAND TOTAL").font = Font(bold=True)
        gf = wsb.cell(rr, 3, grand_final); gf.font = Font(bold=True); gf.number_format = money; gf.fill = lite
        for col, w in zip("ABCDEFGHIJK", [8, 34, 15, 15, 13, 15, 15, 16, 15, 7, 18]):
            wsb.column_dimensions[col].width = w

        # ── One sheet per equipment (full BOM) ───────────────────────────
        used = {"Summary", "Cost Breakdown"}
        for idx, eq in enumerate(req.equipments):
            base = "".join(ch for ch in eq.name if ch not in '[]:*?/\\').strip()[:28] or f"Equipment {idx+1}"
            name = base; n = 2
            while name in used:
                name = f"{base[:26]} {n}"; n += 1
            used.add(name)
            sh = wb.create_sheet(name)
            sh["A1"] = eq.name; sh["A1"].font = Font(bold=True, size=12, color=navy)
            hr = 3
            cols = ["S. No.", "Media", "Item", "Ref / Size", "Qty", "Unit Price", "Total"]
            for c, h in enumerate(cols, start=1):
                cell = sh.cell(hr, c, h); cell.fill = hdr_fill; cell.font = hdr_font; cell.border = border
                cell.alignment = Alignment(horizontal="center")
            sub = 0.0
            for j, row in enumerate(eq.bom, start=1):
                rr = hr + j
                tot = float(row.get("total") or 0); sub += tot
                vals = [j, row.get("media", ""), row.get("item", ""),
                        row.get("ref", row.get("size", "")),
                        row.get("qty", ""), float(row.get("unit_price") or 0), tot]
                for c, v in enumerate(vals, start=1):
                    cell = sh.cell(rr, c, v); cell.border = border
                    if c in (6, 7): cell.number_format = money
            tr = hr + len(eq.bom) + 1
            # ORIGINAL total = unit price x qty (the equipment's own price, BEFORE
            # any combined P&F / designing / negotiation / transport is distributed).
            _q = max(1, int(eq.qty or 1))
            _orig = (float(eq.unit_price) * _q) if eq.unit_price is not None \
                    else (eq.total if eq.total is not None else sub)
            sh.cell(tr, 3, "TOTAL").font = bold
            tc = sh.cell(tr, 7, _orig)
            tc.font = bold; tc.number_format = money
            widths = [8, 16, 40, 18, 8, 14, 16]
            for c, w in enumerate(widths, start=1):
                sh.column_dimensions[openpyxl.utils.get_column_letter(c)].width = w

        safe_company = "".join(ch for ch in (req.company_name or "Client")
                               if ch.isalnum() or ch in " _-").strip().replace(" ", "_") or "Client"
        stamp = datetime.now().strftime("%d%b%Y_%H%M%S")
        out_name = f"Combined_Costing_{safe_company}_{stamp}.xlsx"
        out_path = os.path.join(QUOTES_FOLDER, out_name)
        wb.save(out_path)
        _drv = {"ok": False, "reason": "error", "msg": "Drive status unknown"}
        try:                                   # mirror the costing sheet to Drive too
            from engine.drive_uploader import upload_offer_async, drive_status as _drive_status
            upload_offer_async(out_path, out_name, "combined")
            _drv = _drive_status("combined")
        except Exception as _drv_err:
            print(f"WARN: combined costing Drive upload failed: {_drv_err}")
        return {
            "success":      True,
            "filename":     out_name,
            "download_url": f"/api/download-xlsx/{out_name}",
            "sheets":       len(req.equipments) + 1,
            "grand_total":  grand,
            "drive":        _drv,
        }
    except Exception as e:
        import traceback
        return {"error": str(e), "trace": traceback.format_exc()}


# ── Per-costing Excel: itemised BOM + the cost-summary box ─────────────────
class CostingExcelRequest(BaseModel):
    product: str = "Costing"               # equipment / product name
    project_name: Optional[str] = ""
    company_name: Optional[str] = ""
    bom: List[dict] = []                   # [{media, item, ref|size, qty, unit_price, total}]
    summary: dict = {}                     # subtotal_label, subtotal, pf_pct, pf_amount,
                                           # design_pct, design_amount, neg_pct, neg_amount,
                                           # transport_amount, final_total


@app.post("/api/costing-excel")
def costing_excel(req: CostingExcelRequest):
    """Build a single-sheet .xlsx for one costing: the itemised bill of
    materials, then the Grand Total / P&F / Designing / Negotiation / Transport
    / Final Total summary box. Uniform across every product costing page."""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

        navy = "1A3A5C"
        hdr_fill = PatternFill("solid", fgColor=navy)
        hdr_font = Font(bold=True, color="FFFFFF")
        bold = Font(bold=True)
        grey = PatternFill("solid", fgColor="F8FAFC")
        green = Font(bold=True, color="065F46", size=12)
        green_bg = PatternFill("solid", fgColor="F0FDF4")
        thin = Side(style="thin", color="D0D7DE")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)
        money = '#,##0.00'

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Costing"

        ws["A1"] = req.product or "Costing"
        ws["A1"].font = Font(bold=True, size=14, color=navy)
        ws["A2"] = f"Project: {req.project_name or '—'}"
        ws["A3"] = f"Client: {req.company_name or '—'}"
        ws["A4"] = f"Date: {datetime.now().strftime('%d %b %Y')}"

        r = 6
        # ── Itemised BOM ────────────────────────────────────────────────
        if req.bom:
            ws.cell(r, 1, "BILL OF MATERIALS").font = Font(bold=True, color=navy, size=11)
            r += 1
            cols = ["S. No.", "Media", "Item", "Ref / Size", "Qty", "Unit Price", "Total"]
            for c, h in enumerate(cols, start=1):
                cell = ws.cell(r, c, h); cell.fill = hdr_fill; cell.font = hdr_font
                cell.border = border; cell.alignment = Alignment(horizontal="center")
            r += 1
            # Live formulas so every derived value is visible: per-line
            # Total = Qty(E) × Unit Price(F); BOM TOTAL = SUM over the item rows.
            sub = 0.0
            _item_rows = []
            for i, row in enumerate(req.bom, start=1):
                tot = float(row.get("total") or 0); sub += tot
                qv = row.get("qty", "")
                vals = [i, row.get("media", ""), row.get("item", ""),
                        row.get("ref", row.get("size", "")),
                        qv, float(row.get("unit_price") or 0)]
                for c, v in enumerate(vals, start=1):
                    cell = ws.cell(r, c, v); cell.border = border
                    if c == 6:
                        cell.number_format = money; cell.alignment = Alignment(horizontal="right")
                if isinstance(qv, (int, float)):
                    tcell = ws.cell(r, 7, f"=E{r}*F{r}"); _item_rows.append(r)
                else:
                    tcell = ws.cell(r, 7, tot)
                tcell.border = border; tcell.number_format = money
                tcell.alignment = Alignment(horizontal="right")
                r += 1
            tc = ws.cell(r, 3, "BOM TOTAL"); tc.font = bold
            _bom_total_row = r
            vc = ws.cell(r, 7, f"=SUM(G{_item_rows[0]}:G{_item_rows[-1]})" if _item_rows else sub)
            vc.font = bold; vc.number_format = money
            vc.alignment = Alignment(horizontal="right")
            r += 2

        # ── Cost summary box ────────────────────────────────────────────
        s = req.summary or {}
        ws.cell(r, 1, "COST SUMMARY").font = Font(bold=True, color=navy, size=11)
        r += 1

        import math

        def _num(x):
            try:
                return float(x)
            except (TypeError, ValueError):
                return None

        def _fmt_pct(p):
            return str(int(p)) if float(p) == int(p) else str(p)

        # `formula` (when given) is written instead of the literal amount so the
        # derivation is visible in the cell. Returns the excel row used.
        def _line(label, amount, is_total=False, formula=None):
            nonlocal r
            bg = green_bg if is_total else grey
            lc = ws.cell(r, 1, label)
            lc.font = green if is_total else Font(bold=False)
            lc.fill = bg; lc.border = border
            ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=6)
            for c in range(2, 7):
                ws.cell(r, c).fill = bg; ws.cell(r, c).border = border
            val = formula if formula is not None else (_num(amount) if _num(amount) is not None else 0)
            vc = ws.cell(r, 7, val)
            vc.number_format = money; vc.fill = bg; vc.border = border
            vc.alignment = Alignment(horizontal="right")
            vc.font = green if is_total else bold
            used = r
            r += 1
            return used

        subtotal = _num(s.get("subtotal"))
        sub_label = s.get("subtotal_label") or "Grand Total"
        rows = {}
        if subtotal is not None:
            rows["sub"] = _line(sub_label, s.get("subtotal"))

        # A percentage line becomes "=G{sub}*pct/100" only when the passed amount
        # actually equals subtotal × pct (some products apply the % to a marked-up
        # base, so the shown subtotal isn't the pct base → keep the exact value).
        def _pct_line(amt_key, pct_key, label_fmt):
            amt = _num(s.get(amt_key))
            if amt is None:
                return
            pct = _num(s.get(pct_key)) or 0
            f = None
            if "sub" in rows and subtotal not in (None, 0) and abs(amt - subtotal * pct / 100) < 0.5:
                f = f"=G{rows['sub']}*{_fmt_pct(pct)}/100"
            rows[amt_key] = _line(label_fmt.format(pct=_fmt_pct(pct)), s.get(amt_key), formula=f)

        _pct_line("pf_amount", "pf_pct", "Packaging & Forwarding ({pct} %)")
        _pct_line("design_amount", "design_pct", "Designing ({pct} %)")
        _pct_line("neg_amount", "neg_pct", "Negotiation ({pct} %)")
        if _num(s.get("transport_amount")) is not None:
            rows["transport"] = _line("Transport", s.get("transport_amount"))

        # Final Total: formula only when it reconciles with the lines above
        # (either a straight SUM, or a SUM rounded up to the nearest ₹1,000).
        final = _num(s.get("final_total"))
        ffml = None
        if "sub" in rows and final is not None:
            first, last = rows["sub"], r - 1
            parts = [subtotal or 0] + [(_num(s.get(k)) or 0) for k in
                     ("pf_amount", "design_amount", "neg_amount", "transport_amount")
                     if _num(s.get(k)) is not None]
            ssum = sum(parts)
            if abs(final - (math.ceil(ssum / 1000) * 1000)) < 0.5:
                ffml = f"=CEILING(SUM(G{first}:G{last}),1000)"
            elif abs(final - ssum) < 0.5:
                ffml = f"=SUM(G{first}:G{last})"
        _line("Final Total", s.get("final_total"), is_total=True, formula=ffml)

        for col, w in zip("ABCDEFG", [8, 16, 40, 18, 8, 14, 16]):
            ws.column_dimensions[col].width = w

        safe_company = "".join(ch for ch in (req.company_name or "Client")
                               if ch.isalnum() or ch in " _-").strip().replace(" ", "_") or "Client"
        safe_prod = "".join(ch for ch in (req.product or "Costing")
                            if ch.isalnum() or ch in " _-").strip().replace(" ", "_") or "Costing"
        stamp = datetime.now().strftime("%d%b%Y_%H%M%S")
        out_name = f"Costing_{safe_prod}_{safe_company}_{stamp}.xlsx"
        out_path = os.path.join(QUOTES_FOLDER, out_name)
        wb.save(out_path)
        try:                                   # mirror the cost sheet to Drive too
            from engine.drive_uploader import upload_offer_async
            upload_offer_async(out_path, out_name, req.product or "costing")
        except Exception as _drv_err:
            print(f"WARN: costing Drive upload failed: {_drv_err}")
        return {"success": True, "filename": out_name,
                "download_url": f"/api/download-xlsx/{out_name}"}
    except Exception as e:
        import traceback
        return {"error": str(e), "trace": traceback.format_exc()}


# ── Unified combined offer: one cover/customer/T&C, equipment sections ──────
class CombinedOfferEquipment(BaseModel):
    name: str
    specs: Optional[str] = ""          # one-line specs summary
    spec_rows: List[dict] = []         # [{label, value}] -> Technical Specifications
    qty: int = 1
    unit_price: float = 0
    bom: List[dict] = []               # itemised BOM -> Scope of Supply
    quote_payload: dict = {}           # full standalone QuoteRequest -> narrative Scope of Supply


class CombinedOfferRequest(BaseModel):
    customer: HpuCustomer
    equipments: List[CombinedOfferEquipment] = []
    project_name: Optional[str] = ""
    # Commercial adjustments applied once to the combined grand total.
    pf_pct: float = 0          # Packaging & Forwarding %
    design_pct: float = 0      # Designing %
    neg_pct: float = 0         # Negotiation %
    transport_amt: float = 0   # Transport (flat Rs.)


def _build_spec_comparison_table(docx_path, columns, rows):
    """Rewrite the rendered 'Parameter | Specification' placeholder table into a
    side-by-side comparison: parameters down the left column, one column per
    equipment with the equipment name as the heading. Modifies the placeholder
    in place so it keeps the offer's table style (borders, header shading)."""
    if not columns:
        return
    import copy as _copy
    from docx import Document as _Docx
    from docx.oxml import OxmlElement as _OE
    from docx.oxml.ns import qn as _qn

    doc = _Docx(docx_path)
    table = None
    for t in doc.tables:
        if t.rows and t.rows[0].cells and t.rows[0].cells[0].text.strip().lower() == "parameter":
            table = t
            break
    if table is None:
        return

    ncol = 1 + len(columns)
    grid = table._tbl.find(_qn("w:tblGrid"))
    # Extend to one column per equipment by cloning the last cell of each row
    # (so new cells inherit borders / the header shading).
    guard = 0
    while len(table.columns) < ncol and guard < 40:
        guard += 1
        if grid is not None:
            grid.append(_OE("w:gridCol"))
        for row in table.rows:
            tcs = row._tr.findall(_qn("w:tc"))
            new_tc = _copy.deepcopy(tcs[-1])
            for t_el in new_tc.iter(_qn("w:t")):
                t_el.text = ""
            row._tr.append(new_tc)

    hdr = table.rows[0].cells
    hdr[0].text = "Parameter"
    for j, name in enumerate(columns):
        if 1 + j < len(hdr):
            hdr[1 + j].text = str(name)
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True

    for r in rows:
        cells = table.add_row().cells
        cells[0].text = str(r.get("param", ""))
        for j, v in enumerate(r.get("values", [])):
            if 1 + j < len(cells):
                cells[1 + j].text = str(v)

    table.autofit = True
    doc.save(docx_path)


def _build_narrative_scope_combined(combined_path, equipments, cust_base):
    """Replace the grouped Scope-of-Supply table in the combined offer with the
    full narrative scope of each equipment, reusing the standalone renderer
    (generate_quote_docx) per equipment and splicing in its scope section.
    Safe: any failure leaves the existing grouped scope untouched."""
    import copy as _copy
    from docx import Document as _Docx
    from docx.oxml import OxmlElement as _OE
    from docx.oxml.ns import qn as _qn
    from engine.quote_writer import generate_quote_docx as _gqd
    MARK = "►"

    def _retext(p_el, text):
        runs = p_el.findall(_qn('w:r'))
        if not runs:
            return
        first = runs[0]
        for t in first.findall(_qn('w:t')):
            first.remove(t)
        for r in runs[1:]:
            p_el.remove(r)
        t = _OE('w:t'); t.set(_qn('xml:space'), 'preserve'); t.text = text
        first.append(t)

    blocks = []
    for eq in equipments:
        qp = getattr(eq, "quote_payload", None) or {}
        if not qp:
            continue
        customer = dict(cust_base)
        for k, v in qp.items():
            if k not in ("bom_items", "items"):
                customer[k] = v
        customer["bom_items"] = qp.get("bom_items", [])
        qd = {"quote_seq": 0, "quote_no": "", "date": "", "customer": customer,
              "items": qp.get("items", [])}
        tmp = os.path.join(QUOTES_FOLDER, f"_scope_{abs(hash(eq.name)) % 999999}.docx")
        head_el, paras = None, []
        try:
            _gqd(qd, tmp)
            ed = _Docx(tmp)
            collecting = False
            for p in ed.paragraphs:
                u = (p.text or "").strip().upper()
                sty = (p.style.name if p.style else "") or ""
                if not collecting and u == "SCOPE OF SUPPLY":
                    head_el = _copy.deepcopy(p._p)
                    collecting = True
                    continue
                if collecting:
                    # Stop at the next annexure — including the formal
                    # "ANNEXURE I — SCOPE OF SUPPLY" heading that follows the
                    # body scope (it contains "SCOPE OF SUPPLY", so must NOT be
                    # excluded or we'd swallow it + its intro into every block).
                    if u.startswith("ANNEXURE"):
                        break
                    if not paras and not sty.startswith("Heading 3"):
                        continue   # skip the per-equipment intro line
                    paras.append(_copy.deepcopy(p._p))
        except Exception:
            paras = []
        finally:
            if os.path.exists(tmp):
                try: os.remove(tmp)
                except Exception: pass
        if paras and head_el is not None:
            blocks.append((eq.name, head_el, paras))

    if not blocks:
        return
    doc = _Docx(combined_path)
    target = None
    for t in doc.tables:
        full = "\n".join(c.text for r in t.rows for c in r.cells)
        if MARK in full or "Combustion Air Train" in full or "ENCON Supplied" in full:
            target = t
            break
    if target is None:
        return
    anchor = target._tbl
    for eqname, head_el, paras in blocks:
        h = _copy.deepcopy(head_el)
        _retext(h, (eqname or "Equipment").upper())
        anchor.addprevious(h)
        for pel in paras:
            anchor.addprevious(_copy.deepcopy(pel))
    anchor.getparent().remove(anchor)
    doc.save(combined_path)


@app.post("/api/generate-combined-offer")
def generate_combined_offer(req: CombinedOfferRequest):
    """Render Combined_Offer_Template.docx — one shared cover/customer/T&C,
    a technical row per equipment, and one combined price schedule."""
    try:
        from datetime import datetime as _dt
        from docxtpl import DocxTemplate
        from engine.quote_writer import amount_in_words_indian, _format_inr

        if not req.equipments:
            return {"error": "No equipment supplied."}
        cust = req.customer

        seq = next_quote_seq()
        full_ref = build_enquiry_ref(seq, cust.technical or "", cust.location or "")
        short_ref = full_ref.split(" DT.")[0]
        date_str = full_ref.split(" DT.")[-1] if " DT." in full_ref else _dt.now().strftime("%d/%m/%Y")
        company_address = ", ".join(filter(None, [
            (cust.address or "").strip(), (cust.city or "").strip(),
            (cust.state or "").strip(), (cust.pin or "").strip()]))

        # If two equipment share the same name (e.g. two 60 T verticals), append
        # a differentiating factor in brackets so they can be told apart across
        # the specs, scope and price schedule. Pick the first spec whose value is
        # distinct for each; fall back to a unit number.
        from collections import defaultdict as _dd
        _groups = _dd(list)
        for eq in req.equipments:
            _groups[(eq.name or "").strip()].append(eq)
        _PRIORITY = ["Fuel", "CV of Fuel", "Main Burner", "Combustion Air Blower",
                     "Air Line Size", "Gas Train Size", "Max. Heating Temperature",
                     "Burner Firing Rate", "Fuel Consumption Rate"]
        for _nm, _grp in _groups.items():
            if len(_grp) < 2:
                continue
            _maps = []
            for eq in _grp:
                m = {}
                for s in (eq.spec_rows or []):
                    l = str((s or {}).get("label", "")).strip()
                    v = str((s or {}).get("value", "")).strip()
                    if l and l not in m:
                        m[l] = v
                _maps.append(m)
            _chosen = None
            for lbl in _PRIORITY + sorted({k for m in _maps for k in m}):
                vals = [m.get(lbl, "") for m in _maps]
                if all(vals) and len(set(vals)) == len(vals):
                    _chosen = lbl
                    break
            for _i, eq in enumerate(_grp):
                factor = _maps[_i].get(_chosen, "") if _chosen else f"Unit {_i + 1}"
                eq.name = f"{(eq.name or '').strip()} ({factor})"

        # Price schedule — uses the SAME math as the costing Excel's Cost
        # Breakdown sheet so the offer's grand total matches it exactly:
        #   1. each equipment's sell = unit x (1 + negotiation%), rounded to the
        #      nearest Rs.10,000;  basic = sum of (sell x qty)
        #   2. P&F% + Designing% (of basic) + Transport(=Freight) = additional cost
        #   3. that additional cost is spread across the equipments in proportion
        #      to their price; grand total = basic + additional cost.
        # Packaging & Forwarding, Designing, Negotiation and Transport are NOT
        # shown as separate lines — they are baked into the equipment prices.
        _negp = float(req.neg_pct or 0); _pfp = float(req.pf_pct or 0)
        _desp = float(req.design_pct or 0); _trn = float(req.transport_amt or 0)
        _subs = []
        basic = 0.0
        _sn = 0
        for eq in req.equipments:
            unit = float(eq.unit_price or 0)
            if unit <= 0:
                continue          # skip unconfigured / zero-price equipment — no blank ₹0 line
            _sn += 1
            qty = max(1, int(eq.qty or 1))
            sell = round((unit * (1 + _negp / 100)) / 10000) * 10000   # nearest Rs.10,000
            sub = sell * qty
            basic += sub
            _subs.append((_sn, eq, qty, sub))
        _addl  = basic * _pfp / 100 + basic * _desp / 100 + _trn       # P&F + designing + freight
        _ratio = (_addl / basic) if basic else 0.0
        # amounts kept for the (hidden) context fields / words
        _pf, _des, _neg = basic * _pfp / 100, basic * _desp / 100, 0.0
        price_lines, grand = [], 0.0
        for i, eq, qty, sub in _subs:
            line_total = sub * (1 + _ratio)        # equipment's share of the additional cost
            grand += line_total
            price_lines.append({
                "sno":        f"{i}.",
                "name":       eq.name,
                "qty":        f"{qty:02d} No.",
                "unit_price": _format_inr(line_total / qty),
                "total":      _format_inr(line_total),
            })

        # Scope of Supply (Annexure I) — modelled on the reference offer:
        # per equipment, the BOM is grouped by system (Combustion Air Train,
        # Gas Train, Nitrogen Train, …) with lettered sub-items + makes, then
        # the standard ENCON inclusions (Control Panel internals, I&C, safety)
        # are appended. docxtpl doesn't XML-escape free text reliably here, so
        # escape & < > ourselves (e.g. makes like L&T).
        import xml.sax.saxutils as _saxutils, string as _string
        def _xesc(s):
            return _saxutils.escape(str(s))

        def _system_name(media: str) -> str:
            m = (media or "").strip()
            mu = m.upper()
            mapping = {
                "COMB AIR": "Combustion Air Train",
                "ENCON ITEMS": "ENCON Supplied Equipment",
                "PURGING LINE": "Nitrogen Purging Train",
                "MISC ITEMS": "Miscellaneous Items",
            }
            if mu in mapping:
                return mapping[mu]
            if "PILOT LINE" in mu:
                return "Pilot Gas Train"
            if mu.endswith("LINE"):
                base = m[:-4].strip()          # keep fuel acronym case, e.g. "NG", "COG"
                return base + (" Train" if "GAS" in base.upper() else " Gas Train")
            return m.title() if m else "Items"

        # Standard inclusions present on every ENCON offer (not all are priced
        # BOM lines). Makes follow the reference offer.
        STANDARD_SCOPE = [
            ("Instrumentation & Control", [
                ("Thermocouples", "2 Nos", "Toshniwal / Tempsen / SBI"),
                ("Temperature Transmitters", "2 Nos", "Honeywell / ABB"),
            ]),
            ("Control Panel", [
                ("Isolation Switch", "1 No", "Siemens / ABB"),
                ("Emergency Stop", "1 No", "Standard"),
                ("Ammeter & Voltmeter", "1 Set", "Standard"),
                ("MCB / MPCB", "1 Lot", "Siemens / ABB"),
                ("Contactors & Relays", "1 Lot", "Siemens / ABB"),
                ("Temperature Indicator", "1 No", "Masibus"),
                ("PLC with HMI", "1 No", "Siemens"),
            ]),
            ("Safety Systems", [
                ("Gas Leakage Detection System", "1 Set", "Honeywell"),
                ("Limit Switches", "2 Nos", "BCH"),
            ]),
        ]
        _ltr = _string.ascii_lowercase

        def _subdesc(item, ref, qty, make):
            d = item
            if ref:
                d += f" — {ref}"
            extra = []
            if qty:
                extra.append(f"{qty} No.")
            if make:
                extra.append(make)
            if extra:
                d += "   ·   " + "  ·  ".join(str(x) for x in extra)
            return d

        scope_rows = []
        for eq in req.equipments:
            scope_rows.append({"sno": "", "desc": _xesc((eq.name or "Equipment").upper())})
            # group this equipment's BOM by media, preserving first-seen order
            groups, order = {}, []
            for b in (eq.bom or []):
                item = (b.get("item") or "").strip()
                if not item or item.upper() == "CONTROL PANEL":
                    continue  # control panel is expanded via standard inclusions
                media = (b.get("media") or "Items").strip() or "Items"
                if media not in groups:
                    groups[media] = []
                    order.append(media)
                groups[media].append(b)
            sysno = 0
            for media in order:
                sysno += 1
                scope_rows.append({"sno": f"{sysno}.", "desc": _xesc(_system_name(media))})
                for i, b in enumerate(groups[media]):
                    lbl = _ltr[i] if i < 26 else str(i + 1)
                    desc = _subdesc((b.get("item") or "").strip(),
                                    (b.get("ref") or b.get("size") or "").strip(),
                                    b.get("qty"), (b.get("make") or "").strip())
                    scope_rows.append({"sno": lbl, "desc": _xesc(desc)})
            for sysname, items in STANDARD_SCOPE:
                sysno += 1
                scope_rows.append({"sno": f"{sysno}.", "desc": _xesc(sysname)})
                for i, (it, q, mk) in enumerate(items):
                    scope_rows.append({"sno": _ltr[i], "desc": _xesc(f"{it}   ·   {q}  ·  {mk}")})

        # Technical Specifications — side by side: parameters down the left, one
        # column per equipment (equipment name = column heading). Built directly
        # with python-docx after rendering (dynamic column count), so plain text
        # here — no XML escaping needed.
        spec_columns = []          # equipment names (column headings)
        _eq_maps     = []          # per-equipment {param: value}
        _param_order = []          # union of parameter names, first-seen order
        for eq in req.equipments:
            spec_columns.append(eq.name or "Equipment")
            rows = eq.spec_rows or []
            if not rows and (eq.specs or "").strip():
                rows = [{"label": "Specifications", "value": eq.specs}]
            m = {}
            for s in rows:
                lbl = str((s or {}).get("label", "")).strip()
                if not lbl:
                    continue
                if lbl not in m:
                    m[lbl] = str((s or {}).get("value", "")).strip()
                if lbl not in _param_order:
                    _param_order.append(lbl)
            _eq_maps.append(m)
        # Drop equipment columns that carry no spec data at all — so an
        # equipment that isn't part of this offer (or wasn't configured) doesn't
        # leave a blank column in the Technical Specifications table.
        _kept = [(n, m) for n, m in zip(spec_columns, _eq_maps)
                 if any((v or "").strip() for v in m.values())]
        if _kept:
            spec_columns = [n for n, _ in _kept]
            _eq_maps = [m for _, m in _kept]
        _param_order = []
        for m in _eq_maps:
            for lbl in m:
                if lbl not in _param_order:
                    _param_order.append(lbl)
        spec_rows = [{"param": p, "values": [m.get(p, "") for m in _eq_maps]}
                     for p in _param_order]

        # P&F, Designing, Negotiation AND Transport are all distributed into the
        # equipment prices above (grand), so the grand total is the final figure
        # — no separate transport line.
        _combined_final = grand

        from engine.quote_writer import _supervision_rates
        _sup_mech, _sup_plc = _supervision_rates()

        ctx = {
            "project_name":      req.project_name or "Combined Equipment Offer",
            "subject":           cust.subject or "Offer for Combined Equipment",
            "application":       "Combined Equipment",
            "equipment_name":    "Combined Equipment Offer",
            "company_name":      cust.company or "",
            "company_address":   company_address,
            "email":             cust.email or "",
            "mobile_no":         cust.phone or "",
            "poc_name":          _with_salutation(cust.salutation, cust.name),
            "poc_designation":   cust.designation or "",
            "client_enq_ref":    "",
            "enquiry_ref":       full_ref,
            "enquiry_ref_short": short_ref,
            "enquiry_date_str":  date_str,
            "marketing_person":  _with_salutation(cust.marketing_salutation, cust.marketing),
            "marketing_email":   cust.marketing_email or "",
            "marketing_phone":   cust.marketing_phone or "",
            "technical_person":  _with_salutation(cust.technical_salutation, cust.technical) or _with_salutation(cust.marketing_salutation, cust.marketing),
            "technical_phone":   cust.technical_phone or cust.marketing_phone or "",
            "technical_email":   cust.technical_email or cust.marketing_email or "",
            # technical section + scope of supply + price schedule loops
            "equipments":   [{"name": e.name, "specs": e.specs or ""} for e in req.equipments],
            "scope_rows":   scope_rows,
            "price_lines":  price_lines,
            "grand_total":  _format_inr(grand),
            # P&F, Designing, Negotiation and Transport are all distributed into
            # the equipment prices — none shown as a separate price-schedule line.
            "pf_amount":         _format_inr(_pf),
            "design_amount":     _format_inr(_des),
            "neg_amount":        _format_inr(_neg),
            "transport_amount":  _format_inr(_trn),
            "show_pf":           False,
            "show_design":       False,
            "show_neg":          False,
            "show_transport":    False,
            "final_total":       _format_inr(_combined_final),
            "grand_total_in_words": f"INR. {amount_in_words_indian(round(_combined_final))} ONLY.",
            # Supervision charges — pulled from the price master (component_price_master).
            "supervision_mech":  _sup_mech,
            "supervision_plc":   _sup_plc,
            # shared T&C
            "tnc_prices":             cust.tnc_prices or "",
            "tnc_delivery":           cust.tnc_delivery or "",
            "tnc_gst":                cust.tnc_gst or "",
            "tnc_hsn_code":           cust.tnc_hsn_code or "",
            "tnc_pan_gst":            cust.tnc_pan_gst or "",
            "tnc_payment_terms":      cust.tnc_payment_terms or "",
            "tnc_packing_forwarding": cust.tnc_packing_forwarding or "",
            "tnc_freight":            cust.tnc_freight or "",
            "tnc_transit_insurance":  cust.tnc_transit_insurance or "",
            "tnc_validity":           cust.tnc_validity or "",
            "tnc_inspection":         cust.tnc_inspection or "",
            "tnc_guarantee":          cust.tnc_guarantee or "",
        }

        tpl = DocxTemplate(os.path.join(BASE_DIR, "Combined_Offer_Template.docx"))
        tpl.render(ctx)
        safe_company = "".join(ch for ch in (cust.company or "Client")
                               if ch.isalnum() or ch in " _-").strip().replace(" ", "_") or "Client"
        docx_name = f"Combined_Offer_{safe_company}_{seq}.docx"
        docx_path = os.path.join(QUOTES_FOLDER, docx_name)
        tpl.save(docx_path)
        # Build the side-by-side technical-spec table (one column per equipment).
        _build_spec_comparison_table(docx_path, spec_columns, spec_rows)
        # Replace the grouped Scope of Supply with the full narrative scope per
        # equipment (falls back to the grouped scope if anything goes wrong).
        try:
            _build_narrative_scope_combined(
                docx_path, req.equipments,
                {"company_name": cust.company or "Client",
                 "project_name": req.project_name or "",
                 "subject": cust.subject or "", "address": "", "poc_name": ""})
        except Exception as _scope_err:
            print(f"WARN: narrative scope build failed, keeping grouped scope: {_scope_err}")

        pdf_name = docx_name.replace(".docx", ".pdf")
        pdf_ok = _docx_to_pdf(docx_path, os.path.join(QUOTES_FOLDER, pdf_name))
        try:
            from engine.drive_uploader import upload_offer_async
            upload_offer_async(docx_path, docx_name, "combined")
            if pdf_ok:
                upload_offer_async(os.path.join(QUOTES_FOLDER, pdf_name), pdf_name, "combined")
        except Exception as _drv_err:
            print(f"WARN: drive upload kickoff failed: {_drv_err}")

        try:
            from engine.drive_uploader import drive_status as _drive_status
            _drv = _drive_status("combined")
        except Exception:
            _drv = {"ok": False, "reason": "error", "msg": "Drive status unknown"}

        _log_equipment_quote(cust, "Combined", grand, docx_path)

        return {
            "success":      True,
            "filename":     docx_name,
            "download_url": f"/api/download-quote/{docx_name}",
            "pdf_url":      f"/api/pdf-quote/{pdf_name}" if pdf_ok else None,
            "preview_url":  f"/api/preview-quote/{docx_name}",
            "quote_no":     full_ref,
            "grand_total":  grand,
            "count":        len(req.equipments),
            "drive":        _drv,
        }
    except Exception as e:
        import traceback
        return {"error": str(e), "trace": traceback.format_exc()}


@app.get("/api/next-quote-ref")
def api_next_quote_ref(technical_person: str = "", location: str = ""):
    """Preview the auto-generated enquiry ref for the form.
    Reads (does not consume) the next sequence number."""
    seq = peek_quote_seq()
    return {"seq": seq, "ref": build_enquiry_ref(seq, technical_person, location)}


_FX_CACHE = {}   # {to: {"rate": float, "ts": epoch, "date": str}}

@app.get("/api/fx-rate")
def api_fx_rate(to: str = "USD"):
    """Live INR → `to` exchange rate (default USD). Cached ~6h; falls back to a
    sensible constant if the provider is unreachable so the UI never breaks."""
    import time as _t, json as _j, urllib.request as _u
    to = (to or "USD").upper()
    now = _t.time()
    c = _FX_CACHE.get(to)
    if c and (now - c["ts"] < 6 * 3600):
        return {"from": "INR", "to": to, "rate": c["rate"], "inverse": round(1 / c["rate"], 4),
                "date": c["date"], "cached": True}
    rate = None; date = ""
    try:
        with _u.urlopen("https://open.er-api.com/v6/latest/INR", timeout=6) as resp:
            d = _j.loads(resp.read().decode())
            rate = float(d["rates"][to]); date = d.get("time_last_update_utc", "") or ""
    except Exception:
        rate = None
    fallback = False
    if not rate or rate <= 0:
        rate = {"USD": 0.01198, "EUR": 0.0110, "GBP": 0.0094}.get(to, 0.012)  # ~₹83.5/USD
        date = "fallback"; fallback = True
    _FX_CACHE[to] = {"rate": rate, "ts": now, "date": date}
    return {"from": "INR", "to": to, "rate": rate, "inverse": round(1 / rate, 4),
            "date": date, "cached": False, "fallback": fallback}


def _break_out_transport(docx_path: str, transport: float):
    """In a standalone offer's Annexure III, pull Transport out of the single
    all-inclusive price onto its own line: reduce the equipment row by the
    transport amount and insert a 'Transport' row just before TOTAL (TOTAL is
    unchanged, so the schedule still reconciles). No-op when transport <= 0."""
    try:
        transport = float(transport or 0)
    except (TypeError, ValueError):
        return
    if transport <= 0:
        return
    import copy, re
    from docx import Document
    from engine.quote_writer import _format_inr

    def _num(s):
        s = (s or "").replace(",", "").replace("₹", "").strip()
        try:
            return float(s)
        except ValueError:
            return None

    def _set_cell(cell, text):
        # set text on the first run (preserve its formatting); clear the rest
        if not cell.paragraphs:
            cell.text = text
            return
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = text
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.add_run(text)

    doc = Document(docx_path)
    for t in doc.tables:
        if not t.rows:
            continue
        hdr = " ".join(c.text for c in t.rows[0].cells).lower()
        if "unit price" not in hdr or "total price" not in hdr:
            continue
        total_idx = next((i for i, r in enumerate(t.rows)
                          if any(c.text.strip().upper() == "TOTAL" for c in r.cells)), None)
        if not total_idx:               # None or 0 (TOTAL is never the header row)
            continue
        eq = t.rows[total_idx - 1]
        if len(eq.cells) < 5:
            continue
        tot = _num(eq.cells[4].text)
        if tot is None:
            continue
        qty = 1
        m = re.search(r"\d+", eq.cells[2].text or "")
        if m:
            qty = max(1, int(m.group()))
        new_total = tot - transport
        _set_cell(eq.cells[3], _format_inr(new_total / qty))
        _set_cell(eq.cells[4], _format_inr(new_total))
        # insert a Transport row before TOTAL (clone the equipment row to match style)
        new_tr = copy.deepcopy(eq._tr)
        t.rows[total_idx]._tr.addprevious(new_tr)
        trow = t.rows[total_idx]
        for i, lab in enumerate(["", "Transport", "", "", _format_inr(transport)]):
            if i < len(trow.cells):
                _set_cell(trow.cells[i], lab)
        break
    doc.save(docx_path)


@app.post("/api/generate-quote")
async def generate_quote(req: QuoteRequest):
    try:
        from engine.quote_engine import calculate_quote
        from engine.quote_writer import generate_quote_docx
        from engine.pdf_writer import generate_quote_pdf

        seq = next_quote_seq()
        # Auto-generate enquiry reference: ET{YY}-{seq}-{initials}.
        # Overrides whatever the form sent so the reference always follows
        # the canonical ENCON pattern (ET26-001-JR for FY2026/Jyotirmoy Rabha).
        auto_ref = build_enquiry_ref(seq, req.technical_person or "", req.location or "")
        form_data = {
            "quote_seq": seq,
            "customer": {
                "company_name":    req.company_name,
                "company_city":    req.company_city,
                "company_state":   req.company_state,
                "address":         ", ".join(filter(None, [req.company_address, req.company_city, req.company_state, req.company_pin])),
                "poc_name":        _with_salutation(req.poc_salutation, req.poc_name),
                "poc_designation": req.poc_designation,
                "mobile_no":       req.mobile_no,
                "email":           req.email,
                "project_name":    req.project_name,
                "subject":         req.subject or req.project_name,
                "ref_no":          auto_ref,
                "your_ref":        auto_ref,
                "enquiry_ref":     auto_ref,
                "marketing_person": _with_salutation(req.marketing_salutation, req.marketing_person),
                "marketing_phone": req.marketing_phone,
                "marketing_email": req.marketing_email,
                "technical_person": _with_salutation(req.technical_salutation, req.technical_person),
                "technical_phone": req.technical_phone,
                "technical_email": req.technical_email,
                "gstin":           req.company_gstin,
                "currency":        (req.currency or "INR"),
                "fx_rate":         (req.fx_rate or 0),
                "extra_context":   (req.extra_context or {}),
                # Technical data (for template tech table)
                "ladle_tons":          req.ladle_tons,
                "ladle_dim":           req.ladle_dim,
                "ladle_drawing_no":    req.ladle_drawing_no,
                "refractory_weight_kg": req.refractory_weight_kg,
                "heating_schedule":    req.heating_schedule,
                "fuel_cv":             req.fuel_cv,
                "fuel_consumption":    req.fuel_consumption,
                "burner_model":        req.burner_model,
                "blower_model":        req.blower_model,
                "blower_size":         req.blower_size,
                "blower_capacity":     req.blower_capacity,
                "hydraulic_motor_hp":  req.hydraulic_motor_hp,
                "max_electrical_load": req.max_electrical_load,
                "total_in_words":      req.total_in_words,
                "heating_time":        req.heating_time,
                "fuel_name":           req.fuel_name,
                "burner_capacity_range": req.burner_capacity_range,
                "pumping_unit":        req.pumping_unit,
                "hood_movement":       req.hood_movement,
                "hood_type":           req.hood_type or "up_down",
                "pilot_gas_type":      req.pilot_gas_type or "LPG",
                "ignition_method":     req.ignition_method,
                "num_burners":         req.num_burners,
                "fuel2_cv":            req.fuel2_cv,
                "fuel2_consumption":   req.fuel2_consumption,
                "max_fuel_consumption1": req.max_fuel_consumption1,
                "max_fuel_consumption2": req.max_fuel_consumption2,
                "is_oil":              bool(req.is_oil),
                "is_dual":             bool(req.is_dual),
                "control_mode":        req.control_mode or "automatic",
                "auto_control_type":   req.auto_control_type or "plc",
                "control_valve_type":  req.control_valve_type or "pneumatic",
                "special_auto_ignition": bool(req.special_auto_ignition),
                "special_auto_controls": bool(req.special_auto_controls),
                "vertical_qty":        int(req.vertical_qty or 1),
                "horizontal_qty":      int(req.horizontal_qty or 1),
                "nitrogen_purging":    (req.purging_line or "no").lower() == "yes",
                "hpu_variant":         req.hpu_variant or "Duplex 1",
                "burner_kw_value":     req.burner_kw_value or "",
                "bom_items":           req.bom_items or [],
                # Annexure IV — Terms & Conditions (editable per offer)
                "tnc_prices":             req.tnc_prices or "",
                "tnc_delivery":           req.tnc_delivery or "",
                "tnc_gst":                req.tnc_gst or "",
                "tnc_hsn_code":           req.tnc_hsn_code or "",
                "tnc_pan_gst":            req.tnc_pan_gst or "",
                "tnc_payment_terms":      req.tnc_payment_terms or "",
                "tnc_packing_forwarding": req.tnc_packing_forwarding or "",
                "tnc_freight":            req.tnc_freight or "",
                "tnc_transit_insurance":  req.tnc_transit_insurance or "",
                "tnc_validity":           req.tnc_validity or "",
                "tnc_inspection":         req.tnc_inspection or "",
                "tnc_guarantee":          req.tnc_guarantee or "",
            },
            "items": [item.dict() for item in req.items],
            "gst_percent": req.gst_percent,
            "freight": req.freight,
            "valid_days": req.valid_days,
        }

        quote_data = calculate_quote(form_data)
        # Filename format: {YYYY-MM-DD}_{Customer}_{Product}-{Capacity}T.docx
        # (e.g. '2026-05-30_MagnoSteel_VLPH-10T.docx'). Falls back gracefully
        # if fields are missing.
        _PRODUCT_SHORT = {
            "Vertical Ladle Preheater":   "VLPH",
            "Horizontal Ladle Preheater": "HLPH",
        }
        _first_pt = (req.items[0].product_type if req.items else "") or ""
        _product = _PRODUCT_SHORT.get(_first_pt) or (
            "Tundish" if "Tundish" in _first_pt
            else "Recup" if "Recuperator" in _first_pt
            else (_first_pt.replace(" ", "") or "Offer")
        )
        _safe_company = "".join(ch for ch in (req.company_name or "Client")
                                if ch.isalnum() or ch in " _-").strip().replace(" ", "_") or "Client"
        _date = datetime.now().strftime("%Y-%m-%d")
        _capacity = f"-{int(req.ladle_tons)}T" if req.ladle_tons else ""
        filename = f"{_date}_{_safe_company}_{_product}{_capacity}.docx"
        output_path = os.path.join(QUOTES_FOLDER, filename)
        # Regenerative-burner offers use their own template (blower-style cover
        # letter + regen technical body); everything else uses the VLPH template.
        _regen_tpl = None
        if "regen" in _first_pt.lower():
            # Enrich the regen body's template vars from the raw calc values.
            from engine.quote_writer import amount_in_words_indian, _format_inr
            _ec = dict(req.extra_context or {})
            _rp = int(_ec.get("regen_pairs") or 1)
            _rf = str(_ec.get("regen_fuel") or "Natural Gas")
            _rk = _ec.get("regen_kw") or ""
            _price = float(req.items[0].total) if req.items else 0.0
            _oil = _rf.strip().lower() in {"hsd", "ldo", "hdo", "fo", "sko", "cfo", "lshs", "oil"}
            # Oil fuels use the dedicated oil template; gases use the standard one.
            _base = os.path.dirname(os.path.abspath(__file__))
            _tpl_name = "Regen_Oil_Offer_Template.docx" if _oil else "Regen_Offer_Template.docx"
            _cand = os.path.join(_base, _tpl_name)
            if os.path.exists(_cand):
                _regen_tpl = _cand
            elif os.path.exists(os.path.join(_base, "Regen_Offer_Template.docx")):
                _regen_tpl = os.path.join(_base, "Regen_Offer_Template.docx")   # fallback
            _fword = "OIL" if _oil else ("NG" if "natural" in _rf.lower() else _rf.upper())
            # Gas-train label: strip a trailing "GAS" so "{fuel} GAS TRAIN" doesn't
            # read "BLAST FURNACE GAS GAS TRAIN". NG/PNG keep their names.
            import re as _re_gt
            _gtf = _re_gt.sub(r"\s*GAS\s*$", "", _fword, flags=_re_gt.I).strip() or _fword
            _qtyw = f"{_rp} Pair" + ("s" if _rp > 1 else "")
            # Gas-train supply pressure: NG at 2.1 bar; all other gas fuels at 1000 mm.
            _press_clause = ("Pressure at TOP 2.1 bar and minimum pressure as 1 bar is considered."
                             if _fword == "NG"
                             else "Pressure at 1000 mm is considered.")
            _ec.update({
                "fuel_word": _fword, "gas_train_fuel": _gtf, "is_oil": _oil,
                "pressure_clause": _press_clause,
                "savings_fuel": ("Fuel Oil" if _oil else _rf),
                "fuel_name": _rf, "kw": _rk, "pairs": _rp,
                "burner_count": f"{_rp * 2} Nos", "qty_words": _qtyw,
                "price_line_desc": f"{_qtyw} Regenerative Burner System with PLC",
                "price_inr": _format_inr(_price),
                "price_in_words": "INR " + amount_in_words_indian(_price) + " only.",
            })
            quote_data.setdefault("customer", {})["extra_context"] = _ec
        generate_quote_docx(quote_data, output_path, template_path=_regen_tpl)
        # Regen offers: build the MAKE LIST dynamically from the real BOM — only
        # the item categories actually present (with their makes) are listed.
        if _regen_tpl:
            try:
                from bom.regen_builder import build_regen_df, select_model
                _mkw = select_model(float(_rk)) if _rk else 1000
                _bomdf = build_regen_df(_mkw, num_pairs=_rp, fuel=_rf, db_path=DB_PATH)
                from engine.regen_bom_table import (
                    fill_make_list, fill_temp_control, fill_gas_train,
                    fill_oil_supply, fill_consist_list)
                if not fill_make_list(output_path, _bomdf):
                    print("WARN: regen MAKE LIST table not found in template")
                try:
                    fill_consist_list(output_path, _oil, _gtf)
                except Exception as _cl_err:
                    print(f"WARN: regen consist-list fill failed: {_cl_err}")
                try:
                    fill_temp_control(output_path, _bomdf)
                except Exception as _tc_err:
                    print(f"WARN: regen TEMP CONTROL fill failed: {_tc_err}")
                try:
                    # oil offers -> HPU/oil-line section; gas -> gas-train section
                    if not fill_oil_supply(output_path, _bomdf):
                        fill_gas_train(output_path, _bomdf)
                except Exception as _gt_err:
                    print(f"WARN: regen fuel-supply fill failed: {_gt_err}")
            except Exception as _bom_err:
                print(f"WARN: regen MAKE LIST fill failed: {_bom_err}")
        # Pull Transport onto its own price-schedule line (P&F/designing/
        # negotiation stay baked into the equipment price). No-op if 0.
        try:
            _break_out_transport(output_path, req.transport_amt)
        except Exception as _trn_err:
            print(f"WARN: transport line break-out failed: {_trn_err}")

        # Build the PDF as a faithful render of the .docx via LibreOffice
        # so the PDF and Word offer are visually identical. Fall back to the
        # reportlab generator only if LibreOffice is missing or fails (e.g.
        # local Windows dev box without LibreOffice installed).
        pdf_path = os.path.splitext(output_path)[0] + ".pdf"
        try:
            if not _docx_to_pdf(output_path, pdf_path):
                generate_quote_pdf(quote_data, pdf_path)
        except Exception as _pdf_err:
            print(f"WARN: PDF generation failed for {filename}: {_pdf_err}")

        # Persist the raw quote_data so the PDF can be regenerated later
        # if the .pdf file gets wiped (e.g., container redeploy).
        try:
            import json as _json
            json_path = os.path.splitext(output_path)[0] + ".json"
            with open(json_path, "w", encoding="utf-8") as _jf:
                _json.dump(quote_data, _jf, default=str)
        except Exception as _json_err:
            print(f"WARN: quote-data persist failed for {filename}: {_json_err}")

        # Mirror the docx + pdf to Google Drive (process@encon.in) in a
        # background thread. Drive credentials live in env vars; failures
        # are logged but never break offer generation.
        try:
            from engine.drive_uploader import upload_offer_async
            _product = (req.items[0].product_type if req.items else "")
            upload_offer_async(output_path, filename, _product)
            pdf_filename = f"{os.path.splitext(filename)[0]}.pdf"
            upload_offer_async(pdf_path, pdf_filename, _product)
        except Exception as _drv_err:
            print(f"WARN: drive upload kickoff failed: {_drv_err}")

        # Persist to quotes_log so we can list/re-download past quotes later
        _pt = (req.items[0].product_type if req.items else "") or ""
        _log_quote(
            quote_no=quote_data["quote_no"], ref_no=req.ref_no,
            company_name=req.company_name, poc_name=req.poc_name,
            email=req.email, mobile_no=req.mobile_no, project_name=req.project_name,
            equipment_type=("HLPH" if "horizontal" in _pt.lower() else "VLPH"),
            location=getattr(req, "location", ""), ladle_tons=req.ladle_tons,
            grand_total=quote_data["grand_total"], marketing_person=req.marketing_person,
            technical_person=req.technical_person, file_path=output_path,
        )

        return {
            "success": True,
            "quote_no": quote_data["quote_no"],
            "download_url": f"/api/download-quote/{filename}",
            "preview_url":  f"/api/preview-quote/{filename}",
            "pdf_url":      f"/api/pdf-quote/{filename}",
            "summary": {
                "subtotal": quote_data["subtotal"],
                "total":    quote_data["grand_total"],
            }
        }
    except Exception as e:
        import traceback
        return {"error": str(e), "trace": traceback.format_exc()}


@app.get("/api/price-master/items")
def price_master_items():
    """Return all items grouped by category for the price master page."""
    try:
        conn = sqlite3.connect(DB_PATH)
        rows = conn.execute(
            "SELECT category, item, unit, price, previous_price FROM component_price_master ORDER BY category, item"
        ).fetchall()
        conn.close()
        result = {}
        for cat, item, unit, price, prev in rows:
            result.setdefault(cat, []).append({
                "item": item, "unit": unit or "nos",
                "price": price, "previous_price": prev
            })
        return {"categories": result, "total": len(rows)}
    except Exception as e:
        return {"error": str(e), "categories": {}, "total": 0}


@app.get("/api/price-master/coverage")
def price_master_coverage():
    """Check which BOM items have prices in the DB and which are missing."""
    from bom.vlph_builder import LEGACY_ITEM_SEQUENCE
    from bom.static_items import static_items
    needed = set(LEGACY_ITEM_SEQUENCE) | {item for _, item, _, _ in static_items()}
    needed -= {"RATIO CONTROLLER"}  # excluded from BOM
    try:
        conn = sqlite3.connect(DB_PATH)
        have = {r[0] for r in conn.execute("SELECT item FROM component_price_master").fetchall()}
        conn.close()
    except Exception:
        have = set()
    covered = needed & have
    missing = needed - have
    return {
        "total_in_db":  len(have),
        "total_bom":    len(needed),
        "covered_count": len(covered),
        "missing_count": len(missing),
        "covered":  sorted(covered),
        "missing":  sorted(missing),
        "extra":    sorted(have - needed),
    }


@app.post("/api/upload-pricelist")
async def upload_pricelist(file: UploadFile = File(...)):
    """
    Parses the full ENCON Pricelist WorkBook.
    Updates all master tables: component_price_master, hpu_master,
    burner_pricelist_master, blower_pricelist_master, horizontal_master,
    vertical_master, recuperator_master, gail_gas_burner_master,
    rad_heat_master, rad_heat_tata_master, and burner parts tables.
    """
    try:
        file_path = os.path.join(UPLOAD_FOLDER, file.filename)
        with open(file_path, "wb") as buf:
            shutil.copyfileobj(file.file, buf)

        conn = sqlite3.connect(DB_PATH)
        results = _parse_pricelist_all(file_path, conn)
        clean_duplicate_rates(conn)
        conn.close()

        updated = {t: r for t, r in results.items() if "rows" in r}
        skipped = {t: r["skipped"] for t, r in results.items() if "skipped" in r}
        errors  = {t: r["error"]   for t, r in results.items() if "error"   in r}

        total_rows = sum(r["rows"] for r in updated.values())

        return {
            "success": True,
            "total_rows_loaded": total_rows,
            "tables_updated": {t: r["rows"] for t, r in updated.items()},
            "tables_skipped": skipped,
            "errors": errors,
        }

    except Exception as e:
        import traceback
        return {"error": str(e), "trace": traceback.format_exc()}


_STOCK_CACHE: dict = {}   # keyed by (file_path, mtime)

@app.get("/api/stock-rates")
def get_stock_rates():
    """
    Parse ENCON Stock Summary Excel and return all sections/items with purchase rates.
    Result is cached in memory; re-parsed only if the file changes.
    """
    import glob as glob_mod
    import openpyxl

    # Find stock file
    candidates = (
        glob_mod.glob(os.path.join(UPLOAD_FOLDER, "Stock*.xlsx")) +
        glob_mod.glob(os.path.join(BASE_DIR, "Stock*.xlsx"))
    )
    if not candidates:
        return {"error": "Stock file not found. Upload a file named Stock*.xlsx"}

    stock_path = candidates[0]
    mtime = os.path.getmtime(stock_path)
    cache_key = (stock_path, mtime)
    if cache_key in _STOCK_CACHE:
        return _STOCK_CACHE[cache_key]

    try:
        wb = openpyxl.load_workbook(stock_path, read_only=True, data_only=True)
        ws = wb["Stock Summary"]
    except Exception as e:
        return {"error": f"Cannot open stock file: {e}"}

    sections = []
    cur_section = "Raw Material"
    cur_items = []

    for r in range(9, 5000):
        v = [ws.cell(r, c).value for c in range(1, 6)]
        if all(x is None for x in v):
            break
        # New section header: col A = None, col B = text, col D ≠ 'TOTAL'
        if v[0] is None and v[1] and str(v[3] or "").strip() != "TOTAL":
            if cur_items:
                sections.append({"section": cur_section, "items": cur_items})
            cur_section = str(v[1]).strip()
            cur_items = []
        elif v[0] is not None and isinstance(v[0], (int, float)):
            name = str(v[1]).strip() if v[1] else ""
            rate = v[3]
            cur_items.append({
                "sl":   int(v[0]),
                "name": name,
                "rate": round(float(rate), 4) if isinstance(rate, (int, float)) else None,
            })

    if cur_items:
        sections.append({"section": cur_section, "items": cur_items})

    total_items = sum(len(s["items"]) for s in sections)
    result = {
        "file": os.path.basename(stock_path),
        "sections": sections,
        "total_items": total_items,
    }
    _STOCK_CACHE[cache_key] = result
    return result


# ── Stock → component_price_master mapping ────────────────────────────────────
# Maps pricelist item name (or prefix) → stock item name (exact)
_STOCK_PRICE_MAP = {
    # MS structural
    "M.S. Angle 65,50":       "M.S Angle 50x50x6 mm",
    "M.S. Channel":           "M.S Channel 100x50",
    "M.S.Chanel":             "M.S Channel 100x50",
    "M.S. Chanel":            "M.S Channel 100x50",
    "M.S. Angle 100,100":     "M.S Angle 100X100X10",
    "M.S. Angle 50*6":        "M.S Angle 50x50x6 mm",
    "M.S. Flat":              None,   # no good match
    "M.S. Round":             "M.S Round Dia 16 MM",
    # MS plates / sheets
    "M.S. Plate 5mm":         "M.S PLATE 1500X6300X5 MM",
    "M.S. Plate 8mm":         "M.S Plate 1500X6300X10 MM",
    "M.S. Plate 16mm* 5mm":   "M.S PLATE 1500X6300X5 MM",
    "M.S. Plate 16mm*5mm":    "M.S PLATE 1500X6300X5 MM",
    "M.S. Plate 16mm*10mm":   "M.S Plate 1500X6300X10 MM",
    "M.S. Sheet 2mm":         None,
    "M.S. Sheet  2mm":        None,
    "M.S. Sheet 3mm":         None,
    "M.S. Sheet  3mm":        None,
    "M.S. Sheet 4mm":         None,
    "M.S. Sheet 5mm":         "M.S Chequered Plate 1250X5000X5 MM",
    "M.S. Sheet 8mm":         "M.S Plate 1510x6310x16mm",
    # MS tubes
    'M.S. Tube "B" Class 1.5 in': "M.S ERW Pipe 40 NB",
    'M.S. Tube "C" Class 1.5 in': "M.S ERW Pipe 40 NB",
    "M.S. Tube B Class 1.5 in":   "M.S ERW Pipe 40 NB",
    "M.S. Tube C Class 1.5 in":   "M.S ERW Pipe 40 NB",
    # SS
    "S.S. Sheet 3mm":         "S.S Plate 1500X3000X3 MM",
    "SS Pipe 304 60 X 3mm":   "S.S 304 ERW PIPE 100 MM OD",
    "SS Pipe 304 60x3mm (per mtr)": "S.S 304 ERW PIPE 100 MM OD",
    "SS Pipe 304 76 X 3mm":   "S.S ERW PIPE OD 65 X ID 58 / 57 MM",
    "SS Pipe 304 76x3mm (per mtr)": "S.S ERW PIPE OD 65 X ID 58 / 57 MM",
    "SS Pipe 304 100 X 3mm":  "S.S 304 ERW Pipe 100 NB",
    "SS Pipe 304 100x3mm (per mtr)": "S.S 304 ERW Pipe 100 NB",
    # Refractories
    "Ceramic Fiber":          "Ceramic Fiber 128 Kg/m3",
    "Whyteheat K":            "Whytheat-A",
}

@app.post("/api/stock/sync")
def sync_stock_to_pricelist():
    """
    Reads the stock file, matches items to component_price_master via
    _STOCK_PRICE_MAP, and updates matched rows.  Returns a summary.
    """
    import glob as glob_mod
    import openpyxl

    candidates = (
        glob_mod.glob(os.path.join(UPLOAD_FOLDER, "Stock*.xlsx")) +
        glob_mod.glob(os.path.join(BASE_DIR, "Stock*.xlsx"))
    )
    if not candidates:
        return {"error": "Stock file not found"}
    stock_path = candidates[0]

    try:
        wb = openpyxl.load_workbook(stock_path, read_only=True, data_only=True)
        ws = wb["Stock Summary"]
    except Exception as e:
        return {"error": str(e)}

    # Build {name_lower: rate} lookup from stock
    stock_lookup = {}
    for r in range(9, 5000):
        sl   = ws.cell(r, 1).value
        name = ws.cell(r, 2).value
        rate = ws.cell(r, 4).value
        if sl is None and name is None:
            break
        if isinstance(sl, (int, float)) and name and isinstance(rate, (int, float)):
            stock_lookup[str(name).strip().lower()] = (str(name).strip(), float(rate))
    wb.close()

    conn = sqlite3.connect(DB_PATH)
    updated, skipped = [], []

    # Fetch all pricelist items
    pl_items = conn.execute("SELECT rowid, item, price FROM component_price_master").fetchall()
    for rowid, item, old_price in pl_items:
        stock_name = _STOCK_PRICE_MAP.get(item)
        if stock_name is None:
            if item in _STOCK_PRICE_MAP:
                skipped.append({"item": item, "reason": "manual_skip"})
            continue
        hit = stock_lookup.get(stock_name.lower())
        if not hit:
            skipped.append({"item": item, "stock_name": stock_name, "reason": "not_in_stock"})
            continue
        new_price = round(hit[1], 2)
        if new_price != old_price:
            conn.execute("UPDATE component_price_master SET price=? WHERE rowid=?", (new_price, rowid))
            updated.append({"item": item, "stock_name": hit[0], "old": old_price, "new": new_price})

    conn.commit()

    # ── Cascade: re-run all parsers so every computed amount/selling price
    #    reflects the updated rates (recuperator, HPU, blower, LPH, burner…)
    cascade_ok = False
    cascade_tables = {}
    xl_path = _find_latest_pricebook()
    if xl_path and updated:
        try:
            cascade_tables = _cascade_recalculate(xl_path, conn)
            conn.commit()
            cascade_ok = True
        except Exception as ce:
            cascade_tables = {"error": str(ce)}

    conn.close()
    _STOCK_CACHE.clear()
    return {
        "file": os.path.basename(stock_path),
        "updated": updated,
        "skipped_count": len(skipped),
        "updated_count": len(updated),
        "cascade": cascade_ok,
        "cascade_tables": {k: v for k, v in cascade_tables.items() if isinstance(v, dict)},
    }


@app.post("/api/stock/upload")
async def upload_stock_file(file: UploadFile = File(...)):
    """Accept a new Stock*.xlsx upload and save to BASE_DIR, then auto-sync."""
    import shutil
    # Validate name
    if not file.filename.startswith("Stock") or not file.filename.endswith(".xlsx"):
        return {"error": "File must be named Stock*.xlsx"}
    dest = os.path.join(BASE_DIR, file.filename)
    with open(dest, "wb") as f:
        shutil.copyfileobj(file.file, f)
    _STOCK_CACHE.clear()
    # Auto-sync after upload
    sync_result = sync_stock_to_pricelist()
    return {"uploaded": file.filename, "sync": sync_result}


class ExcelExportRequest(BaseModel):
    equipment_type: str          # "VLPH" | "HLPH" | "Regen"
    customer: dict = {}
    calculations: dict = {}
    bom: list = []
    cost_summary: dict = {}
    pipes: dict = {}
    equipment: dict = {}
    # Commercial buildup (markup / P&F / Designing / Negotiation / Transport /
    # Final) merged into the same sheet below the BOM. Optional — when absent the
    # old 3-row cost summary is written instead.
    commercial: dict = {}
    order_qty: int = 1           # order more than one unit → scales the costing
    currency: str = "INR"        # "USD" → add USD equivalents of the commercial totals
    fx_rate: float = 0           # INR → USD (e.g. 0.012); only used when currency == "USD"


def _regen_basis(item, spec):
    """Describe where a regen BOM line's unit price comes from — for the BOM
    Excel's 'BASIS' column. Matches how bom/regen_pricelist resolves each line."""
    it = item or ""
    s = (spec or "").strip()
    sz = f" {s}" if s and ("NB" in s or "DN" in s) else ""
    checks = [
        # Oil line (Pricelist 'Oil Line' category). Kept first so they win over
        # the generic valve/flow-meter matches below.
        ("Solenoid Valve (Oil",       "Pricelist → Solenoid Valve (Oil Line) 25 NB (MADAS)"),
        ("Globe Type Oil Control",    "Pricelist → Globe Type Oil Control Valve 25 NB (DEMBLA)"),
        ("Oil Flow Meter",            "Pricelist → Oil Flow Meter (HONEYWELL)"),
        ("TT in Oil Line",            "Pricelist → TT in Oil Line (HONEYWELL)"),
        ("PT in Oil Line",            "Pricelist → PT in Oil Line (HONEYWELL)"),
        ("Flexible Hose Pipe",        "Pricelist → Flexible Hose Pipe (Oil Line) 25 NB (BENGAL)"),
        ("Paperless Recorder",        "Pricelist → Paperless Recorder (EUROTHERM)"),
        ("Heating & Pumping Unit",    "HPU calculator → 9 KW (HPD-9, material cost × markup)"),
        ("ID Fan",                    "Pricelist → ID Fan 15 HP (ENCON)"),
        ("Burner with Regenerator", "Regen-with-Burner tab → Burner portion (burner KW) + Regen portion (regen KW)"),
        ("Combustion Blower",        "Internal costing → blower with motor (Alone×1.8 + Motor×1.5)"),
        ("PLC with HMI",             "Pricelist → PLC with HMI (by no. of pairs)"),
        ("Control Panel",            "Pricelist → Control Panel (per KW)"),
        ("NG Gas Train",             "Pricelist → Gas Train (by NG flow band)"),
        ("Gate Valve",               f"Pricelist → Gate Valve{sz} (L&T)"),
        ("Butterfly Valve (Isolation)", f"Pricelist → Butterfly Valve{sz} (L&T)"),
        ("Butterfly Valve",          f"Pricelist → Butterfly Valve{sz} (L&T)"),
        ("Rotary Joint",             f"Pricelist → Rotary Joint{sz} (ENCON)"),
        ("Pressure Switch Low",      "Pricelist → Pressure Switch Low (MADAS)"),
        ("Orifice Plate",            f"Pricelist → Orifice Plate{sz} (ENCON)"),
        ("Pressure Gauge with TNV",  "Pricelist → Pressure Gauge with TNV (BAUMER)"),
        ("DPT (Gas Train)",          "Pricelist → DPT (HONEYWELL)"),
        ("Control Valve",            f"Pricelist → Pneumatic Control Valve{sz} (DEMBLA)"),
        ("Pneumatic Damper",         f"Pricelist → Pneumatic Damper{sz} (ENCON)"),
        ("Manual Damper",            "Pricelist → DAMPER MANUAL (ENCON)"),
        ("Sequence Controller",      "Pricelist → Sequence Controller (LINEAR)"),
        ("Pilot Burner",             "Pricelist → ENCON-PB-LPG-10KW"),
        ("Ignition Transformer",     "Pricelist → Ignition Transformer (DANFOSS)"),
        ("UV Sensor",                "Pricelist → UV Sensor with Air Jacket (LINEAR)"),
        ("Pilot Regulator",          "Pricelist → Gas Regulator 025 NB, 5 Bar (MADAS)"),
        ("Pilot Solenoid Valve",     "Pricelist → Solenoid Valve 15 NB (MADAS)"),
        ("Solenoid Valve",           f"Pricelist → Solenoid Valve{sz} (MADAS, −45%)"),
        ("Ball Valve",               f"Pricelist → Ball Valve{sz} (L&T)"),
        ("Manual Butterfly Valve",   f"Pricelist → Butterfly Valve{sz} (L&T)"),
        ("Shut-Off Valve",           f"Pricelist → Pneumatic Shut Off Valve{sz} (DEMBLA)"),
        ("Gas Control Valve",        f"Pricelist → Pneumatic Control Valve{sz} (DEMBLA)"),
        ("Air Control Valve",        f"Pricelist → Pneumatic Control Valve{sz} (DEMBLA)"),
        ("Flow Meter (DPT)",         f"Pricelist → Flow Meter (DPT){sz}"),
        ("Flow Meter",               f"Pricelist → Flow Meter (DPT){sz}"),
        ("Flexible Hose",            f"Pricelist → Flexible Hose{sz} (BENGAL)"),
        ("Pressure Gauge 0-1000",    "Pricelist → Pressure Gauge with TNV (BAUMER)"),
        ("Pressure Gauge 0-500",     "Pricelist → Pressure Gauge with TNV (HGURU)"),
        ("Thermocouple with TT (Furnace)", "Pricelist → THERMOCOUPLE (TEMPSENS)"),
        ("Thermocouple",             "Pricelist → Thermocouple Small"),
        ("DPT",                      "Pricelist → DPT (HONEYWELL)"),
        ("Gate Valve",               "Code → 6000 KW gas skid"),
        ("Pressure Switch",          "Code → 6000 KW gas skid"),
        ("Pneumatic Shut-Off Valve", "Code → 6000 KW gas skid"),
        ("Manual Cock",              "Code → 6000 KW gas skid"),
    ]
    for needle, basis in checks:
        if needle in it:
            return basis
    return ""


def _vlph_basis(item, media, ref):
    """Describe where a VLPH/HLPH/Tundish BOM line's price comes from — for the
    Excel 'BASIS' column. Branches on MEDIA first (some item names, e.g. ORIFICE
    PLATE / BALL VALVE / SOLENOID VALVE, mean different sources in different
    lines), then on the item name. Mirrors bom/vlph_builder.py's pricing."""
    it = (item or "").upper().strip()
    md = (media or "").upper()
    r  = (ref or "").strip()
    sz = f" ({r})" if r else ""
    if it in ("BOUGHT OUT ITEMS", "ENCON ITEMS", "GRAND TOTAL"):
        return ""
    # ── Calculated structural items (unit price = qty basis × rate) ──────────
    if "FABRICATION" in it:       return "Calculated: MS structure kg × FABRICATION RATE (Pricelist)"
    if "AIR-GAS PIPELINE" in it:  return "Calculated: pipeline kg × PIPELINE RATE (Pricelist)"
    if "PLUMMER BLOCK" in it:     return "Calculated: plummer block kg × ₹170/kg (code)"
    if it.startswith("SHAFT"):    return "Calculated: shaft kg × ₹120/kg (code)"
    if "CERAMIC FIBRE" in it:     return "Calculated: rolls × ceramic rate/roll (vertical_master)"
    # ── Media-disambiguated items ────────────────────────────────────────────
    if "ORIFICE PLATE" in it:
        if "COG" in md or "BFG" in md:  return "Fixed ₹10,000 (code — ENGINEERING SPECIALITY)"
        if "MG" in md or "MIX" in md:   return "Fixed ₹7,000 (code — ENGINEERING SPECIALITY)"
        return f"orifice_plate_master → NB ≥ size{sz} (ENCON)"
    if "SOLENOID VALVE" in it:
        if "PURGING" in md:  return "Fixed ₹5,000 (code — MADAS)"
        return f"solenoidvalve_component_master → NB ≥ size{sz} (MADAS, −45%)"
    if "BALL VALVE" in it:
        if "PURGING" in md:  return "Fixed (code — AUDCO/L&T/LEADER)"
        if "INSTRUMENTS" in it: return "Pricelist → INSTRUMENTS BALL VALVE (L&T)"
        return f"Pricelist → BALL VALVE{sz} (L&T, cheapest)"
    if "PRESSURE REGULATING VALVE" in it:
        if "PURGING" in md:  return "Fixed ₹35,000 (code — NIRMAL)"
        return f"gas_regulator master → select_gas_regulator{sz} (MADAS)"
    if "CHECK VALVE" in it:       return "Fixed ₹3,300 (code — AUDCO/L&T/LEADER)"
    # ── Selector / master-table items ────────────────────────────────────────
    if "GAS TRAIN" in it:         return "gas_train_master → selected by flow band (MADAS)"
    if "BURNER" in it and "ENCON" in it: return "burner_master → selected burner model (ENCON)"
    if it == "BLOWER" or it.startswith("BLOWER"): return "blower_master / blower_pricelist_master → selected model (ENCON)"
    if "HEATING AND PUMPING" in it or it == "HPU": return "HPU selector (hpu_master) → selected KW / variant"
    if it == "AGR" or "AIR OIL REGULATOR" in it: return "agr_master → selected AGR (ENAG)"
    if "FLEXIBLE HOSE" in it:     return f"flexible_hose_master → DN ≥ size{sz} (BENGAL IND.)"
    if "ROTARY JOINT" in it:      return f"rotary_joint_master → selected{sz} (ENCON)"
    if "GATE VALVE" in it:        return f"Pricelist → GATE VALVE{sz} (L&T)"
    if "SHUT OFF VALVE" in it or "SHUT-OFF" in it: return f"Pricelist → PNEUMATIC SHUT OFF VALVE{sz} (AIRA/DEMBLA)"
    if "MOTORIZED CONTROL VALVE" in it: return f"Pricelist → MOTORIZED CONTROL VALVE{sz} (CAIR)"
    if "CONTROL VALVE" in it:     return f"Pricelist → PNEUMATIC / MOTORIZED CONTROL VALVE{sz}"
    if "BUTTERFLY VALVE" in it:   return f"Pricelist → BUTTERFLY VALVE{sz} (L&T) / lt_butterfly_valve_master"
    if "FLOWMETER" in it:         return "Pricelist → FLOWMETER (ELETA)"
    if "SWIVEL ASSEMBLY" in it:   return "vertical_master → swirling / pipeline cost"
    if "PILOT BURNER" in it:      return f"Pricelist → pilot burner model{sz} (ENCON)"
    # ── Everything else = exact Pricelist (component_price_master) item ───────
    return "Pricelist (component_price_master) → exact item"


@app.post("/api/export-excel")
def export_excel(req: ExcelExportRequest):
    import io
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from fastapi.responses import StreamingResponse

    NAVY   = "1A3A5C"
    LIGHT  = "EFF6FF"
    WHITE  = "FFFFFF"
    GREY   = "F8FAFC"
    GREEN  = "065F46"
    GREEN_BG = "F0FDF4"

    def thin():
        s = Side(style="thin", color="E2E8F0")
        return Border(left=s, right=s, top=s, bottom=s)

    wb = Workbook()

    def hdr(ws, row, col, val, bg=NAVY, fg=WHITE, bold=True, size=11, span=None):
        c = ws.cell(row=row, column=col, value=val)
        c.font = Font(bold=bold, color=fg, size=size, name="Calibri")
        c.fill = PatternFill("solid", fgColor=bg)
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        c.border = thin()
        return c

    def cell(ws, row, col, val, bold=False, align="left", bg=WHITE, fg="1E293B", num_fmt=None):
        c = ws.cell(row=row, column=col, value=val)
        c.font = Font(bold=bold, color=fg, size=10, name="Calibri")
        c.fill = PatternFill("solid", fgColor=bg)
        c.alignment = Alignment(horizontal=align, vertical="center")
        c.border = thin()
        if num_fmt:
            c.number_format = num_fmt
        return c

    def section_hdr(ws, r, ncols, text, bg=NAVY, fg=WHITE, size=10):
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=ncols)
        c = ws.cell(row=r, column=1, value=text)
        c.font = Font(bold=True, color=fg, size=size, name="Calibri")
        c.fill = PatternFill("solid", fgColor=bg)
        c.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[r].height = 20
        return c

    def title_row(ws, r, ncols, text, size=14):
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=ncols)
        c = ws.cell(row=r, column=1, value=text)
        c.font = Font(bold=True, color=WHITE, size=size, name="Calibri")
        c.fill = PatternFill("solid", fgColor=NAVY)
        c.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[r].height = 30
        return c

    # ════════════════════════════════════════════════════════════════════════
    #  REGEN — 5 sheets matching HTML tabs exactly
    # ════════════════════════════════════════════════════════════════════════
    if req.equipment_type == "Regen":
        supp = (req.equipment or {}).get("supplementary", {}) if req.equipment else {}
        bs   = supp.get("burner_sizing", {}) if supp else {}
        calc = req.calculations
        cs   = req.cost_summary

        # ── Sheet 1: Process Calcs ────────────────────────────────────────
        ws1 = wb.active
        ws1.title = "Process Calcs"
        for col, w in zip("ABCDE", [28, 22, 22, 14, 18]):
            ws1.column_dimensions[col].width = w

        r1 = 1
        title_row(ws1, r1, 5, "ENCON — Regenerative Burner System — Process Calcs")
        r1 += 1
        ws1.merge_cells(f"A{r1}:E{r1}")
        d = ws1.cell(row=r1, column=1, value=f"Generated: {datetime.now().strftime('%d %b %Y, %H:%M')}")
        d.font = Font(color="64748B", size=9, name="Calibri")
        d.fill = PatternFill("solid", fgColor=LIGHT)
        d.alignment = Alignment(horizontal="right", vertical="center")
        r1 += 2

        # Customer
        if req.customer:
            section_hdr(ws1, r1, 5, "CUSTOMER DETAILS")
            r1 += 1
            for label, val in [
                ("Company",     req.customer.get("company_name","")),
                ("Contact",     req.customer.get("poc_name","")),
                ("Designation", req.customer.get("poc_designation","")),
                ("Mobile",      req.customer.get("mobile_no","")),
                ("Email",       req.customer.get("email","")),
                ("Project",     req.customer.get("project_name","")),
                ("Ref No.",     req.customer.get("ref_no","")),
            ]:
                if val:
                    cell(ws1, r1, 1, label, bold=True, bg=GREY)
                    ws1.merge_cells(f"B{r1}:E{r1}")
                    cell(ws1, r1, 2, val)
                    r1 += 1
            r1 += 1

        # Process Parameters
        section_hdr(ws1, r1, 5, "PROCESS PARAMETERS")
        r1 += 1
        proc_params = [
            ("Material Weight",   calc.get("material_weight_kg",""), "kg"),
            ("Initial Temp (Ti)", calc.get("Ti",""),                 "°C"),
            ("Final Temp (Tf)",   calc.get("Tf",""),                 "°C"),
            ("Temp Rise (ΔT)",    calc.get("delta_T",""),            "°C"),
            ("Specific Heat Cp",  calc.get("Cp",""),                 "kJ/kg·°C"),
            ("Cycle Time",        calc.get("cycle_time_hr",""),      "hr"),
            ("Efficiency",        f"{round(calc.get('efficiency',0)*100)}%", ""),
            ("Heat Required",     calc.get("heat_required_kj",""),   "kJ"),
            ("Required Power",    calc.get("required_kw",""),        "kW"),
            ("No. of Pairs",      calc.get("num_pairs",""),          f"× {calc.get('model_kw',1000)} KW"),
            ("Total KW",          calc.get("total_kw",""),           "KW"),
        ]
        for i, (label, val, unit) in enumerate(proc_params):
            bg = GREY if i % 2 == 0 else WHITE
            cell(ws1, r1, 1, label, bold=True, bg=bg)
            cell(ws1, r1, 2, val, bg=bg, align="right")
            cell(ws1, r1, 3, unit, bg=bg)
            ws1.merge_cells(f"D{r1}:E{r1}")
            cell(ws1, r1, 4, "", bg=bg)
            r1 += 1
        r1 += 1

        # Cost Summary
        section_hdr(ws1, r1, 5, "COST SUMMARY")
        r1 += 1
        for label, val, is_total in [
            ("Total Cost Price",  cs.get("total_cost", 0),    False),
            ("Markup",            cs.get("markup", 0),         False),
            ("Total Selling Price", cs.get("total_selling", 0), True),
        ]:
            bg  = GREEN_BG if is_total else GREY
            fg_ = GREEN    if is_total else "1E293B"
            cell(ws1, r1, 1, label, bold=is_total, bg=bg, fg=fg_)
            ws1.merge_cells(f"B{r1}:D{r1}")
            cell(ws1, r1, 2, "", bg=bg)
            c = ws1.cell(row=r1, column=5, value=val)
            c.font  = Font(bold=is_total, color=fg_, size=11 if is_total else 10, name="Calibri")
            c.fill  = PatternFill("solid", fgColor=bg)
            c.alignment = Alignment(horizontal="right", vertical="center")
            c.number_format = '#,##0.00'
            c.border = thin()
            ws1.row_dimensions[r1].height = 22 if is_total else 18
            r1 += 1

        # ── Sheet 2: Burner Sizing and Costing ───────────────────────────
        ws2 = wb.create_sheet("Burner Sizing and Costing")
        ALL_COLS_BS = list("ABCDEFGHIJKLMNOP")
        for i, w in enumerate([14,14,14,14,12,12,12,12,14,14,14,14,14,14,14,14]):
            ws2.column_dimensions[ALL_COLS_BS[i]].width = w

        r2 = 1
        title_row(ws2, r2, 16, f"Burner + Regenerator Sizing — All KW Rows (Selected: {bs.get('kw','')} KW)")
        r2 += 2

        all_sz  = bs.get("all_sizing", [])
        sel_kw  = bs.get("kw")
        nozzles = supp.get("nozzle_sizing", [])

        # — Regenerator Dimensions table —
        section_hdr(ws2, r2, 16, "REGENERATOR DIMENSIONS (all KW models)")
        r2 += 1
        dim_hdrs = ["KW","Shell (mm)","Retainer (mm)","Refrac. (mm)","L (m)","H (m)","W (m)","Bottom H (m)",
                    "Vol Total (m³)","Vol Eff. (m³)","Vol Refrac. (m³)","Density Castable",
                    "Wt Refrac. (kg)","Loose Density","Vol Balls (m³)","Balls Fill %"]
        for ci, lbl in enumerate(dim_hdrs, 1):
            hdr(ws2, r2, ci, lbl, size=8)
        r2 += 1
        for row_s in all_sz:
            is_sel = row_s.get("kw") == sel_kw
            bg = LIGHT if is_sel else (GREY if r2 % 2 == 0 else WHITE)
            vals = [
                row_s.get("kw",""), row_s.get("shell_thick",""), row_s.get("retainer_thick",""),
                row_s.get("refractory_thick",""), row_s.get("dim_L",""), row_s.get("dim_H",""),
                row_s.get("dim_W",""), row_s.get("bottom_h",""),
                round(row_s["vol_total"],4) if row_s.get("vol_total") is not None else "",
                round(row_s["vol_effective"],4) if row_s.get("vol_effective") is not None else "",
                round(row_s["vol_refractory"],4) if row_s.get("vol_refractory") is not None else "",
                row_s.get("density_castable",""),
                round(row_s["wt_refractory_insulation"],1) if row_s.get("wt_refractory_insulation") is not None else "",
                row_s.get("loose_density_balls",""),
                round(row_s["vol_available_balls"],4) if row_s.get("vol_available_balls") is not None else "",
                f"{row_s['balls_filling_pct']}%" if row_s.get("balls_filling_pct") is not None else "",
            ]
            for ci, v in enumerate(vals, 1):
                cell(ws2, r2, ci, v, bold=is_sel, bg=bg, align="center")
            r2 += 1
        r2 += 1

        # — Weight Breakdown table —
        section_hdr(ws2, r2, 12, "BURNER SIZE — WEIGHT BREAKDOWN (all KW models)")
        r2 += 1
        wt_cols = list("ABCDEFGHIJKL")
        for i, w in enumerate([14,14,14,14,14,14,14,14,14,12,12,12]):
            ws2.column_dimensions[wt_cols[i]].width = w
        wt_hdrs = ["KW","Burner MS (kg)","Burner Refrac. (kg)","Regen MS (kg)","Regen SS (kg)",
                   "Regen Refrac. (kg)","Ceramic Balls (kg)","BB Refrac. (kg)","Grand Total (kg)",
                   "BB Dia Inner (m)","Burner Length (m)","Burner Dia (m)"]
        for ci, lbl in enumerate(wt_hdrs, 1):
            hdr(ws2, r2, ci, lbl, size=8)
        r2 += 1
        for row_s in all_sz:
            is_sel = row_s.get("kw") == sel_kw
            bg = LIGHT if is_sel else (GREY if r2 % 2 == 0 else WHITE)
            def fmt_wt(v): return round(v, 1) if v is not None else ""
            vals = [
                row_s.get("kw",""), fmt_wt(row_s.get("wt_burner_ms")), fmt_wt(row_s.get("wt_burner_refrac")),
                fmt_wt(row_s.get("wt_regen_ms")), fmt_wt(row_s.get("wt_regen_ss")),
                fmt_wt(row_s.get("wt_regen_refrac")), fmt_wt(row_s.get("wt_ceramic_balls")),
                fmt_wt(row_s.get("wt_burner_block_summary")), fmt_wt(row_s.get("wt_total")),
                row_s.get("bb_dia_inner",""), row_s.get("burner_length",""), row_s.get("burner_dia",""),
            ]
            for ci, v in enumerate(vals, 1):
                cell(ws2, r2, ci, v, bold=is_sel, bg=bg, align="center")
            r2 += 1
        r2 += 1

        # — Nozzle Sizing —
        if nozzles:
            section_hdr(ws2, r2, 8, "NOZZLE SIZING")
            r2 += 1
            for ci, lbl in enumerate(["Burner Name","Power (KW)","DN Air In (mm)","Air Speed (m/s)",
                                       "DN Fume Out (mm)","Fume Speed (m/s)","DN NG In (mm)","NG Speed (m/s)"], 1):
                hdr(ws2, r2, ci, lbl, size=9)
            r2 += 1
            for nz in nozzles:
                is_sel = nz.get("power_kw") == sel_kw
                bg = LIGHT if is_sel else (GREY if r2 % 2 == 0 else WHITE)
                for ci, v in enumerate([nz.get("burner_name",""), nz.get("power_kw",""),
                                        nz.get("dn_air_in",""), nz.get("air_speed_ms",""),
                                        nz.get("dn_fume_out",""), nz.get("fume_speed_ms",""),
                                        nz.get("dn_ng_in",""), nz.get("ng_speed_ms","")], 1):
                    cell(ws2, r2, ci, v, bold=is_sel, bg=bg, align="center")
                r2 += 1
            r2 += 1

        # — Costing Consideration: Material Rates —
        legacy_rates = bs.get("legacy_material_rates", [])
        section_hdr(ws2, r2, 5, "COSTING CONSIDERATION — MATERIAL RATES")
        r2 += 1
        for ci, lbl in enumerate(["Material","Wastage","Material Cost (₹/kg)","Labour Cost (₹/kg)","Effective Rate (₹/kg)"], 1):
            hdr(ws2, r2, ci, lbl, size=9)
        r2 += 1
        default_rates = [("MS",0.1,50,25),("SS",0.1,50,25),("Refractory",0.1,56,25),("Ceramic Balls",0.1,125,0)]
        rate_rows = legacy_rates if legacy_rates else [
            {"material":m,"wastage":wa,"mat_cost":mc,"labour_cost":lc} for m,wa,mc,lc in default_rates
        ]
        for i, mr in enumerate(rate_rows):
            bg = GREY if i % 2 == 0 else WHITE
            wa  = mr.get("wastage") or 0
            mc  = mr.get("mat_cost") or 0
            lc  = mr.get("labour_cost") or 0
            eff = round((mc + lc) * (1 + wa), 2)
            for ci, v in enumerate([mr.get("material",""), f"{round(wa*100)}%", mc or "—", lc or "—", eff], 1):
                cell(ws2, r2, ci, v, bg=bg, align="right" if ci > 1 else "left")
            r2 += 1
        r2 += 1

        # — Cost Breakdown all KW —
        section_hdr(ws2, r2, 9, "COST BREAKDOWN — PER UNIT (1 burner + regenerator), ALL KW MODELS")
        r2 += 1
        cd_hdrs = ["KW","Burner MS (₹)","Burner Refrac. (₹)","Regen MS (₹)","Regen SS (₹)",
                   "Regen Refrac. (₹)","Ceramic Balls (₹)","Burner Block (₹)","TOTAL UNIT COST (₹)"]
        for ci, lbl in enumerate(cd_hdrs, 1):
            hdr(ws2, r2, ci, lbl, size=8)
        r2 += 1
        for row_s in all_sz:
            is_sel = row_s.get("kw") == sel_kw
            bg = LIGHT if is_sel else (GREY if r2 % 2 == 0 else WHITE)
            vals = [
                row_s.get("kw",""),
                round(row_s["cost_burner_ms"],2) if row_s.get("cost_burner_ms") is not None else "",
                round(row_s["cost_burner_refrac"],2) if row_s.get("cost_burner_refrac") is not None else "",
                round(row_s["cost_regen_ms"],2) if row_s.get("cost_regen_ms") is not None else "",
                round(row_s["cost_regen_ss"],2) if row_s.get("cost_regen_ss") is not None else "",
                round(row_s["cost_regen_refrac"],2) if row_s.get("cost_regen_refrac") is not None else "",
                round(row_s["cost_ceramic_balls"],2) if row_s.get("cost_ceramic_balls") is not None else "",
                round(row_s["cost_burner_block"],2) if row_s.get("cost_burner_block") is not None else "",
                round(row_s["cost_total"],2) if row_s.get("cost_total") is not None else "",
            ]
            for ci, v in enumerate(vals, 1):
                cell(ws2, r2, ci, v, bold=is_sel, bg=bg, align="center",
                     num_fmt='#,##0.00' if ci > 1 and isinstance(v,(int,float)) else None)
            r2 += 1
        # Totals row
        r2_tot = r2
        ws2.merge_cells(f"A{r2_tot}:H{r2_tot}")
        cell(ws2, r2_tot, 1, f"Selected: {sel_kw} KW — 1 Pair (×2)", bold=True, bg=GREEN_BG, fg=GREEN, align="right")
        cell(ws2, r2_tot, 9, bs.get("total_pair_cost", 0), bold=True, bg=GREEN_BG, fg=GREEN, align="center",
             num_fmt='#,##0.00')

        # ── Sheet 3: Burner Pipe Size ─────────────────────────────────────
        ws3 = wb.create_sheet("Burner Pipe Size")
        PIPE_COLS = list("ABCDEFGHIJKLM")
        for i, w in enumerate([16,14,14,14,12,12,12,12,12,12,12,12,12]):
            ws3.column_dimensions[PIPE_COLS[i]].width = w

        r3 = 1
        title_row(ws3, r3, 13, "Burner Pipe Size — All Gas Types and KW Models")
        r3 += 2

        all_ps  = supp.get("all_pipe_sizes", [])
        ps_sel  = supp.get("pipe_sizes", {})
        sel_ps_kw = bs.get("kw")
        pipe_col_hdrs = ["Burner Size (KW)","Gas Flow (Nm³/hr)","Air Flow (Nm³/hr)","Total Flue (Nm³/hr)",
                         "DN Air (mm)","DN Gas (mm)","DN Flue (mm)",
                         "Area Air (m²)","Area Gas (m²)","Area Flue (m²)",
                         "Vel Gas (m/s)","Vel Air (m/s)","Vel Flue (m/s)"]

        if all_ps:
            # Group by gas type
            gas_types = []
            for row_p in all_ps:
                gt = row_p.get("gas_type","")
                if gt not in gas_types:
                    gas_types.append(gt)
            for gt in gas_types:
                section_hdr(ws3, r3, 13, gt)
                r3 += 1
                for ci, lbl in enumerate(pipe_col_hdrs, 1):
                    hdr(ws3, r3, ci, lbl, size=8)
                r3 += 1
                for row_p in all_ps:
                    if row_p.get("gas_type") != gt:
                        continue
                    is_sel = row_p.get("burner_size_kw") == sel_ps_kw
                    bg = LIGHT if is_sel else (GREY if r3 % 2 == 0 else WHITE)
                    vals = [
                        row_p.get("burner_size_kw",""),
                        round(row_p["gas_flow_nm3hr"],2) if row_p.get("gas_flow_nm3hr") is not None else "",
                        round(row_p["air_flow_nm3hr"],2) if row_p.get("air_flow_nm3hr") is not None else "",
                        round(row_p["flue_flow_nm3hr"],2) if row_p.get("flue_flow_nm3hr") is not None else "",
                        row_p.get("dn_air_mm",""), row_p.get("dn_gas_mm",""), row_p.get("dn_flue_mm",""),
                        round(row_p["area_air_m2"],4) if row_p.get("area_air_m2") is not None else "",
                        round(row_p["area_gas_m2"],4) if row_p.get("area_gas_m2") is not None else "",
                        round(row_p["area_flue_m2"],4) if row_p.get("area_flue_m2") is not None else "",
                        round(row_p["vel_gas_ms"],1) if row_p.get("vel_gas_ms") is not None else "",
                        round(row_p["vel_air_ms"],1) if row_p.get("vel_air_ms") is not None else "",
                        round(row_p["vel_flue_ms"],1) if row_p.get("vel_flue_ms") is not None else "",
                    ]
                    for ci, v in enumerate(vals, 1):
                        cell(ws3, r3, ci, v, bold=is_sel, bg=bg, align="center")
                    r3 += 1
                r3 += 1
        elif ps_sel:
            # Fallback: single KW from pipe_sizes
            section_hdr(ws3, r3, 7, f"Natural Gas (NG) — {ps_sel.get('fuel','')} @ {ps_sel.get('pressure','')}")
            r3 += 1
            for ci, lbl in enumerate(pipe_col_hdrs[:7], 1):
                hdr(ws3, r3, ci, lbl, size=9)
            r3 += 1
            fb_vals = [ps_sel.get("kw",""), ps_sel.get("ng_flow_nm3hr",""), ps_sel.get("air_flow_nm3hr",""),
                       ps_sel.get("flue_flow_nm3hr",""), ps_sel.get("air_line_dn",""),
                       ps_sel.get("gas_line_dn",""), ps_sel.get("flue_line_dn","")]
            for ci, v in enumerate(fb_vals, 1):
                cell(ws3, r3, ci, v, bg=LIGHT, align="center")

        # ── Sheet 4: Blower ───────────────────────────────────────────────
        ws4 = wb.create_sheet("Blower")
        for col, w in zip("ABCDEF", [22, 14, 14, 16, 20, 20]):
            ws4.column_dimensions[col].width = w

        r4 = 1
        bl  = supp.get("blower_selection", {})
        cat = supp.get("blower_catalogue", [])
        title_row(ws4, r4, 6, f"ENCON 40\" WG Blower Selection — {bl.get('kw','')} KW")
        r4 += 2

        section_hdr(ws4, r4, 6, "SELECTED BLOWER")
        r4 += 1
        sel_blower_rows = [
            ("Burner KW",          f"{bl.get('kw','')} KW"),
            ("Selected Model",     bl.get("selected_model","")),
            ("HP",                 bl.get("hp","")),
            ("CFM",                bl.get("cfm","")),
            ("Flow Rate (Nm³/hr)", bl.get("nm3hr","")),
            ("Qty per pair",       bl.get("qty_per_pair","")),
            ("Price w/o Motor ₹",  bl.get("price_without_motor",0)),
            ("Price with Motor ₹", bl.get("price_with_motor",0)),
            ("Costing Price Used ₹", bl.get("costing_price",0)),
        ]
        for i, (lbl, val) in enumerate(sel_blower_rows):
            is_cost = "Costing" in lbl
            bg = GREEN_BG if is_cost else (GREY if i % 2 == 0 else WHITE)
            fg_ = GREEN if is_cost else "1E293B"
            cell(ws4, r4, 1, lbl, bold=is_cost, bg=bg, fg=fg_)
            ws4.merge_cells(f"B{r4}:F{r4}")
            cell(ws4, r4, 2, val, bold=is_cost, bg=bg, fg=fg_, align="right",
                 num_fmt='#,##0.00' if isinstance(val,(int,float)) else None)
            r4 += 1
        r4 += 1

        section_hdr(ws4, r4, 6, "ENCON 40\" WG BLOWER CATALOGUE")
        r4 += 1
        for ci, lbl in enumerate(["Model","HP","CFM","Nm³/hr","Price w/o Motor ₹","Price with Motor ₹"], 1):
            hdr(ws4, r4, ci, lbl, size=9)
        r4 += 1
        for row_b in cat:
            is_sel_b = row_b.get("model","") == bl.get("selected_model","")
            bg = LIGHT if is_sel_b else (GREY if r4 % 2 == 0 else WHITE)
            for ci, v in enumerate([row_b.get("model",""), row_b.get("hp",""), row_b.get("cfm",""),
                                     row_b.get("nm3hr",""), row_b.get("price_without_motor",""),
                                     row_b.get("price_with_motor","")], 1):
                cell(ws4, r4, ci, v, bold=is_sel_b, bg=bg,
                     align="right" if ci > 2 else "left",
                     num_fmt='#,##0.00' if ci > 4 and isinstance(v,(int,float)) else None)
            r4 += 1

        # ── Sheet 5: BOM (with a BASIS column + live formulas) ────────────
        ws5 = wb.create_sheet("BOM")
        for col, w in zip("ABCDEFGHIJK", [6, 16, 26, 11, 20, 7, 13, 14, 13, 15, 42]):
            ws5.column_dimensions[col].width = w

        mk = cs.get("markup", 1.8) or 1.8
        r5 = 1
        title_row(ws5, r5, 11, f"BILL OF MATERIALS — REGENERATIVE BURNER SYSTEM ({calc.get('model_kw','1000')} KW)")
        r5 += 2

        bom_col_hdrs = ["S.No.","SECTION","ITEM NAME","MAKE","SPECIFICATION","QTY",
                        "COST/UNIT ₹","TOTAL COST ₹\n(=Qty×Unit)","SELL/UNIT ₹\n(=Unit×Markup)",
                        f"TOTAL SELLING ₹\n(=Qty×Sell)","BASIS — how the unit price is derived"]
        for ci, lbl in enumerate(bom_col_hdrs, 1):
            hdr(ws5, r5, ci, lbl, size=9)
        ws5.row_dimensions[r5].height = 30
        r5 += 1
        data_start = r5
        # Each BOM section gets its own colour band so line items are grouped
        # visually. Colours are assigned in section-first-appearance order and
        # match the on-screen preview (_SECTION_COLORS in regen_costing.html).
        _sec_palette = ["DBEAFE", "DCFCE7", "FEF3C7", "EDE9FE", "CCFBF1",
                        "FCE7F3", "FFEDD5", "E0F2FE", "FEE2E2", "F3E8FF"]
        _sec_color = {}
        for _rd in req.bom:
            _s = _rd.get("SECTION", "")
            if _s not in _sec_color:
                _sec_color[_s] = _sec_palette[len(_sec_color) % len(_sec_palette)]
        sno = 0
        for i, row_d in enumerate(req.bom):
            bg = _sec_color.get(row_d.get("SECTION", ""), WHITE)
            sno += 1
            cell(ws5, r5, 1, sno, bg=bg, align="center")
            cell(ws5, r5, 2, row_d.get("SECTION",""), bg=bg)
            cell(ws5, r5, 3, row_d.get("ITEM NAME",""), bg=bg)
            cell(ws5, r5, 4, row_d.get("MAKE",""), bg=bg)
            cell(ws5, r5, 5, row_d.get("SPECIFICATION",""), bg=bg)
            cell(ws5, r5, 6, row_d.get("QTY",""), bg=bg, align="right")
            cell(ws5, r5, 7, row_d.get("COST/UNIT",0), bg=bg, align="right", num_fmt='#,##0.00')
            # Derived values as live formulas so the calculation is visible.
            cell(ws5, r5, 8,  f"=F{r5}*G{r5}",     bg=bg, align="right", num_fmt='#,##0.00')
            cell(ws5, r5, 9,  f"=G{r5}*{mk}",      bg=bg, align="right", num_fmt='#,##0.00')
            cell(ws5, r5, 10, f"=F{r5}*I{r5}",     bg=bg, align="right", num_fmt='#,##0.00')
            b = cell(ws5, r5, 11, _regen_basis(row_d.get("ITEM NAME",""), row_d.get("SPECIFICATION","")),
                     bg=bg, fg="475569")
            b.font = Font(color="475569", size=9, italic=True, name="Calibri")
            ws5.row_dimensions[r5].height = 18
            r5 += 1

        # Grand total row — sum formulas over the data rows.
        ws5.merge_cells(f"A{r5}:G{r5}")
        cell(ws5, r5, 1, f"GRAND TOTAL   (Selling = Cost × {mk} markup)", bold=True, bg=GREEN_BG, fg=GREEN, align="right")
        cell(ws5, r5, 8, f"=SUM(H{data_start}:H{r5-1})", bold=True, bg=GREEN_BG, fg=GREEN, align="right", num_fmt='#,##0.00')
        cell(ws5, r5, 9, "", bold=True, bg=GREEN_BG, fg=GREEN)
        cell(ws5, r5, 10, f"=SUM(J{data_start}:J{r5-1})", bold=True, bg=GREEN_BG, fg=GREEN, align="right", num_fmt='#,##0.00')
        cell(ws5, r5, 11, "", bold=True, bg=GREEN_BG, fg=GREEN)
        ws5.row_dimensions[r5].height = 22

        # ── Blower + ID-fan sizing worked example (below the items) ─────────
        fs = supp.get("fan_sizing") if supp else None
        if fs:
            r5 += 2
            section_hdr(ws5, r5, 11,
                f"BLOWER & ID FAN — SIZING EXAMPLE  ({fs.get('num_pairs')} × {fs.get('kw')} KW, "
                f"{'oil' if fs.get('is_oil') else 'gas'})")
            r5 += 1
            def _fline(text):
                nonlocal r5
                ws5.merge_cells(start_row=r5, start_column=1, end_row=r5, end_column=11)
                fc = ws5.cell(row=r5, column=1, value=text)
                fc.font = Font(size=10, name="Consolas", color="1E293B")
                fc.alignment = Alignment(horizontal="left", vertical="center")
                r5 += 1
            def _r0(v): return f"{round(v):,}"
            def _r1(v): return f"{round(v*10)/10:,}"
            if fs.get("is_oil"):
                _fline(f"Oil flow    = KW × 860 ÷ CV × pairs = {fs['kw']} × 860 ÷ {fs['oil_cv']} × {fs['num_pairs']}  =  {_r1(fs['oil_kg'])} kg/hr")
                _fline(f"Air flow    = oil × A/F ({fs['afr']:g})       = {_r1(fs['oil_kg'])} × {fs['afr']:g}  =  {_r0(fs['air_kg'])} kg/hr")
                _fline(f"Blower flow = air ÷ {fs['rho_air']} (air density)   = {_r0(fs['air_kg'])} ÷ {fs['rho_air']}  =  {_r0(fs['comb_air'])} Nm³/hr")
                _fline(f"ID-fan flow = (air + oil) ÷ {fs['rho_flue']} (flue gas) = ({_r0(fs['air_kg'])} + {_r1(fs['oil_kg'])}) ÷ {fs['rho_flue']}  =  {_r0(fs['id_air'])} Nm³/hr")
            else:
                _fline(f"Combustion air = KW × pairs        = {fs['kw']} × {fs['num_pairs']}  =  {_r0(fs['comb_air'])} Nm³/hr")
                _fline(f"Gas flow    = KW × 860 ÷ CV × pairs = {fs['kw']} × 860 ÷ {fs['fuel_cv']} × {fs['num_pairs']}  =  {_r0(fs['gas_flow'])} Nm³/hr")
                _fline(f"ID-fan flow = combustion air + gas = {_r0(fs['comb_air'])} + {_r0(fs['gas_flow'])}  =  {_r0(fs['id_air'])} Nm³/hr")
            _bnote = "   (>60 HP → price ??)" if fs.get("blower_price") is None else ""
            _inote = "   (>60 HP → price ??)" if fs.get("id_price") is None else "   (price mirrored from blower)"
            _fline(f"Blower HP = flow ÷ 1.7 × 40\" ÷ 3200 = {_r0(fs['comb_air'])} ÷ 1.7 × 40 ÷ 3200  =  {_r1(fs['blower_raw_hp'])} HP  →  {fs['blower_hp']:g} HP frame{_bnote}")
            _fline(f"ID fan HP = flow ÷ 1.7 × 36\" ÷ 3200 = {_r0(fs['id_air'])} ÷ 1.7 × 36 ÷ 3200  =  {_r1(fs['id_raw_hp'])} HP  →  {fs['id_hp']:g} HP frame{_inote}")

        # BOM-only download — the derivation tabs were dropped from the UI
        # (we map straight from the Pricelist), so drop them from the workbook too.
        for _name in ("Process Calcs", "Burner Sizing and Costing", "Burner Pipe Size", "Blower"):
            if _name in wb.sheetnames:
                del wb[_name]

        # Save and return
        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        cname = req.customer.get("company_name","").replace(" ","_") if req.customer else "ENCON"
        fname = f"Regen_Costing_{cname or 'ENCON'}_{datetime.now().strftime('%d%b%Y')}.xlsx"
        return StreamingResponse(buf,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f'attachment; filename="{fname}"'})

    # ════════════════════════════════════════════════════════════════════════
    #  BTF / SNSF BRF path
    # ════════════════════════════════════════════════════════════════════════
    if req.equipment_type in ("BTF", "SNSF BRF"):
        ws = wb.active
        ws.title = req.equipment_type
        ws.column_dimensions["A"].width = 36
        ws.column_dimensions["B"].width = 42
        ws.column_dimensions["C"].width = 14
        ws.column_dimensions["D"].width = 12
        ws.column_dimensions["E"].width = 18
        ws.column_dimensions["F"].width = 20
        ws.column_dimensions["G"].width = 12
        ws.column_dimensions["H"].width = 20

        r1 = 1
        ws.merge_cells(f"A{r1}:H{r1}")
        t = ws.cell(row=r1, column=1, value=f"ENCON — {req.equipment_type} Costing Sheet")
        t.font = Font(bold=True, color=WHITE, size=14, name="Calibri")
        t.fill = PatternFill("solid", fgColor=NAVY)
        t.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[r1].height = 30
        r1 += 1

        ws.merge_cells(f"A{r1}:H{r1}")
        d = ws.cell(row=r1, column=1, value=f"Generated: {datetime.now().strftime('%d %b %Y, %H:%M')}")
        d.font = Font(color="64748B", size=9, name="Calibri")
        d.fill = PatternFill("solid", fgColor=LIGHT)
        d.alignment = Alignment(horizontal="right", vertical="center")
        r1 += 2

        # BOM header
        is_brf = req.equipment_type == "SNSF BRF"
        bom_cols = ["SECTION", "ITEM", "QTY", "UNIT", "UNIT PRICE", "COST PRICE",
                    "MARKUP" if is_brf else "", "SELL PRICE"]
        for ci, col in enumerate(bom_cols, 1):
            if col:
                hdr(ws, r1, ci, col, size=9)
        ws.row_dimensions[r1].height = 22
        r1 += 1

        last_section = ""
        for i, row_d in enumerate(req.bom):
            section = row_d.get("SECTION", "")
            if section != last_section:
                last_section = section
                ws.merge_cells(f"A{r1}:H{r1}")
                cell(ws, r1, 1, section, bold=True, bg=LIGHT, fg=NAVY)
                ws.row_dimensions[r1].height = 20
                r1 += 1

            bg = GREY if i % 2 == 0 else WHITE
            cell(ws, r1, 1, "", bg=bg)
            cell(ws, r1, 2, row_d.get("ITEM", ""), bg=bg)
            cell(ws, r1, 3, row_d.get("QTY", ""), bg=bg, align="right")
            cell(ws, r1, 4, row_d.get("UNIT", ""), bg=bg)
            cell(ws, r1, 5, row_d.get("UNIT PRICE", 0), bg=bg, align="right", num_fmt='#,##0')
            cell(ws, r1, 6, row_d.get("COST PRICE", 0), bg=bg, align="right", num_fmt='#,##0')
            if is_brf:
                cell(ws, r1, 7, f'{row_d.get("MARKUP", "")}x', bg=bg, align="center")
            cell(ws, r1, 8, row_d.get("SELL PRICE", 0), bg=bg, align="right", num_fmt='#,##0')
            ws.row_dimensions[r1].height = 18
            r1 += 1
        r1 += 1

        # Cost Summary
        cs = req.cost_summary
        ws.merge_cells(f"A{r1}:H{r1}")
        hdr(ws, r1, 1, "COST SUMMARY", size=10)
        ws.row_dimensions[r1].height = 20
        r1 += 1

        if req.equipment_type == "BTF":
            summary_rows = [
                ("Structure Cost",     cs.get("structure_cost", 0)),
                ("Combustion Cost",    cs.get("combustion_cost", 0)),
                ("Total Cost Price",   cs.get("total_cost", 0)),
                ("Sell Price",         cs.get("sell_price", 0)),
                ("Designing (10%)",    cs.get("designing_10pct", 0)),
                ("Negotiation (10%)",  cs.get("negotiation_10pct", 0)),
                ("Quoted Price",       cs.get("quoted_price", 0)),
            ]
        else:
            summary_rows = [
                ("Main Scope Cost",    cs.get("main_cost", 0)),
                ("Main Scope Sell",    cs.get("main_sell", 0)),
                ("NG Optional Cost",   cs.get("ng_optional_cost", 0)),
                ("NG Optional Sell",   cs.get("ng_optional_sell", 0)),
                ("Client Scope Cost",  cs.get("client_scope_cost", 0)),
                ("Client Scope Sell",  cs.get("client_scope_sell", 0)),
                ("Grand Total Cost",   cs.get("grand_cost", 0)),
                ("Grand Total Sell",   cs.get("grand_sell", 0)),
            ]

        for label, val in summary_rows:
            is_total = "Grand" in label or "Quoted" in label
            bg = GREEN_BG if is_total else GREY
            fg = GREEN if is_total else "1E293B"
            cell(ws, r1, 1, label, bold=is_total, bg=bg, fg=fg)
            ws.merge_cells(f"B{r1}:G{r1}")
            cell(ws, r1, 2, "", bg=bg)
            c = ws.cell(row=r1, column=8, value=val)
            c.font = Font(bold=is_total, color=fg, size=11 if is_total else 10, name="Calibri")
            c.fill = PatternFill("solid", fgColor=bg)
            c.alignment = Alignment(horizontal="right", vertical="center")
            c.number_format = '₹#,##0'
            c.border = thin()
            ws.row_dimensions[r1].height = 20
            r1 += 1

        # ── Additional sheets for BTF (match legacy 5-sheet workbook) ─────
        supp = (req.equipment or {}).get("supplementary", {})
        if supp and req.equipment_type == "BTF":
            def _kv_sheet(ws2, r2, items, col_widths=(36, 24)):
                """Write key-value rows to a sheet."""
                for ci, w in enumerate(col_widths, 1):
                    ws2.column_dimensions[chr(64+ci)].width = w
                return r2

            # ── Sheet: 10 T RHF ──────────────────────────────────────────
            ws_rhf = wb.create_sheet("10 T RHF", 0)
            ws_rhf.column_dimensions["A"].width = 36
            ws_rhf.column_dimensions["B"].width = 18
            ws_rhf.column_dimensions["C"].width = 14
            ws_rhf.column_dimensions["D"].width = 16
            ws_rhf.column_dimensions["E"].width = 14
            r2 = 1
            # Furnace dims
            fd = supp.get("furnace_dimensions", {})
            ws_rhf.merge_cells(f"A{r2}:E{r2}")
            hdr(ws_rhf, r2, 1, "FURNACE DIMENSIONS", size=10); r2 += 1
            cell(ws_rhf, r2, 1, "Furnace Internal Dim.", bold=True, bg=GREY)
            cell(ws_rhf, r2, 2, f'{fd.get("internal_L_mm","")}', bg=WHITE, align="right")
            cell(ws_rhf, r2, 3, f'{fd.get("internal_W_mm","")}', bg=WHITE, align="right")
            cell(ws_rhf, r2, 4, f'{fd.get("internal_H_mm","")} mm', bg=WHITE, align="right")
            r2 += 1
            for fv in supp.get("furnace_volumes", []):
                cell(ws_rhf, r2, 1, fv["item"], bold=True, bg=GREY)
                cell(ws_rhf, r2, 2, fv["value"], bg=WHITE, align="right")
                cell(ws_rhf, r2, 3, fv["unit"], bg=WHITE); r2 += 1
            r2 += 1
            # MS Structure
            ws_rhf.merge_cells(f"A{r2}:E{r2}")
            hdr(ws_rhf, r2, 1, "MS STRUCTURE", size=10); r2 += 1
            for sec in supp.get("ms_structure", []):
                cell(ws_rhf, r2, 1, sec["section"], bold=True, bg=LIGHT, fg=NAVY); r2 += 1
                for it in sec["items"]:
                    bg2 = GREEN_BG if it.get("highlight") else GREY
                    cell(ws_rhf, r2, 1, it["item"], bold=it.get("highlight",False), bg=bg2)
                    cell(ws_rhf, r2, 2, it["value"], bg=bg2, align="right"); r2 += 1
            r2 += 1
            # Ceramic Fibre
            ws_rhf.merge_cells(f"A{r2}:E{r2}")
            hdr(ws_rhf, r2, 1, "CERAMIC FIBRE", size=10); r2 += 1
            for cf in supp.get("ceramic_fibre", []):
                cell(ws_rhf, r2, 1, cf["wall"], bg=GREY)
                cell(ws_rhf, r2, 2, cf["L"], bg=WHITE, align="right")
                cell(ws_rhf, r2, 3, cf["W"], bg=WHITE, align="right")
                cell(ws_rhf, r2, 4, cf["thk"], bg=WHITE, align="right"); r2 += 1
            cfs = supp.get("ceramic_fibre_summary", {})
            cell(ws_rhf, r2, 1, "Total Rolls", bold=True, bg=GREEN_BG)
            cell(ws_rhf, r2, 2, cfs.get("total_rolls",""), bg=GREEN_BG, align="right"); r2 += 2
            # Door
            ws_rhf.merge_cells(f"A{r2}:E{r2}")
            hdr(ws_rhf, r2, 1, "DOOR", size=10); r2 += 1
            hdr(ws_rhf, r2, 1, "Item", size=9); hdr(ws_rhf, r2, 2, "Qty", size=9); hdr(ws_rhf, r2, 3, "Rate", size=9); hdr(ws_rhf, r2, 4, "Total", size=9); r2 += 1
            door_total = 0
            for d in supp.get("door", []):
                cell(ws_rhf, r2, 1, d["item"], bg=GREY)
                cell(ws_rhf, r2, 2, d["qty"], bg=WHITE, align="right")
                cell(ws_rhf, r2, 3, d["rate"], bg=WHITE, align="right", num_fmt='#,##0')
                cell(ws_rhf, r2, 4, d["total"], bg=WHITE, align="right", num_fmt='#,##0')
                door_total += d["total"]; r2 += 1
            cell(ws_rhf, r2, 1, "Total Door", bold=True, bg=GREEN_BG)
            cell(ws_rhf, r2, 4, door_total, bg=GREEN_BG, align="right", num_fmt='#,##0'); r2 += 2
            # Trolley
            ws_rhf.merge_cells(f"A{r2}:E{r2}")
            hdr(ws_rhf, r2, 1, "TROLLEY / BOGIE", size=10); r2 += 1
            hdr(ws_rhf, r2, 1, "Item", size=9); hdr(ws_rhf, r2, 2, "Qty", size=9); hdr(ws_rhf, r2, 3, "Rate", size=9); hdr(ws_rhf, r2, 4, "Total", size=9); r2 += 1
            tr_total = 0
            for t in supp.get("trolley_bogie", []):
                cell(ws_rhf, r2, 1, t["item"], bg=GREY)
                cell(ws_rhf, r2, 2, t["qty"], bg=WHITE, align="right")
                cell(ws_rhf, r2, 3, t["rate"], bg=WHITE, align="right", num_fmt='#,##0')
                cell(ws_rhf, r2, 4, t["total"], bg=WHITE, align="right", num_fmt='#,##0')
                tr_total += t["total"]; r2 += 1
            cell(ws_rhf, r2, 1, "Total Trolley/Bogie", bold=True, bg=GREEN_BG)
            cell(ws_rhf, r2, 4, tr_total, bg=GREEN_BG, align="right", num_fmt='#,##0'); r2 += 2
            # Refractory
            ws_rhf.merge_cells(f"A{r2}:E{r2}")
            hdr(ws_rhf, r2, 1, "REFRACTORY", size=10); r2 += 1
            ref = supp.get("refractory", {})
            for sec_name, sec_key in [("Walls","walls"),("Bogie","bogie")]:
                cell(ws_rhf, r2, 1, sec_name, bold=True, bg=LIGHT, fg=NAVY); r2 += 1
                sec_total = 0
                for br in ref.get(sec_key, []):
                    cell(ws_rhf, r2, 1, br["type"], bg=GREY)
                    cell(ws_rhf, r2, 2, br["qty"], bg=WHITE, align="right")
                    cell(ws_rhf, r2, 3, br.get("cost_per",""), bg=WHITE, align="right")
                    cell(ws_rhf, r2, 4, br["total"], bg=WHITE, align="right", num_fmt='#,##0')
                    sec_total += br["total"]; r2 += 1
                cell(ws_rhf, r2, 1, f"Subtotal {sec_name}", bold=True, bg=GREEN_BG)
                cell(ws_rhf, r2, 4, sec_total, bg=GREEN_BG, align="right", num_fmt='#,##0'); r2 += 1
            for sec_name, sec_key in [("Misc (Mortar)","misc"),("Arch Bricks","arch")]:
                cell(ws_rhf, r2, 1, sec_name, bold=True, bg=LIGHT, fg=NAVY); r2 += 1
                for br in ref.get(sec_key, []):
                    cell(ws_rhf, r2, 1, br["type"], bg=GREY)
                    cell(ws_rhf, r2, 2, br["qty"], bg=WHITE, align="right")
                    cell(ws_rhf, r2, 3, br.get("rate", br.get("cost_per","")), bg=WHITE, align="right")
                    cell(ws_rhf, r2, 4, br["total"], bg=WHITE, align="right", num_fmt='#,##0'); r2 += 1

            # ── Sheet: Calculation ────────────────────────────────────────
            ws_calc = wb.create_sheet("Calculation", 1)
            ws_calc.column_dimensions["A"].width = 28
            ws_calc.column_dimensions["B"].width = 20
            ws_calc.column_dimensions["C"].width = 14
            r2 = 1
            fp = supp.get("furnace_params", {})
            ws_calc.merge_cells(f"A{r2}:C{r2}")
            hdr(ws_calc, r2, 1, "HEAT LOAD", size=10); r2 += 1
            for h in supp.get("heat_load", []):
                is_t = h["item"] in ("Total", "Gross")
                bg2 = GREEN_BG if is_t else GREY
                cell(ws_calc, r2, 1, h["item"], bold=is_t, bg=bg2)
                cell(ws_calc, r2, 2, h["value"], bg=bg2, align="right", num_fmt='#,##0')
                cell(ws_calc, r2, 3, h["unit"], bg=bg2); r2 += 1
            r2 += 1
            ws_calc.merge_cells(f"A{r2}:C{r2}")
            hdr(ws_calc, r2, 1, "FURNACE PARAMETERS", size=10); r2 += 1
            for label, val in [
                ("Capacity", f'{fp.get("furnace_capacity_tph","")} Ton/hr'),
                ("Fuel Consumption", f'{fp.get("fuel_consumption_nm3hr","")} Nm3/hr'),
                ("Air Flow", f'{fp.get("air_flow_nm3hr","")} Nm3/hr'),
                ("CFM", f'{fp.get("cfm","")}'),
                ("Blower HP", f'{fp.get("blower_hp_calc","")} → {fp.get("blower_hp_selected","")} HP'),
                ("Burners", f'{fp.get("no_of_burners","")} ({fp.get("no_of_zones","")} zones)'),
                ("Rating/Zone", f'{fp.get("rating_per_zone_kcal","")} kcal ({fp.get("rating_per_zone_kw","")} kW)'),
            ]:
                cell(ws_calc, r2, 1, label, bold=True, bg=GREY)
                cell(ws_calc, r2, 2, val, bg=WHITE); r2 += 1
            r2 += 1
            ws_calc.merge_cells(f"A{r2}:C{r2}")
            hdr(ws_calc, r2, 1, "PIPE SIZING", size=10); r2 += 1
            ps = supp.get("pipe_sizing", {})
            for label, key in [("Air Zone 1","air_zone1"),("Air Zone 2","air_zone2"),("Gas Zone 1","gas_zone1"),("Gas Zone 2","gas_zone2")]:
                p = ps.get(key, {})
                cell(ws_calc, r2, 1, label, bold=True, bg=GREY)
                cell(ws_calc, r2, 2, f'{p.get("flow_nm3hr","")} Nm3/hr @ {p.get("velocity_ms","")} m/s', bg=WHITE)
                cell(ws_calc, r2, 3, f'd = {p.get("inner_dia_mm","")} mm', bg=WHITE); r2 += 1
            r2 += 1
            ws_calc.merge_cells(f"A{r2}:C{r2}")
            hdr(ws_calc, r2, 1, "GAS CONSUMPTION", size=10); r2 += 1
            for g in supp.get("gas_consumption", []):
                bg2 = GREEN_BG if "Guarantee" in g["item"] else GREY
                cell(ws_calc, r2, 1, g["item"], bold="Guarantee" in g["item"], bg=bg2)
                cell(ws_calc, r2, 2, g["value"], bg=bg2); r2 += 1
            r2 += 1
            ws_calc.merge_cells(f"A{r2}:C{r2}")
            hdr(ws_calc, r2, 1, "GEAR BOX", size=10); r2 += 1
            for g in supp.get("gear_box", []):
                cell(ws_calc, r2, 1, g["item"], bold=True, bg=GREY)
                cell(ws_calc, r2, 2, g["value"], bg=WHITE); r2 += 1

            # ── Sheet: Recuperator ────────────────────────────────────────
            ws_rec = wb.create_sheet("Recuperator", 2)
            ws_rec.column_dimensions["A"].width = 36
            ws_rec.column_dimensions["B"].width = 18
            ws_rec.column_dimensions["C"].width = 14
            r2 = 1
            rcp = supp.get("recuperator", {})
            ws_rec.merge_cells(f"A{r2}:C{r2}")
            hdr(ws_rec, r2, 1, "RECUPERATOR CALCULATION", size=10); r2 += 1
            rec_rows = [
                ("Total Flue Gas", f'{rcp.get("total_flue_gas_nm3hr","")}', "Nm3/Hr"),
                ("Total Mass of Flue Gas", f'{rcp.get("total_mass_flue_gas_kghr","")}', "Kg/Hr"),
                ("Specific Heat of Flue Gas", f'{rcp.get("specific_heat_flue_gas","")}', "Kcal/Kg-C"),
                ("Initial Temp of Flue Gas", f'{rcp.get("initial_temp_flue_gas_c","")}', "°C"),
                ("Final Temp of Flue Gas", f'{rcp.get("final_temp_flue_gas_c","")}', "°C"),
                ("Heat Transfer Coefficient", f'{rcp.get("heat_transfer_coeff","")}', "Kcal/m2-C"),
                ("Volume of Combustion Air", f'{rcp.get("air_volume_nm3hr","")}', "Nm3/Hr"),
                ("Initial Temp of Air", f'{rcp.get("initial_air_temp_c","")}', "°C"),
                ("Final Temp of Air", f'{rcp.get("final_air_temp_c","")}', "°C"),
                ("Specific Heat of Air", f'{rcp.get("specific_heat_air","")}', "Kcal/Kg-C"),
                ("Heat Required By Combustion Air", f'{rcp.get("heat_required_kcal","")}', "Kcal"),
                ("Logarithmic Mean Temp Diff", f'{rcp.get("lmtd_c","")}', "°C"),
                ("Surface Area of Recuperator", f'{rcp.get("surface_area_m2","")}', "mtr2"),
                ("Pipes in Row × Column", f'{rcp.get("pipes_per_row","")} × {rcp.get("pipes_per_column","")}', f'= {rcp.get("pipes_total","")}'),
            ]
            for label, val, unit in rec_rows:
                cell(ws_rec, r2, 1, label, bold=True, bg=GREY)
                cell(ws_rec, r2, 2, val, bg=WHITE, align="right")
                cell(ws_rec, r2, 3, unit, bg=WHITE); r2 += 1
            r2 += 1
            ws_rec.merge_cells(f"A{r2}:C{r2}")
            hdr(ws_rec, r2, 1, "HOT BANK (SS 304)", size=10); r2 += 1
            for label, val, unit in [
                ("Pipe Diameter", rcp.get("pipe_dia_mm",""), "mm"),
                ("Pipe Length", rcp.get("pipe_length_m",""), "mtr"),
                ("Pipe Thickness", rcp.get("pipe_thickness_mm",""), "mm"),
                ("Weight per Pipe", rcp.get("pipe_weight_kg_m",""), "Kg/Mtr"),
                ("Total Weight Hot Bank", rcp.get("hot_bank_weight_kg",""), "Kg"),
            ]:
                cell(ws_rec, r2, 1, label, bg=GREY)
                cell(ws_rec, r2, 2, val, bg=WHITE, align="right")
                cell(ws_rec, r2, 3, unit, bg=WHITE); r2 += 1
            r2 += 1
            ws_rec.merge_cells(f"A{r2}:C{r2}")
            hdr(ws_rec, r2, 1, f'COLD BANK ({rcp.get("cold_bank_material","")})', size=10); r2 += 1
            for label, val, unit in [
                ("Pipe Diameter", rcp.get("cold_bank_dia_mm",""), "mm"),
                ("Pipe Length", rcp.get("cold_bank_length_m",""), "mtr"),
                ("Pipe Thickness", rcp.get("cold_bank_thickness_mm",""), "mm"),
                ("Weight per Pipe", rcp.get("cold_bank_weight_kg_m",""), "Kg/Mtr"),
                ("Total Weight Cold Bank", rcp.get("cold_bank_total_wt_kg",""), "Kg"),
            ]:
                cell(ws_rec, r2, 1, label, bg=GREY)
                cell(ws_rec, r2, 2, val, bg=WHITE, align="right")
                cell(ws_rec, r2, 3, unit, bg=WHITE); r2 += 1
            r2 += 1
            ws_rec.merge_cells(f"A{r2}:C{r2}")
            hdr(ws_rec, r2, 1, "COST OF RECUPERATOR", size=10); r2 += 1
            cost_items = [
                ("Cost of All Pipes", rcp.get("cost_all_pipes",0)),
                ("MS Outer Shell", rcp.get("cost_ms_outer_shell",0)),
                ("MS Combustion Air Inlet", rcp.get("cost_ms_combustion_air_inlet",0)),
                ("MS Channel 150×75×10", rcp.get("cost_ms_channel_150x75",0)),
                ("Angle 65×25 mtr", rcp.get("cost_angle_65",0)),
                ("Angle 75×10 mtr", rcp.get("cost_angle_75",0)),
                ("Angle 50×15 mtr", rcp.get("cost_angle_50",0)),
            ]
            for label, val in cost_items:
                cell(ws_rec, r2, 1, label, bg=GREY)
                cell(ws_rec, r2, 2, val, bg=WHITE, align="right", num_fmt='#,##0')
                cell(ws_rec, r2, 3, "Rs", bg=WHITE); r2 += 1
            cell(ws_rec, r2, 1, "Total Material Cost", bold=True, bg=GREEN_BG)
            cell(ws_rec, r2, 2, rcp.get("cost_total_material",0), bg=GREEN_BG, align="right", num_fmt='#,##0'); r2 += 1
            for label, val in [
                ("Bending of Pipes", rcp.get("cost_pipe_bending",0)),
                ("Welding Rod", rcp.get("cost_welding_rod",0)),
                ("Hole Fabrication", rcp.get("cost_hole_fabrication",0)),
                ("Thermocouple with TT", rcp.get("cost_thermocouple_tt",0)),
            ]:
                cell(ws_rec, r2, 1, label, bg=GREY)
                cell(ws_rec, r2, 2, val, bg=WHITE, align="right", num_fmt='#,##0')
                cell(ws_rec, r2, 3, "Rs", bg=WHITE); r2 += 1
            cell(ws_rec, r2, 1, "Total Cost of Recuperator", bold=True, bg=GREEN_BG, fg=GREEN)
            cell(ws_rec, r2, 2, rcp.get("cost_total_recuperator",0), bg=GREEN_BG, align="right", num_fmt='₹#,##0')

            # ── Sheet: Combustion ─────────────────────────────────────
            comb_mode = cs.get("combustion_mode", "onoff")
            comb_title = "Combustion-Massflow" if comb_mode == "massflow" else "Combustion-ONOFF"
            ws_comb = wb.create_sheet(comb_title, 3)
            ws_comb.column_dimensions["A"].width = 8
            ws_comb.column_dimensions["B"].width = 40
            ws_comb.column_dimensions["C"].width = 10
            ws_comb.column_dimensions["D"].width = 18
            r2 = 1
            ws_comb.merge_cells(f"A{r2}:D{r2}")
            hdr(ws_comb, r2, 1, f"COMBUSTION SYSTEM ({comb_title})", size=10); r2 += 1
            hdr(ws_comb, r2, 1, "S.No.", size=9); hdr(ws_comb, r2, 2, "Control System", size=9)
            hdr(ws_comb, r2, 3, "Qty", size=9); hdr(ws_comb, r2, 4, "Total Price", size=9); r2 += 1
            comb_items = [b for b in req.bom if b.get("SECTION") == "Combustion System"]
            comb_total = 0
            for i, ci in enumerate(comb_items, 1):
                cell(ws_comb, r2, 1, i, bg=GREY, align="right")
                cell(ws_comb, r2, 2, ci.get("ITEM",""), bg=WHITE)
                cell(ws_comb, r2, 3, ci.get("QTY",""), bg=WHITE, align="right")
                cell(ws_comb, r2, 4, ci.get("COST PRICE",0), bg=WHITE, align="right", num_fmt='#,##0')
                comb_total += ci.get("COST PRICE",0); r2 += 1
            cell(ws_comb, r2, 2, "TOTAL", bold=True, bg=GREEN_BG)
            cell(ws_comb, r2, 4, comb_total, bold=True, bg=GREEN_BG, align="right", num_fmt='#,##0')

            # Rename BOM sheet to legacy costing name
            costing_title = "Costing with Mass flow" if comb_mode == "massflow" else "Costing with Pulse Firing"
            ws.title = costing_title

        # ── Additional sheets for SNSF BRF (match legacy 7-sheet workbook) ──
        if supp and req.equipment_type == "SNSF BRF":
            # Sheet: Furnace (2) — full 110 rows from legacy
            ws2 = wb.create_sheet("Furnace (2)", 0)
            for ci, w in enumerate([4, 44, 18, 14, 4, 44, 18, 14, 4, 44], 1):
                ws2.column_dimensions[chr(64+ci)].width = w
            SECTION_KW = ['REHEATING','1.','2.','FLUE DUCT','STRUCTURAL','CASTING','PIPE CALC','PREHEATING']
            r2 = 0
            for row_data in supp.get("furnace_full", []):
                r2 += 1
                first = str(row_data[0] if row_data else "").strip()
                is_sec = any(first.upper().startswith(k) for k in SECTION_KW)
                is_hl = first.startswith("Total") or first.startswith("EFFECTIVE") or first.startswith("Overall")
                if is_sec:
                    bg2 = LIGHT
                elif is_hl:
                    bg2 = GREEN_BG
                else:
                    bg2 = GREY if r2 % 2 == 0 else WHITE
                for ci, v in enumerate(row_data):
                    val = v if v != '' else None
                    is_num = isinstance(val, (int, float))
                    cell(ws2, r2, ci+1, val,
                         bold=is_sec or is_hl,
                         bg=bg2,
                         fg=NAVY if is_sec else ("375623" if is_hl else "1E293B"),
                         align="right" if is_num else "left",
                         num_fmt='#,##0.##' if is_num else None)

            # Sheet: Refractory
            ws3 = wb.create_sheet("Refractory", 1)
            ws3.column_dimensions["A"].width = 44; ws3.column_dimensions["B"].width = 12; ws3.column_dimensions["C"].width = 14; ws3.column_dimensions["D"].width = 12; ws3.column_dimensions["E"].width = 16
            r2 = 1; ws3.merge_cells(f"A{r2}:E{r2}"); hdr(ws3, r2, 1, "REFRACTORY WITH MAHA KOSAL/BHILWARA/TRL", size=10); r2 += 1
            hdr(ws3, r2, 1, "Item", size=9); hdr(ws3, r2, 2, "Qty", size=9); hdr(ws3, r2, 3, "Weight", size=9); hdr(ws3, r2, 4, "Rate", size=9); hdr(ws3, r2, 5, "Cost", size=9); r2 += 1
            ref_total = 0
            for ri in supp.get("refractory_items", []):
                cell(ws3, r2, 1, ri["item"], bg=GREY); cell(ws3, r2, 2, ri["qty"], bg=WHITE, align="right")
                cell(ws3, r2, 3, ri["weight_kg"], bg=WHITE, align="right", num_fmt='#,##0')
                cell(ws3, r2, 4, ri["rate"], bg=WHITE, align="right"); cell(ws3, r2, 5, ri["cost"], bg=WHITE, align="right", num_fmt='#,##0')
                ref_total += ri["cost"]; r2 += 1
            for ri in supp.get("refractory_extra", []):
                cell(ws3, r2, 1, ri["item"], bg=GREY); cell(ws3, r2, 2, ri["qty"], bg=WHITE, align="right")
                cell(ws3, r2, 3, ri["wt"], bg=WHITE, align="right", num_fmt='#,##0')
                cell(ws3, r2, 4, ri["rate"], bg=WHITE, align="right"); cell(ws3, r2, 5, ri["cost"], bg=WHITE, align="right", num_fmt='#,##0')
                ref_total += ri["cost"]; r2 += 1
            cell(ws3, r2, 1, "Total Refractory", bold=True, bg=GREEN_BG); cell(ws3, r2, 5, ref_total, bold=True, bg=GREEN_BG, align="right", num_fmt='#,##0')

            # Sheet: Mild Steel
            ws4 = wb.create_sheet("Mild Steel", 2)
            ws4.column_dimensions["A"].width = 44; ws4.column_dimensions["B"].width = 14; ws4.column_dimensions["C"].width = 12; ws4.column_dimensions["D"].width = 14; ws4.column_dimensions["E"].width = 16
            r2 = 1; ws4.merge_cells(f"A{r2}:E{r2}"); hdr(ws4, r2, 1, "MILD STEEL", size=10); r2 += 1
            hdr(ws4, r2, 1, "Description", size=9); hdr(ws4, r2, 2, "Qty", size=9); hdr(ws4, r2, 3, "Wt/qty", size=9); hdr(ws4, r2, 4, "Total Wt", size=9); hdr(ws4, r2, 5, "Cost", size=9); r2 += 1
            for mi in supp.get("mild_steel", []):
                cell(ws4, r2, 1, mi["item"], bg=GREY); cell(ws4, r2, 2, mi["qty_m"], bg=WHITE, align="right")
                cell(ws4, r2, 3, mi["wt_per_m"], bg=WHITE, align="right"); cell(ws4, r2, 4, mi["total_wt"], bg=WHITE, align="right", num_fmt='#,##0')
                cell(ws4, r2, 5, mi["cost"], bg=WHITE, align="right", num_fmt='#,##0'); r2 += 1
            cell(ws4, r2, 1, "Total Structure", bold=True, bg=GREEN_BG)
            cell(ws4, r2, 4, supp.get("mild_steel_total_wt", 0), bg=GREEN_BG, align="right", num_fmt='#,##0')
            cell(ws4, r2, 5, supp.get("mild_steel_total_cost", 0), bg=GREEN_BG, align="right", num_fmt='#,##0'); r2 += 2
            ws4.merge_cells(f"A{r2}:E{r2}"); hdr(ws4, r2, 1, "PIPELINE & CASTING", size=10); r2 += 1
            for pi in supp.get("pipeline_casting", []):
                cell(ws4, r2, 1, pi["item"], bg=GREY); cell(ws4, r2, 3, pi["wt"], bg=WHITE, align="right")
                cell(ws4, r2, 4, pi["rate"], bg=WHITE, align="right"); cell(ws4, r2, 5, pi["cost"], bg=WHITE, align="right", num_fmt='#,##0'); r2 += 1

            # Sheet: Mass Flow Control
            ws5 = wb.create_sheet("Mass Flow Control", 3)
            ws5.column_dimensions["A"].width = 44; ws5.column_dimensions["B"].width = 10; ws5.column_dimensions["C"].width = 16
            r2 = 1; ws5.merge_cells(f"A{r2}:C{r2}"); hdr(ws5, r2, 1, "MASS FLOW CONTROL SYSTEM", size=10); r2 += 1
            mfc_total = 0
            for z in supp.get("mass_flow_control", []):
                cell(ws5, r2, 1, z["zone"], bold=True, bg=LIGHT, fg=NAVY); r2 += 1
                for it in z["items"]:
                    cell(ws5, r2, 1, it["item"], bg=GREY); cell(ws5, r2, 2, it["qty"], bg=WHITE, align="right")
                    t = it["qty"] * it["price"]; cell(ws5, r2, 3, t, bg=WHITE, align="right", num_fmt='#,##0'); mfc_total += t; r2 += 1
            cell(ws5, r2, 1, "TOTAL", bold=True, bg=GREEN_BG); cell(ws5, r2, 3, mfc_total, bold=True, bg=GREEN_BG, align="right", num_fmt='#,##0')

            # Sheet: Combustion
            ws6 = wb.create_sheet("Combustion", 4)
            ws6.column_dimensions["A"].width = 44; ws6.column_dimensions["B"].width = 10; ws6.column_dimensions["C"].width = 18; ws6.column_dimensions["D"].width = 18
            r2 = 1; ws6.merge_cells(f"A{r2}:D{r2}"); hdr(ws6, r2, 1, "COMBUSTION EQUIPMENT", size=10); r2 += 1
            hdr(ws6, r2, 1, "Description", size=9); hdr(ws6, r2, 2, "Qty", size=9); hdr(ws6, r2, 3, "Unit Price", size=9); hdr(ws6, r2, 4, "Cost Price", size=9); r2 += 1
            comb_total2 = 0
            for ci in supp.get("combustion_items", []):
                t = ci["qty"] * ci["price"]
                cell(ws6, r2, 1, ci["item"], bg=GREY); cell(ws6, r2, 2, ci["qty"], bg=WHITE, align="right")
                cell(ws6, r2, 3, ci["price"], bg=WHITE, align="right", num_fmt='#,##0')
                cell(ws6, r2, 4, t, bg=WHITE, align="right", num_fmt='#,##0'); comb_total2 += t; r2 += 1
            cell(ws6, r2, 1, "TOTAL", bold=True, bg=GREEN_BG); cell(ws6, r2, 4, comb_total2, bold=True, bg=GREEN_BG, align="right", num_fmt='#,##0')

            # Sheet: Recuperator1
            ws7 = wb.create_sheet("Recuperator1", 5)
            ws7.column_dimensions["A"].width = 40; ws7.column_dimensions["B"].width = 18; ws7.column_dimensions["C"].width = 14
            r2 = 1; rcp2 = supp.get("recuperator", {}); rcc = supp.get("recuperator_cost", {})
            ws7.merge_cells(f"A{r2}:C{r2}"); hdr(ws7, r2, 1, rcp2.get("title", "Recuperator Calculation"), size=10); r2 += 1
            for label, val, unit in [
                ("Total Flue Gas", rcp2.get("total_flue_gas_nm3hr",""), "Nm3/Hr"),
                ("Total Mass of Flue Gas", rcp2.get("total_mass_flue_gas_kghr",""), "Kg/Hr"),
                ("Initial Temp of Flue Gas", rcp2.get("initial_temp_flue_gas_c",""), "°C"),
                ("Final Temp of Flue Gas", rcp2.get("final_temp_flue_gas_c",""), "°C"),
                ("Air Volume", rcp2.get("air_volume_nm3hr",""), "Nm3/Hr"),
                ("Heat Required", rcp2.get("heat_required_kcal",""), "Kcal"),
                ("LMTD", rcp2.get("lmtd_c",""), "°C"),
                ("Surface Area", rcp2.get("surface_area_m2",""), "mtr2"),
                ("Pipes", f'{rcp2.get("pipes_per_row","")}×{rcp2.get("pipes_per_column","")} = {rcp2.get("pipes_total","")}', ""),
                ("Hot Bank Weight", rcp2.get("hot_bank_weight_kg",""), "Kg"),
                ("Cold Bank Weight", rcc.get("cold_bank_total_wt_kg",""), "Kg"),
            ]:
                cell(ws7, r2, 1, label, bold=True, bg=GREY); cell(ws7, r2, 2, val, bg=WHITE, align="right"); cell(ws7, r2, 3, unit, bg=WHITE); r2 += 1
            r2 += 1; ws7.merge_cells(f"A{r2}:C{r2}"); hdr(ws7, r2, 1, "COST OF RECUPERATOR", size=10); r2 += 1
            for label, val in [
                ("Cost of All Pipes", rcc.get("cost_all_pipes",0)),("MS Outer Shell", rcc.get("cost_ms_outer_shell",0)),
                ("MS Combustion Air Inlet", rcc.get("cost_ms_combustion_air_inlet",0)),
                ("MS Channel 150×75", rcc.get("cost_ms_channel_150x75",0)),
                ("Angle 65×25", rcc.get("cost_angle_65",0)),("Angle 75×10", rcc.get("cost_angle_75",0)),("Angle 50×15", rcc.get("cost_angle_50",0)),
            ]:
                cell(ws7, r2, 1, label, bg=GREY); cell(ws7, r2, 2, val, bg=WHITE, align="right", num_fmt='#,##0'); cell(ws7, r2, 3, "Rs", bg=WHITE); r2 += 1
            cell(ws7, r2, 1, "Total Material Cost", bold=True, bg=GREEN_BG)
            cell(ws7, r2, 2, rcc.get("cost_total_material",0), bg=GREEN_BG, align="right", num_fmt='#,##0'); r2 += 1
            for label, val in [("Bending",rcc.get("cost_pipe_bending",0)),("Welding Rod",rcc.get("cost_welding_rod",0)),
                ("Hole Fabrication",rcc.get("cost_hole_fabrication",0)),("Thermocouple",rcc.get("cost_thermocouple_tt",0))]:
                cell(ws7, r2, 1, label, bg=GREY); cell(ws7, r2, 2, val, bg=WHITE, align="right", num_fmt='#,##0'); r2 += 1
            cell(ws7, r2, 1, "Total Cost of Recuperator", bold=True, bg=GREEN_BG, fg=GREEN)
            cell(ws7, r2, 2, rcc.get("cost_total_recuperator",0), bg=GREEN_BG, align="right", num_fmt='₹#,##0')

            # Rename BOM sheet
            ws.title = "Breakup"

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        cname = req.customer.get("company_name", "").replace(" ", "_") or "ENCON"
        date_str = datetime.now().strftime("%d%b%Y")
        fname = f"{req.equipment_type}_Costing_{cname}_{date_str}.xlsx"
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f'attachment; filename="{fname}"'})

    # ════════════════════════════════════════════════════════════════════════
    #  VLPH / HLPH path
    # ════════════════════════════════════════════════════════════════════════
    ws = wb.active
    ws.title = req.equipment_type
    ws.column_dimensions["A"].width = 32
    ws.column_dimensions["B"].width = 22
    ws.column_dimensions["C"].width = 22
    ws.column_dimensions["D"].width = 14
    ws.column_dimensions["E"].width = 18
    ws.column_dimensions["F"].width = 18
    ws.column_dimensions["G"].width = 18
    ws.column_dimensions["H"].width = 18

    r = 1
    # ── Title ──────────────────────────────────────────────────────────────
    ws.merge_cells(f"A{r}:H{r}")
    t = ws.cell(row=r, column=1, value=f"ENCON — {req.equipment_type} Costing Sheet")
    t.font = Font(bold=True, color=WHITE, size=14, name="Calibri")
    t.fill = PatternFill("solid", fgColor=NAVY)
    t.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[r].height = 30
    r += 1

    # Date
    ws.merge_cells(f"A{r}:H{r}")
    d = ws.cell(row=r, column=1, value=f"Generated: {datetime.now().strftime('%d %b %Y, %H:%M')}")
    d.font = Font(color="64748B", size=9, name="Calibri")
    d.fill = PatternFill("solid", fgColor=LIGHT)
    d.alignment = Alignment(horizontal="right", vertical="center")
    r += 2

    # ── Customer ──────────────────────────────────────────────────────────
    if req.customer:
        ws.merge_cells(f"A{r}:H{r}")
        hdr(ws, r, 1, "CUSTOMER DETAILS", size=10)
        ws.row_dimensions[r].height = 20
        r += 1
        cust_fields = [
            ("Company",     req.customer.get("company_name","")),
            ("Contact",     req.customer.get("poc_name","")),
            ("Designation", req.customer.get("poc_designation","")),
            ("Mobile",      req.customer.get("mobile_no","")),
            ("Email",       req.customer.get("email","")),
            ("Project",     req.customer.get("project_name","")),
            ("Ref No.",     req.customer.get("ref_no","")),
        ]
        for label, val in cust_fields:
            if val:
                cell(ws, r, 1, label, bold=True, bg=GREY)
                ws.merge_cells(f"B{r}:H{r}")
                cell(ws, r, 2, val, bg=WHITE)
                r += 1
        r += 1

    # ── Process Parameters ────────────────────────────────────────────────
    ws.merge_cells(f"A{r}:H{r}")
    hdr(ws, r, 1, "PROCESS PARAMETERS", size=10)
    ws.row_dimensions[r].height = 20
    r += 1

    calc = req.calculations
    fuel_name = calc.get("fuel1_name", "")
    fuel_type = calc.get("fuel1_type", "ng")
    is_oil = fuel_type in {"hsd", "ldo", "hdo", "fo", "sko", "cfo", "lshs"}
    cv_unit = "kcal/kg" if is_oil else "kcal/Nm³"
    flow_unit = "kg" if is_oil else "Nm³"
    rate_unit = "kg/hr" if is_oil else "Nm³/hr"
    params = [
        ("Initial Temp (Ti)",       f"{calc.get('Ti','')} °C"),
        ("Final Temp (Tf)",         f"{calc.get('Tf','')} °C"),
        ("Refractory Weight",       f"{calc.get('refractory_weight','')} kg"),
        ("Fuel",                    fuel_name),
        ("Fuel CV",                 f"{calc.get('fuel_cv','')} {cv_unit}"),
        ("Time Taken",              f"{calc.get('time_taken_hr','')} hr"),
        ("Heat Load",               f"{calc.get('heat_load_kcal','')} kcal"),
        ("Firing Rate",             f"{calc.get('firing_rate_kcal','')} kcal/hr"),
        ("Fuel Consumption",        f"{calc.get('fuel_consumption_nm3','')} {flow_unit}"),
        ("Calc. Firing Rate",       f"{calc.get('calculated_firing_rate_nm3hr','')} {rate_unit}"),
        ("Design Firing Rate",      f"{calc.get('extra_firing_rate_nm3hr','')} {rate_unit}"),
        ("Equiv. Firing Rate (LPH)", f"{calc.get('equivalent_lph','')} ltr/hr"),
        ("Fuel Density",            f"{calc.get('fuel_density','')} {calc.get('fuel_density_unit','kg/ltr')}"),
    ]
    for i, (label, val) in enumerate(params):
        bg = GREY if i % 2 == 0 else WHITE
        cell(ws, r, 1, label, bold=True, bg=bg)
        ws.merge_cells(f"B{r}:H{r}")
        cell(ws, r, 2, val, bg=bg)
        r += 1
    r += 1

    # ── BOM Table ─────────────────────────────────────────────────────────
    ws.merge_cells(f"A{r}:H{r}")
    hdr(ws, r, 1, "BILL OF MATERIALS", size=10)
    ws.row_dimensions[r].height = 20
    r += 1

    # Columns: MEDIA(A) ITEM NAME(B) REFERENCE(C) QTY(D) MAKE(E) UNIT PRICE(F)
    #          TOTAL(G) BASIS(H)
    ws.column_dimensions["H"].width = 52
    bom_cols = ["MEDIA", "ITEM NAME", "REFERENCE", "QTY", "MAKE", "UNIT PRICE",
                "TOTAL", "BASIS — where the price comes from"]
    for ci, col in enumerate(bom_cols, 1):
        hdr(ws, r, ci, col, size=9)
    ws.row_dimensions[r].height = 22
    r += 1

    # Live formulas so every derived value is visible in the cell (like Regen):
    # per-line TOTAL = QTY(D) × UNIT PRICE(F); GRAND TOTAL = SUM over the item
    # rows. Calculated structural unit prices (weight × rate) are shown as
    # formulas when the weight is in the REFERENCE. The BASIS column names the
    # source (which master / Pricelist / calculation) of every line.
    import re as _re
    def _numf(x):
        return str(int(x)) if float(x) == int(x) else str(round(x, 4))
    _SUBTOTAL = {"BOUGHT OUT ITEMS", "ENCON ITEMS", "GRAND TOTAL"}
    _CALC_KG = ("FABRICATION", "AIR-GAS PIPELINE")   # unit price = kg × rate
    # Classification MUST match bom/vlph_builder.py: ENCON = MEDIA 'ENCON ITEMS'
    # only; Bought Out = everything else (incl. MISC ITEMS) minus the excluded
    # items. (MISC ITEMS are bought-out, not in-house.)
    _BUY_EXCLUDE = {"RATIO CONTROLLER"}
    _ENCON_BG  = "DCFCE7"   # light green  → ENCON (in-house) items
    _BOUGHT_BG = "DBEAFE"   # light blue   → bought-out items
    _item_rows = []
    _bought_rows = []
    _encon_rows = []
    for i, row_d in enumerate(req.bom):
        vals = list(row_d.values())
        media     = str(vals[0]) if len(vals) > 0 else ""
        item_name = str(vals[1]).strip() if len(vals) > 1 else ""
        ref       = str(vals[2]) if len(vals) > 2 else ""
        is_sub = item_name in _SUBTOTAL
        _is_encon = media.strip().upper() == "ENCON ITEMS"
        # Colour ENCON (in-house) vs bought-out distinctly; subtotal rows grey.
        bg = GREY if is_sub else (_ENCON_BG if _is_encon else _BOUGHT_BG)
        qty = vals[3] if len(vals) > 3 else None
        up  = vals[5] if len(vals) > 5 else None
        # cols 1-5 (MEDIA .. MAKE) written as-is
        for ci in range(1, 6):
            v = vals[ci - 1] if ci - 1 < len(vals) else ""
            cell(ws, r, ci, v, bg=bg, align="right" if ci == 4 else "left",
                 num_fmt='#,##0.00' if isinstance(v, (int, float)) and ci == 4 else None)
        # col 6 = UNIT PRICE — a weight×rate formula for calc items whose kg is
        # in the REFERENCE; otherwise the value.
        up_done = False
        if isinstance(up, (int, float)) and up and any(k in item_name.upper() for k in _CALC_KG):
            m = _re.search(r"([\d,]+(?:\.\d+)?)\s*kg", ref, _re.I)
            if m:
                kg = float(m.group(1).replace(",", ""))
                if kg:
                    cell(ws, r, 6, f"={_numf(kg)}*{_numf(round(up / kg, 4))}",
                         bg=bg, align="right", num_fmt='#,##0.00')
                    up_done = True
        if not up_done:
            cell(ws, r, 6, up if up is not None else "", bg=bg, align="right",
                 num_fmt='#,##0.00' if isinstance(up, (int, float)) else None)
        # col 7 = TOTAL
        if not is_sub and isinstance(qty, (int, float)) and isinstance(up, (int, float)):
            cell(ws, r, 7, f"=D{r}*F{r}", bg=bg, align="right", num_fmt='#,##0.00')
            _item_rows.append(r)
            if _is_encon:
                _encon_rows.append(r)
            elif item_name.upper() not in _BUY_EXCLUDE:
                _bought_rows.append(r)
        elif item_name == "GRAND TOTAL" and _item_rows:
            cell(ws, r, 7, f"=SUM(G{_item_rows[0]}:G{_item_rows[-1]})",
                 bold=True, bg=bg, align="right", num_fmt='#,##0.00')
        else:
            v = vals[6] if len(vals) > 6 else ""
            cell(ws, r, 7, v, bold=is_sub, bg=bg, align="right",
                 num_fmt='#,##0.00' if isinstance(v, (int, float)) else None)
        # col 8 = BASIS (source of the line)
        b = cell(ws, r, 8, _vlph_basis(item_name, media, ref), bg=bg, fg="475569")
        b.font = Font(color="475569", size=9, italic=True, name="Calibri")
        ws.row_dimensions[r].height = 18
        r += 1
    r += 1

    # ── Cost Summary (merged: cost breakup + commercial buildup) ──────────
    cs   = req.cost_summary
    comm = req.commercial or {}
    ws.merge_cells(f"A{r}:H{r}")
    hdr(ws, r, 1, "COST SUMMARY", size=10)
    ws.row_dimensions[r].height = 20
    r += 1

    import math
    def _pf(p):
        p = float(p or 0)
        return str(int(p)) if p == int(p) else str(round(p, 4))

    def _sline(label, value, formula=None, is_total=False):
        """One summary row: value/formula in col H, label in col A (B:G merged)."""
        nonlocal r
        bg = GREEN_BG if is_total else GREY
        fg = GREEN if is_total else "1E293B"
        cell(ws, r, 1, label, bold=is_total, bg=bg, fg=fg)
        ws.merge_cells(f"B{r}:G{r}")
        cell(ws, r, 2, "", bg=bg)
        c = ws.cell(row=r, column=8, value=formula if formula is not None else value)
        c.font = Font(bold=is_total, color=fg, size=11 if is_total else 10, name="Calibri")
        c.fill = PatternFill("solid", fgColor=bg)
        c.alignment = Alignment(horizontal="right", vertical="center")
        c.number_format = '₹#,##0.00'
        c.border = thin()
        ws.row_dimensions[r].height = 20
        used = r
        r += 1
        return used

    def _f(x):
        try:    return float(x)
        except (TypeError, ValueError): return None

    bought = _f(cs.get("bought_out_total")) or 0.0
    encon  = _f(cs.get("encon_total")) or 0.0
    grand  = _f(cs.get("grand_total"))

    if comm:
        # Merged commercial buildup. markup applies to bought-out only:
        #   Grand = BoughtOut×markup + ENCON  → then P&F/Designing/Negotiation are
        #   % of Grand, Final = round to nearest ₹1000, Order Total = Final × qty.
        markup = _f(comm.get("markup")) or 1.0
        # Bought Out / ENCON as live formulas over the BOM item rows (col G =
        # line TOTAL). The builder groups all bought-out rows together and all
        # ENCON/MISC rows together, so each is a clean =SUM(range). If a block
        # isn't contiguous, fall back to SUMIF / total−ENCON.
        # Clean SUM over each block's rows; joins contiguous runs with SUM(a:b)
        # and separates non-contiguous runs with '+' (bought-out = top block +
        # the MISC block, since the ENCON block sits between them).
        def _sum_rows(rows):
            if not rows:
                return None
            rows = sorted(rows)
            runs = []
            a = p = rows[0]
            for x in rows[1:]:
                if x == p + 1:
                    p = x
                else:
                    runs.append((a, p)); a = p = x
            runs.append((a, p))
            return "=" + "+".join(f"SUM(G{s}:G{e})" if s != e else f"G{s}" for s, e in runs)
        bought_f = _sum_rows(_bought_rows)
        encon_f  = _sum_rows(_encon_rows)
        r_bought = _sline("Bought Out Total", bought, formula=bought_f)
        r_encon  = _sline("ENCON Total", encon, formula=encon_f)
        # Grand — formula when it reconciles with Bought×markup + ENCON.
        g_formula = None
        if grand is not None and abs(grand - (bought * markup + encon)) < 0.5:
            g_formula = f"=H{r_bought}*{_pf(markup)}+H{r_encon}"
        r_grand = _sline(f"Grand Total (Bought Out × {_pf(markup)} + ENCON)",
                         grand if grand is not None else bought * markup + encon,
                         formula=g_formula)

        def _pct(label_key, amt_key, pct_key, text):
            amt = _f(comm.get(amt_key))
            if amt is None:
                return None
            pct = _f(comm.get(pct_key)) or 0
            f = None
            gval = grand if grand is not None else (bought * markup + encon)
            if gval and abs(amt - gval * pct / 100) < 0.5:
                f = f"=H{r_grand}*{_pf(pct)}/100"
            return _sline(text.format(pct=_pf(pct)), amt, formula=f)

        r_pf  = _pct("pf", "pf_amount", "pf_pct", "Packaging & Forwarding ({pct} %)")
        r_ds  = _pct("design", "design_amount", "design_pct", "Designing ({pct} %)")
        r_ng  = _pct("neg", "neg_amount", "neg_pct", "Negotiation ({pct} %)")
        r_tr  = None
        if _f(comm.get("transport_amount")) is not None:
            r_tr = _sline("Transport", _f(comm.get("transport_amount")))
        # Final Total — SUM(Grand..Transport) rounded to nearest ₹1000.
        final = _f(comm.get("final_total"))
        last_row = r - 1
        parts = [grand if grand is not None else (bought * markup + encon)]
        for k in ("pf_amount", "design_amount", "neg_amount", "transport_amount"):
            if _f(comm.get(k)) is not None:
                parts.append(_f(comm.get(k)))
        ssum = sum(parts)
        ffml = None
        if final is not None and abs(final - (round(ssum / 1000) * 1000)) < 0.5:
            ffml = f"=MROUND(SUM(H{r_grand}:H{last_row}),1000)"
        elif final is not None and abs(final - (math.ceil(ssum / 1000) * 1000)) < 0.5:
            ffml = f"=CEILING(SUM(H{r_grand}:H{last_row}),1000)"
        r_final = _sline("Final Total", final if final is not None else ssum,
                         formula=ffml, is_total=True)
        # Order quantity → order total (× N units).
        oq = int(req.order_qty or comm.get("order_qty") or 1)
        r_ordertotal = None
        if oq > 1:
            r_oq = _sline("Order Quantity (units)", oq)
            ws.cell(r_oq, 8).number_format = '#,##0'
            r_ordertotal = _sline(f"Order Total ({oq} units)", (final or ssum) * oq,
                                  formula=f"=H{r_final}*H{r_oq}", is_total=True)
        # USD conversion (commercial totals only). rate = INR→USD.
        if (req.currency or "").upper() == "USD" and (req.fx_rate or 0) > 0:
            _rate = _pf(req.fx_rate)
            r_rate = _sline(f"Exchange Rate (1 USD = ₹{_pf(round(1 / req.fx_rate, 2))})", req.fx_rate)
            ws.cell(r_rate, 8).number_format = '0.000000'
            _usd_base = r_ordertotal or r_final
            r_usd = _sline("Final Total (USD)", (final or ssum) * (oq if r_ordertotal else 1) * req.fx_rate,
                           formula=f"=H{_usd_base}*H{r_rate}", is_total=True)
            ws.cell(r_usd, 8).number_format = '$#,##0.00'
    else:
        # Fallback (no commercial payload): the plain 3-row breakup. Grand Total
        # is written as its exact value (its markup rule is product-specific).
        _sline("Bought Out Total", bought)
        _sline("ENCON Total", encon)
        _sline("Grand Total", grand if grand is not None else bought + encon, is_total=True)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    equip = req.equipment_type
    cname = req.customer.get("company_name", "").replace(" ", "_") or "ENCON"
    date_str = datetime.now().strftime("%d%b%Y")
    fname = f"{equip}_Costing_{cname}_{date_str}.xlsx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'}
    )


@app.get("/api/download-quote/{filename}")
def download_quote(filename: str):
    file_path = os.path.join(QUOTES_FOLDER, filename)
    if not os.path.exists(file_path):
        return {"error": "File not found"}
    return FileResponse(path=file_path, filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


def _soffice_binary():
    """Locate LibreOffice binary across Linux / macOS / Windows."""
    import shutil as _sh
    for cand in ("soffice", "libreoffice",
                 "/usr/bin/soffice", "/usr/bin/libreoffice",
                 r"C:\Program Files\LibreOffice\program\soffice.exe",
                 r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"):
        if os.path.exists(cand) or _sh.which(cand):
            return _sh.which(cand) or cand
    return None


def _docx_to_pdf(docx_path: str, pdf_path: str) -> bool:
    """Convert docx -> pdf via LibreOffice headless. Returns True on success.

    Each invocation gets its own UserInstallation profile so concurrent
    requests don't collide on the shared lock file.
    """
    soffice = _soffice_binary()
    if not soffice:
        return False

    import subprocess
    import tempfile
    import shutil as _sh

    out_dir = os.path.dirname(pdf_path) or "."
    profile_dir = tempfile.mkdtemp(prefix="lo_profile_")
    try:
        proc = subprocess.run(
            [
                soffice,
                "--headless",
                "--norestore",
                "--nolockcheck",
                f"-env:UserInstallation=file://{profile_dir}",
                "--convert-to", "pdf",
                "--outdir", out_dir,
                docx_path,
            ],
            capture_output=True, text=True, timeout=120,
        )
        if proc.returncode != 0:
            print(f"WARN: soffice convert failed (rc={proc.returncode}): "
                  f"stdout={proc.stdout!r} stderr={proc.stderr!r}")
            return False

        produced = os.path.join(
            out_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
        )
        if produced != pdf_path and os.path.exists(produced):
            _sh.move(produced, pdf_path)
        return os.path.exists(pdf_path)
    except Exception as e:
        print(f"WARN: soffice convert exception: {e}")
        return False
    finally:
        _sh.rmtree(profile_dir, ignore_errors=True)


@app.get("/api/pdf-quote/{filename}")
def pdf_quote(filename: str):
    """Serve the PDF version of a generated offer.

    Each call to /api/generate-quote now writes both a .docx (Word
    template) and a .pdf (built directly with reportlab) into
    QUOTES_FOLDER. This endpoint just serves the pre-generated PDF
    so the client doesn't need any docx->pdf conversion at request
    time. If the PDF doesn't exist (older quote or generator error),
    it builds it on demand from the docx-equivalent quote_data.
    """
    base = os.path.splitext(filename)[0]
    pdf_dst = os.path.join(QUOTES_FOLDER, f"{base}.pdf")

    if os.path.exists(pdf_dst):
        return FileResponse(path=pdf_dst, filename=f"{base}.pdf",
            media_type="application/pdf")

    # Prefer reconverting the existing .docx via LibreOffice so the PDF is
    # a faithful copy of the Word offer.
    docx_path = os.path.join(QUOTES_FOLDER, f"{base}.docx")
    if os.path.exists(docx_path):
        if _docx_to_pdf(docx_path, pdf_dst):
            return FileResponse(path=pdf_dst, filename=f"{base}.pdf",
                media_type="application/pdf")

    # Last-resort fallback: rebuild from persisted quote JSON via reportlab.
    json_path = os.path.join(QUOTES_FOLDER, f"{base}.json")
    if os.path.exists(json_path):
        try:
            import json as _json
            from engine.pdf_writer import generate_quote_pdf
            with open(json_path, encoding="utf-8") as f:
                quote_data = _json.load(f)
            generate_quote_pdf(quote_data, pdf_dst)
            return FileResponse(path=pdf_dst, filename=f"{base}.pdf",
                media_type="application/pdf")
        except Exception as e:
            return {"error": f"PDF build failed: {e}"}

    return {"error": "PDF not available — regenerate the quote to produce it."}


def _stamp_page_numbers(pdf_path: str):
    """Overlay 'Page X / Y' at the bottom-right of every page."""
    from pypdf import PdfReader, PdfWriter
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    import io

    reader = PdfReader(pdf_path)
    total = len(reader.pages)
    writer = PdfWriter()

    for i, page in enumerate(reader.pages, start=1):
        # Build a single-page overlay with the label
        box = page.mediabox
        width  = float(box.width)
        height = float(box.height)
        packet = io.BytesIO()
        c = canvas.Canvas(packet, pagesize=(width, height))
        c.setFont("Helvetica-Bold", 10)
        c.setFillColorRGB(0.10, 0.22, 0.36)   # dark navy
        c.drawRightString(width - 36, 20, f"Page {i} / {total}")
        c.save()
        packet.seek(0)
        stamp = PdfReader(packet).pages[0]
        page.merge_page(stamp)
        writer.add_page(page)

    with open(pdf_path, "wb") as f:
        writer.write(f)


@app.get("/api/preview-quote/{filename}")
def preview_quote(filename: str):
    """Preview a generated offer.

    Serves the PDF **inline** — a faithful LibreOffice render of the final
    .docx — so the preview looks pixel-identical to the Word document (cover
    layout, header logo, table styling, page breaks, etc.). Falls back to a
    simplified mammoth HTML render only when no PDF can be produced (e.g. a dev
    box without LibreOffice installed).
    """
    base = os.path.splitext(filename)[0]
    file_path = os.path.join(QUOTES_FOLDER, filename)
    pdf_dst   = os.path.join(QUOTES_FOLDER, f"{base}.pdf")

    # 1) Faithful PDF preview (build it from the .docx on demand if missing).
    if not os.path.exists(pdf_dst):
        docx_path = os.path.join(QUOTES_FOLDER, f"{base}.docx")
        if os.path.exists(docx_path):
            try:
                _docx_to_pdf(docx_path, pdf_dst)
            except Exception as _e:
                print(f"WARN: preview PDF build failed: {_e}")
    if os.path.exists(pdf_dst):
        return FileResponse(
            path=pdf_dst, media_type="application/pdf",
            headers={"Content-Disposition": f'inline; filename="{base}.pdf"'},
        )

    # 2) Fallback: simplified HTML render (LibreOffice unavailable).
    import mammoth
    if not os.path.exists(file_path):
        return HTMLResponse("<p style='color:red'>File not found</p>", status_code=404)
    with open(file_path, "rb") as f:
        result = mammoth.convert_to_html(f)
    html = f"""<!doctype html><html><head><meta charset="utf-8">
<style>
  body{{font-family:'Calibri','Segoe UI',Arial,sans-serif;font-size:11pt;color:#222;max-width:900px;margin:24px auto;padding:0 28px;line-height:1.4;}}
  h1,h2,h3,h4{{color:#1a3a5c;margin:14px 0 6px;}}
  table{{border-collapse:collapse;margin:10px 0;width:100%;}}
  table td,table th{{border:1px solid #bbb;padding:4px 8px;font-size:10.5pt;vertical-align:top;}}
  table th{{background:#eef2f7;}}
  p{{margin:4px 0;}}
  img{{max-width:100%;}}
  a{{color:#1d4ed8;}}
</style></head><body>{result.value}</body></html>"""
    return HTMLResponse(html)


@app.post("/upload-excel/")
async def upload_excel(request: Request, file: UploadFile = File(...)):
    try:
        filename = file.filename
        table_name = filename.replace(".xlsx", "").strip().lower()
        if VALID_TABLES and table_name not in VALID_TABLES:
            return {"error": f"Invalid table name: {table_name}"}
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        df = pd.read_excel(file_path)
        if df.empty:
            return {"error": "Excel file is empty"}
        df.columns = [col.strip().lower() for col in df.columns]
        client_ip = request.client.host if request.client else "unknown"
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        conn = sqlite3.connect(DB_PATH)
        df.to_sql(table_name, conn, if_exists="replace", index=False)
        conn.execute("""INSERT INTO table_update_log (table_name, updated_at, uploaded_by)
            VALUES (?, ?, ?) ON CONFLICT(table_name) DO UPDATE SET
            updated_at = excluded.updated_at, uploaded_by = excluded.uploaded_by""",
            (table_name, now, client_ip))
        conn.commit()
        conn.close()
        return {"message": f"{table_name} updated successfully!", "updated_at": now}
    except Exception as e:
        return {"error": str(e)}


@app.get("/db/tables")
def get_tables():
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("""SELECT name FROM sqlite_master
            WHERE type='table' AND name NOT LIKE 'sqlite_%' AND name != 'table_update_log'
            ORDER BY name""")
        tables = [row[0] for row in cursor.fetchall()]
        cursor.execute("SELECT table_name, updated_at, uploaded_by FROM table_update_log")
        log = {row[0]: {"updated_at": row[1], "uploaded_by": row[2]} for row in cursor.fetchall()}
        result = []
        for table in tables:
            cursor.execute(f"SELECT COUNT(*) FROM [{table}]")
            count = cursor.fetchone()[0]
            entry = log.get(table, {})
            result.append({"name": table, "rows": count,
                "updated_at": entry.get("updated_at"), "uploaded_by": entry.get("uploaded_by")})
        conn.close()
        return {"tables": result}
    except Exception as e:
        return {"error": str(e)}


@app.get("/db/table/{table_name}")
def get_table_data(table_name: str):
    if VALID_TABLES and table_name not in VALID_TABLES:
        return {"error": "Invalid table name"}
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute(f"SELECT * FROM [{table_name}]")
        rows = cursor.fetchall()
        columns = [desc[0] for desc in cursor.description]
        cursor.execute("SELECT updated_at, uploaded_by FROM table_update_log WHERE table_name = ?", (table_name,))
        log_row = cursor.fetchone()
        conn.close()
        return {"table": table_name, "columns": columns,
            "rows": [list(r) for r in rows], "total": len(rows),
            "updated_at": log_row[0] if log_row else None,
            "uploaded_by": log_row[1] if log_row else None}
    except Exception as e:
        return {"error": str(e)}