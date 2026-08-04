"""
Add the National Energy Efficiency Innovation Award block to the recuperator
offer's cover letter.

The regen and vertical letters both close with the award heading and the
trophy/certificate photo; the recup template had neither — no heading, and the
award image was never in the file. This copies both from the vertical template
(the recup template was derived from it, so the styling matches), inserting them
after the signature block.

Run:  python patch_recup_award.py
      python patch_offer_typography.py     # put the new runs on the size scale

Idempotent — a second run finds the heading already present.
"""

import copy
import os

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Emu
from docx.text.paragraph import Paragraph

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SOURCE = "Offer_Template.docx"           # vertical — has the award block
TARGET = "Recup_Offer_Template.docx"
AWARD_TEXT = "NATIONAL ENERGY EFFICIENCY INNOVATION AWARD"
# Last line of the signature block. The recup letter closes with
# "Email: {{technical_email}}   |   www.encon.co.in" — one line, where the
# vertical has a separate "Website:" line — so match on the email placeholder.
SIGN_OFF = "{{technical_email}}"


def _award_image(doc):
    """(bytes, width, height) of the award photo — the small inline image, not
    the full-page cover logo."""
    best = None
    for shape in doc.inline_shapes:
        w = Emu(shape.width).inches
        if w > 3:
            continue                     # cover logo
        rId = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
        blob = doc.part.related_parts[rId].blob
        if best is None or w < best[1]:
            best = (blob, w, shape.width, shape.height)
    return best


def patch():
    src = Document(os.path.join(BASE_DIR, SOURCE))
    tgt_path = os.path.join(BASE_DIR, TARGET)
    tgt = Document(tgt_path)

    if any(AWARD_TEXT in p.text.upper() for p in tgt.paragraphs):
        print(f"  {TARGET}: award block already present")
        return

    award_para = next((p for p in src.paragraphs if AWARD_TEXT in p.text.upper()), None)
    img = _award_image(src)
    if award_para is None or img is None:
        print(f"  SKIP — award heading or image not found in {SOURCE}")
        return
    blob, _w_in, width, height = img

    anchor = next((p for p in tgt.paragraphs if SIGN_OFF in p.text), None)
    if anchor is None:
        print(f"  SKIP — no '{SIGN_OFF}' line to anchor to in {TARGET}")
        return

    # Heading: clone the vertical's paragraph so the colour/weight carry over.
    new_head = copy.deepcopy(award_para._element)
    anchor._element.addnext(new_head)

    # Image: a fresh paragraph after the heading. add_picture registers the
    # image part properly — a deepcopy would reference a relationship that does
    # not exist in this document.
    spacer = tgt.add_paragraph()
    new_head.addnext(spacer._element)
    run = spacer.add_run()
    import io
    run.add_picture(io.BytesIO(blob), width=width, height=height)

    tgt.save(tgt_path)
    print(f"  OK  {TARGET}: award heading + photo added "
          f"({Emu(width).inches:.2f} x {Emu(height).inches:.2f} in)")


if __name__ == "__main__":
    patch()
