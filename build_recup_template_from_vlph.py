"""Build Recup_Offer_Template.docx by cloning the VLPH offer template
(Offer_Template.docx) and surgically replacing the ladle-preheater
content with recuperator content.

What this script keeps from the VLPH template (the "shell"):
  * Cover page, ENCON banner, cover letter
  * Table of Contents + List of Annexures
  * Section 1 Company Profile (universal)
  * Section 2 About the Client + Client Details + Marketing Person tables
  * Section heading "3. TECHNICAL SPECIFICATIONS"
  * ANNEXURE I (Scope of Supply table shell — content swapped)
  * ANNEXURE II (Exclusions paragraphs — universal)
  * ANNEXURE III (Price Schedule table — rebuilt for single/full toggle)
  * Supervision Charges sub-table
  * ANNEXURE IV (T&C 12-row table — universal)
  * ANNEXURE V (Reference List — universal)
  * ANNEXURE VI (Make List — universal)

What this script replaces:
  * "VERTICAL LADLE PREHEATERS AND DRYERS" -> "RECUPERATOR (Waste Heat
    Recovery System)"
  * 24-row Technical Specs table  -> 15-row Recup Designing Params
  * Material of Construction sub-table inserted ABOVE the Designing
    Params table
  * 3D image placeholders inserted AFTER the Designing Params table
  * VLPH scope-of-supply body (~200 paragraphs of fuel/oil/PLC/hood
    conditionals) -> short recup operation description
  * Annexure I 12-row Scope table content -> recup component scope rows
  * Annexure III Price Schedule -> single/full toggle (matches
    build_recup_template_v2.py output)

After this script the template still expects all VLPH-style
placeholders to be supplied by the backend:
  equipment_name, poc_designation, technical_person, technical_phone,
  technical_email, supervision_mech, supervision_plc, tnc_* etc.

Plus the recup-specific keys introduced here:
  hot_tube_plate_material, hot_tube_material, cold_tube_plate_material,
  cold_tube_material, flue_*, air_*, surface_area_m2, pipe_*,
  image_recup_side/front/top, bom_rows, bought_out_total, encon_total,
  grand_total, grand_total_in_words, price_schedule_style,
  supervision_include/rate/note.
"""
from __future__ import annotations

import os
import shutil
from copy import deepcopy

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.enum.table import WD_ROW_HEIGHT_RULE

SOURCE = "Offer_Template.docx"
TARGET = "Recup_Offer_Template.docx"


# ── Material of Construction layout (cold/hot/others) ──────────────────────
# Mirrors build_recup_template_v2.py's _MOC_ROWS.
_MOC_ROWS = [
    ('header',  'Material of Construction',  None),
    ('section', 'COLD BANK',                 None),
    ('item',    'Tube Plate',                '{{ cold_tube_plate_material }}'),
    ('item',    'Tube',                      '{{ cold_tube_material }}'),
    ('item',    'Duct Inlet',                'Mild Steel'),
    ('item',    'Inlet Collar',              'Mild Steel'),
    ('section', 'HOT BANK',                  None),
    ('item',    'Tube Plate',                '{{ hot_tube_plate_material }}'),
    ('item',    'Tube',                      '{{ hot_tube_material }}'),
    ('item',    'Duct Outlet',               'Mild Steel'),
    ('item',    'Outlet Collar',             'Mild Steel'),
    ('item',    'Supporting Frame',          'Mild Steel'),
    ('item',    'Bottom Duct',               'Mild Steel'),
    ('item',    'Flange',                    'Mild Steel'),
    ('section', 'OTHERS',                    None),
    ('item',    'Flanges',                   'Mild Steel'),
    ('item',    'Air Inlet Duct',            'Mild Steel'),
    ('item',    'Air Outlet Duct',           'Mild Steel'),
    ('item',    'Bottom Air Receiving Box',  'Mild Steel'),
    ('item',    'Matching Flange',           'Mild Steel'),
    ('item',    'Nut, Bolt & Washer',        'Mild Steel'),
    ('item',    'Gasket',                    'Heat Resistant'),
    ('item',    'Supporting Structure',      'Mild Steel'),
]

# ── Recup Designing Parameters layout (replaces VLPH's 24-row spec) ────────
_DESIGNING_PARAMS = [
    ('No. of Units',              '{{ recup_qty }}'),
    ('Flue Gas Flow',             '{{ flue_flow_nm3hr }} Nm³/hr'),
    ('Flue Gas Temperature (In)', '{{ flue_temp_in_C }} °C'),
    ('Flue Gas Temperature (Out)','{{ flue_temp_out_C }} °C'),
    ('Air Volume',                '{{ air_volume_nm3hr }} Nm³/hr'),
    ('Air Temperature (In)',      '{{ air_temp_in_C }} °C'),
    ('Air Temperature (Out)',     '{{ air_temp_out_C }} °C'),
    ('Heat Transfer Area',        '{{ surface_area_m2 }} m²'),
    ('Pipe Outside Diameter',     '{{ pipe_dia_mm }} mm'),
    ('Pipe Length (per bank)',    '{{ pipe_length_m }} m'),
    ('Pipe Wall Thickness',       '{{ pipe_thick_mm }} mm'),
]

# ── Annexure V — Reference List (recup-specific clients) ───────────────────
# Pulled from the ENCON recuperator offer (separate page-V layout the
# user supplied). Format is (S.No., Client, Application). All entries
# are Recuperator deliveries so the Application column is uniform.
_RECUP_CLIENTS = [
    "Jordan Steel, Jordan",
    "Utility Alloys Pvt. Ltd., Coimbatore",
    "CUMI Refractories, Chennai",
    "Sunflag Iron & Steel Company Ltd., Maharashtra",
    "Tata Steel Ltd, Boisar, MH",
    "Modern Steel Ltd, Mandi Gobindgarh",
    "Raj Shree Udyog, Mandi Gobindgarh",
    "Dhiman Iron & Steel Industries, Mandi Gobindgarh",
    "Bharat Ispat Udyog Pvt. Ltd., Mandi Gobindgarh",
    "Tameer Steel Factory Company Ltd., Saudi Arabia",
    "Premier Rolling Mills, Kenya",
    "Gulf Steel Industries, Abu Dhabi",
    "GTB Columbo Corporation, Sri Lanka",
    "Divine Alloys & Power Company Ltd., Jamshedpur",
    "Maadi Steel, Egypt",
    "Steel Rolling Mills Ltd., Uganda",
    "Star Wire, Faridabad",
    "Sadhu Forging Ltd., Faridabad",
    "Hindustan Udyog Ltd., Kolkata",
    "Steel Makers Ltd., Kenya",
]


# ── Annexure I — Scope of Supply (recup) ────────────────────────────────────
_RECUP_SCOPE_ITEMS = [
    'Recuperator (Hot & Cold Bank) with tube plate, tubes, and supporting frame',
    'Material of Construction as per Annexure I (Material of Construction table above)',
    'Air inlet & outlet ducts with matching flanges and gaskets',
    'Combustion air receiving box with all internal baffles',
    'Bottom duct with expansion joints',
    'Pipe holding plates and supporting structure',
    'MS Side Hood (with insulation provision)',
    'Thermocouple (Type-K) for flue gas temperature monitoring',
    'Nut, bolt, washer, and matching flanges',
    'Painting and Packing (heat-resistant outer paint)',
]


def _make_moc_row(table, kind: str, label: str | None, value: str | None):
    new_row = table.add_row()
    cells = new_row.cells
    if kind in ('header', 'section'):
        cells[0].merge(cells[1])
        cells[0].text = label
        for p in cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
    elif kind == 'item':
        cells[0].text = label
        cells[1].text = value or ''
    return new_row._element


def _replace_spec_table_with_designing_params(table) -> None:
    """Wipe table rows (except header) and re-fill with recup Designing
    Parameters. The first row is kept as the 'Parameter | Specification'
    header from VLPH."""
    tbl_xml = table._element
    # Keep header row (index 0), strip the rest.
    for tr in list(tbl_xml.findall(qn('w:tr')))[1:]:
        tbl_xml.remove(tr)
    # Replace header text with a recup-friendly banner.
    head = table.rows[0]
    head.cells[0].text = 'Recuperator Designing Parameters'
    head.cells[1].text = ''
    head.cells[0].merge(head.cells[1])
    for p in head.cells[0].paragraphs:
        for r in p.runs:
            r.bold = True
    # Add new rows.
    for label, value in _DESIGNING_PARAMS:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value


def _insert_moc_table_before(designing_table) -> None:
    """Build a brand-new MoC table (cloning the Designing Params table's
    style) and insert it BEFORE designing_table in the document body."""
    src_xml = designing_table._element
    moc_xml = deepcopy(src_xml)
    # Wipe all rows in clone; we'll add fresh.
    for tr in list(moc_xml.findall(qn('w:tr'))):
        moc_xml.remove(tr)
    # Insert the clone before designing_table.
    src_xml.addprevious(moc_xml)
    # Add a separating paragraph BETWEEN the MoC and Designing Params.
    sep = OxmlElement('w:p')
    src_xml.addprevious(sep)
    # Re-wrap the clone via a Table object so .add_row works.
    from docx.table import Table
    moc_table = Table(moc_xml, designing_table._parent)
    # Header row first
    moc_table.add_row()
    moc_table.add_row()  # we'll merge-fix in _make_moc_row
    # Actually clearer: build rows one at a time
    for tr in list(moc_xml.findall(qn('w:tr'))):
        moc_xml.remove(tr)
    for kind, label, value in _MOC_ROWS:
        _make_moc_row(moc_table, kind, label, value)


def _insert_3d_images_before(designing_table) -> None:
    """Insert '3D Image of the Proposed Recuperator' heading + three
    image placeholders immediately before the Designing Parameters
    table (so they appear between the MoC table and the Designing
    Params table — i.e. right BEFORE the client-given parameters)."""
    anchor = designing_table._element

    def _para(text: str = '', bold: bool = False) -> OxmlElement:
        p = OxmlElement('w:p')
        if text:
            r = OxmlElement('w:r')
            if bold:
                rPr = OxmlElement('w:rPr')
                b = OxmlElement('w:b'); rPr.append(b); r.append(rPr)
            t = OxmlElement('w:t')
            t.text = text
            r.append(t)
            p.append(r)
        return p

    image_block = [
        _para('3D Image of the Proposed Recuperator', bold=True),
        _para('{{ image_recup_side }}'),
        _para(),
        _para('{{ image_recup_front }}'),
        _para(),
        _para('{{ image_recup_top }}'),
        _para(),
    ]
    for el in image_block:
        anchor.addprevious(el)


def _replace_scope_table_with_recup(table) -> None:
    """Replace the Annexure I 12-row Scope of Supply rows with recup
    components. The first row (header) stays untouched.

    VLPH structure for this table was:
      row 0: 'S. No. | Item Description'  (header)
      row 1: equipment name banner row (merged)
      rows 2-11: numbered scope items
    """
    tbl_xml = table._element
    # Drop rows 1+ (keep the 'S. No. | Item Description' header)
    for tr in list(tbl_xml.findall(qn('w:tr')))[1:]:
        tbl_xml.remove(tr)
    # Banner row (merged across both cols) — short label only
    banner = table.add_row()
    banner.cells[0].merge(banner.cells[1])
    banner.cells[0].text = 'Recuperator'
    for p in banner.cells[0].paragraphs:
        for r in p.runs:
            r.bold = True
    # Numbered scope items
    for idx, item in enumerate(_RECUP_SCOPE_ITEMS, start=1):
        row = table.add_row()
        row.cells[0].text = str(idx)
        row.cells[1].text = item


def _delete_vlph_scope_body(doc: Document) -> None:
    """Walk the body looking for the 'SCOPE OF SUPPLY' paragraph (the
    inline one — i.e. NOT the 'ANNEXURE I — SCOPE OF SUPPLY' heading)
    and delete every element from there up to (but not including) the
    'ANNEXURE I' heading. Replace with a short recup operation block."""
    body = doc.element.body
    elements = list(body)

    start_idx = None
    end_idx = None
    for i, el in enumerate(elements):
        if el.tag.split('}')[-1] != 'p':
            continue
        txt = ''.join(t.text or '' for t in el.iter(qn('w:t'))).strip()
        if start_idx is None and txt == 'SCOPE OF SUPPLY':
            start_idx = i
        elif start_idx is not None and txt.startswith('ANNEXURE I'):
            end_idx = i
            break

    if start_idx is None or end_idx is None:
        print('Could not locate VLPH scope body markers — skipping delete.')
        return

    # Remember the anchor (the ANNEXURE I element); we'll add replacement
    # paragraphs before it.
    anchor = elements[end_idx]

    # Delete elements [start_idx .. end_idx)
    for el in elements[start_idx:end_idx]:
        body.remove(el)

    # Insert recup operation block right before the ANNEXURE I anchor.
    def _para(text: str = '', bold: bool = False) -> OxmlElement:
        p = OxmlElement('w:p')
        if text:
            r = OxmlElement('w:r')
            if bold:
                rPr = OxmlElement('w:rPr')
                b = OxmlElement('w:b'); rPr.append(b); r.append(rPr)
            t = OxmlElement('w:t')
            t.text = text
            r.append(t)
            p.append(r)
        return p

    # The 3D-image block is inserted separately (between MoC and
    # Designing Params) by _insert_3d_images_before(). Keep this
    # post-Designing-Params block focused on text only.
    recup_body = [
        _para('SCOPE OF SUPPLY', bold=True),
        _para('Our scope of supply will cover design, engineering, '
              'manufacturing, supply, and supervision for erection & '
              'commissioning of the proposed Recuperator (Waste Heat '
              'Recovery System).'),
        _para('OPERATION', bold=True),
        _para('The Recuperator is a heat-recovery exchanger that uses '
              'the hot flue gases leaving the furnace to preheat the '
              'combustion air being supplied to the burner. The hot '
              'flue gas passes through the hot bank tubes, while '
              'ambient combustion air flows around them in a counter-'
              'cross-flow arrangement. The preheated air reduces fuel '
              'consumption and improves overall furnace efficiency.'),
        _para('Design and material of construction details are listed '
              'in the Material of Construction sub-table above. '
              'Designing parameters and process flows are listed in '
              'the Recuperator Designing Parameters table.'),
    ]
    # addprevious inserts BEFORE the anchor — iterate in normal order so
    # the first item lands earliest in the document.
    for el in recup_body:
        anchor.addprevious(el)

    # Also flip the section title "VERTICAL LADLE PREHEATERS AND DRYERS"
    for el in body.iter(qn('w:p')):
        for t in el.iter(qn('w:t')):
            if t.text and 'VERTICAL LADLE PREHEATERS' in t.text:
                t.text = 'RECUPERATOR (Waste Heat Recovery System)'


def _replace_reference_list(doc: Document) -> None:
    """Replace the Annexure V Reference List (50-row ladle-preheater
    table) with the 20-row recup client list. The table header
    (S. No. | Client | Application) is preserved. Application column
    is set to 'Recuperator' for every entry.

    Also rewrites the intro paragraph above the table from the
    'ladle preheating systems' line to a recup equivalent."""
    # 1. Find the reference list table by header signature.
    table = None
    for t in doc.tables:
        head = [c.text.strip() for c in t.rows[0].cells] if t.rows else []
        if (len(head) >= 3 and head[0].upper() == 'S. NO.'
                and head[1].upper() == 'CLIENT' and 'APPLICATION' in head[2].upper()):
            table = t
            break
    if table is None:
        print('Reference List table not found — skipping replace.')
        return

    tbl_xml = table._element
    # Strip rows 1+; keep the header row.
    for tr in list(tbl_xml.findall(qn('w:tr')))[1:]:
        tbl_xml.remove(tr)

    # Add recup client rows.
    for idx, client in enumerate(_RECUP_CLIENTS, start=1):
        row = table.add_row()
        row.cells[0].text = str(idx)
        row.cells[1].text = client
        row.cells[2].text = 'Recuperator'

    # 2. Rewrite the intro paragraph above the table. VLPH had:
    #   'ENCON has supplied ladle preheating systems to leading steel
    #    and casting houses across India and overseas. A representative
    #    list is below:'
    # Recup intro per the user's screenshot:
    #   'We have supplied Recuperators to various clients, some of
    #    them are listed below:'
    for p in doc.paragraphs:
        if 'ladle preheating systems' in p.text.lower():
            # Replace the entire paragraph text.
            for run in p.runs:
                run.text = ''
            if p.runs:
                p.runs[0].text = (
                    'We have supplied Recuperators to various clients, '
                    'some of them are listed below:'
                )
            else:
                p.add_run(
                    'We have supplied Recuperators to various clients, '
                    'some of them are listed below:'
                )
            break


def _inject_scope_of_supply_paragraph(doc: Document) -> None:
    """Rewrite Annexure I — Scope of Supply with the recup-specific
    intro paragraph + 'RECUPERATOR (Waste Heat Recovery System)'
    sub-heading + descriptive paragraph that the user supplied.

    Final order in Annexure I:
      1. 'ANNEXURE I — SCOPE OF SUPPLY' heading        (unchanged)
      2. Intro:  'The Scope of supply will cover Design,
                  manufacturing and supply of the recuperator
                  as per the specifications provided by the
                  client mentioned below.'
      3. (blank spacer)
      4. 'RECUPERATOR (Waste Heat Recovery System)'   (bold + underline)
      5. Descriptive paragraph (recup operation)
      6. (blank spacer)
      7. T7 Scope of Supply table                     (unchanged)
    """
    body = doc.element.body
    elements = list(body)

    # 1. Locate the 'ANNEXURE I' heading.
    heading_idx = None
    for i, el in enumerate(elements):
        if el.tag.split('}')[-1] != 'p':
            continue
        txt = ''.join(t.text or '' for t in el.iter(qn('w:t'))).strip()
        if 'ANNEXURE I' in txt.upper() and 'SCOPE OF SUPPLY' in txt.upper():
            heading_idx = i
            break
    if heading_idx is None:
        print('Annexure I heading not found — skipping scope-intro injection.')
        return

    # 2. Find the intro paragraph immediately after the heading.
    intro_p = None
    for j in range(heading_idx + 1, len(elements)):
        if elements[j].tag.split('}')[-1] == 'p':
            intro_p = elements[j]
            break
    if intro_p is None:
        return

    # Idempotency: if the intro already starts with our recup text, skip.
    intro_txt = ''.join(t.text or '' for t in intro_p.iter(qn('w:t'))).strip()
    if intro_txt.startswith('The Scope of supply will cover Design, manufacturing'):
        return

    # 3. Replace the intro text.
    NEW_INTRO = (
        'The Scope of supply will cover Design, manufacturing and supply '
        'of the recuperator as per the specifications provided by the '
        'client mentioned below.'
    )
    # Wipe all <w:t> elements inside intro_p, then set the first to our text.
    t_els = list(intro_p.iter(qn('w:t')))
    for t in t_els:
        t.text = ''
    if t_els:
        t_els[0].text = NEW_INTRO
    else:
        # No runs at all — synthesize one.
        r = OxmlElement('w:r')
        t = OxmlElement('w:t'); t.text = NEW_INTRO; r.append(t)
        intro_p.append(r)

    # 4. Insert blank + RECUPERATOR heading + descriptive paragraph
    #    + blank spacer, AFTER intro_p (so they sit between intro and
    #    the Scope of Supply table).
    def _para(text: str = '', *, bold: bool = False, underline: bool = False) -> OxmlElement:
        p = OxmlElement('w:p')
        if text:
            r = OxmlElement('w:r')
            if bold or underline:
                rPr = OxmlElement('w:rPr')
                if bold:
                    rPr.append(OxmlElement('w:b'))
                if underline:
                    u = OxmlElement('w:u')
                    u.set(qn('w:val'), 'single')
                    rPr.append(u)
                r.append(rPr)
            t = OxmlElement('w:t')
            t.text = text
            t.set(qn('xml:space'), 'preserve')
            r.append(t)
            p.append(r)
        return p

    DESCRIPTION = (
        'A waste heat recovery Recuperator of suitable capacity will '
        'be provided, which will preheat the combustion air to a '
        'temperature of about {{ air_temp_out_C }} °C, which in turn '
        'decreases fuel consumption. Recuperator will be of '
        'Convective type having two passes for air and single pass '
        'for flue gas. The flue gas will pass over the bank while '
        'air passes through the tubes. The tubes are provided for '
        'hot Bank and cold Bank. The outer body of the recuperator '
        'shall be fabricated from MS plates of suitable thickness '
        'so that it may sustain thermal stresses developed during '
        'its work. The recuperator will be installed above the ground.'
    )

    # addnext goes immediately after intro_p. Insert in reverse order
    # so the final sequence is: intro_p, spacer, recup-head, desc, spacer2.
    nodes_in_order = [
        _para(),                                                  # spacer
        _para('RECUPERATOR (Waste Heat Recovery System)',
              bold=True, underline=True),
        _para(DESCRIPTION),
        _para(),                                                  # spacer
    ]
    for el in reversed(nodes_in_order):
        intro_p.addnext(el)


def _remove_annexure_vi(doc: Document) -> None:
    """Strip Annexure VI — Make List from the cloned VLPH template.
    Removes:
      * The 'ANNEXURE VI — MAKE LIST' heading paragraph
      * The intro paragraph below it
      * The empty make-list table
      * The 'Annexure VI | Make List | Attached' row from the LIST OF
        ANNEXURES table on the cover page (T2)
    """
    # ── Part A: strip the Annexure VI row from LIST OF ANNEXURES ──────
    for t in doc.tables:
        if not t.rows or len(t.rows[0].cells) < 2:
            continue
        head = [c.text.strip().upper() for c in t.rows[0].cells]
        # Match the 3-col List of Annexures header
        if 'ANNEXURE NO.' in head[0]:
            tbl_xml = t._element
            for tr in list(tbl_xml.findall(qn('w:tr'))):
                row_txt = ''.join(x.text or '' for x in tr.iter(qn('w:t'))).upper()
                # 'ANNEXURE VI' but NOT 'ANNEXURE V ' (don't match Annexure V row)
                if 'ANNEXURE VI' in row_txt and 'MAKE' in row_txt:
                    tbl_xml.remove(tr)
            break

    # ── Part B: strip the body section (heading + intro + table) ─────
    body = doc.element.body
    elements = list(body)

    heading_idx = None
    for i, el in enumerate(elements):
        if el.tag.split('}')[-1] != 'p':
            continue
        txt = ''.join(t.text or '' for t in el.iter(qn('w:t'))).strip()
        if 'ANNEXURE VI' in txt.upper() and 'MAKE' in txt.upper():
            heading_idx = i
            break
    if heading_idx is None:
        return  # idempotent — already removed

    # Remove the heading + everything after it up to (and including)
    # the first table we encounter (the make-list table). Then stop.
    i = heading_idx
    removed_table = False
    while i < len(elements):
        el = elements[i]
        tag = el.tag.split('}')[-1]
        if tag == 'sectPr':
            break  # don't remove the doc's section properties
        body.remove(el)
        if tag == 'tbl':
            removed_table = True
            # Remove one trailing blank paragraph too (if present) for clean spacing.
            if i + 1 < len(elements) and elements[i+1].tag.split('}')[-1] == 'p':
                nxt_txt = ''.join(t.text or '' for t in elements[i+1].iter(qn('w:t'))).strip()
                if not nxt_txt:
                    body.remove(elements[i+1])
            break
        i += 1
    print(f'Removed Annexure VI (table removed={removed_table})')


def _pad_table_rows(doc: Document, min_height_cm: float = 0.75) -> None:
    """Give the Price Schedule and Supervision sub-table a generous
    minimum row height so the printed offer doesn't look cramped.
    Control rows ({%tr if%}, {%tr endif%}, {%tr for%}, etc.) are
    intentionally NOT padded — they get stripped at render time and
    leaving them at default height keeps the source template compact."""
    targets = []
    for t in doc.tables:
        if not t.rows or len(t.rows[0].cells) < 2:
            continue
        head_l = t.rows[0].cells[0].text.strip().upper()
        head_r = t.rows[0].cells[1].text.strip().upper() if len(t.rows[0].cells) > 1 else ''
        # Price Schedule (header row contains 'ITEM DESCRIPTION')
        if 'ITEM DESCRIPTION' in head_r:
            targets.append(t)
        # Supervision sub-table (either VLPH 'Supervision Charges' or
        # our wrapped '{%tr if supervision_include %}' opener)
        elif t.rows[0].cells[0].text.strip().startswith('Supervision Charges'):
            targets.append(t)
        elif t.rows[0].cells[0].text.strip().startswith('{%tr if supervision_include'):
            targets.append(t)

    height = Cm(min_height_cm)
    for tbl in targets:
        for row in tbl.rows:
            first_cell_text = row.cells[0].text.strip() if row.cells else ''
            # Skip control rows so they collapse at render time.
            if first_cell_text.startswith('{%tr') or first_cell_text.startswith('{%p'):
                continue
            row.height = height
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def _rebuild_supervision_subtable(doc: Document) -> None:
    """Replace the Supervision sub-table with a single supervision
    row + Note row, wrapped with {%tr if supervision_include %}.

    Final layout (4 source rows, 2 rendered when supervision_include
    is true, 0 rendered when false):
      r0: {%tr if supervision_include %}
      r1: 'Supervision Charges for Erection and Commissioning'
          | '{{ supervision_rate }}'
      r2: Note row (merged across both columns)
      r3: {%tr endif %}
    """
    target = None
    for t in doc.tables:
        if not t.rows or not t.rows[0].cells:
            continue
        head = t.rows[0].cells[0].text.strip()
        # The VLPH original starts with 'Supervision Charges for
        # Erection' on row 0. Our previously-wrapped version starts
        # with '{%tr if supervision_include %}'. Match either so this
        # function is idempotent on a re-run.
        if head.startswith('Supervision Charges for Erection') or head.startswith('{%tr'):
            target = t
            break
    if target is None:
        print('Supervision sub-table not found — skipping rebuild.')
        return

    tbl_xml = target._element
    # Strip every existing row — we'll add a fresh 4-row layout.
    for tr in list(tbl_xml.findall(qn('w:tr'))):
        tbl_xml.remove(tr)

    NOTE_TEXT = ('Note: To-and-fro fare from Delhi to site, plus boarding, '
                 'lodging, local conveyance, and medical assistance if '
                 'required.')

    # r0: opener control row
    opener = target.add_row()
    opener.cells[0].text = '{%tr if supervision_include %}'

    # r1: supervision rate row
    rate_row = target.add_row()
    rate_row.cells[0].text = 'Supervision Charges for Erection and Commissioning'
    rate_row.cells[1].text = '{{ supervision_rate }}'

    # r2: note row (merged)
    note_row = target.add_row()
    note_row.cells[0].merge(note_row.cells[1])
    note_row.cells[0].text = NOTE_TEXT

    # r3: closer control row
    closer = target.add_row()
    closer.cells[0].text = '{%tr endif %}'


def _rebuild_price_schedule(doc: Document) -> None:
    """Find the 3-row Price Schedule (target by 'ITEM DESCRIPTION' header)
    and expand into single/full toggled rows + supervision + summary."""
    table = None
    for t in doc.tables:
        if not t.rows or len(t.rows[0].cells) < 2:
            continue
        if 'ITEM DESCRIPTION' in t.rows[0].cells[1].text.upper():
            table = t
            break
    if table is None:
        print('Price Schedule table not found — skipping rebuild.')
        return

    tbl_xml = table._element
    # Strip rows 1+ (keep the header 'S. No. | Item Description | Qty | Unit Price | Total Price')
    for tr in list(tbl_xml.findall(qn('w:tr')))[1:]:
        tbl_xml.remove(tr)

    def add(c1, c2, c3, c4, c5, *, bold=False):
        r = table.add_row()
        r.cells[0].text = c1
        r.cells[1].text = c2
        r.cells[2].text = c3
        r.cells[3].text = c4
        r.cells[4].text = c5
        if bold:
            for cell in r.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True

    # Layout matches the VLPH default Annexure III: header + item row(s)
    # + TOTAL row. Single mode = one item (default). Full mode = iterate
    # over bom_rows. No amount-in-words footer — the standalone
    # Supervision sub-table (Mechanical / PLC / Note) lives separately
    # in the template body just below this table.

    # ── style: single ─────────────────────────────────────────────────
    add("{%tr if price_schedule_style == 'single' %}", "", "", "", "")
    add("1.", "Recuperator",
        "{{ recup_qty }}", "{{ recup_unit_price }}", "{{ recup_total_price }}")
    add("", "TOTAL", "", "", "{{ grand_total }}", bold=True)
    add("{%tr endif %}", "", "", "", "")

    # ── style: full ────────────────────────────────────────────────────
    add("{%tr if price_schedule_style == 'full' %}", "", "", "", "")
    add("{%tr for r in bom_rows %}", "", "", "", "")
    add("{{ r.sno }}", "{{ r.item }}", "{{ r.qty }}",
        "{{ r.unit_price }}", "{{ r.total }}")
    add("{%tr endfor %}", "", "", "", "")
    add("", "TOTAL", "", "", "{{ grand_total }}", bold=True)
    add("{%tr endif %}", "", "", "", "")


def main() -> None:
    if not os.path.exists(SOURCE):
        raise SystemExit(f"missing: {SOURCE}")

    # Backup current target if it exists (caller already does this via the
    # _v2_backup file, but be defensive).
    if os.path.exists(TARGET) and not os.path.exists(TARGET + '.bak'):
        shutil.copy(TARGET, TARGET + '.bak')

    shutil.copy(SOURCE, TARGET)
    doc = Document(TARGET)
    print(f'Cloned {SOURCE} -> {TARGET}: '
          f'{len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables')

    # 1. Delete VLPH scope-of-supply body, replace with recup operation
    #    paragraphs + 3D image placeholders.
    _delete_vlph_scope_body(doc)

    # 2. Replace T5 (Tech Specs 24-row) with recup Designing Parameters.
    #    After the previous step, table indices haven't shifted (we only
    #    touched paragraphs), so doc.tables[5] is still the spec table.
    designing_table = doc.tables[5]
    _replace_spec_table_with_designing_params(designing_table)

    # 3. Insert Material of Construction sub-table BEFORE the
    #    Designing Parameters table.
    _insert_moc_table_before(designing_table)

    # 3b. Insert 3D image heading + placeholders between the MoC table
    #     and the Designing Parameters table (so the diagrams appear
    #     right before the client-given parameters).
    _insert_3d_images_before(designing_table)

    # 4. Replace Annexure I scope table content with recup items.
    #    After step 3, MoC took position 5 and Designing Params is now
    #    table 6, so the Scope of Supply table is at 7.
    # Locate by content to be robust:
    scope_table = None
    for t in doc.tables:
        if not t.rows or len(t.rows[0].cells) < 2:
            continue
        head = t.rows[0].cells[0].text.strip().upper()
        if head == 'S. NO.' and 'ITEM DESCRIPTION' not in t.rows[0].cells[1].text.upper():
            # candidate — the Annexure I Scope of Supply uses 'S. No.' header,
            # but so does the Reference List (50 rows) and Make List (1 row).
            # Distinguish by row count: scope is 12 rows.
            if 10 <= len(t.rows) <= 15:
                scope_table = t
                break
    if scope_table is not None:
        _replace_scope_table_with_recup(scope_table)
    else:
        print('Annexure I Scope of Supply table not found — leaving as-is.')

    # 5. Rebuild the Price Schedule (Annexure III) for single/full toggle.
    _rebuild_price_schedule(doc)

    # 5b. Rebuild Supervision sub-table to a single 'Supervision Charges
    #     for Erection and Commissioning' row + Note row, wrapped with
    #     {%tr if supervision_include %} for the Step-4 checkbox.
    _rebuild_supervision_subtable(doc)

    # 5c. Inject the RECUPERATOR descriptive paragraph into Annexure I.
    _inject_scope_of_supply_paragraph(doc)

    # 6. Replace Annexure V Reference List with recup-specific clients.
    _replace_reference_list(doc)

    # 7. Remove Annexure VI Make List (not applicable for recup).
    _remove_annexure_vi(doc)

    # 8. Give Price Schedule + Supervision rows more vertical room so
    #    the printed offer doesn't look cramped.
    _pad_table_rows(doc, min_height_cm=0.75)

    doc.save(TARGET)
    print(f'Saved -> {TARGET}')


if __name__ == '__main__':
    main()
