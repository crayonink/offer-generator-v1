"""One-off: take the Real Ispat reference offer and rewrite the hardcoded
values as Jinja placeholders, producing Recup_Offer_Template.docx.

Reuses existing placeholder names where possible ({{ enquiry_ref }},
{{ company_name }}, {{ tnc_* }}) so engine/quote_writer.py context for
recup can share most of the VLPH builder. Adds recup-specific names for
the Designing Parameters and Price Schedule tables.
"""
from __future__ import annotations

import os
import shutil

from docx import Document
from docx.oxml.ns import qn

SRC = "Encon offer Recuperator_Real Group_13.08.2025 (1) (1).docx"
DST = "Recup_Offer_Template.docx"


def _clear_runs(para):
    """Remove every <w:r> inside a paragraph. Keeps the paragraph
    properties (style, alignment) intact."""
    for r in list(para._element.findall(qn('w:r'))):
        para._element.remove(r)


def set_para_text(para, text):
    """Replace the paragraph's text with `text`, collapsed into ONE run.
    Single-run is required so Jinja {{ x }} markers survive docxtpl
    parsing (a placeholder split across runs is unreplaceable)."""
    # Reuse formatting from the first existing run if any
    rPr = None
    runs = para._element.findall(qn('w:r'))
    if runs:
        first_rPr = runs[0].find(qn('w:rPr'))
        if first_rPr is not None:
            from copy import deepcopy
            rPr = deepcopy(first_rPr)
    _clear_runs(para)
    new_run = para.add_run(text)
    if rPr is not None:
        # remove the auto-added rPr and inject the copy
        existing_rPr = new_run._element.find(qn('w:rPr'))
        if existing_rPr is not None:
            new_run._element.remove(existing_rPr)
        new_run._element.insert(0, rPr)


def set_cell_text(cell, text):
    """Replace the cell's text with `text` in its first paragraph;
    drop any other paragraphs in the cell."""
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    set_para_text(paras[0], text)
    # remove extra paragraphs
    for p in paras[1:]:
        p._element.getparent().remove(p._element)


def set_cell_multiline(cell, lines):
    """Replace cell content with multiple paragraphs (one per line)."""
    # Keep the first paragraph (preserve cell style), wipe its text and add lines.
    paras = cell.paragraphs
    if not paras:
        for ln in lines:
            cell.add_paragraph(ln)
        return
    first = paras[0]
    set_para_text(first, lines[0] if lines else "")
    # remove every other paragraph
    for p in paras[1:]:
        p._element.getparent().remove(p._element)
    # add the remaining lines as new paragraphs in the cell
    for ln in lines[1:]:
        cell.add_paragraph(ln)


def clear_para(para):
    _clear_runs(para)


def main():
    if not os.path.exists(SRC):
        raise SystemExit(f"Source missing: {SRC}")
    shutil.copy(SRC, DST)

    doc = Document(DST)
    paras = doc.paragraphs

    # ── Title/intro block (paragraphs 4–9) ─────────────────────────────
    set_para_text(paras[4], "PROJECT/EQUIPMENT\t:\t{{ project_name }}")
    # paras[5] = "NAME OF EQUIOMENT : Recuperator" — keep static (always recup)
    set_para_text(paras[6], "CLIENT\t:\t{{ company_name }}")
    set_para_text(paras[7], "ENQ. No./DATE\t:\t{{ client_enq_ref }}")
    set_para_text(paras[8], "OUR REF. /DATE\t:\t{{ enquiry_ref }}")
    clear_para(paras[9])  # DT.13/08/2025 — date is already in enquiry_ref

    # ── 'To,' address block (33–38) ─────────────────────────────────────
    set_para_text(paras[34], "{{ company_name }}")
    set_para_text(paras[35], "{{ company_address }}")
    set_para_text(paras[36], "E: {{ email }}")
    set_para_text(paras[37], "Phone: {{ mobile_no }}")
    set_para_text(paras[38], "Kind Attn. - {{ poc_name }}")

    # ── Subject + body letter ──────────────────────────────────────────
    set_para_text(paras[41], "Sub: {{ subject }}")
    set_para_text(
        paras[45],
        "We refer to your e-mail and specification sheet, the requirement "
        "of Recuperator for {{ application }}. ",
    )

    # ── Signature block (52–56) ─────────────────────────────────────────
    set_para_text(paras[52], "{{ marketing_person }}")
    # 53, 54 — keep (company name + address are static here)
    set_para_text(paras[55], "Mob {{ marketing_phone }}")
    set_para_text(paras[56], "Email: {{ marketing_email }} | www.encon.co.in ")

    # ── Terms & Conditions section (144–169) ────────────────────────────
    set_para_text(paras[144], "PRICES\t: {{ tnc_prices }}")
    set_para_text(paras[146], "DELIVERY\t: {{ tnc_delivery }}")
    clear_para(paras[147])  # was the 2nd line of delivery
    set_para_text(paras[149], "GST\t: {{ tnc_gst }}")
    set_para_text(paras[151], "HSN CODE\t: {{ tnc_hsn_code }}")
    set_para_text(paras[152], "PAN NO / GST NO\t: {{ tnc_pan_gst }}")
    set_para_text(paras[154], "TERMS OF PAYMENT\t: {{ tnc_payment_terms }}")
    clear_para(paras[155])
    set_para_text(paras[157], "PACKING \t\t\t: {{ tnc_packing_forwarding }}")
    set_para_text(paras[159], "FREIGHT CHARGES\t: {{ tnc_freight }}")
    set_para_text(paras[161], "TRANSIT INSURANCE\t: {{ tnc_transit_insurance }}")
    set_para_text(paras[163], "VALIDITY\t: {{ tnc_validity }}")
    set_para_text(paras[165], "INSPECTION\t: {{ tnc_inspection }}")
    set_para_text(paras[167], "GUARANTEE\t: {{ tnc_guarantee }}")
    clear_para(paras[168])
    clear_para(paras[169])

    # ── Tables ──────────────────────────────────────────────────────────
    tabs = doc.tables

    # Table 0: header ref row
    t0 = tabs[0]
    set_cell_text(t0.cell(1, 0), "{{ client_enq_ref }}")
    set_cell_text(t0.cell(1, 1), "{{ enquiry_ref_short }}")
    set_cell_text(t0.cell(1, 2), "{{ enquiry_date_str }}")

    # Table 3: client + marketing details (4x1, multi-paragraph cells)
    t3 = tabs[3]
    set_cell_multiline(t3.cell(1, 0), [
        "{{ company_name }}",
        "{{ company_address }}",
        "Kind Attn. - {{ poc_name }}",
        "E: {{ email }}",
        "Phone: {{ mobile_no }}",
    ])
    set_cell_multiline(t3.cell(3, 0), [
        "{{ marketing_person }}",
        "E. {{ marketing_email }}",
        "Mob {{ marketing_phone }}",
    ])

    # Table 4: Recuperator Designing Parameters
    t4 = tabs[4]
    # Column 1 is the value column. Column 2 is the unit (already correct).
    set_cell_text(t4.cell(1, 1), "{{ flue_flow_nm3hr }}")
    set_cell_text(t4.cell(2, 1), "{{ flue_temp_in_C }}")
    set_cell_text(t4.cell(3, 1), "{{ flue_temp_out_C }}")
    set_cell_text(t4.cell(4, 1), "{{ air_volume_nm3hr }}")
    set_cell_text(t4.cell(5, 1), "{{ air_temp_in_C }}")
    set_cell_text(t4.cell(6, 1), "{{ air_temp_out_C }}")
    set_cell_text(t4.cell(7, 1), "{{ surface_area_m2 }}")
    # r8 = "Hot Bank (CS Boiler Grade)" header — keep
    set_cell_text(t4.cell(9, 1),  "{{ pipe_dia_mm }}")
    set_cell_text(t4.cell(10, 1), "{{ pipe_length_m }}")
    set_cell_text(t4.cell(11, 1), "{{ pipe_thick_mm }}")
    # r12 = "Cold Bank" header — keep
    set_cell_text(t4.cell(13, 1), "{{ pipe_dia_mm }}")
    set_cell_text(t4.cell(14, 1), "{{ pipe_length_m }}")
    set_cell_text(t4.cell(15, 1), "{{ pipe_thick_mm }}")

    # Table 5: Price Schedule (3x5)
    t5 = tabs[5]
    # Row 1: 1. | <desc> | <qty> | <unit price> | <total>
    set_cell_text(t5.cell(1, 1), "Recuperator for {{ application }}")
    set_cell_text(t5.cell(1, 2), "{{ recup_qty }}")
    set_cell_text(t5.cell(1, 3), "{{ recup_unit_price }}")
    set_cell_text(t5.cell(1, 4), "{{ recup_total_price }}")
    # Row 2: amount in words across the merged cells + total
    set_cell_text(t5.cell(2, 0), "{{ recup_total_in_words }}")
    set_cell_text(t5.cell(2, 4), "{{ recup_total_price }}")

    doc.save(DST)
    print(f"Wrote {DST}")


if __name__ == "__main__":
    main()
