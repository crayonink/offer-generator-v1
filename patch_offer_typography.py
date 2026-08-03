"""
Force every offer template onto Calibri and exactly three sizes: 15, 12, 10.

The templates were already Calibri at run and document-default level, but a
handful of runs sat off the scale — the cover page's 22 pt letterhead, 18 pt
"OFFER FOR", 16 pt equipment name, the 13.5 pt award line, a stray 11 pt, and
runs with no explicit size at all. Those are snapped on:

    >= 14 pt  ->  15   (headings and cover display type)
    11-13.9   ->  12   (sub-headings, table headers)
    otherwise ->  10   (body copy, including unset)

The document default is pinned to Calibri 10 as well, so anything that picks up
a size by inheritance rather than an explicit run property lands on the scale.

Run:  python patch_offer_typography.py

Idempotent — a second run finds every size already on the scale.
"""

import os

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATES = [
    "Regen_Offer_Template.docx",
    "Regen_Oil_Offer_Template.docx",
    "Regen_Dual_Offer_Template.docx",
    "Offer_Template.docx",
]

FONT_NAME = "Calibri"
SIZE_HEADING, SIZE_SUB, SIZE_BODY = 15, 12, 10


def _target(pt):
    if pt is None:
        return SIZE_BODY
    if pt >= 14:
        return SIZE_HEADING
    if pt >= 11:
        return SIZE_SUB
    return SIZE_BODY


def _fix_run(run):
    """Calibri + a size on the scale. Theme attributes are dropped: w:asciiTheme
    beats w:ascii in Word, so a themed run would ignore the typeface we set."""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for key in list(rFonts.attrib):
        if key.split("}")[-1].lower().endswith("theme"):
            del rFonts.attrib[key]
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT_NAME)
    want = _target(run.font.size.pt if run.font.size else None)
    run.font.size = Pt(want)
    return want


def _pin_defaults(doc):
    """Calibri 10 as the document default, so inherited text lands on scale."""
    dd = doc.styles.element.find(qn("w:docDefaults"))
    if dd is None:
        return
    rprd = dd.find(qn("w:rPrDefault"))
    if rprd is None:
        return
    rPr = rprd.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        rprd.append(rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for key in list(rFonts.attrib):
        if key.split("}")[-1].lower().endswith("theme"):
            del rFonts.attrib[key]
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT_NAME)
    for tag, val in (("w:sz", str(SIZE_BODY * 2)), ("w:szCs", str(SIZE_BODY * 2))):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.append(el)
        el.set(qn("w:val"), val)


def _fix_container(container, counts):
    """Every run in a body / header / footer, tables included."""
    for p in container.paragraphs:
        for r in p.runs:
            counts[_fix_run(r)] += 1
    for tbl in container.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        counts[_fix_run(r)] += 1


def patch(path):
    d = Document(path)
    counts = {SIZE_HEADING: 0, SIZE_SUB: 0, SIZE_BODY: 0}
    _fix_container(d, counts)
    # Headers and footers carry the ref line and the address strip; they sat at
    # 8 pt, which is off the three-size scale.
    for sec in d.sections:
        for part in (sec.header, sec.footer,
                     sec.first_page_header, sec.first_page_footer,
                     sec.even_page_header, sec.even_page_footer):
            if part is not None:
                _fix_container(part, counts)
    # Styles too, so nothing inherits an off-scale size.
    for s in d.styles:
        el = getattr(s, "element", None)
        rPr = el.find(qn("w:rPr")) if el is not None else None
        if rPr is None:
            continue
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is not None:
            for key in list(rFonts.attrib):
                if key.split("}")[-1].lower().endswith("theme"):
                    del rFonts.attrib[key]
            for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                rFonts.set(qn(attr), FONT_NAME)
        sz = rPr.find(qn("w:sz"))
        if sz is not None:
            try:
                sz.set(qn("w:val"), str(_target(int(sz.get(qn("w:val"))) / 2) * 2))
            except (TypeError, ValueError):
                pass
    _pin_defaults(d)
    d.save(path)
    print(f"  OK  {os.path.basename(path)} — 15pt:{counts[SIZE_HEADING]} "
          f"12pt:{counts[SIZE_SUB]} 10pt:{counts[SIZE_BODY]}")


if __name__ == "__main__":
    for name in TEMPLATES:
        p = os.path.join(BASE_DIR, name)
        if os.path.exists(p):
            patch(p)
        else:
            print(f"  SKIP {name} — not found")
