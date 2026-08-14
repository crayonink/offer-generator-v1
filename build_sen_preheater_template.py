"""
build_sen_preheater_template.py — v4
Builds SEN_Preheater_Offer_Template.docx from Burner_Offer_Template.docx.
Cleanly formats tables, TOC, Technical Specifications, Scope Narrative, and Price Schedule.
"""
from __future__ import annotations
import os, shutil
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor

BASE   = os.path.dirname(os.path.abspath(__file__))
SOURCE = os.path.join(BASE, "Burner_Offer_Template.docx")
TARGET = os.path.join(BASE, "SEN_Preheater_Offer_Template.docx")

shutil.copy2(SOURCE, TARGET)
doc = Document(TARGET)

# ── Helper ────────────────────────────────────────────────────────────────────
def _set_cell(cell, text, bold=False, size=9, bg=None, color=None, align=None):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Arial"
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if bg:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), bg)
        tcPr.append(shd)

def _new_para(text, bold=False, underline=False, size=9.5, italic=False, color=None):
    """Create a bare w:p element with font, size, spacing, and styling."""
    p_el = OxmlElement("w:p")
    pPr  = OxmlElement("w:pPr")
    sp   = OxmlElement("w:spacing")
    sp.set(qn("w:after"), "80")  # 80 dxa = 4pt
    pPr.append(sp)
    
    rPr  = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rPr.append(rFonts)
    if bold:
        rPr.append(OxmlElement("w:b"))
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    if italic:
        rPr.append(OxmlElement("w:i"))
    if color:
        c_el = OxmlElement("w:color")
        c_el.set(qn("w:val"), color)
        rPr.append(c_el)
    
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(size * 2)))
    rPr.append(szCs)
    
    p_el.append(pPr)
    r = OxmlElement("w:r")
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    if text.startswith(" ") or text.endswith(" "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    p_el.append(r)
    return p_el

# ── TABLE 1: Table of Contents ────────────────────────────────────────────────
toc = doc.tables[1]
for tr in list(toc._element.findall(qn("w:tr")))[4:]:
    toc._element.remove(tr)
toc_data = [
    ("1.", "COMPANY PROFILE"),
    ("2.", "ABOUT THE CLIENT"),
    ("3.", "TECHNICAL SPECIFICATIONS"),
    ("4.", "LIST OF ANNEXURES"),
]
for i, (a, b) in enumerate(toc_data):
    if i < len(toc.rows):
        _set_cell(toc.rows[i].cells[0], a, bold=True)
        _set_cell(toc.rows[i].cells[1], b, bold=True)

# ── TABLE 2: Annexures ─────────────────────────────────────────────────────────
ann = doc.tables[2]
ann_data = [
    ("Annexure - I",   "SCOPE OF SUPPLY",   "Attached"),
    ("Annexure \u2013 II",  "PRICE SCHEDULE",    "Attached"),
    ("Annexure - III", "TERMS & CONDITIONS", "Attached"),
    ("Annexure - IV",  "REFERENCE LIST",     "Attached"),
]
for i, (a, b, c_) in enumerate(ann_data):
    if i + 1 < len(ann.rows):
        _set_cell(ann.rows[i+1].cells[0], a)
        _set_cell(ann.rows[i+1].cells[1], b)
        _set_cell(ann.rows[i+1].cells[2], c_)
    else:
        nr = ann.add_row()
        _set_cell(nr.cells[0], a); _set_cell(nr.cells[1], b); _set_cell(nr.cells[2], c_)

# ── TABLE 5: Tech Spec ─────────────────────────────────────────────────────────
spec = doc.tables[5]
SEN_ROWS = [
    ("Heating Medium",                  "{{ fuel_name }} Fired Burners"),
    ("Quantity",                        "{{ sen_quantity }}"),
    ("Minimum Temperature",             "Ambient to 1200\u00B0C"),
    ("Heating Time",                    "{{ heating_time }}"),
    ("{{ gas_line }} Burner Capacity",  "{{ burner_capacity_sen }}"),
    ("No. of Burners",                  "{{ num_burners_sen }}"),
    ("Ignition of Burner",              "Manual"),
]
# Header - Merge cells
if spec.rows:
    hdr = spec.rows[0]
    hdr.cells[0].merge(hdr.cells[1])
    _set_cell(hdr.cells[0], "TECHNICAL SPECIFICATIONS: SEN PREHEATER", bold=True, size=10, bg="1A3A5C", color=RGBColor(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER)

# Data rows
existing = list(spec.rows[1:])
for i, (lbl, val) in enumerate(SEN_ROWS):
    bg_color = "F8FAFC" if i % 2 == 1 else "FFFFFF"
    if i < len(existing):
        _set_cell(existing[i].cells[0], lbl, bold=True, size=9, bg="EEF2F7")
        _set_cell(existing[i].cells[1], val, size=9, bg=bg_color)
    else:
        nr = spec.add_row()
        _set_cell(nr.cells[0], lbl, bold=True, size=9, bg="EEF2F7")
        _set_cell(nr.cells[1], val, size=9, bg=bg_color)

# Remove excess rows
for row in existing[len(SEN_ROWS):]:
    row._element.getparent().remove(row._element)

# ── TABLE 6: Price Schedule ────────────────────────────────────────────────────
def _rebuild_price_schedule(table):
    tbl_xml = table._element
    for tr in list(tbl_xml.findall(qn("w:tr")))[1:]:
        tbl_xml.remove(tr)

    def add(c1, c2, c3, c4, c5, *, bold=False, bg=None):
        r = table.add_row()
        vals = (c1, c2, c3, c4, c5)
        aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT]
        for idx, (cell, txt) in enumerate(zip(r.cells, vals)):
            _set_cell(cell, txt, bold=bold, size=9, bg=bg, align=aligns[idx])

    # Header styling
    _set_cell(table.rows[0].cells[0], "S. No.", bold=True, size=9.5, bg="1A3A5C", color=RGBColor(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(table.rows[0].cells[1], "Item Description", bold=True, size=9.5, bg="1A3A5C", color=RGBColor(255,255,255), align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_cell(table.rows[0].cells[2], "Qty", bold=True, size=9.5, bg="1A3A5C", color=RGBColor(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(table.rows[0].cells[3], "Unit Price (INR)", bold=True, size=9.5, bg="1A3A5C", color=RGBColor(255,255,255), align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(table.rows[0].cells[4], "Total Price (INR)", bold=True, size=9.5, bg="1A3A5C", color=RGBColor(255,255,255), align=WD_ALIGN_PARAGRAPH.RIGHT)

    add("1.", "SEN Preheater ({{ gas_line }} Line)", "{{ item_qty }}", "{{ unit_price }}", "{{ total_price }}")
    add("", "TOTAL", "", "", "{{ grand_total }}", bold=True, bg="EEF2F7")

    col_widths = [Cm(1.2), Cm(6.0), Cm(2.0), Cm(3.5), Cm(4.0)]
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

_rebuild_price_schedule(doc.tables[6])
p_words = _new_para("(Rupees {{ total_in_words }} Only)", bold=True, size=9.5, italic=True)
doc.tables[6]._element.addnext(p_words)

# ── Paragraphs: replace scope section ─────────────────────────────────────────
paras = list(doc.paragraphs)
scope_start = None
ann2_idx    = None
for i, p in enumerate(paras):
    if "scope_intro" in p.text and scope_start is None:
        scope_start = i
    if "ANNEXURE II" in p.text.upper() and "PRICE" in p.text.upper():
        ann2_idx = i
        break

if scope_start is not None and ann2_idx is not None:
    body = doc.element.body
    to_remove = [p._element for p in paras[scope_start:ann2_idx]]
    ann2_el = paras[ann2_idx]._element

    new_paras = [
        # Scope intro
        _new_para("{{ pipeline_scope_text }}", size=9.5),
        _new_para(""),
        # OBJECTIVE
        _new_para("OBJECTIVE", bold=True, underline=True, size=10, color="1A3A5C"),
        _new_para("{{ sen_objective }}", size=9.5),
        _new_para(""),
        # SCOPE OF SUPPLY
        _new_para("SCOPE OF SUPPLY", bold=True, underline=True, size=10, color="1A3A5C"),
        _new_para("Our scope of supply will cover design, engineering, manufacture supply, supervision of commissioning & erection of the following main components:", size=9.5),
        _new_para(""),
        # ENCON BURNERS
        _new_para("{{ sen_burners_heading }}", bold=True, underline=True, size=9.5, color="1A3A5C"),
        _new_para("{{ sen_burners_body }}", size=9.5),
        _new_para(""),
        # GAS LINE
        _new_para("{{ gas_line }} LINE FOR BURNERS", bold=True, underline=True, size=9.5, color="1A3A5C"),
        _new_para("The {{ gas_line }} line for main burners shall be routed from the main gas train and will consist of the following main components:", size=9.5),
        # gas line items loop
        _new_para("{%p for x in fuel1_line_items %}"),
        _new_para("\u2022  {{ x.item }}", size=9.5),
        _new_para("{%p endfor %}"),
        _new_para(""),
        # AIR LINE
        _new_para("AIR LINE FOR MAIN BURNERS:", bold=True, underline=True, size=9.5, color="1A3A5C"),
        # air line items loop
        _new_para("{%p for x in air_pipeline_items %}"),
        _new_para("\u2022  {{ x.item }}", size=9.5),
        _new_para("{%p endfor %}"),
        _new_para(""),
        # TROLLEY
        _new_para("TROLLEY", bold=True, underline=True, size=9.5, color="1A3A5C"),
        _new_para("{{ sen_trolley_text }}", size=9.5),
        _new_para(""),
        # ROLLER RACK
        _new_para("ROLLER SUPPORTED RACK WITH HANDLE", bold=True, underline=True, size=9.5, color="1A3A5C"),
        _new_para("{{ sen_roller_rack_text }}", size=9.5),
        _new_para(""),
    ]

    for np_el in new_paras:
        ann2_el.addprevious(np_el)

    for el in to_remove:
        try:
            el.getparent().remove(el)
        except Exception:
            pass

# ── Rename section headings ────────────────────────────────────────────────────
for p in doc.paragraphs:
    txt = p.text.strip()
    if txt == "ENCON BURNER":
        for run in p.runs:
            run.text = ""
        p.add_run("SEN PREHEATER")
    elif "ANNEXURE I" in txt.upper() and ("SCOPE" in txt.upper() or "EXCLUSION" in txt.upper() or "BURNER" in txt.upper()):
        for run in p.runs:
            run.text = ""
        p.add_run("ANNEXURE I \u2014 SCOPE OF SUPPLY: SEN PREHEATER")

# ── Save ───────────────────────────────────────────────────────────────────────
doc.save(TARGET)
print(f"Saved: {TARGET}")

# ── Verify ─────────────────────────────────────────────────────────────────────
from docxtpl import DocxTemplate
tpl = DocxTemplate(TARGET)
try:
    v = sorted(tpl.get_undeclared_template_variables())
    print(f"Variables ({len(v)}): {v}")
except Exception as e:
    print(f"Jinja error: {e}")
