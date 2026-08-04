"""
Give the VLPH-derived offers the same layout treatment as the regen offer.

Covers the vertical and the recuperator — the recup template was built from the
vertical one (see build_recup_template_from_vlph.py), so both carry the same
heading hierarchy and this keys off the built-in styles rather than the font
signature the regen patch has to detect:

    Heading 1  (16 pt navy)   TABLE OF CONTENTS, 1. COMPANY PROFILE,
                              2. ABOUT THE CLIENT, 3. TECHNICAL SPECIFICATIONS,
                              ANNEXURE I … VI
    Heading 2  (13 pt orange) ABOUT US, SCOPE OF SUPPLY, …
    Heading 3  (11 pt navy)   CLIENT DETAILS, LADLE HOOD, …

What it does, matching the regen offer:

0. Drop the "COVER LETTER" heading — it sat directly above "To," and added
   nothing; nothing in the contents list refers to it.
1. Page break before "To," — the cover block and the letter ran together, so
   "To," sat alone at the foot of page 1 with the address overleaf.
2. Page break before every Heading 1, so each numbered section and annexure
   opens its own page.
3. Sections bound with keep-with-next so one that will not fit starts on the
   next page instead of splitting. Skipped for sections too tall to fit a page
   anyway (INSTRUMENTATION & CONTROL and friends), where forcing it would just
   waste a page before splitting regardless.
4. cantSplit on every table row, so no table breaks mid-row.
5. The whole recipient block and the whole signature block bold.
6. Calibri throughout on three sizes: 15 pt Heading 1, 12 pt Heading 2/3 and
   table headers, 10 pt body copy. Cover-page display type keeps its sizes.

Run:  python patch_vlph_layout.py

Idempotent.
"""

import os

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATES = ["Offer_Template.docx", "Recup_Offer_Template.docx"]

LETTER_START = "To,"
LETTER_END = "NATIONAL ENERGY EFFICIENCY INNOVATION AWARD"
LETTER_BOLD = [
    "To,", "{{company_name}}", "Kind Attn:", "{{company_address}}",
    "Contact No.:", "Email: {{email}}",
    "Regards,", "{{technical_person}}", "ENCON Thermal Engineers Pvt. Ltd.",
    "53 K.M Stone", "Mob: {{technical_phone}}", "Email: {{technical_email}}",
    "Website:",
]

FONT_NAME = "Calibri"
SIZE_H1 = 15
SIZE_SUB = 12
SIZE_BODY = 10
# A section longer than this cannot fit a page, so binding it would only push a
# blank page ahead of a split that happens anyway.
MAX_BOUND_BLOCKS = 16


def _style_of(p):
    return (p.style.name or "") if p.style is not None else ""


def _heading_level(p):
    s = _style_of(p)
    if s.startswith("Heading "):
        try:
            return int(s.split()[1])
        except (IndexError, ValueError):
            return None
    return None


def _set_page_break_before(p):
    pPr = p._element.get_or_add_pPr()
    for old in pPr.findall(qn("w:pageBreakBefore")):
        pPr.remove(old)
    pPr.append(OxmlElement("w:pageBreakBefore"))


def _no_split_rows(table):
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        if trPr.find(qn("w:cantSplit")) is None:
            trPr.append(OxmlElement("w:cantSplit"))


def _set_font(run, size_pt=None):
    run.font.name = FONT_NAME
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT_NAME)
    if size_pt is not None:
        run.font.size = Pt(size_pt)


COVER_HEADING = "COVER LETTER"


def patch(path):
    d = Document(path)

    # ── 0. drop the redundant COVER LETTER heading ───────────────────────────
    dropped = 0
    for p in list(d.paragraphs):
        if p.text.strip().upper() == COVER_HEADING:
            p._element.getparent().remove(p._element)
            dropped += 1

    by_el = {p._element: p for p in d.paragraphs}

    # ── 1 + 2. page breaks ───────────────────────────────────────────────────
    breaks = 0
    for p in d.paragraphs:
        if p.text.strip() == LETTER_START or _heading_level(p) == 1:
            _set_page_break_before(p)
            breaks += 1

    # ── 3. bind each section; 4. no mid-row table splits ─────────────────────
    groups, current = [], None
    for el in d.element.body:
        p = by_el.get(el)
        if p is not None and _heading_level(p):
            if current:
                groups.append(current)
            current = [el]
        elif current is not None and (el.tag == qn("w:tbl") or p is not None):
            current.append(el)
    if current:
        groups.append(current)

    bound = 0
    for g in groups:
        if len(g) > MAX_BOUND_BLOCKS:
            continue                       # taller than a page — let it flow
        bound += 1
        for i, el in enumerate(g):
            last = i == len(g) - 1 and i != 0     # a lone heading still binds
            if el.tag == qn("w:tbl"):
                t = Table(el, d)
                for ri, row in enumerate(t.rows):
                    if last and ri == len(t.rows) - 1:
                        continue
                    for cell in row.cells:
                        for cp in cell.paragraphs:
                            cp.paragraph_format.keep_with_next = True
            else:
                par = Paragraph(el, d)
                par.paragraph_format.keep_together = True
                if not last:
                    par.paragraph_format.keep_with_next = True
    for t in d.tables:
        _no_split_rows(t)

    # A list item that has deeper items under it is a heading for them, so it
    # keeps-with-next — "GAS TRAIN" was stranding at the foot of a page with
    # its sub-bullets overleaf. Applies inside the long sections too, which are
    # not bound as a whole.
    def _list_level(p):
        pPr = p._element.find(qn("w:pPr"))
        if pPr is None:
            return None
        numPr = pPr.find(qn("w:numPr"))
        if numPr is None:
            return None
        ilvl = numPr.find(qn("w:ilvl"))
        return int(ilvl.get(qn("w:val"))) if ilvl is not None else 0

    paras = d.paragraphs
    parents = 0
    for i, p in enumerate(paras[:-1]):
        lvl = _list_level(p)
        nxt = _list_level(paras[i + 1])
        # a bullet with a deeper bullet under it, or the lead-in line above a list
        if (lvl is not None and nxt is not None and nxt > lvl) or \
           (lvl is None and nxt is not None and p.text.strip()):
            p.paragraph_format.keep_with_next = True
            parents += 1

    # ── 5. cover-letter emphasis ─────────────────────────────────────────────
    bolded = 0
    for p in d.paragraphs:
        t = p.text.strip()
        if t.startswith(LETTER_END):
            break
        if any(frag in t for frag in LETTER_BOLD):
            for run in p.runs:
                run.font.bold = True
            bolded += 1

    # ── 6. typography ────────────────────────────────────────────────────────
    retyped = 0
    started = False
    for p in d.paragraphs:
        t = p.text.strip()
        if t == LETTER_START:
            started = True
        lvl = _heading_level(p)
        if not started or t.startswith(LETTER_END):
            size = None                    # cover display type keeps its size
        elif lvl == 1:
            size = SIZE_H1
        elif lvl in (2, 3):
            size = SIZE_SUB
        else:
            size = SIZE_BODY
        for r in p.runs:
            _set_font(r, size_pt=size)
            retyped += 1
    for tbl in d.tables:
        for ri, row in enumerate(tbl.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        _set_font(r, size_pt=SIZE_SUB if ri == 0 else SIZE_BODY)
                        retyped += 1

    d.save(path)
    print(f"  OK  {os.path.basename(path)} — {breaks} page breaks; "
          f"{bound}/{len(groups)} sections bound; {bolded} letter lines bold; "
          f"{retyped} runs re-typed; {parents} list parents kept with children; "
          f"{dropped} cover heading(s) dropped")


if __name__ == "__main__":
    for name in TEMPLATES:
        patch(os.path.join(BASE_DIR, name))
