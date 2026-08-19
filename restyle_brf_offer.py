"""Give the BRF offer the vertical offer's typography.

The content was right and the document looked nothing like ours. It came out of
a Word file set in Times New Roman on US Letter with quarter-inch margins and
one flat black heading level; the vertical offer is Calibri on A4 with proper
margins, a two-level heading hierarchy and ENCON's orange.

What the vertical offer does, and what this copies:

    body          Calibri 10pt, colour inherited
    Heading 1     Calibri 15pt bold, 1F2937   the major sections
    Heading 2     Calibri 12pt bold, E87722   everything under them
    page          A4, margins 0.75 / 0.75 / 1.00 / 1.00

Which BRF heading is major and which is not is a judgement, so it is written
down in MAJOR rather than guessed from the text at run time.

Run:  python restyle_brf_offer.py
Idempotent — running twice changes nothing the second time.
"""
import os
import re
import shutil
import sys

from docx import Document
from docx.shared import Inches, Pt, RGBColor

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DEFAULT_TARGET = os.path.join(BASE_DIR, "BRF_Offer_Template.docx")

FONT = "Calibri"
SLATE = RGBColor(0x1F, 0x29, 0x37)
ORANGE = RGBColor(0xE8, 0x77, 0x22)

BODY_PT = 10.0
H1_PT = 15.0
H2_PT = 12.0

PAGE_W, PAGE_H = Inches(8.27), Inches(11.69)          # A4
MARGINS = dict(left=Inches(0.75), right=Inches(0.75),
               top=Inches(1.00), bottom=Inches(1.00))

# The sections that carry their own weight. Everything else that is currently
# a heading becomes a subheading under one of these.
MAJOR = {
    "SCOPE OF SUPPLY",
    "REFRACTORY LINING",
    "MATERIAL HANDLING EQUIPMENT HYDRAULIC PUSHER",
    "PRICE SCHEDULE",
    "SUPERVISION CHARGES FOR ERECTION & COMMISSIONING",
    "EXCLUSION",
    "TERMS AND CONDITIONS",
}
# Letterhead lines that are styled as headings but are not headings at all.
NOT_A_HEADING = re.compile(
    r"^(297, Sector|E-mail:|Tel:|ENCON YOUR ANSWER)", re.I)


def _is_major(text):
    t = re.sub(r"\s+", " ", text).strip().upper()
    if t in MAJOR:
        return True
    # "Basic Parameters For {{ capacity }} Ton/Hr..." opens the technical part
    return t.startswith("BASIC PARAMETERS")


def restyle_run(run, size_pt, colour=None, bold=None):
    run.font.name = FONT
    # East-Asian and complex-script names have to be set too, or Word keeps
    # the old face for anything it does not consider plain Latin.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts")
    if rfonts is not None:
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
                + attr, FONT)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if colour is not None:
        run.font.color.rgb = colour
    if bold is not None:
        run.bold = bold


def main():
    target = sys.argv[1] if len(sys.argv) > 1 else DEFAULT_TARGET
    if not os.path.exists(target):
        print(f"not found: {target}")
        return 1
    shutil.copyfile(target, target + ".bak")
    doc = Document(target)

    for sec in doc.sections:
        sec.page_width, sec.page_height = PAGE_W, PAGE_H
        sec.left_margin = MARGINS["left"]
        sec.right_margin = MARGINS["right"]
        sec.top_margin = MARGINS["top"]
        sec.bottom_margin = MARGINS["bottom"]

    h1 = h2 = body = 0
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        style = p.style.name if p.style is not None else ""
        heading = style.startswith("Heading") and text and not NOT_A_HEADING.match(text)

        if heading and _is_major(text):
            for r in p.runs:
                restyle_run(r, H1_PT, SLATE, True)
            h1 += 1
        elif heading:
            for r in p.runs:
                restyle_run(r, H2_PT, ORANGE, True)
            h2 += 1
        else:
            for r in p.runs:
                # A run already sized for emphasis keeps its size; only the
                # face and the flat black are brought into line.
                size = r.font.size.pt if r.font.size else BODY_PT
                if size < BODY_PT:
                    size = BODY_PT
                restyle_run(r, size)
                try:
                    if r.font.color is not None and r.font.color.rgb is not None \
                            and str(r.font.color.rgb) == "000000":
                        r.font.color.rgb = SLATE
                except Exception:
                    pass
            body += 1

    # The cover page. The vertical offer sets the company name at 15pt in
    # ENCON orange; this one had it at 27pt in a different orange, with the
    # tagline in red and the phone line larger than the address above it.
    COVER_FIX = {
        "ENCON Thermal Engineers (P) Ltd": (H1_PT, ORANGE, True),
        "ENERGY – CONSERVATION": (H2_PT, ORANGE, True),
        "Techno-Commercial": (H1_PT, SLATE, True),
        "Tel:": (BODY_PT, None, None),
    }
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        for needle, (size, colour, bold) in COVER_FIX.items():
            if needle in text:
                for r in p.runs:
                    restyle_run(r, size, colour, bold)
                break

    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        restyle_run(r, r.font.size.pt if r.font.size else BODY_PT)

    doc.save(target)
    print(f"{os.path.basename(target)}: A4, {FONT}; "
          f"{h1} major headings, {h2} subheadings, {body} body paragraphs")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
