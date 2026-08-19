"""Turn the Suraj Products offer into BRF_Offer_Template.docx.

The content is already right and already ENCON's; what it lacks is the seams.
Every figure the app now computes becomes a placeholder, and the price table
becomes one the writer can grow a row at a time.

Run once. Idempotent in the sense that it always rebuilds from the source
document, so re-running after editing the source picks the edits up.
"""
import os
import re
import shutil

from docx import Document

BASE = r"c:\Users\RUPA\Downloads\offer_generator_structure"
SRC = os.path.join(BASE, "Encon offer BRF 60TPH_Suraj Products Ltd.30.07.2026.docx")
OUT = os.path.join(BASE, "BRF_Offer_Template.docx")

# Longest first, so "12,300 mm" is not eaten by a shorter pattern.
SUBS = [
    # ── the parties ────────────────────────────────────────────────
    ("Suraj Products Ltd.", "{{ company_name }}"),
    ("Rajganpur,", "{{ company_address }}"),
    ("Odisha, India,", "{{ company_city_state }}"),
    ("Mr. Yogesh Dalmia", "{{ poc_name }}"),
    ("dalmiayk@gmail.com", "{{ email }}"),
    ("+91 783072787", "{{ mobile_no }}"),
    # ── the references ─────────────────────────────────────────────
    ("Mail enquiry 20/07/2026", "{{ your_ref }}"),
    ("E-Mail dt. 20.07.2026", "{{ your_ref }}"),
    ("ENCON.04026.192 /FBD/JR", "{{ ref_no }}"),
    ("ENCON-04026.192/FBD/JR", "{{ ref_no }}"),
    ("Dated -30/07/2026", "Dated - {{ quote_date }}"),
    ("30/07/2026", "{{ quote_date }}"),
    # ── who signs it ───────────────────────────────────────────────
    ("Jyotirmay Rabha", "{{ marketing_person }}"),
    ("+91 8099702129", "{{ marketing_phone }}"),
    ("west@encon.co.in", "{{ marketing_email }}"),
    # ── the furnace, all of it computed ────────────────────────────
    ("60Ton/Hr.", "{{ capacity }} Ton/Hr."),
    ("60Tons/Hr. Pusher type", "{{ capacity }} Tons/Hr. Pusher type"),
    ("60 Tons per hour", "{{ capacity }} Tons per hour"),
    ("60Tons/hr", "{{ capacity }} Tons/hr"),
    ("Cap: 60Tons/hr", "Cap: {{ capacity }} Tons/hr"),
    ("130mm2 x 12000mm long , 150mm2 x 12000mm long", "{{ billet_size }}"),
    ("17,000 mm", "{{ eff_length_mm }} mm"),
    ("12,300 mm", "{{ eff_width_mm }} mm"),
    ("20,800 mm", "{{ overall_length_mm }} mm"),
    ("13,800 mm", "{{ overall_width_mm }} mm"),
    ("8,600 kcal/ Nm3", "{{ cv }} kcal/Nm\u00b3"),
    ("3 nos 100 HP, 40\" w.g. (2W+1S)", "{{ blower_desc }}"),
    ("Flue Gaas: 31000 Nm3/hr.", "Flue gas: {{ flue_gas_nm3hr }} Nm\u00b3/hr"),
    ("Five zones, Two-Soaking Zone; Three Heating Zone",
     "{{ zone_count }} zones, {{ zone_split }}"),
]

# The per-zone burner counts, in the order the document names them.
ZONE_HEADINGS = ["SOAKING ZONE-1", "SOAKING ZONE-2",
                 "HEATING ZONE-1", "HEATING ZONE-2", "HEATING ZONE-3"]


def _replace_in_paragraph(p, old, new):
    """Replace across runs, keeping the first run's formatting."""
    full = "".join(r.text for r in p.runs)
    if old not in full:
        return False
    full = full.replace(old, new)
    for r in p.runs[1:]:
        r.text = ""
    if p.runs:
        p.runs[0].text = full
    return True


# The two email addresses are hyperlinks. python-docx's paragraph.runs does not
# return runs inside a hyperlink, so those are replaced in the package XML.
HYPERLINK_SUBS = [("dalmiayk@gmail.com", "{{ email }}"),
                  ("west@encon.co.in", "{{ marketing_email }}")]


def _templatise_hyperlinks(path):
    import zipfile
    tmp = path + ".tmp"
    zin = zipfile.ZipFile(path, "r")
    zout = zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED)
    for item in zin.infolist():
        data = zin.read(item.filename)
        if item.filename.endswith(".xml"):
            txt = data.decode("utf-8")
            for old, new in HYPERLINK_SUBS:
                txt = txt.replace(old, new)
            data = txt.encode("utf-8")
        zout.writestr(item, data)
    zin.close()
    zout.close()
    shutil.move(tmp, path)


def main():
    shutil.copyfile(SRC, OUT)
    doc = Document(OUT)

    hits = {}
    targets = list(doc.paragraphs)
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                targets.extend(cell.paragraphs)

    for old, new in SUBS:
        for p in targets:
            if _replace_in_paragraph(p, old, new):
                hits[old] = hits.get(old, 0) + 1

    # The burner count on each zone heading's following line.
    zone_i = 0
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip().upper()
        if t in ZONE_HEADINGS and zone_i < len(ZONE_HEADINGS):
            zone_i += 1
            for q in doc.paragraphs[i + 1:i + 3]:
                if "ENCON GAS BURNERS" in (q.text or "").upper():
                    _replace_in_paragraph(
                        q, q.text.strip(),
                        f"ENCON GAS BURNERS: {{{{ zone{zone_i}_burners }}}} Nos.")
                    hits[f"zone{zone_i}_burners"] = 1
                    break

    doc.save(OUT)
    _templatise_hyperlinks(OUT)

    print(f"wrote {os.path.basename(OUT)}")
    missed = [o for o, _ in SUBS if o not in hits]
    print(f"  {len(hits)} placeholders inserted, {len(missed)} patterns not found")
    for m in missed:
        print(f"    not found: {m[:60]}")


if __name__ == "__main__":
    main()
