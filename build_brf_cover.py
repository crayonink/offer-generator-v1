"""Rebuild the BRF offer's cover page in the vertical offer's shape.

What came out of the source document was a left-aligned list of tab-separated
labels under a bare "Techno-Commercial / OFFER / FOR" — no logo colour, no
hierarchy, the company name buried at the bottom in the same weight as an
address line.

The vertical offer opens differently, and this copies it: a centred title
block that names the equipment and the company in ENCON orange, the tagline
and contact lines beneath it in grey, and then the project details as a
two-column table rather than tabs that drift when a value gets longer.

    OFFER FOR                                15pt bold  1F2937
    {{ equipment_name }}                     15pt bold  E87722
    ENCON Thermal Engineers (P) Ltd          15pt bold  E87722
    Your Answer to the Continued Need...     12pt       6B7280
    address / telephone                      10pt       6B7280
    sales@encon.co.in  •  www.encon.co.in    10pt       E87722

    +-----------------------+---------------------------+
    | Project / Equipment   | {{ equipment_name }}      |
    | Project Name          | {{ project_name }}        |
    | Client                | {{ company_name }}        |
    | Enquiry No. / Date    | {{ your_ref }}            |
    | Our Ref. / Date       | {{ ref_no }}              |
    +-----------------------+---------------------------+

Run:  python build_brf_cover.py
Run it after build_brf_offer_template.py and before restyle_brf_offer.py.
"""
import os
import shutil
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DEFAULT_TARGET = os.path.join(BASE_DIR, "BRF_Offer_Template.docx")

FONT = "Calibri"
SLATE = RGBColor(0x1F, 0x29, 0x37)
ORANGE = RGBColor(0xE8, 0x77, 0x22)
GREY = RGBColor(0x6B, 0x72, 0x80)
RULE = "D1D5DB"

# The centred block: text, size, colour, bold
TITLE_BLOCK = [
    ("OFFER FOR", 15, SLATE, True),
    ("{{ equipment_name }}", 15, ORANGE, True),
    ("", 10, None, None),
    ("ENCON Thermal Engineers (P) Ltd", 15, ORANGE, True),
    ("Your Answer to the Continued Need for Energy Conservation", 12, GREY, False),
    ("297, Sector-21B  \u2022  Faridabad \u2013 121 001  \u2022  Haryana, India", 10, GREY, False),
    ("Tel: +91 (129) 4047847, 2439458", 10, GREY, False),
    ("sales@encon.co.in   \u2022   www.encon.co.in", 10, ORANGE, False),
    ("", 10, None, None),
]

DETAILS = [
    ("Project / Equipment", "{{ equipment_name }}"),
    ("Project Name", "{{ project_name }}"),
    ("Client", "{{ company_name }}"),
    ("Enquiry No. / Date", "{{ your_ref }}"),
    ("Our Ref. / Date", "{{ ref_no }}"),
]

# The cover runs from the "Techno-Commercial" line to the tagline, inclusive.
FIRST = "Techno-Commercial"
LAST = "ENCON YOUR ANSWER TO THE CONTINUED NEED"


def _style(run, size, colour, bold):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rf)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rf.set(qn("w:" + attr), FONT)
    run.font.size = Pt(size)
    if colour is not None:
        run.font.color.rgb = colour
    if bold is not None:
        run.bold = bold


def _borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    b = tcPr.makeelement(qn("w:tcBorders"), {})
    for edge in ("top", "left", "bottom", "right"):
        el = b.makeelement(qn("w:" + edge), {})
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), RULE)
        b.append(el)
    tcPr.append(b)


def main():
    target = sys.argv[1] if len(sys.argv) > 1 else DEFAULT_TARGET
    if not os.path.exists(target):
        print("not found: " + target)
        return 1
    doc = Document(target)
    paras = doc.paragraphs

    start = end = None
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if start is None and t.startswith(FIRST):
            start = i
        if start is not None and LAST in t.upper():
            end = i
            break
    if start is None or end is None:
        print("cover already rebuilt, or its bounds have moved")
        return 0

    shutil.copyfile(target, target + ".bak")
    anchor = paras[end]._p                      # build in front of the tagline

    # One of the paragraphs about to be deleted ends the cover's own section.
    # Losing it would merge the cover into the covering letter, and the letter
    # would inherit the cover's deliberately blank header — which is precisely
    # the fault the recuperator header patch was written for. Keep it and hang
    # it on the paragraph that closes the new cover.
    carried_sect = None
    for p in paras[start:end + 1]:
        found = p._p.find(qn("w:pPr") + "/" + qn("w:sectPr"))
        if found is None:
            pPr = p._p.find(qn("w:pPr"))
            found = pPr.find(qn("w:sectPr")) if pPr is not None else None
        if found is not None:
            carried_sect = found
            break

    for text, size, colour, bold in TITLE_BLOCK:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _style(p.add_run(text), size, colour, bold)
        anchor.addprevious(p._p)

    table = doc.add_table(rows=len(DETAILS), cols=2)
    table.autofit = True
    for row, (label, value) in zip(table.rows, DETAILS):
        for cell, text, bold, colour in ((row.cells[0], label, True, SLATE),
                                         (row.cells[1], value, False, None)):
            para = cell.paragraphs[0]
            _style(para.add_run(text), 10, colour, bold)
            _borders(cell)
    anchor.addprevious(table._tbl)

    tail = doc.add_paragraph()
    if carried_sect is not None:
        carried_sect.getparent().remove(carried_sect)
        tail._p.get_or_add_pPr().append(carried_sect)
    anchor.addprevious(tail._p)

    # the old cover goes, the tagline with it — the block above says it better
    for p in paras[start:end + 1]:
        p._p.getparent().remove(p._p)

    doc.save(target)
    print("cover rebuilt: %d centred lines and a %d-row detail table"
          % (len(TITLE_BLOCK), len(DETAILS)))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
