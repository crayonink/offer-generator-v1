"""
Targeted pagination and figure fixes for the regen offer templates.

Each one was asked for specifically — this is deliberately NOT the
break-before-every-heading treatment, which was tried and reverted:

1. Page break before "To," — the cover block and the letter ran together, so
   "To, / <client> / Kind Attn / address" sat crammed under the letterhead at
   the foot of page 1 and the letter resumed on page 2 mid-address at
   "Contact No.:". The whole letter now fits on one page.

2. Page break before "TECHNICAL OFFER" — it and REGENERATIVE BURNERS were
   trailing at the bottom of the letter page, stranded from their body text on
   the next page.

3. Close the gap between "TECHNICAL OFFER" and "REGENERATIVE BURNERS" — the
   second heading carried 14 pt space-before, leaving a blank line between the
   two headings.

4. Figure captions sit ABOVE their diagram, so both are set keep-with-next —
   "Fig 2: Parts of Regenerative Burner" was stranding at the foot of a page
   with its diagram overleaf. Fig 2 also gets a page break, so each diagram
   has its own page.

5. Scale both figure diagrams to 80%. They were 6.20 in wide (92% of the text
   width) and 4.30 / 3.33 in tall, so the pair alone overran a page.

6. Page break before "CLIENT DETAILS", promoted from the 11 pt Heading-3
   sub-label to the 15 pt section heading the rest of the document uses, so it
   reads as a page opener. Its siblings MARKETING PERSON DETAILS and TECHNICAL
   DATA stay as sub-labels within that page.

7. cantSplit on every savings-table row, so a table never breaks mid-row. On a
   dual-fuel offer the second table is cloned from the first at generation
   time, so it inherits this too — see engine/regen_bom_table.add_oil_savings_table,
   which also keeps the pair on one page.

Run once against the gas and oil templates, then rebuild the dual template
(it is derived from the gas one):

    python patch_regen_pagination.py
    python build_regen_dual_template.py

Idempotent.
"""

import os

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATES = ["Regen_Offer_Template.docx", "Regen_Oil_Offer_Template.docx"]

PAGE_BREAK_BEFORE = ["To,", "TECHNICAL OFFER", "CLIENT DETAILS", "EXCLUSIONS",
                     "PRICE SCHEDULE", "TERMS AND CONDITIONS"]
# SUPERVISION CHARGES is a heading in its own right but belongs with the price
# schedule, so it does not start a new keep-together group — the two share a page.
MERGE_WITH_PREVIOUS = {"SUPERVISION CHARGES FOR ERECTION & COMMISSIONING"}
# Matched on a prefix, because the heading carries template placeholders:
# "ENCON REGENERATIVE {{ fuel_word }} BURNERS – {{ burner_count }}". It opens
# the scope narrative and was crammed under the TECHNICAL DATA table.
PAGE_BREAK_BEFORE_PREFIX = ("ENCON REGENERATIVE", "Fig 2:")
CLOSE_GAP_BEFORE = "REGENERATIVE BURNERS"
# CLIENT DETAILS opens its own page, so it is promoted from the 11 pt
# Heading-3 sub-label to the document's 15 pt section-heading treatment (the
# same one TECHNICAL OFFER and ENERGY SAVING use). Its siblings MARKETING
# PERSON DETAILS / TECHNICAL DATA stay as sub-labels within that page.
PROMOTE_TO_SECTION = "CLIENT DETAILS"
SECTION_PT = 15.0
SECTION_COLOR = "1F2937"
# The whole recipient block and the whole signature block go bold. Matched on a
# distinctive fragment; only applied before the award line, so the letterhead
# "ENCON Thermal Engineers (P) Ltd" on the cover is left alone.
LETTER_BOLD = [
    "To,", "{{company_name}}", "Kind Attn:", "{{company_address}}",
    "Contact No.:", "Email: {{email}}",
    "Regards,", "{{technical_person}}", "ENCON Thermal Engineers Pvt. Ltd.",
    "53 K.M Stone", "Mob: {{technical_phone}}", "Email: {{technical_email}}",
    "Website:",
]
LETTER_END = "NATIONAL ENERGY EFFICIENCY INNOVATION AWARD"
FIG1_CAPTION = "Fig 1:"
# Captions sit ABOVE their diagram, so they must keep-with-next — otherwise a
# caption strands at the foot of a page with its figure overleaf.
FIG_CAPTIONS = ("Fig 1:", "Fig 2:")
# An absolute target (80% of the original 6.20 in) rather than a scale factor,
# so re-running the script doesn't shrink the diagrams again. Aspect ratio kept.
FIGURE_TARGET_WIDTH_IN = 4.96
FIGURE_MIN_WIDTH_IN = 4.0    # only the big diagrams, not the logo/award images
# Typography: Calibri throughout, on three sizes only.
FONT_NAME = "Calibri"
SIZE_HEADING = 15    # section headings
SIZE_SUB = 12        # sub-labels, figure captions, table header rows
SIZE_BODY = 10       # body copy, bullets, table data


def _set_page_break_before(p):
    pPr = p._element.get_or_add_pPr()
    for old in pPr.findall(qn("w:pageBreakBefore")):
        pPr.remove(old)
    pPr.append(OxmlElement("w:pageBreakBefore"))


def _no_split_rows(table):
    """Stop a table breaking mid-row across a page."""
    n = 0
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        if trPr.find(qn("w:cantSplit")) is None:
            trPr.append(OxmlElement("w:cantSplit"))
        n += 1
    return n


def _set_font(run, name=FONT_NAME, size_pt=None):
    """Force the typeface (incl. the complex-script / east-asian slots, which
    Word falls back to if only w:ascii is set) and optionally the size."""
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)


def _normalise_type(doc):
    """Calibri everywhere, on three sizes: 15 pt section headings, 12 pt
    sub-labels and table headers, 10 pt body copy.

    The cover-page display type (OFFER FOR / equipment name / letterhead) and
    the award line keep their own sizes — shrinking a 22 pt letterhead to 15 pt
    would wreck the title page — but they are still switched to Calibri.
    """
    counts = {SIZE_HEADING: 0, SIZE_SUB: 0, SIZE_BODY: 0, "display": 0}

    # Anything before the letter starts is cover-page display type.
    body_started = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "To,":
            body_started = True
        if not body_started or t.startswith(LETTER_END):
            for r in p.runs:
                _set_font(r)
                counts["display"] += 1
            continue
        if _is_section_heading(p):
            size = SIZE_HEADING
        elif p.style is not None and (p.style.name or "").startswith("Heading"):
            size = SIZE_SUB
        elif t.startswith(FIG_CAPTIONS):
            size = SIZE_SUB
        else:
            size = SIZE_BODY
        for r in p.runs:
            _set_font(r, size_pt=size)
            counts[size] += 1

    for tbl in doc.tables:
        for ri, row in enumerate(tbl.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    # header row reads as a sub-label; the data rows are body
                    size = SIZE_SUB if ri == 0 else SIZE_BODY
                    for r in p.runs:
                        _set_font(r, size_pt=size)
                        counts[size] += 1
    return counts


def _is_section_heading(p):
    """A real section heading — 15 pt, bold, 1F2937."""
    r = p.runs[0] if p.runs else None
    if r is None or not p.text.strip():
        return False
    if not r.font.size or round(r.font.size.pt, 1) != SECTION_PT:
        return False
    if not r.font.bold:
        return False
    col = r.font.color
    return bool(col and col.rgb == RGBColor.from_string(SECTION_COLOR))


def _keep_section_together(blocks, parent):
    """Bind a section (heading + body + tables) so Word carries it as one unit.

    Every block but the last keeps-with-next, so a section that will not fit in
    the space left simply starts on the next page instead of splitting. Every
    section here is well under a page, so nothing gets forced.
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for i, el in enumerate(blocks):
        # The heading always keeps with what follows: a section that is only a
        # heading (TECHNICAL OFFER, whose body is the next heading) would
        # otherwise be left alone on a page.
        last_block = i == len(blocks) - 1 and i != 0
        if el.tag == qn("w:tbl"):
            t = Table(el, parent)
            _no_split_rows(t)
            for ri, row in enumerate(t.rows):
                if last_block and ri == len(t.rows) - 1:
                    continue
                for cell in row.cells:
                    for cp in cell.paragraphs:
                        cp.paragraph_format.keep_with_next = True
        else:
            p = Paragraph(el, parent)
            p.paragraph_format.keep_together = True      # no mid-paragraph split
            if not last_block:
                p.paragraph_format.keep_with_next = True


def patch(path):
    d = Document(path)
    done, gap, kept = [], False, 0
    promoted = False
    after_fig1 = False
    for p in d.paragraphs:
        t = p.text.strip()
        if t in PAGE_BREAK_BEFORE and t not in done:
            _set_page_break_before(p)
            done.append(t)
        elif t.startswith(PAGE_BREAK_BEFORE_PREFIX) and t[:20] not in done:
            _set_page_break_before(p)
            done.append(t[:20])
        if t == PROMOTE_TO_SECTION and not promoted:
            # Match the 15 pt section headings so it reads as a page opener.
            for run in p.runs:
                run.font.size = Pt(SECTION_PT)
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(SECTION_COLOR)
            promoted = True
        elif t == CLOSE_GAP_BEFORE and not gap:
            # Sits directly under TECHNICAL OFFER — no blank line between them.
            p.paragraph_format.space_before = Pt(0)
            gap = True
        if t.startswith(FIG_CAPTIONS):
            # Caption sits above its diagram, so it must travel with it.
            p.paragraph_format.keep_with_next = True
            kept += 1

    # Recipient + signature blocks fully bold.
    bolded = 0
    for p in d.paragraphs:
        t = p.text.strip()
        if t.startswith(LETTER_END):
            break
        if any(frag in t for frag in LETTER_BOLD):
            for run in p.runs:
                run.font.bold = True
            bolded += 1

    scaled = 0
    target = Inches(FIGURE_TARGET_WIDTH_IN)
    for shape in d.inline_shapes:
        if Emu(shape.width).inches < FIGURE_MIN_WIDTH_IN:
            continue                      # logo / award image — leave alone
        ratio = shape.height / shape.width
        shape.width = target
        shape.height = Emu(int(target * ratio))
        scaled += 1

    # Savings tables must never break mid-row; on a dual offer the second table
    # is cloned from this one at generation time and inherits the property.
    nosplit = 0
    for t in d.tables:
        if t.rows and "Savings In" in t.rows[0].cells[0].text:
            nosplit += _no_split_rows(t)

    # ── Keep each section on one page ────────────────────────────────────────
    # Walk the body in order, group heading + body + tables, and bind each group
    # so a section that will not fit starts on the next page instead of
    # splitting. PRICE SCHEDULE and SUPERVISION CHARGES are bound as one group.
    by_el = {p._element: p for p in d.paragraphs}
    groups, current = [], None
    for el in d.element.body:
        p = by_el.get(el)
        is_head = p is not None and _is_section_heading(p)
        if is_head and p.text.strip() not in MERGE_WITH_PREVIOUS:
            if current:
                groups.append(current)
            current = [el]
        elif current is not None and (el.tag == qn("w:tbl") or p is not None):
            current.append(el)
    if current:
        groups.append(current)
    for g in groups:
        _keep_section_together(g, d)

    # Typography last, so it also covers anything the steps above touched.
    tc = _normalise_type(d)

    missing = [x for x in PAGE_BREAK_BEFORE if x not in done]
    note = f"  (not found: {', '.join(missing)})" if missing else ""
    print(f"  OK  {os.path.basename(path)} — breaks: {', '.join(done) or 'none'};"
          f" gap closed: {gap};"
          f" captions kept with figure: {kept}; figures scaled: {scaled};"
          f" CLIENT DETAILS promoted: {promoted}; savings rows no-split: {nosplit};"
          f" letter lines bolded: {bolded}; sections bound: {len(groups)};"
          f" runs re-typed: {sum(tc.values())}{note}")
    d.save(path)


if __name__ == "__main__":
    for name in TEMPLATES:
        patch(os.path.join(BASE_DIR, name))
    print("Now run: python build_regen_dual_template.py")
