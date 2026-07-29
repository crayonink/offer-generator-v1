"""
Patch the regen offer templates' PRICE SCHEDULE so it supports an order
quantity (buying 2, 3 … complete regen systems).

Before, the row bound Unit Price, Total Price and TOTAL all to the same
{{ price_inr }}, and the Qty column to {{ qty_words }} ("2 Pairs") — so the
schedule could only ever quote one system, and Unit == Total.

After:
    Qty          {{ qty_label }}   e.g. "03 Sets"
    Unit Price   {{ price_inr }}   price of ONE system
    Total Price  {{ price_total }} unit x qty
    TOTAL        {{ price_total }}

Run once against the gas and oil templates, then rebuild the dual template
(it is derived from the gas one):

    python patch_regen_price_schedule_qty.py
    python build_regen_dual_template.py

Idempotent — re-running it is a no-op.
"""

import os

from docx import Document

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATES = ["Regen_Offer_Template.docx", "Regen_Oil_Offer_Template.docx"]

# (row, col) -> (old placeholder, new placeholder)
SWAPS = {
    (1, 2): ("{{ qty_words }}", "{{ qty_label }}"),
    (1, 4): ("{{ price_inr }}", "{{ price_total }}"),
    (2, 4): ("{{ price_inr }}", "{{ price_total }}"),
}


def _price_table(doc):
    for t in doc.tables:
        if t.rows and any("Unit Price" in c.text for c in t.rows[0].cells):
            return t
    return None


def patch(path):
    d = Document(path)
    t = _price_table(d)
    if t is None:
        print(f"  SKIP {os.path.basename(path)} — no price schedule table")
        return
    changed = 0
    for (ri, ci), (old, new) in SWAPS.items():
        if ri >= len(t.rows):
            continue
        cell = t.rows[ri].cells[ci]
        for p in cell.paragraphs:
            for run in p.runs:
                if old in (run.text or ""):
                    run.text = run.text.replace(old, new)
                    changed += 1
    if changed:
        d.save(path)
    print(f"  {os.path.basename(path)}: {changed} cell(s) repointed"
          f"{' (already patched)' if not changed else ''}")


if __name__ == "__main__":
    for name in TEMPLATES:
        patch(os.path.join(BASE_DIR, name))
    print("Now run: python build_regen_dual_template.py")
