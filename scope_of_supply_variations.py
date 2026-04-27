"""
Generate one Word document showing the full Scope of Supply
(prose paragraphs + component tables) for every combination of:

    Hood movement: Up and Down / Swivelling - Manual / Swivelling - Geared
    Temperature control: Manual / Automatic-PLC / Automatic-PLC+AGR /
                         Automatic-PID

12 combinations total. Useful for proof-reading wording variants
side-by-side.

Usage:
    python scope_of_supply_variations.py [server_url]

Default server is the deployed automation.encon.co.in. Edit OTHER_INPUTS
at the top to change ladle / fuel / temperature etc.
"""
from __future__ import annotations
import json, os, sys, ssl, urllib.request

DEFAULT_BASE = "https://automation.encon.co.in"
OUTPUT_DIR   = "scope_of_supply_outputs"
OUTPUT_FILE  = "ScopeOfSupply_AllVariations.docx"

# Everything except the two axes we vary (hood + control mode)
OTHER_INPUTS = dict(
    mode          = "calc",
    ladle_tons    = 60,
    Ti            = 40,
    Tf            = 1200,
    time_taken_hr = 2,
    refractory_weight       = 21500,
    refractory_heat_factor  = 0.25,
    efficiency    = 0.52,
    fuel_cv       = 9000,
    fuel1_type    = "ng",
    fuel1_cv      = 9000,
    fuel2_type    = "none",
    fuel2_cv      = 0,
    blower_pressure        = "28",
    pressure_gauge_vendor  = "baumer",
    butterfly_valve_vendor = "lt_lever",
    shutoff_valve_vendor   = "aira",
    control_valve_vendor   = "dembla",
    hpu_variant            = "Duplex 1",
    pilot_burner           = "lpg_10",
    pilot_line_fuel        = "lpg",
    manual_pilot_burner    = "yes",
    pipeline_weight_kg     = 1000,
    purging_line           = "no",
    num_burners            = 1,
    ms_structure_kg_override = 0,
    ceramic_rolls_override   = 0,
    special_auto_ignition    = True,    # so pilot sections render
    special_auto_controls    = False,
)

HOODS = [
    ("Up and Down (hydraulic)", "up_down"),
    ("Swivelling - Manual",     "swivel_manual"),
    ("Swivelling - Geared",     "swivel_geared"),
]

CONTROLS = [
    ("Manual",         "manual",    "plc"),
    ("PLC",            "automatic", "plc"),
    ("PLC with AGR",   "automatic", "plc_agr"),
    ("PID",            "automatic", "pid"),
]

# Default fuel set covers both gas-train categories:
#   - NG  : packaged MADAS gas-train (representative of NG / LPG / RLNG)
#   - COG : discrete components, no rotary joint
#   - MG  : discrete components with rotary joint
#   - BFG : discrete components with double-block shut-off
# (label, fuel_type, fuel_cv kcal/Nm3) -- typical values, just for routing.
FUELS = [
    ("NG",  "ng",  9000),
    ("COG", "cog", 4500),
    ("MG",  "mg",  2500),
    ("BFG", "bg",   900),
]


# ─────────────── HTTP ───────────────
def _ctx():
    c = ssl.create_default_context()
    c.check_hostname = False
    c.verify_mode    = ssl.CERT_NONE
    return c


def _post_calc(base_url, payload):
    url  = base_url.rstrip("/") + "/api/vlph-calculate"
    data = json.dumps(payload).encode("utf-8")
    req  = urllib.request.Request(url, data=data,
        headers={"Content-Type": "application/json"}, method="POST")
    try:
        with urllib.request.urlopen(req, context=_ctx(), timeout=120) as resp:
            return json.loads(resp.read().decode("utf-8"))
    except urllib.request.HTTPError as e:
        return {"error": f"HTTP {e.code}: {e.read().decode('utf-8', 'ignore')[:200]}"}
    except Exception as e:
        return {"error": str(e)}


# ─────────────── Generate prose + lists for one variation ───────────────
def _build_variation(base_url, hood_label, hood_type, ctrl_label, control_mode, auto_control_type,
                     fuel_label="NG", fuel_type="ng", fuel_cv=9000):
    """Returns dict:
        prose_blocks    : [(heading, text), ...]
        gas_items       : [str, ...]
        air_items       : [str, ...]
        pilot_items     : [str, ...]
        gas_train_label : str (e.g. 'NATURAL GAS TRAIN FOR MAIN BURNER')
        pilot_label     : str (e.g. 'LPG LINE FOR PILOT BURNER')
        temp_items      : [str, ...] (numbered list content)
        op_seq_text     : str
    """
    payload = dict(OTHER_INPUTS,
                   hood_type=hood_type,
                   control_mode=control_mode,
                   auto_control_type=auto_control_type,
                   fuel1_type=fuel_type,
                   fuel1_cv=fuel_cv,
                   fuel_cv=fuel_cv)
    data = _post_calc(base_url, payload)
    if "error" in data:
        return {"error": data["error"]}

    bom_items = []
    for r in data.get("bom", []):
        item = (r.get("ITEM NAME") or "").strip()
        if not item or item in ("BOUGHT OUT ITEMS", "ENCON ITEMS", "GRAND TOTAL"):
            continue
        bom_items.append({
            "item":  item,
            "make":  (r.get("MAKE") or "ENCON"),
            "media": (r.get("MEDIA") or ""),
            "ref":   (r.get("REFERENCE") or ""),
        })

    # Re-use the actual offer-generator logic so prose / lists match exactly
    # what the customer would see in a generated offer.
    from engine.pdf_writer import _split_bom, _prose_blocks, _product_kind
    from engine.quote_writer import (
        _temp_control_items_for_mode,
        _operational_sequence_text,
    )

    # Pretty fuel name for the BURNER prose (e.g. "Natural Gas", "Coke Oven Gas")
    pretty_fuel = {
        "ng":  "Natural Gas",
        "lpg": "LPG",
        "rlng":"RLNG",
        "cog": "Coke Oven Gas",
        "mg":  "Mixed Gas",
        "bg":  "Blast Furnace Gas",
    }.get(fuel_type, fuel_label)

    items_for_kind = [{"product_type": "Vertical Ladle Preheater"}]
    customer = dict(
        hood_type=hood_type,
        control_mode=control_mode,
        auto_control_type=auto_control_type,
        fuel_name=pretty_fuel,
        pilot_gas_type=OTHER_INPUTS["pilot_line_fuel"].upper(),
        is_oil=False,
        is_dual=False,
        special_auto_ignition=OTHER_INPUTS["special_auto_ignition"],
        bom_items=bom_items,
    )

    scope = _split_bom(bom_items)
    prose = list(_prose_blocks(customer, scope, control_mode,
                                bool(scope.get("purging")),
                                product_kind=_product_kind(items_for_kind)))

    return {
        "hood_label":      hood_label,
        "control_label":   ctrl_label,
        "fuel_label":      fuel_label,
        "prose_blocks":    prose,
        "gas_items":       scope["gas_main"],
        "air_items":       scope["air"],
        "pilot_items":     scope["pilot"],
        "gas_train_label": scope.get("gas_main_label", "GAS TRAIN"),
        "gas_train_intro": scope.get("gas_main_intro",
                                     "Gas train will be supplied for firing of Burner, consisting of the following components:"),
        "pilot_label":     scope.get("pilot_label", "PILOT LINE"),
        "purging_items":   scope["purging"],
        "temp_items":      [d["item"] for d in _temp_control_items_for_mode(control_mode, auto_control_type)],
        "op_seq_text":     _operational_sequence_text(control_mode, auto_control_type),
    }


# ─────────────── Word doc ───────────────
def _write_doc(variations: list, out_path: str):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2)
        s.top_margin  = s.bottom_margin = Cm(2)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Scope of Supply -- All Variations")
    r.bold = True
    r.font.size = Pt(18)

    fuels_in_doc = sorted({v.get("fuel_label", "?") for v in variations if "error" not in v})
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(
        f"Ladle: {OTHER_INPUTS['ladle_tons']} T  |  "
        f"Ti->Tf: {OTHER_INPUTS['Ti']}->{OTHER_INPUTS['Tf']} C  |  "
        f"Fuels: {', '.join(fuels_in_doc)}  |  "
        f"{len(variations)} combinations"
    )
    sr.italic = True
    sr.font.size = Pt(10)

    for v in variations:
        if "error" in v:
            doc.add_page_break()
            doc.add_heading(f"{v.get('hood_label','?')} -- {v.get('control_label','?')}  [ERROR]", level=1)
            doc.add_paragraph(v["error"])
            continue

        doc.add_page_break()
        h = doc.add_heading(
            f"{v.get('fuel_label','?')}  |  {v['hood_label']}  |  Temperature Control: {v['control_label']}",
            level=1)

        # Prose blocks (Steel Structure, Ladle Hood, Hood Movement, Burner, etc.)
        for heading, body in v["prose_blocks"]:
            doc.add_heading(heading, level=2)
            if isinstance(body, (list, tuple)):
                for para in body:
                    doc.add_paragraph(para)
            else:
                doc.add_paragraph(body)

        # Combustion Air Line
        doc.add_heading("COMBUSTION AIR LINE", level=2)
        doc.add_paragraph("The airline will consist of the following items:")
        for x in v["air_items"] or ["(none)"]:
            doc.add_paragraph(x, style="List Bullet")

        # Main fuel gas train
        if v["gas_items"]:
            doc.add_heading(v["gas_train_label"], level=2)
            doc.add_paragraph(v["gas_train_intro"])
            for x in v["gas_items"]:
                doc.add_paragraph(x, style="List Bullet")

        # Pilot line
        if v["pilot_items"]:
            doc.add_heading(v["pilot_label"], level=2)
            doc.add_paragraph(
                "We shall be supplying a gas train for Pilot Burner, which will "
                "supply required gas flow and pressure to the pilot burner for "
                "ignition of the main burner. The pilot line will consist of "
                "the following main components:")
            for x in v["pilot_items"]:
                doc.add_paragraph(x, style="List Bullet")

        # Purging line (only when present)
        if v["purging_items"]:
            doc.add_heading("NITROGEN PURGING LINE", level=2)
            for x in v["purging_items"]:
                doc.add_paragraph(x, style="List Bullet")

        # Temperature Control System
        doc.add_heading("TEMPERATURE CONTROL SYSTEM", level=2)
        doc.add_paragraph(
            "To control and maintain the temperature of the ladle accurately, "
            "the thermocouple will be fitted in the Ladle at suitable location. "
            "Temperature control system will consist of the following main "
            "components:")
        for i, item in enumerate(v["temp_items"], start=1):
            doc.add_paragraph(f"{i}. {item}", style="List Number")

        # Operational Sequence (mode-specific)
        doc.add_heading("OPERATIONAL SEQUENCE", level=2)
        doc.add_paragraph(v["op_seq_text"])

    doc.save(out_path)


# ─────────────── main ───────────────
def main():
    base_url = sys.argv[1] if len(sys.argv) > 1 else DEFAULT_BASE
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    print(f"Server: {base_url}")
    total = len(HOODS) * len(CONTROLS) * len(FUELS)
    print(f"Generating {total} variations -> {OUTPUT_DIR}/{OUTPUT_FILE}")

    variations = []
    for fuel_label, fuel_type, fuel_cv in FUELS:
        for hood_label, hood_type in HOODS:
            for ctrl_label, control_mode, auto_control_type in CONTROLS:
                tag = f"{fuel_label} / {hood_label} / {ctrl_label}"
                print(f"  [{len(variations)+1}/{total}] {tag} ...")
                v = _build_variation(base_url,
                                     hood_label, hood_type,
                                     ctrl_label, control_mode, auto_control_type,
                                     fuel_label=fuel_label, fuel_type=fuel_type,
                                     fuel_cv=fuel_cv)
                v.setdefault("hood_label", hood_label)
                v.setdefault("control_label", ctrl_label)
                v.setdefault("fuel_label", fuel_label)
                variations.append(v)

    out_path = os.path.join(OUTPUT_DIR, OUTPUT_FILE)
    _write_doc(variations, out_path)
    print(f"\nSaved: {out_path}  ({os.path.getsize(out_path):,} bytes)")


if __name__ == "__main__":
    main()
