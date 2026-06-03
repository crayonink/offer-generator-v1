"""Build PU_Offer_Template.docx by cloning the VLPH offer template
(Offer_Template.docx) and surgically replacing the ladle-preheater
content with stand-alone Pumping Unit content.

Sister script of build_hpu_template_from_vlph.py — identical shell,
the only differences are equipment labels ('Pumping Unit' instead
of 'Pumping Unit') and the rendered docx target.

Template placeholders introduced (filled by /api/generate-pu-quote):
  pu_variant, pu_kw, pu_lph, pu_qty  — same value source as the HPU
  template's hpu_* keys; the offer endpoint pours customer.hpu_* into
  these too so a single context dict feeds either template.
"""
from __future__ import annotations

import os
import shutil
from copy import deepcopy

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.enum.table import WD_ROW_HEIGHT_RULE

SOURCE = "Offer_Template.docx"
TARGET = "PU_Offer_Template.docx"


# ── PU Tech Specs table (replaces VLPH's 24-row preheater spec) ──────────
# Keys are 'hpu_*' rather than 'pu_*' so the same render context dict
# feeds either HPU or PU templates without endpoint-side branching.
_HPU_SPECS = [
    ('Equipment',           '{{ equipment_name }}'),
    ('Variant',             '{{ hpu_variant }}'),
    ('Motor Capacity',      '{{ hpu_kw }} kW'),
    ('Oil Flow Rate',       '{{ hpu_lph }} LPH'),
    ('Quantity',            '{{ hpu_qty }} No.'),
]


# ── Annexure I — Scope of Supply (PU accessories) ─────────────────────────
# Intentionally empty: previous defaults were fabricated by the AI rather
# than sourced from ENCON catalogs. Populate with real (description, qty)
# tuples — or wire to a master table — when the catalog list is provided.
_HPU_SCOPE_ITEMS = []


def _replace_spec_table_with_hpu(table) -> None:
    """Wipe the VLPH 24-row tech-data table and re-fill with PU specs.
    Keep the header row from VLPH; relabel it 'PU Specifications'."""
    tbl_xml = table._element
    for tr in list(tbl_xml.findall(qn('w:tr')))[1:]:
        tbl_xml.remove(tr)
    head = table.rows[0]
    head.cells[0].text = 'PU Specifications'
    head.cells[1].text = ''
    head.cells[0].merge(head.cells[1])
    for p in head.cells[0].paragraphs:
        for r in p.runs:
            r.bold = True
    for label, value in _HPU_SPECS:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value


def _replace_scope_table_with_hpu(table) -> None:
    """Replace the Annexure I 12-row Scope of Supply with HPU accessories.
    VLPH layout: header row, banner row (merged), 10 scope rows.
    HPU layout : header row, banner row (merged), 15 scope rows."""
    tbl_xml = table._element
    for tr in list(tbl_xml.findall(qn('w:tr')))[1:]:
        tbl_xml.remove(tr)

    # Banner row (merged) — short label
    banner = table.add_row()
    banner.cells[0].merge(banner.cells[1])
    banner.cells[0].text = 'Pumping Unit'
    for p in banner.cells[0].paragraphs:
        for r in p.runs:
            r.bold = True

    # Numbered scope rows. The table is 2-col 'S. No. | Description'; the
    # original VLPH layout has no Qty column. To keep table-width parity,
    # fold quantity into the description: "Item — Qty".
    for idx, (desc, qty) in enumerate(_HPU_SCOPE_ITEMS, start=1):
        row = table.add_row()
        row.cells[0].text = str(idx)
        row.cells[1].text = f'{desc}  ({qty})'


def _delete_vlph_scope_body(doc: Document) -> None:
    """Walk the body looking for the inline 'SCOPE OF SUPPLY' paragraph
    (NOT the Annexure heading) and delete every element from there up
    to (but not including) the 'ANNEXURE I' heading. Replace with a
    short HPU description block."""
    body = doc.element.body
    elements = list(body)

    start_idx, end_idx = None, None
    for i, el in enumerate(elements):
        if el.tag.split('}')[-1] != 'p':
            continue
        txt = ''.join(t.text or '' for t in el.iter(qn('w:t'))).strip()
        if start_idx is None and txt == 'SCOPE OF SUPPLY':
            start_idx = i
        elif start_idx is not None and txt.startswith('ANNEXURE I'):
            end_idx = i
            break

    if start_idx is None or end_idx is None:
        print('Could not locate VLPH scope body markers — skipping delete.')
        return

    for el in elements[start_idx:end_idx]:
        body.remove(el)
    # No replacement description block — the previous wording was
    # AI-fabricated rather than sourced from ENCON's catalog. The
    # HPU Specifications table speaks for itself.

    # Flip the section title. In the source template this heading is a
    # Jinja conditional:
    #   {% if is_tundish %}TUNDISH{% elif is_horizontal %}HORIZONTAL LADLE
    #   {% else %}VERTICAL LADLE{% endif %} PREHEATERS AND DRYERS
    # We match on the trailing 'PREHEATERS AND DRYERS' which is stable,
    # then wipe all runs and stamp the HPU heading.
    for el in body.iter(qn('w:p')):
        t_els = list(el.iter(qn('w:t')))
        joined = ''.join((t.text or '') for t in t_els).upper()
        if 'PREHEATERS AND DRYERS' in joined:
            for t in t_els:
                t.text = ''
            if t_els:
                t_els[0].text = 'PUMPING UNIT'


def _inject_scope_intro_paragraph(doc: Document) -> None:
    """Rewrite Annexure I's intro paragraph to HPU-specific wording."""
    body = doc.element.body
    elements = list(body)

    heading_idx = None
    for i, el in enumerate(elements):
        if el.tag.split('}')[-1] != 'p':
            continue
        txt = ''.join(t.text or '' for t in el.iter(qn('w:t'))).strip()
        if 'ANNEXURE I' in txt.upper() and 'SCOPE OF SUPPLY' in txt.upper():
            heading_idx = i
            break
    if heading_idx is None:
        return

    intro_p = None
    for j in range(heading_idx + 1, len(elements)):
        if elements[j].tag.split('}')[-1] == 'p':
            intro_p = elements[j]
            break
    if intro_p is None:
        return

    NEW_INTRO = (
        'The Scope of supply covers Design, manufacturing and supply '
        'of the Pumping Unit.'
    )
    t_els = list(intro_p.iter(qn('w:t')))
    if t_els and t_els[0].text and t_els[0].text.startswith(NEW_INTRO[:40]):
        return  # idempotent
    for t in t_els:
        t.text = ''
    if t_els:
        t_els[0].text = NEW_INTRO
    else:
        r = OxmlElement('w:r')
        t = OxmlElement('w:t'); t.text = NEW_INTRO; r.append(t)
        intro_p.append(r)


def _strip_project_name_from_cover_box(doc: Document) -> None:
    """Remove the 'Project Name' row from the cover-box table.
    The HPU offer doesn't have a separate project name beyond the
    equipment name itself, so the duplicate row is just noise.

    The target is the first 4-row table whose row 0 label is
    'Project / Equipment'. We drop the row whose first cell starts
    with 'Project Name'.
    """
    target = None
    for t in doc.tables:
        if not t.rows or not t.rows[0].cells:
            continue
        if t.rows[0].cells[0].text.strip().startswith('Project / Equipment'):
            target = t
            break
    if target is None:
        print('Cover-box table not found — skipping Project Name strip.')
        return
    tbl_xml = target._element
    for tr in list(tbl_xml.findall(qn('w:tr'))):
        first_cell_txt = ''.join(t.text or '' for t in tr.iter(qn('w:t'))).strip()
        if first_cell_txt.lower().startswith('project name'):
            tbl_xml.remove(tr)
            print('Stripped Project Name row from cover-box table.')
            return


def _remove_annexure_ii(doc: Document) -> None:
    """Strip Annexure II — Exclusions from the cloned VLPH template.
    Annexure II is a bullet list (no table), so we remove paragraphs
    from the 'ANNEXURE II' heading up to (but NOT including) the next
    'ANNEXURE ...' heading. Also strips the matching row from the
    List of Annexures table on the cover page.
    """
    # ── Part A: strip the Annexure II row from List of Annexures ──────
    for t in doc.tables:
        if not t.rows or len(t.rows[0].cells) < 2:
            continue
        head = [c.text.strip().upper() for c in t.rows[0].cells]
        if 'ANNEXURE NO.' in head[0]:
            tbl_xml = t._element
            for tr in list(tbl_xml.findall(qn('w:tr'))):
                row_txt = ''.join(x.text or '' for x in tr.iter(qn('w:t'))).upper()
                # Match 'ANNEXURE II' + 'EXCLUSIONS' but NOT 'ANNEXURE III'.
                if ('ANNEXURE II' in row_txt
                        and 'ANNEXURE III' not in row_txt
                        and 'EXCLUS' in row_txt):
                    tbl_xml.remove(tr)
            break

    # ── Part B: strip the body block (heading + intro + bullets) ──────
    body = doc.element.body
    elements = list(body)

    heading_idx = None
    for i, el in enumerate(elements):
        if el.tag.split('}')[-1] != 'p':
            continue
        txt = ''.join(t.text or '' for t in el.iter(qn('w:t'))).strip().upper()
        # 'ANNEXURE II — EXCLUSIONS' but not 'ANNEXURE III'
        if (txt.startswith('ANNEXURE II')
                and not txt.startswith('ANNEXURE III')
                and 'EXCLUS' in txt):
            heading_idx = i
            break
    if heading_idx is None:
        return  # idempotent — already removed

    # Walk forward until we hit the next 'ANNEXURE …' heading, deleting
    # everything in between (the bullets live as plain paragraphs).
    i = heading_idx
    removed = 0
    while i < len(elements):
        el = elements[i]
        tag = el.tag.split('}')[-1]
        if tag == 'sectPr':
            break
        if tag == 'p' and i != heading_idx:
            txt = ''.join(t.text or '' for t in el.iter(qn('w:t'))).strip().upper()
            if txt.startswith('ANNEXURE'):
                break  # next annexure — stop here, leave it intact
        body.remove(el)
        removed += 1
        i += 1
    print(f'Removed Annexure II (elements removed={removed})')


def _remove_annexure_i_scope_table(doc: Document) -> None:
    """Strip the Annexure I scope-of-supply table itself. The list of
    accessories is empty (see _HPU_SCOPE_ITEMS = []), so an S.No. /
    Item-Description header with only an equipment-name banner row
    adds no information.

    Identified by content: it's the table that has a 'Heating and
    Pumping Unit' (or 'Hydraulic Pumping Unit') banner row stamped
    by _replace_scope_table_with_hpu(). Identifying by content avoids
    a false-match on the 5-row Section TOC ('S. No. | Section')."""
    SCOPE_BANNERS = {'PUMPING UNIT', 'HYDRAULIC PUMPING UNIT'}
    target = None
    for t in doc.tables:
        for row in t.rows:
            cell_text = row.cells[0].text.strip().upper() if row.cells else ''
            if cell_text in SCOPE_BANNERS:
                target = t
                break
        if target is not None:
            break
    if target is None:
        return
    tbl_xml = target._element
    parent = tbl_xml.getparent()
    nxt = tbl_xml.getnext()
    parent.remove(tbl_xml)
    if nxt is not None and nxt.tag.split('}')[-1] == 'p':
        nxt_txt = ''.join(t.text or '' for t in nxt.iter(qn('w:t'))).strip()
        if not nxt_txt:
            parent.remove(nxt)
    print('Removed empty Annexure I scope table.')


def _rename_annexure_i_heading(doc: Document) -> None:
    """Extend 'ANNEXURE I — SCOPE OF SUPPLY' to
    'ANNEXURE I — SCOPE OF SUPPLY: PUMPING UNIT'."""
    body = doc.element.body
    for el in body.iter(qn('w:p')):
        t_els = list(el.iter(qn('w:t')))
        joined = ''.join((t.text or '') for t in t_els)
        joined_u = joined.upper()
        if 'ANNEXURE I' in joined_u and 'SCOPE OF SUPPLY' in joined_u:
            # Skip if already extended.
            if 'PUMPING UNIT' in joined_u:
                return
            # Only the *standalone* Annexure I heading paragraph, not
            # the List-of-Annexures row (which is inside a table cell —
            # those are also reached by body.iter, so filter by length).
            if len(joined.strip()) > 60:
                continue
            new_text = 'ANNEXURE I — SCOPE OF SUPPLY: PUMPING UNIT'
            for t in t_els:
                t.text = ''
            if t_els:
                t_els[0].text = new_text
            print('Renamed Annexure I heading.')
            return


def _remove_supervision_table(doc: Document) -> None:
    """Strip the 3-row Supervision Charges sub-table (Mechanical + PLC
    + Note) that sits just below the Price Schedule. HPU offers don't
    quote supervision separately."""
    target = None
    for t in doc.tables:
        if not t.rows or not t.rows[0].cells:
            continue
        head = t.rows[0].cells[0].text.strip()
        # Match either the original VLPH heading or the {%tr%}-wrapped
        # supervision_include variant introduced by build_recup_template.
        if head.startswith('Supervision Charges') or head.startswith('{%tr if supervision_include'):
            target = t
            break
    if target is None:
        return  # idempotent — already removed
    tbl_xml = target._element
    parent = tbl_xml.getparent()
    # Drop the table itself + one trailing blank paragraph (if any) so
    # the offer doesn't leave a yawning gap where the table used to be.
    nxt = tbl_xml.getnext()
    parent.remove(tbl_xml)
    if nxt is not None and nxt.tag.split('}')[-1] == 'p':
        nxt_txt = ''.join(t.text or '' for t in nxt.iter(qn('w:t'))).strip()
        if not nxt_txt:
            parent.remove(nxt)
    print('Removed Supervision Charges sub-table.')


def _renumber_annexures_after_removal(doc: Document) -> None:
    """After Annexure II is removed, the body still has 'ANNEXURE III'
    (Price Schedule) and 'ANNEXURE IV' (T&Cs) headings, and the
    cover-page List of Annexures lists them under those numbers.
    Shift them down so the document reads I, II, III consistently.

    Replacement order is critical:
      1) III -> II  (do this first; it does not touch IV)
      2) IV  -> III (safe; previous step already consumed the III)
    Doing IV first would create a new 'III' that step 2 would then
    rename to 'II', collapsing both into II.

    Each paragraph's runs are joined, replaced, and the result is
    stamped back into the first run so headings that Word split
    across multiple <w:t> elements are handled correctly.
    """
    renames = [
        ('ANNEXURE III', 'ANNEXURE II'),
        ('Annexure III', 'Annexure II'),
        ('ANNEXURE IV',  'ANNEXURE III'),
        ('Annexure IV',  'Annexure III'),
    ]

    def _rewrite(p_el) -> bool:
        t_els = list(p_el.iter(qn('w:t')))
        if not t_els:
            return False
        joined = ''.join((t.text or '') for t in t_els)
        new = joined
        for old, repl in renames:
            new = new.replace(old, repl)
        if new == joined:
            return False
        for t in t_els:
            t.text = ''
        t_els[0].text = new
        return True

    body = doc.element.body
    count = 0
    for p_el in body.iter(qn('w:p')):
        if _rewrite(p_el):
            count += 1
    print(f'Renumbered annexures III->II and IV->III in {count} paragraphs.')


def _scrub_reference_list_mention(doc: Document) -> None:
    """The Company Profile blurb ends with '... a representative
    reference list is included in Annexure V.' Once Annexure V is
    removed that sentence dangles — rewrite the paragraph to drop the
    trailing reference."""
    needle = 'a representative reference list is included in Annexure V'
    for p in doc.paragraphs:
        if needle.lower() not in p.text.lower():
            continue
        # Join all runs, snip the dangling tail, blank the runs and
        # stamp the cleaned text into the first run.
        full = ''.join(r.text or '' for r in p.runs)
        idx = full.lower().find('— ' + needle.lower())
        if idx == -1:
            idx = full.lower().find(needle.lower())
        if idx == -1:
            continue
        cleaned = full[:idx].rstrip(' —-').rstrip()
        if not cleaned.endswith('.'):
            cleaned += '.'
        for r in p.runs:
            r.text = ''
        if p.runs:
            p.runs[0].text = cleaned
        else:
            p.add_run(cleaned)
        print('Scrubbed dangling Annexure V mention from Company Profile.')
        return


def _remove_annexure_section(doc: Document, roman: str, label_keyword: str) -> None:
    """Generic Annexure remover. Strips:
      * The matching row from the cover-page 'List of Annexures' table
        (matched by 'ANNEXURE <roman>' + the label keyword).
      * The body block: heading paragraph + every element after it up
        to (but NOT including) the next 'ANNEXURE …' heading or
        section properties marker.

    `roman` is e.g. 'V' / 'VI'. `label_keyword` is e.g. 'REFERENCE' /
    'MAKE' — used to disambiguate from same-prefix annexures (e.g.
    'ANNEXURE V' must not match 'ANNEXURE VI').
    """
    target_prefix = f'ANNEXURE {roman}'

    def _is_target_heading(txt_upper: str) -> bool:
        if not txt_upper.startswith(target_prefix):
            return False
        # Reject longer-roman matches (e.g. when looking for V, reject VI).
        rest = txt_upper[len(target_prefix):].lstrip()
        if rest and rest[0].isalpha() and rest[0] in 'IVX':
            return False
        return label_keyword in txt_upper

    # ── Part A: strip the row from List of Annexures ──────────────────
    for t in doc.tables:
        if not t.rows or len(t.rows[0].cells) < 2:
            continue
        head = [c.text.strip().upper() for c in t.rows[0].cells]
        if 'ANNEXURE NO.' in head[0]:
            tbl_xml = t._element
            for tr in list(tbl_xml.findall(qn('w:tr'))):
                row_txt = ''.join(x.text or '' for x in tr.iter(qn('w:t'))).upper()
                if _is_target_heading(row_txt):
                    tbl_xml.remove(tr)
            break

    # ── Part B: strip the body section ────────────────────────────────
    body = doc.element.body
    elements = list(body)

    heading_idx = None
    for i, el in enumerate(elements):
        if el.tag.split('}')[-1] != 'p':
            continue
        txt = ''.join(t.text or '' for t in el.iter(qn('w:t'))).strip().upper()
        if _is_target_heading(txt):
            heading_idx = i
            break
    if heading_idx is None:
        return  # idempotent

    i = heading_idx
    removed = 0
    while i < len(elements):
        el = elements[i]
        tag = el.tag.split('}')[-1]
        if tag == 'sectPr':
            break
        if tag == 'p' and i != heading_idx:
            txt = ''.join(t.text or '' for t in el.iter(qn('w:t'))).strip().upper()
            if txt.startswith('ANNEXURE'):
                break
        body.remove(el)
        removed += 1
        i += 1
    print(f'Removed {target_prefix} ({label_keyword}) — elements removed={removed}')


def _pad_table_rows(doc: Document, min_height_cm: float = 0.75) -> None:
    """Give the Price Schedule and Supervision sub-table some breathing
    room so the printed HPU offer doesn't look cramped."""
    targets = []
    for t in doc.tables:
        if not t.rows or len(t.rows[0].cells) < 2:
            continue
        head_r = t.rows[0].cells[1].text.strip().upper() if len(t.rows[0].cells) > 1 else ''
        head_l = t.rows[0].cells[0].text.strip()
        if 'ITEM DESCRIPTION' in head_r:
            targets.append(t)
        elif head_l.startswith('Supervision Charges'):
            targets.append(t)

    height = Cm(min_height_cm)
    for tbl in targets:
        for row in tbl.rows:
            first_cell_text = row.cells[0].text.strip() if row.cells else ''
            if first_cell_text.startswith('{%tr') or first_cell_text.startswith('{%p'):
                continue
            row.height = height
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def main() -> None:
    if not os.path.exists(SOURCE):
        raise SystemExit(f"missing: {SOURCE}")

    if os.path.exists(TARGET) and not os.path.exists(TARGET + '.bak'):
        shutil.copy(TARGET, TARGET + '.bak')

    shutil.copy(SOURCE, TARGET)
    doc = Document(TARGET)
    print(f'Cloned {SOURCE} -> {TARGET}: '
          f'{len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables')

    # 1. Delete VLPH scope-of-supply body; insert HPU description block.
    _delete_vlph_scope_body(doc)

    # 2. Replace T5 (Tech Specs 24-row) with HPU specs.
    designing_table = doc.tables[5]
    _replace_spec_table_with_hpu(designing_table)

    # 3. Replace Annexure I scope table with HPU accessories. Locate by
    #    row count (10..15 rows, 'S. No.' header, not the Reference List
    #    50-row table and not the Make List).
    scope_table = None
    for t in doc.tables:
        if not t.rows or len(t.rows[0].cells) < 2:
            continue
        head_l = t.rows[0].cells[0].text.strip().upper()
        head_r = t.rows[0].cells[1].text.strip().upper()
        if head_l == 'S. NO.' and 'ITEM DESCRIPTION' not in head_r:
            if 10 <= len(t.rows) <= 15:
                scope_table = t
                break
    if scope_table is not None:
        _replace_scope_table_with_hpu(scope_table)
    else:
        print('Annexure I Scope of Supply table not found — leaving as-is.')

    # 4. Rewrite the intro paragraph above the (about-to-be-removed)
    #    Scope of Supply table.
    _inject_scope_intro_paragraph(doc)

    # 4b. Extend the Annexure I heading to include the equipment label.
    _rename_annexure_i_heading(doc)

    # 4c. Remove the empty Annexure I scope table entirely — with
    #     _HPU_SCOPE_ITEMS=[] it has only header + banner row, no value.
    _remove_annexure_i_scope_table(doc)

    # 5. Drop the redundant 'Project Name' row from the cover-box
    #    table (Project/Equipment + Client + Enquiry No. is enough).
    _strip_project_name_from_cover_box(doc)

    # 5b. Strip Annexure II — Exclusions (HPU offers don't need it).
    _remove_annexure_ii(doc)

    # 5c. Strip Annexure V — Reference List of Clients (no curated
    #     HPU client list yet).
    _remove_annexure_section(doc, 'V',  'REFERENCE')

    # 5d. Strip Annexure VI — Make List (irrelevant for stand-alone HPU).
    _remove_annexure_section(doc, 'VI', 'MAKE')

    # 5e. Scrub the dangling "...representative reference list is
    #     included in Annexure V" tail-sentence in the Company Profile
    #     section, since Annexure V was just removed.
    _scrub_reference_list_mention(doc)

    # 5f. Strip the Supervision Charges sub-table (Mechanical + PLC +
    #     Note). HPU offers don't quote erection / commissioning
    #     supervision separately.
    _remove_supervision_table(doc)

    # 5g. Annexure II was removed in 5b. Renumber the remaining body
    #     headings + List-of-Annexures rows so the offer reads
    #     I -> II -> III instead of I -> III -> IV.
    _renumber_annexures_after_removal(doc)

    # 6. Give the Price Schedule + Supervision rows some breathing room.
    _pad_table_rows(doc, min_height_cm=0.75)

    doc.save(TARGET)
    print(f'Saved -> {TARGET}')


if __name__ == '__main__':
    main()
