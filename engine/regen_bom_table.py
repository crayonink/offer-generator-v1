"""Populate the regen offer's BILL OF MATERIAL table dynamically from the
actual BOM produced by ``bom.regen_builder.build_regen_df``.

The offer template ships with a static placeholder BOM table (header row
"Description | Qty."). At generation time we clear its body and re-emit it
faithfully from the computed BOM: one merged/bold row per BOM section, then an
item row (description + spec) and its real total quantity for every line.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as WA
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = "1F2937"     # body text colour (matches the rest of the offer)
SEC  = "EEF2F7"     # section-row shading


def _set_fill(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexc)
    tcPr.append(shd)


def _style_para(p, align):
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(2); pf.space_after = Pt(2); pf.line_spacing = 1.0


def _qty_label(q):
    try:
        q = int(round(float(q)))
    except (TypeError, ValueError):
        return str(q)
    return f"{q} No." if q == 1 else f"{q} Nos."


def _desc(name, spec):
    name = (name or "").strip()
    spec = (spec or "").strip()
    return f"{name} — {spec}" if spec else name


def _add_section_row(tbl):
    row = tbl.add_row()
    cell = row.cells[0].merge(row.cells[1])
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_fill(cell, SEC)
    p = cell.paragraphs[0]; _style_para(p, WA.LEFT)
    return cell, p


def _add_item_row(tbl):
    row = tbl.add_row()
    for c in row.cells:
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return row.cells[0], row.cells[1]


def _run(p, text, *, bold=False, size=12):
    r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = RGBColor.from_string(DARK)
    return r


def fill_temp_control(doc_path, df):
    """Regenerate the TEMPERATURE CONTROL bullet list from the BOM.

    Replaces the static component bullets (PLC / thermocouple / flow meters /
    control valves) with lines driven by the real BOM — each shows its actual
    quantity, and the PLC model follows the selection. Items absent from the BOM
    are dropped. Returns True if the bullet block was found and rewritten.
    """
    import copy, re as _re

    d = Document(doc_path)
    paras = d.paragraphs
    hi = next((i for i, p in enumerate(paras)
               if p.text.strip() == "TEMPERATURE CONTROL"), None)
    if hi is None:
        return False

    # the component bullets = the consecutive list paragraphs after the intro
    bullets, started = [], False
    for i in range(hi + 1, len(paras)):
        pPr = paras[i]._element.find(qn("w:pPr"))
        is_bullet = pPr is not None and pPr.find(qn("w:numPr")) is not None
        if is_bullet:
            started = True
            bullets.append(paras[i])
        elif started:
            break
    if not bullets:
        return False

    def _lookup(section, needle):
        sub = df[df["SECTION"] == section]
        for _, r in sub.iterrows():
            if needle in str(r["ITEM NAME"]).lower():
                return int(round(float(r["QTY"]))), r
        return None, None

    # PLC model from its BOM spec (e.g. "Siemens S7-1200/1500 with touch panel")
    _pq, _pr = _lookup("CONTROLS", "plc with hmi")
    plc_model = ""
    if _pr is not None:
        m = _re.search(r"S7[-\s]?[\d/]+", str(_pr.get("SPECIFICATION", "")))
        plc_model = m.group(0).replace(" ", "") if m else ""
    plc_phrase = (f"PLC {plc_model} with HMI" if plc_model else "PLC with HMI")

    is_oil = any(str(s).upper().startswith("OIL") for s in df["SECTION"].unique())
    if is_oil:
        SPEC = [
            ("CONTROLS",     "plc with hmi",        plc_phrase),
            ("TEMP CONTROL", "thermocouple",        "Thermocouple with Temperature Transmitter"),
            ("TEMP CONTROL", "air flow meter",      "Orifice with DPT Volumetric Flow Meter for Air"),
            ("TEMP CONTROL", "oil flow meter",      "Flow Meter in Oil Line"),
            ("TEMP CONTROL", "air control valve",   "Pneumatic Air Control Valve"),
            ("TEMP CONTROL", "oil control valve",   "Pneumatic Control Valve in Oil Line"),
            ("TEMP CONTROL", "pneumatic damper",    "Pneumatic Flue Control Valve"),
        ]
    else:
        SPEC = [
            ("CONTROLS",     "plc with hmi",      plc_phrase),
            ("TEMP CONTROL", "thermocouple",      "Thermocouple with Temperature Transmitter"),
            ("TEMP CONTROL", "air flow meter",    "Orifice with DPT Volumetric Flow Meter for Air"),
            ("TEMP CONTROL", "gas flow meter",    "Orifice with DPT Volumetric Flow Meter for Gas"),
            ("TEMP CONTROL", "air control valve", "Pneumatic Air Control Valve"),
            ("TEMP CONTROL", "gas control valve", "Pneumatic Gas Control Valve"),
            ("TEMP CONTROL", "pneumatic damper",  "Pneumatic Flue Control Valve"),
        ]
    lines = []
    for section, needle, phrase in SPEC:
        q, _ = _lookup(section, needle)
        if q:
            lines.append(phrase)   # no quantity in the offer scope
    if not lines:
        return False

    template = bullets[0]._element      # clone for list formatting
    prev = bullets[-1]._element         # insert new bullets after the old block
    for txt in lines:
        nb = copy.deepcopy(template)
        runs = nb.findall(qn("w:r"))
        if runs:
            ts = runs[0].findall(qn("w:t"))
            if ts:
                ts[0].text = txt
                for extra in ts[1:]:
                    extra.getparent().remove(extra)
            for r in runs[1:]:
                r.getparent().remove(r)
        prev.addnext(nb)
        prev = nb
    for p in bullets:
        p._element.getparent().remove(p._element)

    d.save(doc_path)
    return True


def fill_consist_list(doc_path, is_oil=False, gas_train_label="NG"):
    """Rebuild the cover-letter 'consisting of' scope list per fuel.

    Always: Regenerative Burners, Temperature Control, Furnace Pressure Control,
    Blower for Combustion Air, ID Fan for Suction of Flue Gas, Panel with PLC &
    HMI. Gas fuels also get '{fuel} Gas Train' (oil fuels have no gas train).
    Idempotent — removes any existing list first.
    """
    import copy
    from docx.shared import Pt, Inches
    from docx.text.paragraph import Paragraph

    d = Document(doc_path)
    anchor = next((p for p in d.paragraphs
                   if p.text.strip().startswith("We are pleased to enclose")), None)
    if anchor is None:
        return False
    body = next((p for p in d.paragraphs
                 if p.text.strip().startswith("This is with reference")), anchor)

    points = ["Regenerative Burners", "Temperature Control"]
    if not is_oil:
        points.append(f"{gas_train_label} Gas Train")
    points += ["Furnace Pressure Control", "Blower for Combustion Air",
               "ID Fan for Suction of Flue Gas", "Panel with PLC & HMI"]
    LEAD = "The offer broadly comprises the following:"
    KEYS = ("Regenerative Burners", "Temperature Control", "Gas Train",
            "Furnace Pressure Control", "Blower for Combustion", "ID Fan",
            "Panel with PLC")

    # remove any existing list (idempotent)
    for p in list(d.paragraphs):
        t = p.text.strip()
        if t == LEAD or (len(t) > 2 and t[0].isdigit() and t[1] == "."
                         and any(k in t for k in KEYS)):
            p._element.getparent().remove(p._element)

    anchor = next(p for p in d.paragraphs
                  if p.text.strip().startswith("We are pleased to enclose"))

    def _after(prev_el, text, indent=None, sa=2):
        el = copy.deepcopy(body._element)
        for r in el.findall(qn("w:r")):
            r.getparent().remove(r)
        prev_el.addnext(el)
        p = Paragraph(el, body._parent)
        run = p.add_run(text)
        pf = p.paragraph_format
        pf.space_before = Pt(0); pf.space_after = Pt(sa); pf.line_spacing = 1.0
        if indent is not None:
            pf.left_indent = Inches(indent)
        return el

    prev = _after(anchor._element, LEAD, sa=3)
    for i, pt in enumerate(points, 1):
        prev = _after(prev, f"{i}. {pt}", indent=0.35, sa=1)
    d.save(doc_path)
    return True


def fill_oil_supply(doc_path, df):
    """Populate the oil offer's HEATING & PUMPING UNIT bullet list from the BOM.

    Lists the OIL AUXILIARY (Heating & Pumping Unit) and OIL LINE — BURNER
    components with quantities. No-op if the BOM has no oil sections (gas offer).
    """
    import copy

    # HPU first, then the oil line to the burners (exclude the ID Fan — it has
    # its own SUCTION BLOWER FOR FLUE GAS section in the offer).
    oil_rows = []
    for section in ("OIL AUXILIARY", "OIL LINE — BURNER"):
        sub = df[df["SECTION"] == section]
        sub = sub[~sub["ITEM NAME"].str.contains("ID Fan", case=False, na=False)]
        oil_rows += list(sub.iterrows())
    if not oil_rows:
        return False

    lines = []
    for _, r in oil_rows:
        q = int(round(float(r["QTY"])))
        name = str(r["ITEM NAME"]).strip()
        lines.append(name)   # no quantity in the offer scope
    if not lines:
        return False

    d = Document(doc_path)
    paras = d.paragraphs
    hi = next((i for i, p in enumerate(paras)
               if p.text.strip() == "HEATING & PUMPING UNIT"), None)
    if hi is None:
        return False
    bullets, started = [], False
    for i in range(hi + 1, len(paras)):
        pPr = paras[i]._element.find(qn("w:pPr"))
        is_b = pPr is not None and pPr.find(qn("w:numPr")) is not None
        if is_b:
            started = True
            bullets.append(paras[i])
        elif started:
            break
    if not bullets:
        return False

    template = bullets[0]._element
    prev = bullets[-1]._element
    for txt in lines:
        nb = copy.deepcopy(template)
        runs = nb.findall(qn("w:r"))
        if runs:
            ts = runs[0].findall(qn("w:t"))
            if ts:
                ts[0].text = txt
                for extra in ts[1:]:
                    extra.getparent().remove(extra)
            for r in runs[1:]:
                r.getparent().remove(r)
        prev.addnext(nb)
        prev = nb
    for p in bullets:
        p._element.getparent().remove(p._element)

    d.save(doc_path)
    return True


def fill_gas_train(doc_path, df):
    """Rewrite the main-burner gas-train bullet list for fuels that have NO
    packaged gas train (BFG / COG / Producer Gas): list the BOM's itemized
    components with quantities. Packaged fuels (NG/PNG — a single 'Gas Train'
    BOM row) keep the standard breakdown bullets untouched.
    """
    import copy

    gt = df[df["SECTION"] == "GAS TRAIN"]
    if gt.empty:
        return False
    # packaged train present -> leave the static bullets as-is
    if any("gas train" in str(n).lower() for n in gt["ITEM NAME"]):
        return False

    lines = []
    for _, r in gt.iterrows():
        q = int(round(float(r["QTY"])))
        name = str(r["ITEM NAME"]).strip()
        lines.append(name)   # no quantity / DN in the offer scope
    if not lines:
        return False

    d = Document(doc_path)
    paras = d.paragraphs
    hi = next((i for i, p in enumerate(paras)
               if "GAS TRAIN FOR MAIN BURNERS" in p.text), None)
    if hi is None:
        return False
    bullets, started = [], False
    for i in range(hi + 1, len(paras)):
        pPr = paras[i]._element.find(qn("w:pPr"))
        is_b = pPr is not None and pPr.find(qn("w:numPr")) is not None
        if is_b:
            started = True
            bullets.append(paras[i])
        elif started:
            break
    if not bullets:
        return False

    template = bullets[0]._element
    prev = bullets[-1]._element
    for txt in lines:
        nb = copy.deepcopy(template)
        runs = nb.findall(qn("w:r"))
        if runs:
            ts = runs[0].findall(qn("w:t"))
            if ts:
                ts[0].text = txt
                for extra in ts[1:]:
                    extra.getparent().remove(extra)
            for r in runs[1:]:
                r.getparent().remove(r)
        prev.addnext(nb)
        prev = nb
    for p in bullets:
        p._element.getparent().remove(p._element)

    d.save(doc_path)
    return True


def _makelist_category(name, is_oil=False):
    """Map a BOM item name to a clean MAKE-LIST display category, so many BOM
    lines collapse into one make-list row (e.g. every Ball Valve → 'Ball Valve').
    """
    n = (name or "").lower()
    rules = [
        ("burner with regenerator", "Regen Oil Burner" if is_oil else "Regen Gas Burner"),
        ("pilot burner",            "Pilot Burner"),
        ("sequence controller",     "Burner Sequence Controller"),
        ("burner controller",       "Burner Sequence Controller"),
        ("ignition transformer",    "Ignition Transformer"),
        ("uv sensor",               "UV Sensor"),
        ("pilot regulator",         "Pressure Regulator"),
        ("pressure regulator",      "Pressure Regulator"),
        ("solenoid",                "Solenoid Valve"),
        ("flexible hose",           "Flexible Hose / Pipe"),
        ("butterfly",               "Butterfly Valve"),
        ("ball valve",              "Ball Valve"),
        ("shut-off",                "Pneumatic Shut Off Valve"),
        ("shut off",                "Pneumatic Shut Off Valve"),
        ("oil control valve",       "Oil Control Valve"),
        ("control valve",           "Pneumatic / Control Valve"),
        ("oil flow meter",          "Oil Flow Meter"),
        ("flow meter",              "DPT / Flow Meter"),
        ("dpt",                     "DPT / Flow Meter"),
        ("orifice",                 "Orifice Plate"),
        ("pressure switch",         "Pressure Switch"),
        ("pressure gauge",          "Pressure Gauge"),
        ("in oil line",             "Oil Line Instrumentation"),
        ("thermocouple",            "Thermocouple with TT"),
        ("transmitter",             "Temperature Transmitter"),
        ("damper",                  "Damper"),
        ("blower",                  "Combustion Air Blower"),
        ("id fan",                  "ID Fan – Suction Blower"),
        ("plc",                     "PLC with HMI"),
        ("control panel",           "Control Panel"),
        ("gas train",               "Gas Train Components"),
        ("paperless recorder",      "Paperless Recorder"),
        ("heating & pumping",       "Heating & Pumping Unit"),
        ("pumping unit",            "Pumping Unit"),
    ]
    for key, label in rules:
        if key in n:
            return label
    return (name or "").strip()


def fill_make_list(doc_path, df):
    """Rewrite the MAKE LIST table (header 'ITEMS | MAKE') so it lists ONLY the
    item categories present in BOM ``df``, each with the make(s) actually used.

    Returns True if the table was found and filled, else False.
    """
    from collections import OrderedDict

    d = Document(doc_path)
    tbl = None
    for t in d.tables:
        if (len(t.columns) == 2 and t.rows
                and t.rows[0].cells[0].text.strip().upper() == "ITEMS"):
            tbl = t
            break
    if tbl is None:
        return False

    # category -> distinct makes, in BOM order
    is_oil = any(str(s).upper().startswith("OIL") for s in df["SECTION"].unique())
    cats = OrderedDict()
    for _, row in df.iterrows():
        make = str(row.get("MAKE", "") or "").strip()
        if not make:
            continue
        cat = _makelist_category(row["ITEM NAME"], is_oil)
        cats.setdefault(cat, [])
        if make not in cats[cat]:
            cats[cat].append(make)

    for r in list(tbl.rows[1:]):
        r._tr.getparent().remove(r._tr)

    for cat, makes in cats.items():
        ic, mc = _add_item_row(tbl)
        ip = ic.paragraphs[0]; _style_para(ip, WA.LEFT)
        _run(ip, cat, bold=True)
        mp = mc.paragraphs[0]; _style_para(mp, WA.LEFT)
        _run(mp, " / ".join(makes))

    d.save(doc_path)
    return True


def fill_bom_table(doc_path, df):
    """Rewrite the BILL OF MATERIAL table in ``doc_path`` from BOM DataFrame ``df``.

    Returns True if the table was found and filled, else False.
    """
    d = Document(doc_path)

    tbl = None
    for t in d.tables:
        if (len(t.columns) == 2 and t.rows
                and t.rows[0].cells[0].text.strip().lower().startswith("description")):
            tbl = t
            break
    if tbl is None:
        return False

    # keep the styled header row, drop every existing body row
    for r in list(tbl.rows[1:]):
        r._tr.getparent().remove(r._tr)

    # faithful dump: section header row, then its item rows, in BOM order
    for section in df["SECTION"].drop_duplicates().tolist():
        _, sp = _add_section_row(tbl)
        _run(sp, str(section), bold=True)
        sub = df[df["SECTION"] == section]
        for _, row in sub.iterrows():
            dc, qc = _add_item_row(tbl)
            dp = dc.paragraphs[0]; _style_para(dp, WA.LEFT)
            _run(dp, _desc(row["ITEM NAME"], row.get("SPECIFICATION", "")))
            qp = qc.paragraphs[0]; _style_para(qp, WA.CENTER)
            _run(qp, _qty_label(row["QTY"]))

    d.save(doc_path)
    return True
