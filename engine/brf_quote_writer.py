"""Write the billet reheating furnace offer.

Fills BRF_Offer_Template.docx from what the app computed, and rebuilds the
price schedule as the itemised list the breakup produces rather than a single
lump sum.

Text is replaced across all the w:t nodes of a paragraph rather than run by
run, because Word splits a line into runs wherever formatting or a spell-check
mark changes, and two of the addresses are hyperlinks whose runs python-docx's
paragraph.runs does not return at all.
"""
from __future__ import annotations

import os
from copy import deepcopy

from docx import Document
from docx.shared import Pt

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE_PATH = os.path.join(BASE_DIR, "BRF_Offer_Template.docx")

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _text_nodes(par):
    return par._p.findall(f".//{W}t")


def _fill_paragraph(par, mapping):
    """Replace every placeholder in one paragraph, across its text nodes."""
    nodes = _text_nodes(par)
    if not nodes:
        return 0
    joined = "".join(n.text or "" for n in nodes)
    if "{{" not in joined:
        return 0
    out, hits = joined, 0
    for key, val in mapping.items():
        for token in (f"{{{{ {key} }}}}", f"{{{{{key}}}}}"):
            if token in out:
                out = out.replace(token, str(val))
                hits += 1
    if not hits:
        return 0
    # Everything lands in the first node; the rest are emptied so the text is
    # not duplicated. The first node carries the paragraph's formatting.
    nodes[0].text = out
    for n in nodes[1:]:
        n.text = ""
    return hits


def _all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for sec in doc.sections:
        for part in (sec.header, sec.footer,
                     sec.even_page_header, sec.even_page_footer,
                     sec.first_page_header, sec.first_page_footer):
            if part is None:
                continue
            for p in part.paragraphs:
                yield p
            for tb in part.tables:
                for row in tb.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p


def _inr(v):
    """Indian grouping, two decimals, as the offer writes its figures."""
    try:
        v = float(v)
    except (TypeError, ValueError):
        return str(v)
    neg = v < 0
    whole, frac = divmod(round(abs(v) * 100), 100)
    s = str(whole)
    if len(s) > 3:
        head, tail = s[:-3], s[-3:]
        parts = []
        while len(head) > 2:
            parts.insert(0, head[-2:])
            head = head[:-2]
        if head:
            parts.insert(0, head)
        s = ",".join(parts + [tail])
    return f"{'-' if neg else ''}{s}.{frac:02d}"


def _price_table(doc):
    """The price schedule: the table whose header names a unit price."""
    for tb in doc.tables:
        head = " ".join((c.text or "").upper() for c in tb.rows[0].cells)
        if "UNIT PRICE" in head and "DESCRIPTION" in head:
            return tb
    return None


def _fill_row(row, values, bold=False):
    for cell, val in zip(row.cells, values):
        para = cell.paragraphs[0]
        for extra in cell.paragraphs[1:]:
            extra._p.getparent().remove(extra._p)
        for r in list(para.runs):
            r._r.getparent().remove(r._r)
        run = para.add_run(str(val))
        run.bold = bold
        run.font.size = Pt(10)


def write_price_schedule(doc, rows, sell_total, currency="INR"):
    """Replace the single lump-sum line with the itemised list.

    rows are the breakup's: [group, item, qty, uom, unit, cost, markup, sell].
    The customer sees the selling price, never the cost or the markup.
    """
    tb = _price_table(doc)
    if tb is None or len(tb.rows) < 2:
        return 0
    template_row = tb.rows[1]
    written = 0
    for i, r in enumerate(rows, start=1):
        _, item, qty, uom, _unit, _cost, _mk, sell = r
        if not sell:
            continue                      # a line carried at nil is not offered
        written += 1
        new = deepcopy(template_row._tr)
        template_row._tr.addprevious(new)
        from docx.table import _Row
        _fill_row(_Row(new, tb),
                  [f"{written}.", item,
                   f"{qty:,.0f} {uom}" if qty else uom,
                   _inr(sell / qty) if qty else _inr(sell),
                   _inr(sell)])
    # the original row becomes the total
    _fill_row(template_row,
              ["", f"Total ({currency})", "", "", _inr(sell_total)], bold=True)
    return written


def generate_brf_quote_docx(data: dict, output_path: str,
                            template_path: str | None = None) -> str:
    """data — customer details plus the computed blocks from /api/brf-calculate."""
    doc = Document(template_path or TEMPLATE_PATH)

    cust = data.get("customer") or {}
    fur = data.get("furnace") or {}
    duty = ((data.get("sizing") or {}).get("duty")) or {}
    zones = ((data.get("sizing") or {}).get("zones")) or []
    comb = data.get("combustion_calc") or {}
    recup = data.get("recuperator_calc") or {}
    breakup = data.get("breakup") or {}

    zone_count = int(duty.get("zone_count") or 0)
    fired = zones[:zone_count]

    blower_n = int(comb.get("blower_count") or 0)
    blower_model = comb.get("blower_model") or ""
    blower_hp = comb.get("blower_installed_hp") or 0
    blower_desc = (f"{blower_n} no{'s' if blower_n != 1 else ''} "
                   f"{blower_model}, {blower_hp:,.0f} HP".strip(", ")
                   if blower_model else f"{blower_n} nos")

    billet = (f"{fur.get('billet_width_mm', 0):,.0f}mm² x "
              f"{fur.get('billet_length_mm', 0):,.0f}mm long")

    mapping = {
        # the parties
        "company_name": cust.get("company_name", ""),
        "company_address": cust.get("company_address", ""),
        "company_city_state": ", ".join(
            x for x in (cust.get("company_city"), cust.get("company_state")) if x),
        "poc_name": cust.get("poc_name", ""),
        "email": cust.get("email", ""),
        "mobile_no": cust.get("mobile_no", ""),
        "your_ref": cust.get("your_ref", ""),
        "ref_no": cust.get("ref_no", ""),
        "quote_date": cust.get("quote_date", ""),
        "marketing_person": cust.get("marketing_person", ""),
        "marketing_phone": cust.get("marketing_phone", ""),
        "marketing_email": cust.get("marketing_email", ""),
        # the furnace
        "capacity": f"{fur.get('furnace_capacity_tph', 0):,.0f}",
        "billet_size": billet,
        "eff_length_mm": f"{fur.get('effective_length_mm', 0):,.0f}",
        "eff_width_mm": f"{fur.get('effective_width_mm', 0):,.0f}",
        "overall_length_mm": f"{fur.get('overall_length_mm', 0):,.0f}",
        "overall_width_mm": f"{fur.get('overall_width_mm', 0):,.0f}",
        "cv": f"{duty.get('cv_kcal_nm3', 0):,.0f}",
        "blower_desc": blower_desc,
        "flue_gas_nm3hr": f"{recup.get('flue_gas_nm3hr', 0):,.0f}",
        "zone_count": _words(zone_count),
        "zone_split": _zone_split(fired),
    }
    for i in range(1, 6):
        z = fired[i - 1] if i - 1 < len(fired) else None
        mapping[f"zone{i}_burners"] = (z or {}).get("burner_count", "") if z else ""

    filled = sum(_fill_paragraph(p, mapping) for p in _all_paragraphs(doc))
    lines = write_price_schedule(doc, breakup.get("rows") or [],
                                 breakup.get("sell_total") or 0.0,
                                 cust.get("currency") or "INR")

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    return output_path


_NUM_WORDS = {1: "One", 2: "Two", 3: "Three", 4: "Four", 5: "Five",
              6: "Six", 7: "Seven", 8: "Eight", 9: "Nine", 10: "Ten"}


def _words(n):
    return _NUM_WORDS.get(int(n or 0), str(n))


def _zone_split(fired):
    """"Two Soaking Zones; Three Heating Zones", from the zone names."""
    soak = sum(1 for z in fired if "soak" in str(z.get("name", "")).lower())
    heat = sum(1 for z in fired if "heat" in str(z.get("name", "")).lower())
    if not soak and not heat:            # generic names — split as the sheet does
        soak = min(2, len(fired))
        heat = len(fired) - soak
    parts = []
    if soak:
        parts.append(f"{_words(soak)} Soaking Zone{'s' if soak != 1 else ''}")
    if heat:
        parts.append(f"{_words(heat)} Heating Zone{'s' if heat != 1 else ''}")
    return "; ".join(parts)
