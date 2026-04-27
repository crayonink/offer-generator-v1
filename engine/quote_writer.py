import os
from export.word_writer import generate_word_offer

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE_PATH = os.path.join(BASE_DIR, "Offer_Template.docx")
TUNDISH_TEMPLATE_PATH = os.path.join(BASE_DIR, "Tundish_Offer_Template.docx")


# ── MAKE LIST (static) ────────────────────────────────────────────────────
# Fixed vendor-make table that ships in every VLPH / HLPH offer. Per
# business rule the list does not vary with the BOM contents.
STATIC_MAKE_LIST = [
    {"item": "MILD STEEL",                    "make": "SAIL, JINDAL"},
    {"item": "PRESSURE GAUGE WITH TNV",       "make": "WIKA"},
    {"item": "PRESSURE SWITCH",               "make": "MADAS"},
    {"item": "BALL VALVE",                    "make": "AUDCO/ L&T/ LEADER"},
    {"item": "PNEUMATIC CONTROL VALVE",       "make": "MADAS"},
    {"item": "BUTTERFLY VALVE",               "make": "AUDCO/ L&T/ LEADER"},
    {"item": "ROTARY JOINT",                  "make": "ENCON"},
    {"item": "MIX GAS TRAIN",                 "make": "MADAS"},
    {"item": "FLEXIBALE HOSE PIPE",           "make": "BIL/FLEXIBLE"},
    {"item": "THERMOCOUPLE",                  "make": "TEMPSENS"},
    {"item": "COMPENSATING LEAD",             "make": "TEMPSENS"},
    {"item": "LIMIT SWITCHES",                "make": "BCH"},
    {"item": "CONTROL PANEL",                 "make": "ENCON"},
    {"item": "HYDRAULIC POWER PACK & CYLINDER", "make": "VARITECH"},
    {"item": "TEMPERATURE TRANSMITTER",       "make": "HONEYWELL"},
    {"item": "BURNER",                        "make": "ENCON"},
    {"item": "PILOT BURNER",                  "make": "ENCON"},
    {"item": "BLOWER",                        "make": "ENCON"},
    {"item": "BEARING",                       "make": "FAG/SKF"},
    {"item": "UV SENSOR",                     "make": "LINEAR"},
    {"item": "IGNITION TRANSFORMER",          "make": "COFI/DANFOSS"},
    {"item": "SEQUENCE CONTROLLER",           "make": "LINEAR"},
    {"item": "MOTOR",                         "make": "ABB"},
    {"item": "GEARBOX",                       "make": "POWERTEK"},
    {"item": "P.PID",                         "make": "HONEYWELL"},
    {"item": "RATIO CONTROLLER",              "make": "HONEYWELL"},
    {"item": "AIR GAS RAGULATOR",             "make": "MADAS"},
]


# ── Indian-English number-to-words (lakh / crore system) ──────────────────
_ONES = ("", "ONE", "TWO", "THREE", "FOUR", "FIVE", "SIX", "SEVEN", "EIGHT",
         "NINE", "TEN", "ELEVEN", "TWELVE", "THIRTEEN", "FOURTEEN", "FIFTEEN",
         "SIXTEEN", "SEVENTEEN", "EIGHTEEN", "NINETEEN")
_TENS = ("", "", "TWENTY", "THIRTY", "FORTY", "FIFTY", "SIXTY", "SEVENTY",
         "EIGHTY", "NINETY")


def _two_digits(n: int) -> str:
    if n < 20:
        return _ONES[n]
    t, o = divmod(n, 10)
    return _TENS[t] + ("-" + _ONES[o] if o else "")


def _three_digits(n: int) -> str:
    h, r = divmod(n, 100)
    parts = []
    if h:
        parts.append(_ONES[h] + " HUNDRED")
    if r:
        parts.append(_two_digits(r))
    return " ".join(parts)


def amount_in_words_indian(amount) -> str:
    """Format an integer rupee amount in Indian English words (lakh / crore)."""
    try:
        n = int(round(float(amount)))
    except (TypeError, ValueError):
        return ""
    if n == 0:
        return "ZERO"
    crore, rem  = divmod(n, 10000000)
    lakh,  rem  = divmod(rem, 100000)
    thou,  rem  = divmod(rem, 1000)
    parts = []
    if crore:
        parts.append(_three_digits(crore) + " CRORE")
    if lakh:
        parts.append(_two_digits(lakh) + " LAKH")
    if thou:
        parts.append(_two_digits(thou) + " THOUSAND")
    if rem:
        parts.append(_three_digits(rem))
    return " ".join(parts).strip()


def _supervision_rates() -> tuple:
    """Look up supervision-charge rates from component_price_master.
    Returns (mech_rate_str, plc_rate_str) formatted as Indian rupee strings."""
    import sqlite3
    db_path = os.path.join(BASE_DIR, "vlph.db")
    rates = {"mech": 12500, "plc": 15000}  # safe defaults
    try:
        conn = sqlite3.connect(db_path)
        for key, item in (("mech", "SUPERVISION CHARGE - MECHANICAL"),
                          ("plc",  "SUPERVISION CHARGE - PLC & INSTRUMENTATION")):
            row = conn.execute(
                "SELECT price FROM component_price_master WHERE item=? LIMIT 1",
                (item,)
            ).fetchone()
            if row:
                rates[key] = float(row[0])
        conn.close()
    except Exception:
        pass
    return f"{rates['mech']:,.2f}", f"{rates['plc']:,.2f}"


def _combine_dual(val1, val2) -> str:
    """Combine two fuel values into one string for single-column display.
    If val2 is empty, returns val1 only."""
    v1 = (val1 or "").strip()
    v2 = (val2 or "").strip()
    if v2:
        return f"{v1}\n{v2}"
    return v1


def _fuel_label(fuel_name: str) -> str:
    """Extract the first fuel name from 'NG & MG' or 'Natural Gas & Furnace Oil'."""
    if "&" in fuel_name:
        return fuel_name.split("&")[0].strip()
    return fuel_name.strip()


def _fuel_label2(fuel_name: str) -> str:
    """Extract the second fuel name from 'NG & MG'."""
    if "&" in fuel_name:
        return fuel_name.split("&")[1].strip()
    return ""


import re as _re_pilot


def _rewrite_pilot_name(item: str, pilot_gas: str) -> str:
    """The BOM carries one fixed pilot-burner model. Rewrite its name so the
    MAKE LIST reflects the fuel the user actually picked for the pilot.
    Only rewrite rows whose model starts with 'ENCON-PB' / 'ENCON PB' — do NOT
    match items like 'BALL VALVE (Pilot Burner)' that merely mention the pilot."""
    if not item:
        return item
    s = item.strip()
    upper = s.upper()
    if not (upper.startswith("ENCON-PB") or upper.startswith("ENCON PB")):
        return s
    # Detect KW rating by a number directly followed by KW (word-boundary).
    m = _re_pilot.search(r"\b(\d+)\s*KW\b", upper)
    rating = f"{m.group(1)} KW" if m else "100 KW"
    return f"ENCON-PB {pilot_gas} {rating}"


def _build_equipment_name(customer, quote_data):
    """Derive equipment name from items product_type or customer fields."""
    # Check items for product type
    for item in quote_data.get("items", []):
        pt = (item.get("product_type") or "").lower()
        if "tundish" in pt:
            tons = customer.get("ladle_tons")
            return f"Tundish Preheater – {tons} Ton" if tons else "Tundish Preheater"
        elif "horizontal" in pt:
            tons = customer.get("ladle_tons")
            return f"Horizontal Ladle Preheater – {tons} Ton" if tons else "Horizontal Ladle Preheater"
    # Default: Vertical Ladle Preheater
    tons = customer.get("ladle_tons")
    if tons:
        return f"Vertical Ladle Preheater – {tons} Ton"
    return customer.get("project_name") or ""


# Oil fuels where the oil is pre-heated separately (no heating element in
# the pumping unit). Anything else oil-based gets the heater (HSD, SKO, HDO,
# CFO etc.).
PUMPING_UNIT_ONLY_FUEL_NAMES = {"LDO", "FO", "FURNACE OIL", "LSHS"}


# ─── Annexure I scope-of-supply dynamic strings ────────────────────────────
def _ignition_scope_text(special_auto_ignition: bool, pilot_gas_type: str | None) -> str:
    if not special_auto_ignition:
        return "Manual ignition of the Burner"
    pilot = (pilot_gas_type or "LPG").upper()
    return f"Firing & Ignition of the Burner is Automatic through {pilot} fired Pilot Burner"


def _temp_control_scope_text(control_mode: str | None, auto_control_type: str | None) -> str:
    if (control_mode or "").lower() == "manual":
        return "Manual temperature control"
    act = (auto_control_type or "plc").lower()
    if act == "plc":     return "Temperature Control System with PLC"
    if act == "plc_agr": return "Temperature Control System with PLC + AGR"
    if act == "pid":     return "Temperature Control System with PID Controller"
    return "Temperature Control System"


def _flow_meter_scope_text(is_oil: bool, is_dual: bool) -> str:
    if is_oil or is_dual:
        return "Flow meter and Control valve in Oil and Air line"
    return "Control valve on Air line"


def _pipeline_scope_text(is_oil: bool, is_dual: bool) -> str:
    if is_oil or is_dual:
        return "Interconnecting pipelines from Pumping unit and blower to burner"
    return "Interconnecting pipelines from gas train and blower to burner"


def _pumping_unit_block(fuel_name: str, is_oil: bool, is_dual: bool):
    """Decide heading + intro + bullets for the oil-side pumping section.

    Returns (heading, intro, bullets_list). Heading is empty when neither
    is_oil nor is_dual is set, so the template can hide the section.
    """
    if not (is_oil or is_dual):
        return "", "", []

    name = (fuel_name or "").upper().strip()
    pumping_only = any(token in name for token in PUMPING_UNIT_ONLY_FUEL_NAMES)
    if pumping_only:
        heading = "PUMPING UNIT"
        intro = ("To supply fuel oil to the above burner at the requisite "
                 "pressure, we will supply a Pumping Unit consisting of "
                 "the following:")
        bullets = [
            "2 Nos. oil pumps each fitted with suitable electric motor.",
            "1 No. each Duplex type coarse and fine filter for the cold and hot oil side respectively.",
            "1 No. Pressure regulating valve.",
            "1 No. Pressure gauge.",
        ]
    else:
        heading = "HEATING & PUMPING UNIT"
        intro = ("To supply fuel oil to the above burner at the requisite "
                 "pressure and temperature, we will supply a complete Heating "
                 "& Pumping Unit consisting of the following:")
        bullets = [
            "2 Nos. oil pumps each fitted with suitable electric motor.",
            "1 No. each Duplex type coarse and fine filter for the cold and hot oil side respectively.",
            "1 No. Pressure regulating valve.",
            "1 No. each of Pressure gauge & Temperature gauge.",
            "Electric heater with thermostat.",
        ]
    return heading, intro, bullets


def _temp_control_items_for_mode(control_mode: str, auto_control_type: str) -> list:
    """Return the static numbered list of Temperature Control System items
    for the given control mode. Drives the {%p for x in temp_control_items %}
    loop in the Word template (and the PDF builder). Replaces the previous
    BOM-derived list so every mode has fixed, well-worded content."""
    cm  = (control_mode or "automatic").lower()
    act = (auto_control_type or "plc").lower()
    if cm == "manual":
        items = [
            "Thermocouple with temperature transmitter",
            "Air-Gas Ratio (AGR) regulator on the gas line",
        ]
    elif act == "pid":
        items = [
            "P.PID temperature controller",
            "Ratio Controller",
            "Thermocouple with temperature transmitter",
            "Air-Gas Ratio (AGR) regulator on the gas line",
        ]
    elif act == "plc_agr":
        items = [
            "PLC with HMI",
            "Thermocouple with temperature transmitter",
            "Air-Gas Ratio (AGR) regulator on the gas line",
        ]
    else:
        # Default: PLC
        items = [
            "PLC with HMI",
            "Thermocouple with temperature transmitter",
            "Orifice plate fitted with mass flow transmitter on gas line",
            "Orifice plate fitted with differential pressure transmitter on air line",
        ]
    return [{"item": x, "ref": ""} for x in items]


def _operational_sequence_text(control_mode: str, auto_control_type: str) -> str:
    """Mode-specific wording for the OPERATIONAL SEQUENCE paragraph in the
    offer document. Mirrored in pdf_writer.py."""
    cm  = (control_mode or "automatic").lower()
    act = (auto_control_type or "plc").lower()
    if cm == "manual":
        return ("The temperature of the ladle will be monitored manually through the "
                "temperature indicator fitted on the panel. The operator will open / "
                "close the air and gas control valves to maintain the desired "
                "temperature profile of the ladle as per the heating schedule.")
    if act == "pid":
        return ("The temperature of the ladle will be controlled automatically through "
                "a PID controller. The thermocouple fitted in the ladle will sense the "
                "temperature and feed it to the PID controller. The PID controller will "
                "modulate the air control valve via the Air-Gas Ratio regulator to "
                "maintain the air/gas ratio as the temperature rises / falls to the "
                "set values.")
    if act == "plc_agr":
        return ("The temperature of the ladle will be controlled automatically through "
                "P.L.C. The thermocouple fitted in the ladle will sense the temperature "
                "and signal the P.L.C. The P.L.C will modulate the air control valve "
                "through the Air-Gas Ratio regulator and accordingly the gas flow will "
                "be controlled to maintain the air/gas ratio as the temperature "
                "decreases / increases to the set values.")
    # Default: PLC
    return ("The temperature of the ladle will be controlled automatically through "
            "P.L.C. The thermocouple fitted in the ladle will sense the temperature "
            "and will give signal to the P.L.C. The P.L.C will send a signal to the "
            "control valve fitted on the airline; the air control valve will be "
            "modulated and accordingly the gas flow will be controlled, maintaining "
            "the mass-flow air/gas ratio as the temperature decreases / increases to "
            "the set values.")


def _control_system_sections(bom_items: list) -> dict:
    """Group BOM rows into the lists the offer doc renders.

    Buckets:
      - gas_pipeline_items   - main fuel-line rows (combined; for single-fuel
                                 offers, == fuel1_line_items).
      - fuel1_line_items     - rows whose MEDIA matches fuel 1 (e.g. "NG LINE").
      - fuel2_line_items     - rows whose MEDIA matches fuel 2 (e.g. "MG LINE")
                                 -- empty for single-fuel offers.
      - air_pipeline_items   - combustion-air rows that AREN'T pilot/UV
                                 accessories (those go to the pilot bucket).
      - pilot_pipeline_items - dedicated pilot-line rows (LPG/NG PILOT LINE)
                                 PLUS air-side accessories whose name ends
                                 in "(Pilot Burner)" or "(UV LINE)" --
                                 those belong with the pilot equipment, not
                                 the main air line.
      - nitrogen_purging_items - rows with MEDIA == "PURGING LINE".
      - temp_control_items   - rows with MEDIA == "MISC ITEMS".

    Fuel-line items are split by MEDIA value: the first distinct fuel-line
    MEDIA seen goes to fuel1_line_items, the second to fuel2_line_items.

    Each list item is {"item": ..., "ref": ...} for easy templating.
    """
    import re as _re
    def _clean_name(name: str) -> str:
        """Strip sizing / flow info baked into the item name itself, so the
        offer never shows things like 'GAS TRAIN 500 NM3/Hr' or
        'COMPENSATOR - 250 NB F150#' in the bullet list."""
        n = (name or "").strip()
        # 'GAS TRAIN 500 NM3/Hr ...' -> 'GAS TRAIN'
        n = _re.sub(r"^(GAS\s+TRAIN)\b.*$", r"\1", n, flags=_re.IGNORECASE)
        # Strip trailing ' - <anything>' (catches '- 250 NB F150#', '- 80 NB',
        # '- RANGE- 0-1600 mBAR', etc.)
        n = _re.sub(r"\s+[\-–—]\s+.*$", "", n)
        return n.strip()

    def _fmt(x):
        return {"item": _clean_name(x.get("item", "")), "ref": ""}

    gas, air, pilot, temp, fuel1, fuel2, purging = [], [], [], [], [], [], []
    fuel1_media = None  # The first distinct fuel-line MEDIA we encounter.
    for x in bom_items:
        media = (x.get("media") or "").strip().upper()
        item  = (x.get("item") or "").strip()
        if not item:
            continue
        upper          = item.upper()
        is_pilot_named = "(PILOT BURNER)" in upper or "(UV LINE)" in upper
        is_pilot_media = media.endswith(" PILOT LINE")

        if is_pilot_named or is_pilot_media:
            pilot.append(_fmt(x))
        elif media == "COMB AIR":
            air.append(_fmt(x))
        elif media == "MISC ITEMS":
            temp.append(_fmt(x))
        elif media == "PURGING LINE":
            purging.append(_fmt(x))
        elif media.endswith(" LINE"):
            gas.append(_fmt(x))
            if fuel1_media is None:
                fuel1_media = media
                fuel1.append(_fmt(x))
            elif media == fuel1_media:
                fuel1.append(_fmt(x))
            else:
                fuel2.append(_fmt(x))
        # BOUGHT OUT ITEMS, ENCON ITEMS etc. are ignored here.

    # Strip the trailing " LINE" suffix to surface the short fuel name (NG, MG, BG, COG, LDO...)
    fuel1_label = fuel1_media[:-5] if fuel1_media and fuel1_media.endswith(" LINE") else (fuel1_media or "")
    # The second fuel media is whichever we collected into fuel2; pick from any item.
    fuel2_media = ""
    for x in bom_items:
        m = (x.get("media") or "").strip().upper()
        if m.endswith(" LINE") and m != "PURGING LINE" and m != fuel1_media:
            fuel2_media = m
            break
    fuel2_label = fuel2_media[:-5] if fuel2_media.endswith(" LINE") else fuel2_media

    return {
        "gas_pipeline_items":     gas,
        "fuel1_line_items":       fuel1,
        "fuel2_line_items":       fuel2,
        "fuel1_line_label":       fuel1_label,
        "fuel2_line_label":       fuel2_label,
        "air_pipeline_items":     air,
        "pilot_pipeline_items":   pilot,
        "nitrogen_purging_items": purging,
        "temp_control_items":     temp,
    }


def generate_quote_docx(quote_data: dict, output_path: str):
    customer = quote_data["customer"]
    sup_mech, sup_plc = _supervision_rates()

    # Determine vertical vs horizontal UNIT price from items (per-set sum).
    unit_vertical   = 0.0
    unit_horizontal = 0.0
    for item in quote_data.get("items", []):
        pt = (item.get("product_type") or "").lower()
        if "vertical" in pt:
            unit_vertical += item.get("total", 0)
        elif "horizontal" in pt:
            unit_horizontal += item.get("total", 0)
        else:
            unit_vertical += item.get("total", 0)   # default to vertical

    # Annexure III qty multipliers (Step-4 inputs); qty=0 means absent type.
    qty_v = int(customer.get("vertical_qty") or 0)
    qty_h = int(customer.get("horizontal_qty") or 0)
    # If neither was specified, fall back to legacy behaviour (1 vertical set).
    if qty_v == 0 and qty_h == 0:
        qty_v = 1
    total_vertical   = unit_vertical   * qty_v
    total_horizontal = unit_horizontal * qty_h

    def _qty_label(n: int) -> str:
        if n <= 0: return ""
        return f"{n:02d} {'Set' if n == 1 else 'Sets'}"

    # Map to template variable names (must match {{...}} in Offer_Template.docx)
    context = {
        # Customer / contact
        "project_name":          customer.get("project_name") or "",
        "subject":               customer.get("subject") or customer.get("project_name") or "",
        "company_name":          customer.get("company_name") or "",
        "company_city":          customer.get("company_city") or "",
        "company_state":         customer.get("company_state") or "",
        "company_address":       customer.get("address") or "",
        "mobile_no":             customer.get("mobile_no") or "",
        "email":                 customer.get("email") or "",
        "gstin":                 customer.get("gstin") or "",
        "poc_name":              customer.get("poc_name") or "",
        "poc_designation":       customer.get("poc_designation") or "",
        # Reference / enquiry
        "quote_no":              quote_data.get("quote_no", ""),
        "date":                  quote_data.get("date", ""),
        "ref_no":                customer.get("ref_no", ""),
        "your_ref":              customer.get("your_ref") or customer.get("ref_no", ""),
        "enquiry_ref":           customer.get("enquiry_ref") or customer.get("ref_no", ""),
        "marketing_person":      customer.get("marketing_person", ""),
        "marketing_phone":       customer.get("marketing_phone", ""),
        "marketing_email":       customer.get("marketing_email", ""),
        "technical_person":      customer.get("technical_person", ""),
        "technical_phone":       customer.get("technical_phone", ""),
        "technical_email":       customer.get("technical_email", ""),
        # Pricing
        "unit_price_vertical":   f"{unit_vertical:,.2f}",
        "unit_price_horizontal": f"{unit_horizontal:,.2f}" if unit_horizontal else "N/A",
        "total_price_vertical":  f"{total_vertical:,.2f}",
        "total_price_horizontal": f"{total_horizontal:,.2f}" if total_horizontal else "N/A",
        "qty_label_vertical":    _qty_label(qty_v),
        "qty_label_horizontal":  _qty_label(qty_h),
        "total_in_words": (
            customer.get("total_in_words")
            or amount_in_words_indian(total_vertical + total_horizontal) + " ONLY"
        ),
        "equipment_name": _build_equipment_name(customer, quote_data),
        # Supervision-charge rates (from component_price_master)
        "supervision_mech": sup_mech,
        "supervision_plc":  sup_plc,
        "subtotal":              f"₹ {quote_data.get('subtotal', 0):,.0f}",
        "grand_total":           f"₹ {quote_data.get('grand_total', 0):,.0f}",
        "valid_days":            quote_data.get("valid_days", 30),
        "items":                 quote_data.get("items", []),
        # Technical data (populates the Tech Data table in the template)
        "ladle_tons":            customer.get("ladle_tons") or "",
        "ladle_dim":             customer.get("ladle_dim") or "",
        "ladle_drawing_no":      customer.get("ladle_drawing_no") or "",
        "refractory_weight_kg":  customer.get("refractory_weight_kg") or "",
        "heating_schedule":      customer.get("heating_schedule") or "",
        "fuel_cv":               _combine_dual(customer.get("fuel_cv"), customer.get("fuel2_cv")),
        "fuel_consumption":      _combine_dual(customer.get("fuel_consumption"), customer.get("fuel2_consumption")),
        "burner_model":          customer.get("burner_model") or "",
        # Display name shown in the 'ENCON GAS/OIL/DUAL BURNER' body paragraphs.
        # Prefixed per the ENCON pricelist: oil -> 'ENCON 7A', gas -> 'ENCON -G 7A',
        # dual -> 'ENCON DUAL- 7A'. Falls back to burner_model if empty.
        "burner_display_model":  (
            customer.get("burner_display_model")
            or (
                customer.get("burner_model", "").replace("ENCON ", "ENCON -G ", 1)
                if (not bool(customer.get("is_oil")) and not bool(customer.get("is_dual")))
                else customer.get("burner_model", "").replace("ENCON ", "ENCON DUAL- ", 1)
                     if bool(customer.get("is_dual"))
                else customer.get("burner_model", "")
            )
        ),
        "burner_firing_rate":    (
            customer.get("burner_firing_rate")
            or customer.get("fuel_consumption", "")
        ),
        # Just the kW value, with any 'Fuel name: ' prefix stripped — used in the
        # ENCON Burner body line ("with a firing rate of 2,148 kW.").
        "burner_kw": (
            (customer.get("fuel_consumption") or "").split(":", 1)[-1].strip()
        ),
        "blower_model":          customer.get("blower_model") or "",
        "blower_size":           customer.get("blower_size") or "",
        "blower_capacity":       customer.get("blower_capacity") or "",
        "hydraulic_motor_hp":    customer.get("hydraulic_motor_hp") or "",
        "max_electrical_load":   customer.get("max_electrical_load") or "",
        # Extra tech-data fields for PLC template
        "heating_time":          customer.get("heating_time") or "",
        "fuel_name":             customer.get("fuel_name") or "",
        "burner_capacity_range": customer.get("burner_capacity_range") or "",
        "pumping_unit":          customer.get("pumping_unit") or "",
        "hood_movement":         customer.get("hood_movement") or "Vertical Swiveling through bearing mechanism.",
        "hood_type":             customer.get("hood_type") or "up_down",
        # Pilot fields are blanked when Auto Ignition is not selected so the
        # tech-data post-processor (_strip_empty_tech_rows) drops the
        # "Pilot Burner Fuel" and "Firing & Ignition of Burner" rows.
        "pilot_gas_type":        (customer.get("pilot_gas_type") or "LPG")
                                  if bool(customer.get("special_auto_ignition")) else "",
        "ignition_method":       (customer.get("ignition_method")
                                  or "Automatic Through LPG Fired Pilot Burner")
                                  if bool(customer.get("special_auto_ignition")) else "",
        # Tundish-specific tech-data placeholders (pre-heating station)
        "fuel1_label":           _fuel_label(customer.get("fuel_name", "")),
        "fuel2_label":           _fuel_label2(customer.get("fuel_name", "")),
        "num_burners":           customer.get("num_burners") or "",
        "fuel1_cv":              customer.get("fuel_cv") or "",
        "fuel2_cv":              customer.get("fuel2_cv") or "",
        "max_temperature":       customer.get("heating_schedule") or "",
        "cycle_time":            customer.get("heating_time") or "",
        "lifting_lowering":      customer.get("hood_movement") or "With Hydraulic cylinder & Power pack",
        "firing_rate_fuel1":     customer.get("fuel_consumption") or "",
        "max_consumption_fuel1": customer.get("max_fuel_consumption1") or "",
        "firing_rate_fuel2":     customer.get("fuel2_consumption") or "",
        "max_consumption_fuel2": customer.get("max_fuel_consumption2") or "",
        "combustion_blower":     "Centrifugal type",
        "motor_power_pack":      customer.get("hydraulic_motor_hp") or "10 HP",
        # Tundish drying station — defaults to same as pre-heating (user edits in Word)
        "max_temperature_dry":       customer.get("max_temperature_dry") or customer.get("heating_schedule") or "",
        "cycle_time_dry":            customer.get("cycle_time_dry") or customer.get("heating_time") or "",
        "firing_rate_fuel1_dry":     customer.get("firing_rate_fuel1_dry") or customer.get("fuel_consumption") or "",
        "max_consumption_fuel1_dry": customer.get("max_consumption_fuel1_dry") or customer.get("max_fuel_consumption1") or "",
        "firing_rate_fuel2_dry":     customer.get("firing_rate_fuel2_dry") or customer.get("fuel2_consumption") or "",
        "max_consumption_fuel2_dry": customer.get("max_consumption_fuel2_dry") or customer.get("max_fuel_consumption2") or "",
        "blower_size_dry":           customer.get("blower_size_dry") or customer.get("blower_capacity") or "",
        "max_electrical_load_dry":   customer.get("max_electrical_load_dry") or customer.get("max_electrical_load") or "",
        # Fuel-type flags drive {%p if is_oil %} / {%p if is_gas %} blocks
        # in the scope-of-supply section of the template.
        # Mutually-exclusive burner-type flags: exactly one of is_oil/is_gas/is_dual is true.
        "is_dual":               bool(customer.get("is_dual")),
        "is_oil":                bool(customer.get("is_oil")) and not bool(customer.get("is_dual")),
        "is_gas":                not bool(customer.get("is_oil")) and not bool(customer.get("is_dual")),
        # Product-type flags drive {%p if is_vertical %} / {%p if is_horizontal %}
        # blocks that wrap the two scope-of-supply variants in the template.
        # Derived from items[*].product_type — vertical wins by default for
        # dual-product quotes (rare).
        "is_vertical":   any("vertical"   in (it.get("product_type") or "").lower()
                              for it in quote_data.get("items", [])),
        "is_horizontal": any("horizontal" in (it.get("product_type") or "").lower()
                              for it in quote_data.get("items", [])),
        # Label used in 'On the main ___ pipeline:' heading. Dual fuel keeps 'oil'
        # since the oil-line components (flow meter, manual ball valve, etc.) still apply.
        "fuel_line_label":       (
            "gas" if (not bool(customer.get("is_oil")) and not bool(customer.get("is_dual")))
            else "oil"
        ),
        # Full PIPELINE sentence tail — 'gas pipeline from gas train to burner' for
        # gas-only, otherwise 'LDO pipeline from LDO pumping unit to burner'.
        "fuel_delivery_text":    (
            "gas pipeline from gas train to burner"
            if (not bool(customer.get("is_oil")) and not bool(customer.get("is_dual")))
            else "LDO pipeline from LDO pumping unit to burner"
        ),
        # ─── Scope of Supply (Annexure I) parameters ─────────────────────────
        # Fuel tag next to 'Burner with Burner block & mounting Plate –'
        "fuel_short": (
            "NG"     if (not bool(customer.get("is_oil")) and not bool(customer.get("is_dual")))
            else "LDO/NG" if bool(customer.get("is_dual"))
            else "LDO"
        ),
        # Burner product name shown in row 3 of the scope table
        "burner_scope_name": (
            "ENCON Dual-Fuel Burner" if bool(customer.get("is_dual"))
            else "ENCON Gas Burner"  if not bool(customer.get("is_oil"))
            else "ENCON IIP Film Burner"
        ),
        # Row 5 — fuel delivery unit
        "pumping_scope_text": (
            "Gas train with pressure regulating and control valves"
            if (not bool(customer.get("is_oil")) and not bool(customer.get("is_dual")))
            else "LDO pumping unit with micro valve"
        ),
        # Row 6 (vertical column) — hood mechanism
        "hood_scope_vertical": (
            "Swiveling mechanism through bearing with mechanical locking"
            if (customer.get("hood_type") in ("swivel", "swivel_manual", "swivel_geared"))
            else "Hydraulic Cylinder for lifting and Lowering of the system"
        ),
        # Control-mode flags drive {%p if is_manual %} / {%p if is_automatic %}
        # for temperature-control and control-panel scope sections.
        "is_manual":             customer.get("control_mode") == "manual",
        "is_automatic":          customer.get("control_mode") != "manual",
        # Special Requirements flags — drive Pilot Burner sections in the
        # template ({%p if special_auto_ignition %}). When auto-ignition is
        # not requested, the entire pilot-burner + pilot-line scope is hidden.
        "special_auto_ignition": bool(customer.get("special_auto_ignition")),
        "special_auto_controls": bool(customer.get("special_auto_controls")),
        # Nitrogen Purging block — only when user toggled "purging_line=yes" on Step 4.
        "nitrogen_purging":      bool(customer.get("nitrogen_purging")),
        # Raw control fields, exposed for inline Jinja in the Tech Specs table
        "control_mode":          (customer.get("control_mode") or "automatic"),
        "auto_control_type":     (customer.get("auto_control_type") or "plc"),
        # MAKE LIST is a fixed 27-row vendor table — same for vertical and
        # horizontal preheaters. Per business rule it does NOT vary with
        # the BOM contents.
        "make_list":             STATIC_MAKE_LIST,
        # Control-system sections driven by the BOM's MEDIA column.
        # Each list is [{item, ref}, ...]; the template renders one bullet per entry.
        # For dual-fuel offers the fuel-line items are split into
        # fuel1_line_items / fuel2_line_items by the BOM's MEDIA value.
        **_control_system_sections(customer.get("bom_items") or []),
        # Override temp_control_items with mode-specific static list (does NOT
        # come from BOM). PLC, PLC+AGR, PID and Manual each get their own.
        "temp_control_items": _temp_control_items_for_mode(
            customer.get("control_mode"), customer.get("auto_control_type")),
        # Mode-specific OPERATIONAL SEQUENCE paragraph wording.
        "operational_sequence_text": _operational_sequence_text(
            customer.get("control_mode"), customer.get("auto_control_type")),
        # Annexure I scope-of-supply dynamic strings
        "ignition_scope_text":     _ignition_scope_text(
            bool(customer.get("special_auto_ignition")),
            customer.get("pilot_gas_type")),
        "temp_control_scope_text": _temp_control_scope_text(
            customer.get("control_mode"), customer.get("auto_control_type")),
        "flow_meter_scope_text":   _flow_meter_scope_text(
            bool(customer.get("is_oil")), bool(customer.get("is_dual"))),
        "pipeline_scope_text":     _pipeline_scope_text(
            bool(customer.get("is_oil")), bool(customer.get("is_dual"))),
        # Annexure I section headers — qty is user-editable on Step 4
        "vertical_qty":   customer.get("vertical_qty") or 1,
        "horizontal_qty": customer.get("horizontal_qty") or 1,
        # Pumping-unit section (heading flips between "PUMPING UNIT" and
        # "HEATING & PUMPING UNIT" by fuel; only renders when is_oil/is_dual).
        **(lambda h, i, b: {
            "pumping_unit_heading": h,
            "pumping_unit_intro":   i,
            "pumping_unit_items":   [{"item": x} for x in b],
        })(*_pumping_unit_block(
            customer.get("fuel_name"),
            bool(customer.get("is_oil")),
            bool(customer.get("is_dual")),
        )),
    }

    # Use tundish-specific template if the product type indicates tundish
    is_tundish = any(
        "tundish" in (item.get("product_type") or "").lower()
        for item in quote_data.get("items", [])
    )
    template = TUNDISH_TEMPLATE_PATH if is_tundish and os.path.exists(TUNDISH_TEMPLATE_PATH) else TEMPLATE_PATH
    buffer = generate_word_offer(template, context)
    with open(output_path, "wb") as f:
        f.write(buffer.read())

    # Post-process: drop any tech-data row whose value cell ended up empty.
    _strip_empty_tech_rows(output_path)

    # Post-process: append BOM items (item, make) to the MAKE LIST table.
    _append_make_list(output_path, context.get("make_list", []))

    # Post-process: drop the unused column in the Scope of Supply (Annexure I)
    # table — VLPH-only quotes should not show the Horizontal column and vice
    # versa. total_vertical and total_horizontal were computed above.
    _prune_scope_columns(
        output_path,
        has_vertical=(total_vertical > 0),
        has_horizontal=(total_horizontal > 0),
    )


def _append_make_list(docx_path: str, items: list):
    """Find the table whose first row is ['ITEM', 'MAKE'] and append one
    row per BOM item with item name + vendor make."""
    if not items:
        return
    from docx import Document
    from copy import deepcopy
    from docx.oxml.ns import qn

    doc = Document(docx_path)
    target_table = None
    for t in doc.tables:
        if len(t.rows) >= 1 and len(t.rows[0].cells) >= 2:
            cells = [c.text.strip() for c in t.rows[0].cells]
            # Detect both 2-col ("ITEM","MAKE") and 3-col ("S. No.","ITEM","MAKE") layouts
            if cells[:2] == ["ITEM", "MAKE"] or cells[1:3] == ["ITEM", "MAKE"]:
                target_table = t
                break
    if target_table is None:
        return

    header_row = target_table.rows[0]
    seen = set()  # de-dupe by item name (case-insensitive)

    import re as _re
    # Known vendor names that may leak into item names
    KNOWN_VENDORS = {"BAUMER", "HGURU", "WIKA", "CAIR", "DEMBLA", "AIRA",
                     "MADAS", "ENCON", "L&T", "LT", "AUDCO", "LEADER",
                     "HONEYWELL", "BENGAL IND.", "BENGAL", "ROTEX", "VFLEX",
                     "SB INTERNATIONAL", "SKF", "FAG", "ABB", "BB",
                     "CROMPTON", "MASIBUS", "MURUGAPPA", "UNIFRAX",
                     "SWITZER", "DANFOSS", "JINDAL", "TATA",
                     "MAHARASHTRA SEAMLESS", "VARITECH", "JSD", "VANAZ",
                     "LINEAR SYSTEM", "THIRD PARTY"}
    # Location/usage qualifiers we want to drop
    LOC_QUALIFIERS = {"PILOT BURNER", "UV LINE", "PILOT", "UV"}

    def _clean_item(name: str, make_str: str) -> str:
        n = name.strip()
        # 1. 'GAS TRAIN <flow> NM3/Hr' → 'GAS TRAIN'
        n = _re.sub(r"^(GAS TRAIN)\s.*$", r"\1", n, flags=_re.IGNORECASE)
        # 2. Strip ALL trailing parentheticals that look like vendors or
        #    location qualifiers. Keep technical qualifiers like (Ball),
        #    (Globe), (Butterfly), (Lever), (Gear).
        while True:
            m = _re.search(r"\s*\(([^)]+)\)\s*$", n)
            if not m:
                break
            inside = m.group(1).strip()
            inside_u = inside.upper()
            is_vendor = (
                inside_u == make_str.upper()
                or inside_u in KNOWN_VENDORS
                or inside_u in LOC_QUALIFIERS
            )
            if not is_vendor:
                break
            n = n[:m.start()].rstrip()
        return n

    # Detect column layout: 3-col table = ["S. No.", "ITEM", "MAKE"], 2-col = ["ITEM", "MAKE"]
    header_cells = [c.text.strip().upper() for c in target_table.rows[0].cells]
    has_serial = (len(header_cells) >= 3 and header_cells[0] in ("S. NO.", "S.NO.", "S NO", "SR. NO.", "SR.NO."))

    serial = 0
    for entry in items:
        item = (entry.get("item") or "").strip()
        make = (entry.get("make") or "ENCON").strip() or "ENCON"
        if not item:
            continue
        item = _clean_item(item, make)
        key = item.lower()
        if key in seen:
            continue
        seen.add(key)
        serial += 1

        # Clone the header row so formatting matches, then overwrite text
        new_row = deepcopy(header_row._element)
        for t in new_row.iter(qn("w:t")):
            t.text = ""
        cells = new_row.findall(qn("w:tc"))

        if has_serial and len(cells) >= 3:
            fills = [(cells[0], str(serial)), (cells[1], item), (cells[2], make)]
        elif len(cells) >= 2:
            fills = [(cells[0], item), (cells[1], make)]
        else:
            fills = []

        for cell, text in fills:
            t_elements = list(cell.iter(qn("w:t")))
            if t_elements:
                t_elements[0].text = text
                t_elements[0].set(qn("xml:space"), "preserve")
        target_table._element.append(new_row)

    doc.save(docx_path)


def _strip_empty_tech_rows(docx_path: str):
    """In the rendered offer, find the Technical Data table and remove rows
    whose value cell is blank (meaning the placeholder resolved to empty)."""
    from docx import Document
    doc = Document(docx_path)

    # Identify the tech-data table by looking for the unique label 'Refractory Weight'.
    tech_labels = {
        "Ladle Dimensions", "Ladle Type", "Reference TS",
        "Refractory Weight", "Weight of refractory lining",
        "Heating Schedule", "Heating Temperature", "Heating time",
        "Calorific Value of Fuel", "Calorific value of LDO", "Calorific value",
        "Fuel Consumption", "Firing Rate", "Firing Rate (R1)",
        "Burner Size & Capacity", "Burner Capacity", "Burner Capacity (R1)",
        "ENCON Burner", "Combustion Air Blower",
        "Blower Size", "Capacity of Blower", "Blower Capacity / Motor rating",
        "LDO Pumping Unit", "Pumping Unit",
        "Movement of Hood",
        "Motor recommended for Power Pack", "Maximum Electrical Load",
        "Maximum Electrical Load Required",
        # Auto-ignition-only rows (removed when special_auto_ignition is False)
        "Pilot Burner Fuel", "Firing & Ignition of Burner",
    }

    for table in doc.tables:
        labels_in_table = {row.cells[0].text.strip() for row in table.rows if len(row.cells) >= 2}
        if not (labels_in_table & {"Refractory Weight", "Weight of refractory lining",
                                    "Heating Schedule", "Heating Temperature"}):
            continue
        # This is the tech-data table — drop blank-value rows we own.
        rows_to_remove = []
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            label = row.cells[0].text.strip()
            value = row.cells[1].text.strip()
            # 'Centrifugal type, ' alone (no model) is also empty
            if label in tech_labels and (not value or value == "Centrifugal type,"):
                rows_to_remove.append(row)
        for row in rows_to_remove:
            row._element.getparent().remove(row._element)
        break

    doc.save(docx_path)


def _prune_scope_columns(docx_path: str, has_vertical: bool, has_horizontal: bool):
    """Annexure I scope-of-supply has two columns: Vertical and Horizontal.
    If the quote is for only one of them, drop the other column so the
    offer shows just the relevant scope."""
    if has_vertical and has_horizontal:
        return  # both needed, nothing to do
    if not has_vertical and not has_horizontal:
        return  # nothing to base decision on

    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(docx_path)

    # Find the scope table by its header row text.
    target = None
    for t in doc.tables:
        if len(t.rows) < 2 or len(t.rows[0].cells) != 2:
            continue
        header = " | ".join(c.text.strip() for c in t.rows[0].cells)
        if "Vertical Ladle" in header and "Horizontal Ladle" in header:
            target = t
            break
    if target is None:
        return

    col_to_drop = 1 if has_vertical else 0   # 1 = horizontal, 0 = vertical
    for row in list(target.rows):
        tcs = row._element.findall(qn('w:tc'))
        if col_to_drop < len(tcs):
            tcs[col_to_drop].getparent().remove(tcs[col_to_drop])

    # Also shrink the grid so the remaining column stretches across.
    tblgrid = target._element.find(qn('w:tblGrid'))
    if tblgrid is not None:
        grid_cols = tblgrid.findall(qn('w:gridCol'))
        if col_to_drop < len(grid_cols):
            tblgrid.remove(grid_cols[col_to_drop])

    doc.save(docx_path)
