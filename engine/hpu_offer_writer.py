"""
Standalone HPU (Hydraulic / Oil Pumping Unit) offer generator.

Produces a Word .docx with:
  - cover / header (customer + ENCON contacts)
  - HPU technical specs table (variant, kW, fuel, LPH)
  - standard accessories list
  - scope of supply
  - commercial terms / T&Cs

Self-contained (uses python-docx directly, no template) so the HPU offer
is independent of the VLPH/HLPH/Tundish Offer_Template.docx.
"""

from datetime import datetime

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from engine.quote_writer import amount_in_words_indian, _format_inr


NAVY = RGBColor(0x1A, 0x3A, 0x5C)
GREY = RGBColor(0x64, 0x74, 0x8B)


def _set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_para(doc, text, *, bold=False, size=10, color=None, align=None, space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Calibri"
    if color is not None:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def _add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 13, 3: 11}
    return _add_para(doc, text, bold=True, size=sizes.get(level, 11), color=NAVY, space_after=6)


def _kv_table(doc, rows, col_widths=(Cm(5.5), Cm(10))):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        c1, c2 = table.rows[i].cells
        c1.width, c2.width = col_widths
        c1.text = ""
        c2.text = ""
        p1 = c1.paragraphs[0].add_run(k)
        p1.bold = True
        p1.font.size = Pt(10)
        p2 = c2.paragraphs[0].add_run(str(v))
        p2.font.size = Pt(10)
        _set_cell_bg(c1, "F1F5F9")
    return table


def _accessories_for_hpu(variant: str, kw: float, fuel: str, lph: float):
    """
    Standard accessories shipped with every HPU. Quantities depend on
    Simplex vs Duplex (Duplex = 2x pumps, 2x motors).
    """
    is_duplex = variant.lower().startswith("duplex")
    pump_qty = 2 if is_duplex else 1
    motor_qty = 2 if is_duplex else 1

    return [
        ("1",  "Oil tank (MS construction with cleaning door, breather, level indicator, drain plug)", "1 No."),
        ("2",  f"Gear pump suitable for {fuel} service",                                              f"{pump_qty} No(s)."),
        ("3",  f"TEFC motor — {kw:g} kW × 1440 RPM × 415V / 3-Ph / 50 Hz",                            f"{motor_qty} No(s)."),
        ("4",  "Coupling with guard (pump-motor)",                                                    f"{pump_qty} No(s)."),
        ("5",  "Suction strainer 150 micron",                                                         f"{pump_qty} No(s)."),
        ("6",  "Pressure relief valve",                                                               f"{pump_qty} No(s)."),
        ("7",  "Pressure gauge with isolation valve (0–10 kg/cm²)",                                   "1 No."),
        ("8",  "Duplex line strainer 25 micron (changeover type)" if is_duplex
               else "Inline strainer 25 micron",                                                      "1 No."),
        ("9",  "Non-return valve",                                                                    f"{pump_qty} No(s)."),
        ("10", "Pressure switch (low-pressure cut-off)",                                              "1 No."),
        ("11", "Float-type oil level switch (low-level cut-off)",                                     "1 No."),
        ("12", "Inter-connecting piping inside the unit",                                             "Lot"),
        ("13", "Common skid base frame with mounting bolts",                                          "1 No."),
        ("14", "Control panel (start / stop / interlock / indication)",                              "1 No."),
        ("15", "Set of foundation bolts & matching flanges",                                          "1 Lot"),
    ]


def _scope_table(doc, accessories):
    table = doc.add_table(rows=1 + len(accessories), cols=3)
    table.style = "Light Grid Accent 1"
    table.autofit = False
    widths = (Cm(1.4), Cm(12), Cm(2.6))
    hdr = table.rows[0].cells
    for c, txt, w in zip(hdr, ("S. No.", "Description", "Quantity"), widths):
        c.width = w
        c.text = ""
        r = c.paragraphs[0].add_run(txt)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(c, "1A3A5C")
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (sn, desc, qty) in enumerate(accessories, start=1):
        row = table.rows[i].cells
        for c, w in zip(row, widths):
            c.width = w
        row[0].text = sn
        row[1].text = desc
        row[2].text = qty
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for c in row:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)


def _commercial_terms():
    return [
        ("Price Basis",         "Ex-Works, ENCON Goa. Prices firm till the validity of this offer."),
        ("Taxes & Duties",      "GST @ 18% extra as applicable, to be paid against documentary proof."),
        ("Packing & Forwarding", "Included in the quoted price (standard export-worthy packing)."),
        ("Freight",             "Extra at actuals, on door-delivery basis, against your written confirmation."),
        ("Transit Insurance",   "To be arranged and borne by the customer."),
        ("Payment Terms",       "30% advance against PO. 60% against pro-forma invoice before dispatch. "
                                "10% within 30 days of receipt of material at site."),
        ("Delivery",            "8–10 weeks from the date of receipt of techno-commercially clear PO "
                                "and confirmed advance payment."),
        ("Inspection",          "At our works prior to dispatch. Customer / customer's representative "
                                "is welcome to witness on prior intimation."),
        ("Guarantee / Warranty", "12 months from the date of commissioning or 18 months from dispatch, "
                                 "whichever is earlier, against manufacturing defects under normal use."),
        ("Validity of Offer",   "30 days from the date of this offer."),
    ]


def generate_hpu_quote_docx(payload: dict, output_path: str) -> dict:
    """
    Build a stand-alone HPU offer Word document.

    payload keys:
        customer: dict with name, company, address, city, state, pin, gstin,
                  ref_no, subject, marketing, technical, etc.
        hpu_variant, hpu_kw, fuel_type, fuel_lph, qty
        unit_price, total_price
    """
    cust = payload.get("customer") or {}
    variant = payload["hpu_variant"]
    kw      = float(payload["hpu_kw"])
    fuel    = payload["fuel_type"]
    lph     = float(payload["fuel_lph"])
    qty     = int(payload.get("qty") or 1)
    unit    = float(payload["unit_price"])
    total   = float(payload["total_price"])

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)

    # ── Header band ──
    _add_para(doc, "ENCON THERMAL ENGINEERS PVT. LTD.",
              bold=True, size=18, color=NAVY,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _add_para(doc, "Plot No. L-31, Verna Industrial Estate, Verna, Goa - 403722",
              size=9, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _add_para(doc, "Tel: +91-832-2783100 | Email: encon@encon.in | Web: www.encon.co.in",
              size=9, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # ── Ref / date band ──
    today = datetime.now().strftime("%d %B %Y")
    _kv_table(doc, [
        ("Ref. No.", cust.get("ref_no") or "—"),
        ("Date",     today),
    ])
    doc.add_paragraph()

    # ── To / addressed-to block ──
    _add_para(doc, "To,", bold=True, size=10)
    salutation = cust.get("salutation") or ""
    addr_lines = [
        f"{salutation} {cust.get('name','')}".strip(),
        cust.get("designation",""),
        cust.get("company",""),
        cust.get("address",""),
        ", ".join(filter(None, [cust.get("city",""), cust.get("state",""), cust.get("pin","")])),
    ]
    for line in addr_lines:
        if line:
            _add_para(doc, line, size=10, space_after=2)

    if cust.get("subject"):
        doc.add_paragraph()
        p = doc.add_paragraph()
        r1 = p.add_run("Subject: ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(cust["subject"])
        r2.font.size = Pt(10)

    _add_para(doc, "Dear Sir / Madam,", size=10, space_after=4)
    _add_para(doc,
        "Thank you for your enquiry. We are pleased to submit our offer for the "
        f"supply of one (01) no. {variant} Hydraulic Pumping Unit "
        f"(HPU) of {kw:g} kW capacity, suitable for {fuel} service at "
        f"{lph:g} LPH, complete with standard accessories as detailed below.",
        size=10, space_after=8)

    # ── 1. Technical specifications ──
    _add_heading(doc, "1.  Technical Specifications", level=2)
    _kv_table(doc, [
        ("HPU Variant",        variant),
        ("Motor Capacity",     f"{kw:g} kW"),
        ("Fuel",               fuel),
        ("Oil Flow Rate",      f"{lph:g} LPH"),
        ("Quantity Offered",   f"{qty} No."),
        ("Construction",       "MS skid-mounted with pre-wired control panel"),
        ("Electrical Supply",  "415 V / 3 Phase / 50 Hz, 4-Wire"),
        ("Operating Pressure", "Up to 7 kg/cm² (adjustable)"),
    ])
    doc.add_paragraph()

    # ── 2. Standard accessories / scope of supply ──
    _add_heading(doc, "2.  Scope of Supply", level=2)
    _add_para(doc,
        f"The following items are included as standard with each {variant} "
        f"HPU offered above:",
        size=10, space_after=6)
    accessories = _accessories_for_hpu(variant, kw, fuel, lph)
    _scope_table(doc, accessories)
    doc.add_paragraph()

    # ── 3. Pricing ──
    _add_heading(doc, "3.  Pricing", level=2)
    _kv_table(doc, [
        ("Unit Price (Ex-Works)",    f"INR {_format_inr(unit)}/-"),
        ("Quantity",                 f"{qty} No."),
        ("Total Price (Ex-Works)",   f"INR {_format_inr(total)}/-"),
        ("Amount in Words",          f"Rupees {amount_in_words_indian(total)} Only"),
    ])
    doc.add_paragraph()

    # ── 4. Commercial / T&Cs ──
    _add_heading(doc, "4.  Commercial Terms & Conditions", level=2)
    table = doc.add_table(rows=len(_commercial_terms()), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(_commercial_terms()):
        c1, c2 = table.rows[i].cells
        c1.width, c2.width = Cm(4.5), Cm(11.5)
        c1.text = ""
        c2.text = ""
        r1 = c1.paragraphs[0].add_run(k)
        r1.bold = True
        r1.font.size = Pt(10)
        _set_cell_bg(c1, "F1F5F9")
        r2 = c2.paragraphs[0].add_run(v)
        r2.font.size = Pt(10)
    doc.add_paragraph()

    # ── 5. Exclusions ──
    _add_heading(doc, "5.  Exclusions", level=2)
    for item in (
        "Civil & foundation work for HPU skid.",
        "Power & water supply at battery limit.",
        "Inter-connecting oil piping from HPU to user equipment beyond the unit's outlet flange.",
        "Anything not explicitly listed under Scope of Supply.",
    ):
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.size = Pt(10)
    doc.add_paragraph()

    # ── Closing / signature ──
    _add_para(doc,
        "We trust the above is in line with your requirement. Please feel free to "
        "revert for any clarification.",
        size=10, space_after=10)

    _add_para(doc, "Thanking you,", size=10, space_after=2)
    _add_para(doc, "Yours faithfully,", size=10, space_after=2)
    _add_para(doc, "For ENCON THERMAL ENGINEERS PVT. LTD.",
              bold=True, size=10, color=NAVY, space_after=18)

    mkt_sal = cust.get("marketing_salutation") or ""
    mkt = (cust.get("marketing") or "").strip()
    if mkt:
        _add_para(doc, f"{mkt_sal} {mkt}".strip(), bold=True, size=10, space_after=1)
    if cust.get("marketing_email"):
        _add_para(doc, cust["marketing_email"], size=9, color=GREY, space_after=1)
    if cust.get("marketing_phone"):
        _add_para(doc, cust["marketing_phone"], size=9, color=GREY, space_after=1)

    doc.save(output_path)
    return {
        "filename": output_path,
        "total":    total,
        "unit":     unit,
    }
