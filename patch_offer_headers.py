"""
Put the letterhead — logo and "Ref: {{enquiry_ref}}" — on every page of the
recuperator offer, the way the regen offer already does.

Two separate faults left most pages bare:

  1. The letter section had no header of its own. The recup template is built
     in three sections — cover page, covering letter + company introduction,
     then the annexures — and only the third one carried a header reference.
     The second inherited the cover page's deliberately blank header, so the
     whole covering letter and the ENCON introduction printed with no logo and
     no reference number.

  2. evenAndOddHeaders was switched on document-wide, but only the cover-page
     section defined an even-page header, and that one is empty. Every
     even-numbered page in the document therefore fell back to it — losing the
     header AND the address footer even inside the annexures, which is why the
     branding appeared to come and go page by page. The regen template has the
     setting off; this turns it off here too, so the one primary header and
     footer apply to every page.

The cover page is left alone: its blank header is intentional, and matches
regen, where the same effect comes from a title-page header.

Run:  python patch_offer_headers.py                       # the recup template
      python patch_offer_headers.py Combined_Offer_Template.docx ...

Both faults are shared by the combined, blower, burner and HPU templates —
hence the filename argument — but this only touches what it is asked to.

Idempotent — a second run reports there is nothing to do.
"""

import os
import sys

from docx import Document
from docx.enum.section import WD_HEADER_FOOTER
from docx.oxml.ns import qn
from docx.shared import Emu, Twips

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATES = ["Recup_Offer_Template.docx"]


def _branded_section(doc):
    """The section whose primary header carries the logo and the ref — the one
    every other section should be borrowing its letterhead from."""
    for sec in doc.sections:
        ref = sec._sectPr.get_headerReference(WD_HEADER_FOOTER.PRIMARY)
        if ref is None:
            continue
        part = doc.part.related_parts[ref.rId]
        if "<w:drawing" in part.element.xml:
            return sec
    return None


def _fit_letterhead_width(doc, sec):
    """Make the letterhead table exactly as wide as the text column.

    The rule under the logo is the table's bottom border, so its length is the
    table's width. Recup and Combined carry a 9360-twip table on a 9740-twip
    text column — inherited from an older page setup — and the two cell widths
    inside it (1800 + 7760) do not even add up to that, with autofit left on to
    paper over the difference. The result is a letterhead rule that stops short
    of the right margin while the orange rule under every ANNEXURE heading runs
    the full width, and a reference number that floats in from the edge instead
    of ending flush with the rule. Every other template is already 9740 fixed.

    Returns a description of what changed, or None.
    """
    tbl = sec.header.tables[0] if sec.header.tables else None
    if tbl is None:
        return None
    text_w = int(round((sec.page_width - sec.left_margin - sec.right_margin) / 635))
    tblPr = tbl._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    logo_w = int(tbl.columns[0].width / 635) if tbl.columns[0].width else 1600
    ref_w = text_w - logo_w
    if (tblW is not None and tblW.get(qn("w:w")) == str(text_w)
            and int(tbl.columns[1].width or 0) / 635 == ref_w
            and tbl.autofit is False):
        return None

    if tblW is not None:
        tblW.set(qn("w:w"), str(text_w))
        tblW.set(qn("w:type"), "dxa")
    tbl.autofit = False                    # w:tblLayout fixed — honour the widths
    tbl.columns[0].width = Twips(logo_w)
    tbl.columns[1].width = Twips(ref_w)
    for row in tbl.rows:
        row.cells[0].width = Twips(logo_w)
        row.cells[1].width = Twips(ref_w)
    return f"letterhead width {text_w} twips ({logo_w} logo + {ref_w} ref), fixed"


def patch(target):
    path = os.path.join(BASE_DIR, target)
    doc = Document(path)
    done = []

    branded = _branded_section(doc)
    if branded is None:
        print(f"  {target}: no header carries the logo — nothing to copy from")
        return
    hdr_rid = branded._sectPr.get_headerReference(WD_HEADER_FOOTER.PRIMARY).rId
    ftr_ref = branded._sectPr.get_footerReference(WD_HEADER_FOOTER.PRIMARY)
    ftr_rid = ftr_ref.rId if ftr_ref is not None else None

    # 1. One header and footer for every page, not one for odd and a blank for even.
    settings = doc.settings.element
    even_odd = settings.find(qn("w:evenAndOddHeaders"))
    if even_odd is not None:
        settings.remove(even_odd)
        done.append("turned off even/odd headers")

    # 2. The letterhead rule runs the full width of the text column, so it
    #    lines up with the orange rules under the headings below it.
    fitted = _fit_letterhead_width(doc, branded)
    if fitted:
        done.append(fitted)

    # 3. Every section after the cover page uses the branded letterhead.
    for i, sec in enumerate(doc.sections):
        if i == 0:
            continue                      # cover page stays deliberately bare
        sp = sec._sectPr
        if sp.get_headerReference(WD_HEADER_FOOTER.PRIMARY) is None:
            sp.add_headerReference(WD_HEADER_FOOTER.PRIMARY, hdr_rid)
            done.append(f"section {i}: added the letterhead")
        if ftr_rid and sp.get_footerReference(WD_HEADER_FOOTER.PRIMARY) is None:
            sp.add_footerReference(WD_HEADER_FOOTER.PRIMARY, ftr_rid)
            done.append(f"section {i}: added the address footer")
        # A header sitting 0 mm from the paper edge prints the logo into the
        # trim. Match the section the header was taken from.
        if sec.header_distance != branded.header_distance:
            sec.header_distance = Emu(branded.header_distance)
            done.append(f"section {i}: header distance now "
                        f"{round(branded.header_distance / 36000, 1)} mm")

    if not done:
        print(f"  {target}: already on every page")
        return
    doc.save(path)
    for line in done:
        print(f"  {target}: {line}")


if __name__ == "__main__":
    for name in (sys.argv[1:] or TEMPLATES):
        patch(name)
