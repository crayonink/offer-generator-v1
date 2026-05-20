"""Stage 1: rebuild Recup_Offer_Template.docx with the merged Viraj + Real
Ispat layout.

This script runs ON TOP of the existing Recup_Offer_Template.docx (which
was built by build_recup_template.py from the Real Ispat reference). It
patches in:

  Stage 1a (this run):
    - Material of Construction sub-section in Annexure I, added as new
      rows at the top of the existing Designing Parameters table.
      Cold Bank / Hot Bank / Others — each component lists its material
      via Jinja placeholders so Step 2's Hot/Cold material dropdowns
      flow through.

  Stage 1b (next commit):
    - 3D recuperator image placeholders.

  Stage 1c (next commit):
    - Annexure III rebuilt as a full-BOM iterable price schedule with
      optional supervision-charges row + summary totals.

The added Material of Construction rows use these new Jinja placeholders:
  {{ hot_tube_plate_material }}    Hot bank tube plate material
  {{ hot_tube_material }}          Hot bank tube spec
  {{ cold_tube_plate_material }}   Cold bank tube plate material
  {{ cold_tube_material }}         Cold bank tube spec
The backend (engine/recup_writer.py) will populate these from the
hot_bank_material / cold_bank_material picks on Step 2.
"""
from __future__ import annotations

import os
from copy import deepcopy

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

TEMPLATE_PATH = "Recup_Offer_Template.docx"


# ── Material of Construction layout ────────────────────────────────────────
# Each row is (kind, label, value).
#   kind = 'header'  -> full-width banner (e.g. 'Material of Construction')
#   kind = 'section' -> sub-header (e.g. 'COLD BANK')
#   kind = 'item'    -> two-col label / material row
_MOC_ROWS = [
    ('header',  'Material of Construction',  None),
    ('section', 'COLD BANK',                 None),
    ('item',    'Tube Plate',                '{{ cold_tube_plate_material }}'),
    ('item',    'Tube',                      '{{ cold_tube_material }}'),
    ('item',    'Duct Inlet',                'Mild Steel'),
    ('item',    'Inlet Collar',              'Mild Steel'),
    ('section', 'HOT BANK',                  None),
    ('item',    'Tube Plate',                '{{ hot_tube_plate_material }}'),
    ('item',    'Tube',                      '{{ hot_tube_material }}'),
    ('item',    'Duct Outlet',               'Mild Steel'),
    ('item',    'Outlet Collar',             'Mild Steel'),
    ('item',    'Supporting Frame',          'Mild Steel'),
    ('item',    'Bottom Duct',               'Mild Steel'),
    ('item',    'Flange',                    'Mild Steel'),
    ('section', 'OTHERS',                    None),
    ('item',    'Flanges',                   'Mild Steel'),
    ('item',    'Air Inlet Duct',            'Mild Steel'),
    ('item',    'Air Outlet Duct',           'Mild Steel'),
    ('item',    'Bottom Air Receiving Box',  'Mild Steel'),
    ('item',    'Matching Flange',           'Mild Steel'),
    ('item',    'Nut, Bolt & Washer',        'Mild Steel'),
    ('item',    'Gasket',                    'Heat Resistant'),
    ('item',    'Supporting Structure',      'Mild Steel'),
]


def _make_row(table, kind: str, label: str | None, value: str | None):
    """Append a new row to `table` and return its <w:tr> XML element with
    the cell text set. Uses the existing-row formatting as a template so
    fonts and borders carry over."""
    new_row = table.add_row()
    cells = new_row.cells
    if kind == 'header':
        # Merge all three cells into one banner
        cells[0].merge(cells[1]).merge(cells[2])
        cells[0].text = label
        # Bold the header
        for p in cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
    elif kind == 'section':
        cells[0].merge(cells[1]).merge(cells[2])
        cells[0].text = label
        for p in cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
    elif kind == 'item':
        cells[0].text = label
        cells[1].merge(cells[2])
        cells[1].text = value or ''
    return new_row._element


def add_material_of_construction(doc: Document) -> None:
    """Add Material of Construction rows to the top of Table 4 (the
    existing Recuperator Designing Parameters table). Idempotent."""
    target_table = doc.tables[4]
    tbl_xml = target_table._element

    # Idempotency: if first row already says 'Material of Construction'
    # we've already patched this template.
    if target_table.rows and 'Material of Construction' in target_table.rows[0].cells[0].text:
        print("Material of Construction rows already present — skipping.")
        return

    # Build the rows at the end of the table first (because python-docx
    # only knows how to .add_row at the end), then move them to the top
    # via XML.
    new_row_elements = []
    for kind, label, value in _MOC_ROWS:
        tr = _make_row(target_table, kind, label, value)
        new_row_elements.append(tr)

    # Find the body element (<w:tbl>) and rearrange: move the new rows
    # before the very first row (the 'Recuperator Designing Parameters'
    # banner row).
    all_rows = list(tbl_xml.findall(qn('w:tr')))
    first_existing_row = all_rows[0]   # 'Recuperator Designing Parameters' banner
    # The rows we just added are the last len(new_row_elements) rows.
    n_new = len(new_row_elements)
    rows_to_move = all_rows[-n_new:]
    for tr in rows_to_move:
        tbl_xml.remove(tr)
    # Re-insert before the original banner row.
    for tr in rows_to_move:
        first_existing_row.addprevious(tr)


def split_material_of_construction(doc: Document) -> None:
    """Split Table 4 (which currently contains MoC rows 0..N followed by
    Designing Parameters rows N..end) into TWO separate tables, with an
    empty paragraph between them.

    After this:
      doc.tables[4] = Material of Construction (was rows 0..N of old T4)
      doc.tables[5] = Designing Parameters    (was rows N..end of old T4)
      doc.tables[6] = Price Schedule          (was doc.tables[5])

    Idempotent: skip if a separate Designing Parameters table already
    exists (i.e. there's already a table whose first row says
    'Recuperator Designing Parameters')."""
    designing_banner = "Recuperator Designing Parameters"
    for ti, t in enumerate(doc.tables):
        if t.rows and designing_banner in t.rows[0].cells[0].text and ti != 4:
            # already split (the banner sits as the first row of a non-T4)
            print(f"Already split — Designing Parameters lives in Table {ti}.")
            return

    table4 = doc.tables[4]
    tbl4_xml = table4._element

    # Walk rows and find the split point: the row whose first cell says
    # 'Recuperator Designing Parameters'.
    rows = list(tbl4_xml.findall(qn('w:tr')))
    split_idx = None
    for i, tr in enumerate(rows):
        cell_text = ''.join(t.text or '' for t in tr.iter(qn('w:t')))
        if designing_banner in cell_text:
            split_idx = i
            break
    if split_idx is None:
        print("No Designing Parameters banner found inside Table 4 — nothing to split.")
        return

    # Clone the whole table, then in the CLONE keep only rows from split_idx
    # onward (Designing Params). In the ORIGINAL, drop those same rows.
    new_tbl = deepcopy(tbl4_xml)
    new_rows = list(new_tbl.findall(qn('w:tr')))
    for i, tr in enumerate(new_rows):
        if i < split_idx:
            new_tbl.remove(tr)
    for i, tr in enumerate(rows):
        if i >= split_idx:
            tbl4_xml.remove(tr)

    # Insert an empty paragraph + the new table right after the (now MoC-only)
    # Table 4.
    sep_p = OxmlElement('w:p')
    tbl4_xml.addnext(new_tbl)
    tbl4_xml.addnext(sep_p)
    print(f"Split done: Table 4 = {split_idx} MoC rows; new Table 5 = "
          f"{len(rows) - split_idx} Designing Params rows.")


def rebuild_price_schedule(doc: Document) -> None:
    """Rebuild the Price Schedule table as a docxtpl iterable row over
    `bom_rows` + optional supervision row + summary totals +
    amount-in-words footer.

    Target by CONTENT (find the table whose header row contains
    'ITEM DESCRIPTION') so the function still works after the MoC split
    shuffles table indices."""
    table = None
    for t in doc.tables:
        if not t.rows or len(t.rows[0].cells) < 2:
            continue
        if 'ITEM DESCRIPTION' in t.rows[0].cells[1].text.upper():
            table = t
            break
    if table is None:
        print("Price Schedule table not found — skipping rebuild.")
        return
    tbl_xml = table._element
    # Strip every existing row except the header (row 0) so a re-run
    # cleanly regenerates the iterable + supervision + summary block.
    for tr in list(tbl_xml.findall(qn('w:tr')))[1:]:
        tbl_xml.remove(tr)

    def add_row(c1: str, c2: str, c3: str, c4: str, c5: str, *, bold_total: bool = False):
        r = table.add_row()
        cells = r.cells
        cells[0].text = c1
        cells[1].text = c2
        cells[2].text = c3
        cells[3].text = c4
        cells[4].text = c5
        if bold_total:
            for cell in cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True

    # ── Style: 'single' ─────────────────────────────────────────────
    # Real Ispat-style single-line price (e.g. "Recuperator for 50 TPH
    # Furnace — 01 No. — 28,12,000.00"). Toggled by
    # price_schedule_style == 'single'.
    add_row("{%tr if price_schedule_style == 'single' %}", "", "", "", "")
    add_row("1.", "Recuperator for {{ application }}",
            "{{ recup_qty }}", "{{ recup_unit_price }}", "{{ recup_total_price }}")
    add_row("{%tr endif %}", "", "", "", "")

    # ── Style: 'full' (default) ─────────────────────────────────────
    # Full BOM line-item table — iterates each row in bom_rows.
    add_row("{%tr if price_schedule_style == 'full' %}", "", "", "", "")
    add_row("{%tr for r in bom_rows %}", "", "", "", "")
    add_row("{{ r.sno }}", "{{ r.item }}", "{{ r.qty }}",
            "{{ r.unit_price }}", "{{ r.total }}")
    add_row("{%tr endfor %}", "", "", "", "")

    # Optional supervision-charges row (only meaningful in full mode).
    add_row("{%tr if supervision_include %}", "", "", "", "")
    add_row("", "Supervision Charges for Erection & Commissioning "
            "(Erection by Client, Supervision by ENCON)", "",
            "{{ supervision_rate }}", "{{ supervision_note }}")
    add_row("{%tr endif %}", "", "", "", "")

    # Summary rows (full mode only).
    add_row("", "Bought-out Items Total", "", "", "{{ bought_out_total }}")
    add_row("", "ENCON Items Total",      "", "", "{{ encon_total }}")
    add_row("", "GRAND TOTAL",            "", "", "{{ grand_total }}", bold_total=True)
    add_row("{%tr endif %}", "", "", "", "")  # end 'full' block

    # 4. Amount-in-words footer — merge the first 4 cells into one so
    #    the long text doesn't repeat across the row. Last cell keeps
    #    the numeric grand total.
    footer = table.add_row()
    f_cells = footer.cells
    merged = f_cells[0].merge(f_cells[1]).merge(f_cells[2]).merge(f_cells[3])
    merged.text = "{{ grand_total_in_words }}"
    f_cells[4].text = "{{ grand_total }}"
    for p in merged.paragraphs:
        for run in p.runs:
            run.bold = True


def main() -> None:
    if not os.path.exists(TEMPLATE_PATH):
        raise SystemExit(f"missing: {TEMPLATE_PATH}")

    doc = Document(TEMPLATE_PATH)
    print(f"Loaded template: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

    add_material_of_construction(doc)
    print(f"Added {len(_MOC_ROWS)} Material of Construction rows above Designing Parameters")

    split_material_of_construction(doc)

    rebuild_price_schedule(doc)
    print("Annexure III Price Schedule rebuilt as iterable + supervision + summary")

    doc.save(TEMPLATE_PATH)
    print(f"Saved -> {TEMPLATE_PATH}")


if __name__ == "__main__":
    main()
