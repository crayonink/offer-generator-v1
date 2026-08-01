"""
Style the regen offer templates' SUPERVISION CHARGES table to match the rest of
the document.

It was the only unstyled table in the offer: no shaded label column, no explicit
font size (so it rendered smaller than every other table), and no vertical
alignment — text sat cramped against the cell borders while TECHNICAL DATA and
TERMS AND CONDITIONS above and below it were cleanly banded.

This applies the same treatment the T&C table uses:
  • label column shaded F3F4F6, bold
  • 11 pt on every run
  • vAlign center so the two-line "(PLC and Instrumentation)" row sits straight
  • the trailing Note row spans the table and stays italic, unshaded

Run once against the gas and oil templates, then rebuild the dual template
(it is derived from the gas one):

    python patch_regen_supervision_style.py
    python build_regen_dual_template.py

Idempotent — re-running it just re-applies the same properties.
"""

import os

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATES = ["Regen_Offer_Template.docx", "Regen_Oil_Offer_Template.docx"]

LABEL_FILL = "F3F4F6"     # same grey the T&C table's label column uses
FONT_PT = 11


def _shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _supervision_table(doc):
    for t in doc.tables:
        if t.rows and "Supervision Charges" in t.rows[0].cells[0].text:
            return t
    return None


def patch(path):
    d = Document(path)
    t = _supervision_table(d)
    if t is None:
        print(f"  SKIP {os.path.basename(path)} — no supervision table")
        return
    for r in t.rows:
        # The Note row is one merged cell spanning the table — leave it italic
        # and unshaded, just size it with the rest.
        is_note = r.cells[0].text.strip().lower().startswith("note:")
        for ci, c in enumerate(r.cells):
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ci == 0 and not is_note:
                _shade(c, LABEL_FILL)
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(FONT_PT)
                    if ci == 0 and not is_note:
                        run.font.bold = True
    d.save(path)
    print(f"  OK  {os.path.basename(path)} — {len(t.rows)} rows styled")


if __name__ == "__main__":
    for name in TEMPLATES:
        patch(os.path.join(BASE_DIR, name))
    print("Now run: python build_regen_dual_template.py")
