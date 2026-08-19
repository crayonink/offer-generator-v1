"""Put the letterhead on every page of the BRF offer, as regen and recup have it.

The BRF template came from a Word document that carried no header at all — four
sections, every one of them empty — so the logo and the reference number
appeared on no page but the first, where they are part of the body text.

This builds the same header the recuperator template uses: a two-cell table,
the ENCON logo on the left and "Ref: {{ ref_no }}" on the right, on the letter
and annexure sections. The cover page keeps its blank header deliberately,
which is what regen and recup do too.

evenAndOddHeaders is turned off while we are here. When it is on and only some
sections define an even-page header, every even-numbered page falls back to
whatever that is — usually nothing — and the branding comes and goes page by
page. That is the fault this same script's recuperator sibling was written for.

Run:  python patch_brf_offer_header.py
      python patch_brf_offer_header.py Some_Other_Template.docx

Idempotent — a second run reports there is nothing to do.
"""
import os
import shutil
import sys
import zipfile

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DEFAULT_TARGET = os.path.join(BASE_DIR, "BRF_Offer_Template.docx")
LOGO_SOURCE = os.path.join(BASE_DIR, "Recup_Offer_Template.docx")
REF_PLACEHOLDER = "Ref: {{ ref_no }}"
LOGO_WIDTH = Inches(1.55)


def extract_logo(source_docx, out_path):
    """Pull the header logo out of a template that already has one."""
    with zipfile.ZipFile(source_docx) as z:
        headers = [n for n in z.namelist() if n.startswith("word/header")]
        wanted = set()
        for h in headers:
            xml = z.read(h).decode("utf-8", "replace")
            if "<w:drawing>" not in xml:
                continue
            rels = f"word/_rels/{os.path.basename(h)}.rels"
            if rels not in z.namelist():
                continue
            rel_xml = z.read(rels).decode("utf-8", "replace")
            for part in rel_xml.split("<Relationship")[1:]:
                if "image" in part and 'Target="' in part:
                    tgt = part.split('Target="', 1)[1].split('"', 1)[0]
                    wanted.add("word/" + tgt.replace("../", ""))
        if not wanted:
            return None
        name = sorted(wanted)[0]
        with open(out_path, "wb") as fh:
            fh.write(z.read(name))
    return out_path


def _has_header(section):
    hdr = section.header
    return bool(hdr.tables) or any(p.text.strip() for p in hdr.paragraphs)


def build_header(section, logo_path):
    """Two cells: the logo, and the reference, on one baseline."""
    hdr = section.header
    hdr.is_linked_to_previous = False
    for tb in list(hdr.tables):
        tb._tbl.getparent().remove(tb._tbl)
    for p in list(hdr.paragraphs):
        for r in list(p.runs):
            r._r.getparent().remove(r._r)

    table = hdr.add_table(rows=1, cols=2, width=section.page_width
                          - section.left_margin - section.right_margin)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    left, right = table.rows[0].cells

    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if logo_path and os.path.exists(logo_path):
        lp.add_run().add_picture(logo_path, width=LOGO_WIDTH)

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = rp.add_run(REF_PLACEHOLDER)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

    # No cell borders — the header is a layout device, not a table.
    for cell in (left, right):
        tcPr = cell._tc.get_or_add_tcPr()
        borders = tcPr.makeelement(qn("w:tcBorders"), {})
        for edge in ("top", "left", "bottom", "right"):
            el = borders.makeelement(qn(f"w:{edge}"), {qn("w:val"): "nil"})
            borders.append(el)
        tcPr.append(borders)


def disable_even_odd(doc):
    """One header for every page, rather than a blank one on even pages."""
    settings = doc.settings.element
    changed = False
    for tag in ("w:evenAndOddHeaders",):
        el = settings.find(qn(tag))
        if el is not None:
            settings.remove(el)
            changed = True
    return changed


def main():
    target = sys.argv[1] if len(sys.argv) > 1 else DEFAULT_TARGET
    if not os.path.exists(target):
        print(f"not found: {target}")
        return 1

    tmp_logo = os.path.join(BASE_DIR, "_brf_header_logo.png")
    logo = extract_logo(LOGO_SOURCE, tmp_logo)
    if not logo:
        print(f"no logo found in {os.path.basename(LOGO_SOURCE)}")
        return 1

    shutil.copyfile(target, target + ".bak")
    doc = Document(target)

    # The cover page keeps its blank header on purpose.
    todo = [s for s in doc.sections[1:] if not _has_header(s)]
    if not todo and not disable_even_odd(doc):
        os.remove(tmp_logo)
        print("nothing to do — the header is already there")
        return 0

    for section in doc.sections[1:]:
        build_header(section, logo)
    disable_even_odd(doc)
    doc.save(target)
    os.remove(tmp_logo)

    print(f"letterhead on {len(doc.sections) - 1} of {len(doc.sections)} sections "
          f"of {os.path.basename(target)}; the cover page is left blank")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
