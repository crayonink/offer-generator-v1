"""
Build Regen_Dual_Offer_Template.docx — the regen offer template for DUAL-FUEL
(one gas + one oil) systems.

A dual-fuel regen offer needs BOTH fuel-supply sections:
  • "{{ gas_train_fuel }} GAS TRAIN FOR MAIN BURNERS"  (from the gas template)
  • "HEATING & PUMPING UNIT"                            (from the oil template)

Rather than maintain a third hand-edited docx, this script derives the dual
template from Regen_Offer_Template.docx (the gas one) and inserts an HPU section
after the gas-train block. Every inserted paragraph is a deepcopy of a paragraph
already in that same document, so list numbering and styles stay valid — a
cross-document copy would carry numId references the target doesn't define.

The gas-train sentence is also re-pointed from {{ fuel_word }} to
{{ gas_train_fuel }}: on a dual offer fuel_word reads "DUAL FUEL", which would
make it say "gas train for DUAL FUEL" instead of "gas train for NG".

Run:  python build_regen_dual_template.py
"""

import copy
import os

from docx import Document
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(BASE_DIR, "Regen_Offer_Template.docx")
OUT = os.path.join(BASE_DIR, "Regen_Dual_Offer_Template.docx")

GAS_TRAIN_HEADING = "GAS TRAIN FOR MAIN BURNERS"
PILOT_HEADING = "GAS LINE FOR PILOT BURNER"

HPU_HEADING = "HEATING & PUMPING UNIT"
HPU_INTRO = (
    "To supply fuel oil to the above burners at the requisite pressure and "
    "temperature, we will supply one Heating & Pumping Unit along with the oil "
    "line to the burners, consisting of the following main components:"
)
# Placeholder bullets — fill_oil_supply() rewrites these from the real BOM at
# generation time; they only need to exist so the block can be found.
HPU_BULLETS = ["Heating & Pumping Unit", "Solenoid Valve (Oil Line)",
               "Ball Valve (Oil Line)", "Flexible Hose Pipe (Oil Line)"]


def _set_text(para, text):
    """Replace a paragraph's text, keeping the first run's formatting."""
    runs = para.runs
    if not runs:
        para.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def _is_bullet(para):
    pPr = para._element.find(qn("w:pPr"))
    return pPr is not None and pPr.find(qn("w:numPr")) is not None


def build():
    d = Document(SRC)
    paras = d.paragraphs

    gi = next(i for i, p in enumerate(paras) if GAS_TRAIN_HEADING in p.text)
    pi = next(i for i, p in enumerate(paras) if PILOT_HEADING in p.text)

    # Gas-train sentence: {{ fuel_word }} -> {{ gas_train_fuel }} so a dual
    # offer reads "gas train for NG", not "gas train for DUAL FUEL".
    intro = paras[gi + 1]
    for r in intro.runs:
        if "fuel_word" in r.text:
            r.text = r.text.replace("{{ fuel_word }}", "{{ gas_train_fuel }}")

    # Templates for the clones: the gas-train heading, its intro, and one of its
    # bullets — all from this same document.
    tpl_heading, tpl_intro = paras[gi]._element, intro._element
    tpl_bullet = next(p._element for p in paras[gi + 2:pi] if _is_bullet(p))

    # Insert the HPU block between the gas-train block and the pilot line.
    anchor = paras[pi]._element
    prev = None

    def _insert(tpl, text):
        nonlocal prev
        el = copy.deepcopy(tpl)
        if prev is None:
            anchor.addprevious(el)
        else:
            prev.addnext(el)
        prev = el
        from docx.text.paragraph import Paragraph
        _set_text(Paragraph(el, paras[pi]._parent), text)

    _insert(tpl_heading, HPU_HEADING)
    _insert(tpl_intro, HPU_INTRO)
    for b in HPU_BULLETS:
        _insert(tpl_bullet, b)

    d.save(OUT)
    print(f"wrote {OUT}")

    # Verify both fuel-supply sections are present and findable by the fillers.
    chk = Document(OUT)
    texts = [p.text.strip() for p in chk.paragraphs]
    for needed in (GAS_TRAIN_HEADING, HPU_HEADING, PILOT_HEADING):
        hit = any(needed in t for t in texts)
        print(f"  {'OK ' if hit else 'MISSING'} {needed}")


if __name__ == "__main__":
    build()
