"""
Fix the regen offer templates' page layout and cover-letter emphasis.

Two problems, both visible on a generated offer:

1. PAGE BREAKS — the whole 12-page document carried exactly ONE page break
   (before MAKE LIST). Every other section started wherever the previous one
   happened to end, so e.g. the National Energy Award block stranded at the
   bottom of the cover letter. Every section heading now starts a fresh page.

   Headings are identified by their formatting signature — 15 pt, bold,
   colour 1F2937 — which is exactly the set of real section headings and
   excludes the 11 pt Heading-3 sub-labels (Client Details / Marketing Person
   Details / Technical Data) that belong with their tables.

   A heading immediately following another heading does NOT break, otherwise
   "TECHNICAL OFFER" would sit alone on a page with "REGENERATIVE BURNERS"
   pushed to the next one.

2. COVER-LETTER BOLD — the entire recipient block and the entire signature
   block were bold, so nothing stood out. Only the addressee, the REF/DATE
   labels, the subject and the signatory's name stay bold now; the street
   address, phone, email and website lines go regular.

Run once against the gas and oil templates, then rebuild the dual template
(it is derived from the gas one):

    python patch_regen_layout.py
    python build_regen_dual_template.py

Idempotent — re-running re-applies the same properties.
"""

import os

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATES = ["Regen_Offer_Template.docx", "Regen_Oil_Offer_Template.docx"]

HEADING_PT = 15.0
HEADING_COLOR = RGBColor(0x1F, 0x29, 0x37)

# Cover-letter lines that should NOT be bold (matched on a distinctive fragment).
LETTER_UNBOLD = [
    "Kind Attn:",
    "{{company_address}}",
    "Contact No.:",
    "Email: {{email}}",
    "Regards,",
    "ENCON Thermal Engineers Pvt. Ltd.",
    "53 K.M Stone",
    "Mob: {{technical_phone}}",
    "Email: {{technical_email}}",
    "Website:",
]
LETTER_END = "NATIONAL ENERGY EFFICIENCY INNOVATION AWARD"


def _is_heading(p):
    r = p.runs[0] if p.runs else None
    if r is None or not p.text.strip():
        return False
    if not r.font.size or round(r.font.size.pt, 1) != HEADING_PT:
        return False
    if not r.font.bold:
        return False
    col = r.font.color
    return bool(col and col.rgb == HEADING_COLOR)


def _set_page_break_before(p, on=True):
    pPr = p._element.get_or_add_pPr()
    for old in pPr.findall(qn("w:pageBreakBefore")):
        pPr.remove(old)
    if on:
        pPr.append(OxmlElement("w:pageBreakBefore"))
    # MAKE LIST carried a manual <w:br type="page"/> inside its own runs. Left
    # in place alongside pageBreakBefore it would break twice and leave a blank
    # page, so drop the inline one — the attribute now owns the break.
    for br in p._element.findall(".//" + qn("w:br")):
        if br.get(qn("w:type")) == "page":
            br.getparent().remove(br)


def patch(path):
    d = Document(path)
    paras = d.paragraphs

    # ── 1. page break before every section heading ───────────────────────────
    # Walk the document BODY, not just the paragraphs — a table between two
    # headings (MAKE LIST -> table -> PRICE SCHEDULE) makes them adjacent in
    # doc.paragraphs, which would wrongly suppress the second one's break.
    from docx.text.paragraph import Paragraph

    by_el = {p._element: p for p in paras}
    breaks = 0
    prev_was_heading = False
    for el in d.element.body:
        if el.tag == qn("w:tbl"):
            prev_was_heading = False       # a table separates the two headings
            continue
        p = by_el.get(el)
        if p is None or not p.text.strip():
            continue                       # blank/image paragraph — carry state
        if _is_heading(p):
            # Don't break when the previous line was itself a heading, or this
            # heading would land alone on a page.
            _set_page_break_before(p, not prev_was_heading)
            if not prev_was_heading:
                breaks += 1
            prev_was_heading = True
        else:
            prev_was_heading = False

    # ── 2. tone down the cover letter ────────────────────────────────────────
    unbolded = 0
    for p in paras:
        t = p.text.strip()
        if t.startswith(LETTER_END):
            break                          # letter ends here
        if any(frag in t for frag in LETTER_UNBOLD):
            for run in p.runs:
                run.font.bold = False
            unbolded += 1

    d.save(path)
    print(f"  OK  {os.path.basename(path)} — {breaks} page breaks, "
          f"{unbolded} letter lines un-bolded")


if __name__ == "__main__":
    for name in TEMPLATES:
        patch(os.path.join(BASE_DIR, name))
    print("Now run: python build_regen_dual_template.py")
