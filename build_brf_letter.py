"""Tidy the BRF offer's covering letter into the vertical offer's shape.

What the source document left behind:

  * one sentence of the letter — "We are enclosing here with our most
    competitive..." — set at 15pt bold, so it read as a heading in the middle
    of a paragraph of prose;
  * the address split over two lines, which printed the city twice when the
    address field and the city field said the same thing;
  * labels that did not agree with each other: "Email." with a full stop,
    "Cont.No. :" with a space before the colon, "Mob" with no colon at all;
  * a full-width three-column YOUR REF / OUR REF / DATE table pushing the date
    to the far edge of the page, when the cover already names all three.

The vertical offer's letter is ten-point throughout, one address line, matched
labels, and only OUR REF and DATE — the enquiry is on the cover.

Run:  python build_brf_letter.py
Run it after build_brf_cover.py and before restyle_brf_offer.py.
Idempotent.
"""
import os
import shutil
import sys

from docx import Document
from docx.shared import Pt

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DEFAULT_TARGET = os.path.join(BASE_DIR, "BRF_Offer_Template.docx")

BODY_PT = 10.0

# text that starts a paragraph -> what it should say instead. None deletes it.
REWRITE = {
    "Kind Attn.:": "Kind Attn.: {{ poc_name }}{{ poc_designation_paren }}",
    "Email. {{ email }}": "Contact No.: {{ mobile_no }}",
    "Cont.No. :": "Email: {{ email }}",
    "Regards": "Regards,",
    "Mob {{ marketing_phone }}": "Mob: {{ marketing_phone }}",
    "{{ company_city_state }}": None,
}

# Prose that must not be sized like a heading, matched on its opening words.
FORCE_BODY = ("We are enclosing here with",)


W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _set(par, text):
    """Replace a paragraph's text, hyperlinks included.

    par.runs skips runs inside a w:hyperlink, so writing through it left the
    old hyperlink text in place and appended the new text after it — which is
    how "Email. ...@jindalsteel.inContact No.: 9893496476" happened.
    """
    nodes = par._p.findall(".//" + W + "t")
    if nodes:
        nodes[0].text = text
        for n in nodes[1:]:
            n.text = ""
    else:
        par.add_run(text)


def main():
    target = sys.argv[1] if len(sys.argv) > 1 else DEFAULT_TARGET
    if not os.path.exists(target):
        print("not found: " + target)
        return 1
    shutil.copyfile(target, target + ".bak")
    doc = Document(target)

    changed = 0
    for p in list(doc.paragraphs):
        t = (p.text or "").strip()
        if not t:
            continue
        for needle, replacement in REWRITE.items():
            if t.startswith(needle):
                if replacement is None:
                    p._p.getparent().remove(p._p)
                else:
                    _set(p, replacement)
                changed += 1
                break
        else:
            if t.startswith(FORCE_BODY):
                for r in p.runs:
                    r.font.size = Pt(BODY_PT)
                    r.bold = False
                changed += 1

    # The reference table loses its YOUR REF column; the cover carries it.
    for tb in doc.tables:
        head = [(c.text or "").strip().upper() for c in tb.rows[0].cells]
        if head[:3] == ["YOUR REF:", "OUR REF:", "DATE:"]:
            for row in tb.rows:
                cell = row.cells[0]
                cell._tc.getparent().remove(cell._tc)
            changed += 1
            break

    doc.save(target)
    print("covering letter tidied: %d paragraphs and cells changed" % changed)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
