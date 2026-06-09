"""Build Combined_Offer_Template.docx by cloning the VLPH offer template
and turning it into a UNIFIED multi-equipment offer:

  * ONE cover page, customer block, company profile, T&C (kept from VLPH).
  * Technical Specifications section -> a {%tr for eq in equipments %} table
    that renders one row per equipment (name + key specs).
  * Price Schedule (Annexure) -> a {%tr for line in price_lines %} table that
    lists one priced line per equipment + a GRAND TOTAL.

Rendered self-contained via docxtpl from /api/generate-combined-offer.
Modelled on build_hpu_template_from_vlph.py (minimal shell, Annexures II/V/VI
+ supervision removed).

Placeholders introduced:
  equipments  -> [{ name, specs }]            (technical section)
  price_lines -> [{ sno, name, qty, unit_price, total }]   (price schedule)
  grand_total                                  (combined total)
  + the shared customer block + tnc_* (same names as the recup/hpu offers)
"""
from __future__ import annotations

import os
import shutil

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.enum.table import WD_ROW_HEIGHT_RULE

SOURCE = "Offer_Template.docx"
TARGET = "Combined_Offer_Template.docx"
HEADING = "COMBINED EQUIPMENT OFFER"


def _delete_vlph_scope_body(doc: Document) -> None:
    body = doc.element.body
    elements = list(body)
    start_idx = end_idx = None
    for i, el in enumerate(elements):
        if el.tag.split('}')[-1] != 'p':
            continue
        txt = ''.join(t.text or '' for t in el.iter(qn('w:t'))).strip()
        if start_idx is None and txt == 'SCOPE OF SUPPLY':
            start_idx = i
        elif start_idx is not None and txt.startswith('ANNEXURE I'):
            end_idx = i
            break
    if start_idx is not None and end_idx is not None:
        for el in elements[start_idx:end_idx]:
            body.remove(el)
    for el in body.iter(qn('w:p')):
        t_els = list(el.iter(qn('w:t')))
        joined = ''.join((t.text or '') for t in t_els).upper()
        if 'PREHEATERS AND DRYERS' in joined:
            for t in t_els:
                t.text = ''
            if t_els:
                t_els[0].text = HEADING


def _replace_spec_table_with_equipment_loop(table) -> None:
    """Wipe the 24-row tech-data table; make it loop one row per equipment."""
    tbl_xml = table._element
    for tr in list(tbl_xml.findall(qn('w:tr')))[1:]:
        tbl_xml.remove(tr)
    head = table.rows[0]
    head.cells[0].text = 'Equipment'
    head.cells[1].text = 'Specifications'
    for cell in head.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    def add(c0, c1):
        row = table.add_row()
        row.cells[0].text = c0
        row.cells[1].text = c1
        return row

    add('{%tr for eq in equipments %}', '')
    add('{{ eq.name }}', '{{ eq.specs }}')
    add('{%tr endfor %}', '')


def _strip_project_name_from_cover_box(doc: Document) -> None:
    target = None
    for t in doc.tables:
        if not t.rows or not t.rows[0].cells:
            continue
        if t.rows[0].cells[0].text.strip().startswith('Project / Equipment'):
            target = t
            break
    if target is None:
        return
    tbl_xml = target._element
    for tr in list(tbl_xml.findall(qn('w:tr'))):
        first = ''.join(t.text or '' for t in tr.iter(qn('w:t'))).strip()
        if first.lower().startswith('project name'):
            tbl_xml.remove(tr)
            return


def _remove_annexure_ii(doc: Document) -> None:
    for t in doc.tables:
        if not t.rows or len(t.rows[0].cells) < 2:
            continue
        if 'ANNEXURE NO.' in t.rows[0].cells[0].text.strip().upper():
            tbl_xml = t._element
            for tr in list(tbl_xml.findall(qn('w:tr'))):
                row_txt = ''.join(x.text or '' for x in tr.iter(qn('w:t'))).upper()
                if ('ANNEXURE II' in row_txt and 'ANNEXURE III' not in row_txt and 'EXCLUS' in row_txt):
                    tbl_xml.remove(tr)
            break
    body = doc.element.body
    elements = list(body)
    heading_idx = None
    for i, el in enumerate(elements):
        if el.tag.split('}')[-1] != 'p':
            continue
        txt = ''.join(t.text or '' for t in el.iter(qn('w:t'))).strip().upper()
        if txt.startswith('ANNEXURE II') and not txt.startswith('ANNEXURE III') and 'EXCLUS' in txt:
            heading_idx = i
            break
    if heading_idx is None:
        return
    i = heading_idx
    while i < len(elements):
        el = elements[i]
        tag = el.tag.split('}')[-1]
        if tag == 'sectPr':
            break
        if tag == 'p' and i != heading_idx:
            txt = ''.join(t.text or '' for t in el.iter(qn('w:t'))).strip().upper()
            if txt.startswith('ANNEXURE'):
                break
        body.remove(el)
        i += 1


def _remove_annexure_i_scope_table(doc: Document) -> None:
    target = None
    for t in doc.tables:
        if not t.rows or len(t.rows[0].cells) < 2:
            continue
        head_l = t.rows[0].cells[0].text.strip().upper()
        head_r = t.rows[0].cells[1].text.strip().upper()
        if head_l == 'S. NO.' and 'ITEM DESCRIPTION' not in head_r and 10 <= len(t.rows) <= 15:
            target = t
            break
    if target is None:
        return
    tbl_xml = target._element
    parent = tbl_xml.getparent()
    nxt = tbl_xml.getnext()
    parent.remove(tbl_xml)
    if nxt is not None and nxt.tag.split('}')[-1] == 'p':
        if not ''.join(t.text or '' for t in nxt.iter(qn('w:t'))).strip():
            parent.remove(nxt)


def _remove_supervision_table(doc: Document) -> None:
    target = None
    for t in doc.tables:
        if not t.rows or not t.rows[0].cells:
            continue
        head = t.rows[0].cells[0].text.strip()
        if head.startswith('Supervision Charges') or head.startswith('{%tr if supervision_include'):
            target = t
            break
    if target is None:
        return
    tbl_xml = target._element
    parent = tbl_xml.getparent()
    nxt = tbl_xml.getnext()
    parent.remove(tbl_xml)
    if nxt is not None and nxt.tag.split('}')[-1] == 'p':
        if not ''.join(t.text or '' for t in nxt.iter(qn('w:t'))).strip():
            parent.remove(nxt)


def _remove_annexure_section(doc: Document, roman: str, label_keyword: str) -> None:
    target_prefix = f'ANNEXURE {roman}'

    def _is_target(txt_upper: str) -> bool:
        if not txt_upper.startswith(target_prefix):
            return False
        rest = txt_upper[len(target_prefix):].lstrip()
        if rest and rest[0].isalpha() and rest[0] in 'IVX':
            return False
        return label_keyword in txt_upper

    for t in doc.tables:
        if not t.rows or len(t.rows[0].cells) < 2:
            continue
        if 'ANNEXURE NO.' in t.rows[0].cells[0].text.strip().upper():
            tbl_xml = t._element
            for tr in list(tbl_xml.findall(qn('w:tr'))):
                row_txt = ''.join(x.text or '' for x in tr.iter(qn('w:t'))).upper()
                if _is_target(row_txt):
                    tbl_xml.remove(tr)
            break
    body = doc.element.body
    elements = list(body)
    heading_idx = None
    for i, el in enumerate(elements):
        if el.tag.split('}')[-1] != 'p':
            continue
        txt = ''.join(t.text or '' for t in el.iter(qn('w:t'))).strip().upper()
        if _is_target(txt):
            heading_idx = i
            break
    if heading_idx is None:
        return
    i = heading_idx
    while i < len(elements):
        el = elements[i]
        tag = el.tag.split('}')[-1]
        if tag == 'sectPr':
            break
        if tag == 'p' and i != heading_idx:
            txt = ''.join(t.text or '' for t in el.iter(qn('w:t'))).strip().upper()
            if txt.startswith('ANNEXURE'):
                break
        body.remove(el)
        i += 1


def _scrub_reference_list_mention(doc: Document) -> None:
    needle = 'a representative reference list is included in Annexure V'
    for p in doc.paragraphs:
        if needle.lower() not in p.text.lower():
            continue
        full = ''.join(r.text or '' for r in p.runs)
        idx = full.lower().find('— ' + needle.lower())
        if idx == -1:
            idx = full.lower().find(needle.lower())
        if idx == -1:
            continue
        cleaned = full[:idx].rstrip(' —-').rstrip()
        if not cleaned.endswith('.'):
            cleaned += '.'
        for r in p.runs:
            r.text = ''
        if p.runs:
            p.runs[0].text = cleaned
        else:
            p.add_run(cleaned)
        return


def _rename_annexure_i_heading(doc: Document) -> None:
    body = doc.element.body
    for el in body.iter(qn('w:p')):
        t_els = list(el.iter(qn('w:t')))
        joined = ''.join((t.text or '') for t in t_els)
        ju = joined.upper()
        if 'ANNEXURE I' in ju and 'SCOPE OF SUPPLY' in ju:
            if 'COMBINED' in ju:
                return
            if len(joined.strip()) > 60:
                continue
            for t in t_els:
                t.text = ''
            if t_els:
                t_els[0].text = 'ANNEXURE I — PRICE SCHEDULE'
            return


def _renumber_annexures_after_removal(doc: Document) -> None:
    renames = [('ANNEXURE III', 'ANNEXURE II'), ('Annexure III', 'Annexure II'),
               ('ANNEXURE IV', 'ANNEXURE III'), ('Annexure IV', 'Annexure III')]

    def _rewrite(p_el) -> None:
        t_els = list(p_el.iter(qn('w:t')))
        if not t_els:
            return
        joined = ''.join((t.text or '') for t in t_els)
        new = joined
        for old, repl in renames:
            new = new.replace(old, repl)
        if new == joined:
            return
        for t in t_els:
            t.text = ''
        t_els[0].text = new

    for p_el in doc.element.body.iter(qn('w:p')):
        _rewrite(p_el)


def _rebuild_price_schedule(doc: Document) -> None:
    """Annexure III Price Schedule -> one line per equipment + GRAND TOTAL."""
    table = None
    for t in doc.tables:
        if not t.rows or len(t.rows[0].cells) < 2:
            continue
        if 'ITEM DESCRIPTION' in t.rows[0].cells[1].text.upper():
            table = t
            break
    if table is None:
        print('Price Schedule table not found.')
        return
    tbl_xml = table._element
    for tr in list(tbl_xml.findall(qn('w:tr')))[1:]:
        tbl_xml.remove(tr)

    def add(c1, c2, c3, c4, c5, *, bold=False):
        r = table.add_row()
        for cell, txt in zip(r.cells, (c1, c2, c3, c4, c5)):
            cell.text = txt
        if bold:
            for cell in r.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True

    add('{%tr for line in price_lines %}', '', '', '', '')
    add('{{ line.sno }}', '{{ line.name }}', '{{ line.qty }}', '{{ line.unit_price }}', '{{ line.total }}')
    add('{%tr endfor %}', '', '', '', '')
    add('', 'GRAND TOTAL', '', '', '{{ grand_total }}', bold=True)
    # Commercial adjustments applied once to the combined total -> Final Total.
    # Each row is shown only when its amount is non-zero.
    add("{%tr if show_pf %}", '', '', '', '')
    add('', 'Packaging & Forwarding', '', '', '{{ pf_amount }}')
    add("{%tr endif %}", '', '', '', '')
    add("{%tr if show_design %}", '', '', '', '')
    add('', 'Designing', '', '', '{{ design_amount }}')
    add("{%tr endif %}", '', '', '', '')
    add("{%tr if show_neg %}", '', '', '', '')
    add('', 'Negotiation', '', '', '{{ neg_amount }}')
    add("{%tr endif %}", '', '', '', '')
    add("{%tr if show_transport %}", '', '', '', '')
    add('', 'Transport', '', '', '{{ transport_amount }}')
    add("{%tr endif %}", '', '', '', '')
    add('', 'FINAL TOTAL', '', '', '{{ final_total }}', bold=True)


def _pad_table_rows(doc: Document, min_height_cm: float = 0.7) -> None:
    height = Cm(min_height_cm)
    for t in doc.tables:
        if not t.rows or len(t.rows[0].cells) < 2:
            continue
        if 'ITEM DESCRIPTION' in t.rows[0].cells[1].text.strip().upper() or \
           t.rows[0].cells[1].text.strip() == 'Specifications':
            for row in t.rows:
                first = row.cells[0].text.strip() if row.cells else ''
                if first.startswith('{%'):
                    continue
                row.height = height
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def _fill_annexure_i_scope(doc: Document) -> None:
    """Annexure I — Scope of Supply: replace the VLPH component list with a
    loop over the combined equipment (so the scope reads as the equipment
    being supplied), keeping the annexure itself intact."""
    target = None
    for t in doc.tables:
        if not t.rows or len(t.rows[0].cells) < 2:
            continue
        head_l = t.rows[0].cells[0].text.strip().upper()
        head_r = t.rows[0].cells[1].text.strip().upper()
        if head_l == 'S. NO.' and 'ITEM DESCRIPTION' not in head_r and 10 <= len(t.rows) <= 15:
            target = t
            break
    if target is None:
        print('Annexure I scope table not found — skipping.')
        return
    tbl_xml = target._element
    for tr in list(tbl_xml.findall(qn('w:tr')))[1:]:
        tbl_xml.remove(tr)

    def add(c0, c1):
        r = target.add_row()
        r.cells[0].text = c0
        r.cells[1].text = c1

    add('{%tr for eq in equipments %}', '')
    add('{{ loop.index }}.', '{{ eq.name }}')
    add('{%tr endfor %}', '')


def main() -> None:
    if not os.path.exists(SOURCE):
        raise SystemExit(f"missing: {SOURCE}")
    if os.path.exists(TARGET) and not os.path.exists(TARGET + '.bak'):
        shutil.copy(TARGET, TARGET + '.bak')
    shutil.copy(SOURCE, TARGET)
    doc = Document(TARGET)
    print(f'Cloned {SOURCE} -> {TARGET}: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables')

    # Full proper-offer structure (like the VLPH/HLPH offers): keep Scope of
    # Supply, Exclusions, Reference List, Make List and Supervision. Only the
    # technical-specs table, the Annexure-I scope list and the price schedule
    # are made equipment-driven.
    _delete_vlph_scope_body(doc)
    _replace_spec_table_with_equipment_loop(doc.tables[5])
    _fill_annexure_i_scope(doc)
    _strip_project_name_from_cover_box(doc)
    _rebuild_price_schedule(doc)
    _pad_table_rows(doc)

    doc.save(TARGET)
    print(f'Saved -> {TARGET}')


if __name__ == '__main__':
    main()
