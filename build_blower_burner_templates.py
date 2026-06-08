"""Build Blower_Offer_Template.docx and Burner_Offer_Template.docx by
cloning the VLPH offer template (Offer_Template.docx) and surgically
replacing the ladle-preheater content with stand-alone Blower / Burner
content.

Modelled on build_hpu_template_from_vlph.py (same minimal shell: specs
table + single-line Price Schedule + T&Cs, no fabricated scope list) plus
the recuperator-style single-line Price Schedule rebuild so the offer can
be rendered self-contained via docxtpl from /api/generate-blower-quote and
/api/generate-burner-quote.

What is kept from the VLPH template:
  * Cover page, ENCON banner, cover letter
  * Table of Contents + List of Annexures (renumbered)
  * Section 1 Company Profile, Section 2 About the Client
  * Section heading "3. TECHNICAL SPECIFICATIONS"
  * ANNEXURE III (Price Schedule — rebuilt single-line)
  * ANNEXURE IV (T&C 12-row table)

What is replaced / removed (same as the HPU template):
  * Equipment heading -> "BLOWER" / "ENCON BURNER"
  * 24-row Technical Specs table -> short equipment spec table
  * VLPH scope-of-supply body -> removed
  * Annexure I scope table -> removed (no fabricated accessory list)
  * Annexure II (Exclusions), V (Reference List), VI (Make List) -> removed
  * Supervision Charges sub-table -> removed

Template placeholders introduced (filled by the generate endpoints):
  Common:  equipment_name, item_qty, unit_price, total_price, grand_total
  Blower:  blower_model, blower_hp, blower_airflow, blower_pressure
  Burner:  burner_model, burner_fuel, burner_capacity
"""
from __future__ import annotations

import os
import shutil

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.enum.table import WD_ROW_HEIGHT_RULE

SOURCE = "Offer_Template.docx"


# ── Per-equipment configuration ────────────────────────────────────────────
# heading      : replaces the VLPH 'PREHEATERS AND DRYERS' section title
# specs        : (label, '{{ placeholder }}') rows for the Tech Specs table
# price_label  : Annexure III price-schedule item description (a placeholder
#                so the rendered name matches the equipment chosen)
_EQUIPMENT = {
    "blower": {
        "target": "Blower_Offer_Template.docx",
        "heading": "BLOWER",
        "specs": [
            ('Equipment',       '{{ equipment_name }}'),
            ('Model',           '{{ blower_model }}'),
            ('Motor Rating',    '{{ blower_hp }} HP'),
            ('Air Flow',        '{{ blower_airflow }} Nm³/hr'),
            ('Static Pressure', '{{ blower_pressure }}'),
            ('Quantity',        '{{ item_qty }}'),
        ],
        "specs_header": "Blower Specifications",
    },
    "burner": {
        "target": "Burner_Offer_Template.docx",
        "heading": "ENCON BURNER",
        "specs": [
            ('Equipment',        '{{ equipment_name }}'),
            ('Model',            '{{ burner_model }}'),
            ('Fuel',             '{{ burner_fuel }}'),
            ('Firing Capacity',  '{{ burner_capacity }}'),
            ('Quantity',         '{{ item_qty }}'),
        ],
        "specs_header": "Burner Specifications",
    },
}


def _replace_spec_table(table, specs, header_label: str) -> None:
    """Wipe the VLPH 24-row tech-data table and re-fill with the equipment
    specs. Keep the header row from VLPH; relabel it."""
    tbl_xml = table._element
    for tr in list(tbl_xml.findall(qn('w:tr')))[1:]:
        tbl_xml.remove(tr)
    head = table.rows[0]
    head.cells[0].text = header_label
    head.cells[1].text = ''
    head.cells[0].merge(head.cells[1])
    for p in head.cells[0].paragraphs:
        for r in p.runs:
            r.bold = True
    for label, value in specs:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value


def _delete_vlph_scope_body(doc: Document, heading: str) -> None:
    """Delete every element from the inline 'SCOPE OF SUPPLY' paragraph up
    to (but not including) the 'ANNEXURE I' heading, then flip the section
    title 'PREHEATERS AND DRYERS' to the equipment heading."""
    body = doc.element.body
    elements = list(body)

    start_idx, end_idx = None, None
    for i, el in enumerate(elements):
        if el.tag.split('}')[-1] != 'p':
            continue
        txt = ''.join(t.text or '' for t in el.iter(qn('w:t'))).strip()
        if start_idx is None and txt == 'SCOPE OF SUPPLY':
            start_idx = i
        elif start_idx is not None and txt.startswith('ANNEXURE I'):
            end_idx = i
            break

    if start_idx is not None and end_idx is not None:
        for el in elements[start_idx:end_idx]:
            body.remove(el)
    else:
        print('Could not locate VLPH scope body markers — skipping delete.')

    # Flip the section title. The VLPH heading is a Jinja conditional ending
    # in 'PREHEATERS AND DRYERS'; match that stable tail and stamp the new
    # equipment heading.
    for el in body.iter(qn('w:p')):
        t_els = list(el.iter(qn('w:t')))
        joined = ''.join((t.text or '') for t in t_els).upper()
        if 'PREHEATERS AND DRYERS' in joined:
            for t in t_els:
                t.text = ''
            if t_els:
                t_els[0].text = heading


def _strip_project_name_from_cover_box(doc: Document) -> None:
    """Remove the 'Project Name' row from the cover-box table."""
    target = None
    for t in doc.tables:
        if not t.rows or not t.rows[0].cells:
            continue
        if t.rows[0].cells[0].text.strip().startswith('Project / Equipment'):
            target = t
            break
    if target is None:
        return
    tbl_xml = target._element
    for tr in list(tbl_xml.findall(qn('w:tr'))):
        first_cell_txt = ''.join(t.text or '' for t in tr.iter(qn('w:t'))).strip()
        if first_cell_txt.lower().startswith('project name'):
            tbl_xml.remove(tr)
            return


def _remove_annexure_ii(doc: Document) -> None:
    """Strip Annexure II — Exclusions (heading + bullets) and its
    List-of-Annexures row."""
    for t in doc.tables:
        if not t.rows or len(t.rows[0].cells) < 2:
            continue
        head = [c.text.strip().upper() for c in t.rows[0].cells]
        if 'ANNEXURE NO.' in head[0]:
            tbl_xml = t._element
            for tr in list(tbl_xml.findall(qn('w:tr'))):
                row_txt = ''.join(x.text or '' for x in tr.iter(qn('w:t'))).upper()
                if ('ANNEXURE II' in row_txt and 'ANNEXURE III' not in row_txt
                        and 'EXCLUS' in row_txt):
                    tbl_xml.remove(tr)
            break

    body = doc.element.body
    elements = list(body)
    heading_idx = None
    for i, el in enumerate(elements):
        if el.tag.split('}')[-1] != 'p':
            continue
        txt = ''.join(t.text or '' for t in el.iter(qn('w:t'))).strip().upper()
        if (txt.startswith('ANNEXURE II') and not txt.startswith('ANNEXURE III')
                and 'EXCLUS' in txt):
            heading_idx = i
            break
    if heading_idx is None:
        return
    i = heading_idx
    while i < len(elements):
        el = elements[i]
        tag = el.tag.split('}')[-1]
        if tag == 'sectPr':
            break
        if tag == 'p' and i != heading_idx:
            txt = ''.join(t.text or '' for t in el.iter(qn('w:t'))).strip().upper()
            if txt.startswith('ANNEXURE'):
                break
        body.remove(el)
        i += 1


def _remove_annexure_i_scope_table(doc: Document) -> None:
    """Strip the Annexure I scope-of-supply table (no fabricated accessory
    list). Identified by its 'S. No. | Item Description'-less header with a
    10–15 row body (avoids matching the 5-row Section TOC)."""
    target = None
    for t in doc.tables:
        if not t.rows or len(t.rows[0].cells) < 2:
            continue
        head_l = t.rows[0].cells[0].text.strip().upper()
        head_r = t.rows[0].cells[1].text.strip().upper()
        if head_l == 'S. NO.' and 'ITEM DESCRIPTION' not in head_r and 10 <= len(t.rows) <= 15:
            target = t
            break
    if target is None:
        return
    tbl_xml = target._element
    parent = tbl_xml.getparent()
    nxt = tbl_xml.getnext()
    parent.remove(tbl_xml)
    if nxt is not None and nxt.tag.split('}')[-1] == 'p':
        if not ''.join(t.text or '' for t in nxt.iter(qn('w:t'))).strip():
            parent.remove(nxt)


def _remove_supervision_table(doc: Document) -> None:
    """Strip the Supervision Charges sub-table below the Price Schedule."""
    target = None
    for t in doc.tables:
        if not t.rows or not t.rows[0].cells:
            continue
        head = t.rows[0].cells[0].text.strip()
        if head.startswith('Supervision Charges') or head.startswith('{%tr if supervision_include'):
            target = t
            break
    if target is None:
        return
    tbl_xml = target._element
    parent = tbl_xml.getparent()
    nxt = tbl_xml.getnext()
    parent.remove(tbl_xml)
    if nxt is not None and nxt.tag.split('}')[-1] == 'p':
        if not ''.join(t.text or '' for t in nxt.iter(qn('w:t'))).strip():
            parent.remove(nxt)


def _remove_annexure_section(doc: Document, roman: str, label_keyword: str) -> None:
    """Generic Annexure remover (body block + List-of-Annexures row)."""
    target_prefix = f'ANNEXURE {roman}'

    def _is_target(txt_upper: str) -> bool:
        if not txt_upper.startswith(target_prefix):
            return False
        rest = txt_upper[len(target_prefix):].lstrip()
        if rest and rest[0].isalpha() and rest[0] in 'IVX':
            return False
        return label_keyword in txt_upper

    for t in doc.tables:
        if not t.rows or len(t.rows[0].cells) < 2:
            continue
        head = [c.text.strip().upper() for c in t.rows[0].cells]
        if 'ANNEXURE NO.' in head[0]:
            tbl_xml = t._element
            for tr in list(tbl_xml.findall(qn('w:tr'))):
                row_txt = ''.join(x.text or '' for x in tr.iter(qn('w:t'))).upper()
                if _is_target(row_txt):
                    tbl_xml.remove(tr)
            break

    body = doc.element.body
    elements = list(body)
    heading_idx = None
    for i, el in enumerate(elements):
        if el.tag.split('}')[-1] != 'p':
            continue
        txt = ''.join(t.text or '' for t in el.iter(qn('w:t'))).strip().upper()
        if _is_target(txt):
            heading_idx = i
            break
    if heading_idx is None:
        return
    i = heading_idx
    while i < len(elements):
        el = elements[i]
        tag = el.tag.split('}')[-1]
        if tag == 'sectPr':
            break
        if tag == 'p' and i != heading_idx:
            txt = ''.join(t.text or '' for t in el.iter(qn('w:t'))).strip().upper()
            if txt.startswith('ANNEXURE'):
                break
        body.remove(el)
        i += 1


def _scrub_reference_list_mention(doc: Document) -> None:
    """Drop the dangling '...reference list is included in Annexure V'
    sentence in the Company Profile once Annexure V is removed."""
    needle = 'a representative reference list is included in Annexure V'
    for p in doc.paragraphs:
        if needle.lower() not in p.text.lower():
            continue
        full = ''.join(r.text or '' for r in p.runs)
        idx = full.lower().find('— ' + needle.lower())
        if idx == -1:
            idx = full.lower().find(needle.lower())
        if idx == -1:
            continue
        cleaned = full[:idx].rstrip(' —-').rstrip()
        if not cleaned.endswith('.'):
            cleaned += '.'
        for r in p.runs:
            r.text = ''
        if p.runs:
            p.runs[0].text = cleaned
        else:
            p.add_run(cleaned)
        return


def _rename_annexure_i_heading(doc: Document, equipment_heading: str) -> None:
    """Extend 'ANNEXURE I — SCOPE OF SUPPLY' with the equipment label."""
    body = doc.element.body
    for el in body.iter(qn('w:p')):
        t_els = list(el.iter(qn('w:t')))
        joined = ''.join((t.text or '') for t in t_els)
        joined_u = joined.upper()
        if 'ANNEXURE I' in joined_u and 'SCOPE OF SUPPLY' in joined_u:
            if equipment_heading.upper() in joined_u:
                return
            if len(joined.strip()) > 60:
                continue
            new_text = f'ANNEXURE I — SCOPE OF SUPPLY: {equipment_heading}'
            for t in t_els:
                t.text = ''
            if t_els:
                t_els[0].text = new_text
            return


def _renumber_annexures_after_removal(doc: Document) -> None:
    """After Annexure II removal, shift III->II and IV->III."""
    renames = [
        ('ANNEXURE III', 'ANNEXURE II'),
        ('Annexure III', 'Annexure II'),
        ('ANNEXURE IV',  'ANNEXURE III'),
        ('Annexure IV',  'Annexure III'),
    ]

    def _rewrite(p_el) -> bool:
        t_els = list(p_el.iter(qn('w:t')))
        if not t_els:
            return False
        joined = ''.join((t.text or '') for t in t_els)
        new = joined
        for old, repl in renames:
            new = new.replace(old, repl)
        if new == joined:
            return False
        for t in t_els:
            t.text = ''
        t_els[0].text = new
        return True

    body = doc.element.body
    for p_el in body.iter(qn('w:p')):
        _rewrite(p_el)


def _rebuild_price_schedule(doc: Document) -> None:
    """Rebuild the 3-row Price Schedule (Annexure III, 'ITEM DESCRIPTION'
    header) into a single-line item + TOTAL using docxtpl placeholders."""
    table = None
    for t in doc.tables:
        if not t.rows or len(t.rows[0].cells) < 2:
            continue
        if 'ITEM DESCRIPTION' in t.rows[0].cells[1].text.upper():
            table = t
            break
    if table is None:
        print('Price Schedule table not found — skipping rebuild.')
        return

    tbl_xml = table._element
    for tr in list(tbl_xml.findall(qn('w:tr')))[1:]:
        tbl_xml.remove(tr)

    def add(c1, c2, c3, c4, c5, *, bold=False):
        r = table.add_row()
        for cell, txt in zip(r.cells, (c1, c2, c3, c4, c5)):
            cell.text = txt
        if bold:
            for cell in r.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True

    add("1.", "{{ equipment_name }}", "{{ item_qty }}",
        "{{ unit_price }}", "{{ total_price }}")
    add("", "TOTAL", "", "", "{{ grand_total }}", bold=True)


def _pad_table_rows(doc: Document, min_height_cm: float = 0.75) -> None:
    """Give the Price Schedule rows some breathing room."""
    height = Cm(min_height_cm)
    for t in doc.tables:
        if not t.rows or len(t.rows[0].cells) < 2:
            continue
        head_r = t.rows[0].cells[1].text.strip().upper()
        if 'ITEM DESCRIPTION' in head_r:
            for row in t.rows:
                first = row.cells[0].text.strip() if row.cells else ''
                if first.startswith('{%tr') or first.startswith('{%p'):
                    continue
                row.height = height
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def build_one(key: str) -> None:
    cfg = _EQUIPMENT[key]
    target = cfg["target"]

    if os.path.exists(target) and not os.path.exists(target + '.bak'):
        shutil.copy(target, target + '.bak')

    shutil.copy(SOURCE, target)
    doc = Document(target)
    print(f'Cloned {SOURCE} -> {target}: '
          f'{len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables')

    _delete_vlph_scope_body(doc, cfg["heading"])
    _replace_spec_table(doc.tables[5], cfg["specs"], cfg["specs_header"])
    _rename_annexure_i_heading(doc, cfg["heading"])
    _remove_annexure_i_scope_table(doc)
    _strip_project_name_from_cover_box(doc)
    _remove_annexure_ii(doc)
    _remove_annexure_section(doc, 'V',  'REFERENCE')
    _remove_annexure_section(doc, 'VI', 'MAKE')
    _scrub_reference_list_mention(doc)
    _remove_supervision_table(doc)
    _renumber_annexures_after_removal(doc)
    _rebuild_price_schedule(doc)
    _pad_table_rows(doc, min_height_cm=0.75)

    doc.save(target)
    print(f'Saved -> {target}')


def main() -> None:
    if not os.path.exists(SOURCE):
        raise SystemExit(f"missing: {SOURCE}")
    for key in _EQUIPMENT:
        build_one(key)


if __name__ == '__main__':
    main()
