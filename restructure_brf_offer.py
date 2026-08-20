"""Give the BRF offer the vertical offer's *structure*, not just its fonts.

The earlier pass restyled type and left the document's architecture alone, so
it still read as the supplier's Word file in ENCON colours. Set the two side by
side and the vertical offer has a skeleton the BRF one simply did not have:

    vertical (Offer_Template.docx)          BRF, before this script
    ─────────────────────────────────       ────────────────────────
    cover                                   cover
    covering letter                         covering letter
    TABLE OF CONTENTS + annexure list       —
    1. COMPANY PROFILE  (6 subsections)     —
    2. ABOUT THE CLIENT (2 tables)          —
    3. TECHNICAL SPECIFICATIONS (table)     loose "Label : value" lines
    scope, as Heading 2 / Heading 3         everything flat at Heading 1
    ANNEXURE I…VI, each on its own page     unnumbered, no page breaks
    ANNEXURE V  reference list of clients   —
    ANNEXURE VI make list                   —
    terms as a two-column table             a run of loose sentences

Rather than retype any of that, the shared parts are lifted out of the vertical
template itself — paragraphs, tables, borders and all — so they cannot drift
apart. The tables BRF needs but does not have are cloned from their vertical
counterparts and refilled, which is what keeps the borders, shading and cell
padding identical instead of merely similar.

The commercial terms are BRF's own; only their presentation changes, from
sentences to the same labelled table the vertical offer uses. Where the
vertical offer takes a value from the quote ({{ tnc_prices }} and friends) so
does this, so both read from one place.

Run:  python restructure_brf_offer.py
Idempotent — a second run finds the contents page already there and stops.

Order matters. The full chain is
    build_brf_offer_template.py → build_brf_cover.py → build_brf_letter.py
    → restructure_brf_offer.py → patch_brf_offer_header.py
    → restyle_brf_offer.py
"""
import copy
import os
import re
import shutil
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.shared import Pt

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TARGET = os.path.join(BASE_DIR, "BRF_Offer_Template.docx")
SOURCE = os.path.join(BASE_DIR, "Offer_Template.docx")      # the vertical offer

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


# ── reading the source ────────────────────────────────────────────────────────

def body_children(doc):
    return [el for el in doc.element.body
            if el.tag in (f"{W}p", f"{W}tbl")]


def para_text(el):
    if el.tag != f"{W}p":
        return ""
    return " ".join("".join(t.text or "" for t in el.findall(f".//{W}t")).split())


def harvest(doc, start, end=None):
    """Every body element from the paragraph reading `start` up to `end`."""
    kids = body_children(doc)
    try:
        i = next(n for n, el in enumerate(kids) if para_text(el) == start)
    except StopIteration:
        raise SystemExit(f"not found in the vertical offer: {start!r}")
    if end is None:
        j = len(kids)
    else:
        j = next((n for n, el in enumerate(kids[i + 1:], i + 1)
                  if para_text(el) == end), len(kids))
    out = [copy.deepcopy(el) for el in kids[i:j]]
    # A harvested block must not bring the source document's section
    # definitions with it. A w:sectPr names header and footer parts by
    # relationship id, and those ids mean something else here — rId12 is a
    # footer over there and a hyperlink to our own website on this side, so
    # Word follows it to a part that is not a footer and gives up on the file.
    for el in out:
        for sectPr in el.findall(f".//{W}sectPr"):
            sectPr.getparent().remove(sectPr)
    return out


def port_images(elements, src_doc, dst_doc):
    """Carry the pictures across, and repoint them at the copy.

    A deep-copied drawing still names the relationship id it had in the source
    document. Left alone that id points at nothing here, and Word will not open
    the file. Each image part is added to this document and the id rewritten to
    whatever it is called on this side.
    """
    src, dst = src_doc.part, dst_doc.part
    moved = 0
    for el in elements:
        for blip in el.iter(f"{'{http://schemas.openxmlformats.org/drawingml/2006/main}'}blip"):
            rid = blip.get(qn("r:embed"))
            if not rid or rid not in src.rels:
                continue
            blip.set(qn("r:embed"), dst.relate_to(src.rels[rid].target_part, RT.IMAGE))
            moved += 1
    return moved


# ── writing ───────────────────────────────────────────────────────────────────

def _drop_rpr(run):
    """add_run only writes an rPr when it has something to say; often it has not."""
    existing = run._r.find(qn("w:rPr"))
    if existing is not None:
        run._r.remove(existing)


def cell_text(cell, text, bold=None):
    """Set a cell's text, keeping the formatting its first run already had."""
    para = cell.paragraphs[0]
    for extra in cell.paragraphs[1:]:
        extra._p.getparent().remove(extra._p)
    keep = copy.deepcopy(para.runs[0]._r.find(qn("w:rPr"))) if para.runs else None
    for r in list(para.runs):
        r._r.getparent().remove(r._r)
    run = para.add_run(str(text))
    if keep is not None:
        _drop_rpr(run)
        run._r.insert(0, copy.deepcopy(keep))
    if bold is not None:
        run.bold = bold
    return run


def refill(table, rows, header=True):
    """Resize a cloned table to `rows` and write them, header row kept as is."""
    from docx.table import _Row
    body_start = 1 if header else 0
    if len(table.rows) <= body_start:
        raise SystemExit("cloned table has no body row to copy")
    pattern = table.rows[body_start]._tr
    while len(table.rows) > body_start:
        tr = table.rows[body_start]._tr
        tr.getparent().remove(tr)
    for values in rows:
        tr = copy.deepcopy(pattern)
        table._tbl.append(tr)
        row = _Row(tr, table)
        for cell, val in zip(row.cells, values):
            cell_text(cell, val)
    return table


def make_heading(par, text, level=1, page_break=False):
    """Turn a paragraph into a real Word heading, so the outline means something."""
    keep = copy.deepcopy(par.runs[0]._r.find(qn("w:rPr"))) if par.runs else None
    for r in list(par.runs):
        r._r.getparent().remove(r._r)
    par.style = f"Heading {level}"
    run = par.add_run(text)
    if keep is not None:
        _drop_rpr(run)
        run._r.insert(0, copy.deepcopy(keep))
    pPr = par._p.get_or_add_pPr()
    old = pPr.find(qn("w:pageBreakBefore"))
    if old is not None:
        pPr.remove(old)
    if page_break:
        el = pPr.makeelement(qn("w:pageBreakBefore"), {})
        pPr.insert(0, el)
    return par


def new_paragraph_after(doc, el):
    """An empty paragraph immediately after `el`, as a python-docx Paragraph."""
    from docx.text.paragraph import Paragraph
    p = el.makeelement(qn("w:p"), {})
    el.addnext(p)
    return Paragraph(p, doc)


def find_para(doc, predicate):
    for p in doc.paragraphs:
        if predicate(" ".join((p.text or "").split())):
            return p
    return None


# ── what goes where ───────────────────────────────────────────────────────────

CONTENTS = [
    ("1", "Company Profile"),
    ("2", "About the Client"),
    ("3", "Technical Specifications"),
    ("4", "Scope of Supply"),
    ("5", "Annexures"),
]

ANNEXURES = [
    ("Annexure I", "Price Schedule", "Attached"),
    ("Annexure II", "Exclusions", "Attached"),
    ("Annexure III", "Terms & Conditions", "Attached"),
    ("Annexure IV", "Reference List of Clients", "Attached"),
    ("Annexure V", "Make List", "Attached"),
]

# The commercial terms are the ones the offer already carried. Only the shape
# changes — sentences become labelled rows. The values the quote sets come
# through the same placeholders the vertical offer uses, so a term edited on
# the form lands in both documents; the rest are BRF's own wording, which no
# form asks about.
TERMS = [
    ("Prices", "{{ tnc_prices }}"),
    ("Delivery",
     "Drawings for approval within 5 weeks of your order and advance. "
     "The erection schedule will be worked out mutually against your scope of "
     "erection work."),
    ("GST", "{{ tnc_gst }}"),
    ("Terms of Payment",
     "30% along with the purchase order. 70% against inspection of material at "
     "our works, before dispatch."),
    ("Packing & Forwarding", "4% of the order value, charged extra."),
    ("Freight Charges",
     "Transportation of all materials from our factory to your site is extra, "
     "at actuals."),
    ("Erection & Commissioning",
     "Erection and commissioning are in your scope; supervision is by ENCON at "
     "INR 12,500 per man per day, excluding to-and-fro fare, boarding and "
     "lodging."),
    ("Client's Scope",
     "Material supplied by us to be received into your store for safe custody "
     "and storage. Electricity, water and crane facilities to be provided free "
     "of charge for the erection work."),
    ("Validity", "{{ tnc_validity }}"),
    ("Guarantee",
     "18 months from the date of supply, against any manufacturing defect."),
    ("Taxes & Duties",
     "All taxes and duties as applicable, and to your account."),
    ("Force Majeure",
     "We shall not be responsible for delay in completion arising from the "
     "non-availability of items in your scope, or from causes beyond our "
     "reasonable control."),
]

SUPERVISION = [
    ("Supervision Charges for Erection and Commissioning",
     "INR 12,500 per man per day"),
    ("Note", "To-and-fro fare from Delhi to site, plus boarding and lodging at "
             "site, are to your account."),
]

# The basic-parameter block: every line of it reads "Label : value", and the
# whole run becomes the two-column table the vertical offer uses.
PARAM_LINE = re.compile(r"^(.*?)\s*:\s*(.*)$")

# The source document wrote its temperatures with a superscript zero for the
# degree sign. Flattened to plain text that zero joins the number, and the
# furnace reads as running at 12500 C rather than 1250 C. Ten times the
# temperature is not a typo a reader forgives, so the degree sign goes back.
DEGREES = re.compile(r"^(\d+)0C$")


def tidy(value):
    m = DEGREES.match(value)
    if m:
        return f"{m.group(1)} °C"
    return value.replace("upto1%", "up to 1%")


def main():
    if not os.path.exists(TARGET):
        print(f"not found: {TARGET}")
        return 1
    if not os.path.exists(SOURCE):
        print(f"not found: {SOURCE}")
        return 1

    doc = Document(TARGET)
    if find_para(doc, lambda t: t.upper() == "TABLE OF CONTENTS"):
        print("nothing to do — the offer already has the vertical structure")
        return 0

    shutil.copyfile(TARGET, TARGET + ".bak")
    src = Document(SOURCE)
    body = doc.element.body

    # ── 1. the technical specification table ─────────────────────────────────
    anchor = find_para(doc, lambda t: t.startswith("Basic Parameters"))
    if anchor is None:
        print("the basic-parameters heading is missing; run build_brf_offer_template first")
        return 1

    params, node = [], anchor._p.getnext()
    while node is not None and node.tag == f"{W}p":
        text = para_text(node)
        if not text:
            node = node.getnext()
            continue
        m = PARAM_LINE.match(text)
        if not m or not m.group(1):
            break                                  # the run of parameters ended
        params.append((m.group(1).strip(), tidy(m.group(2).strip())))
        nxt = node.getnext()
        body.remove(node)
        node = nxt

    from docx.table import Table
    spec_tbl = copy.deepcopy(next(t for t in src.tables
                                  if t.rows[0].cells[0].text.strip() == "Parameter")._tbl)
    anchor._p.addnext(spec_tbl)
    refill(Table(spec_tbl, doc), params)

    make_heading(anchor, "3. TECHNICAL SPECIFICATIONS", 1, page_break=True)
    sub = new_paragraph_after(doc, anchor._p)
    make_heading(sub, "BILLET REHEATING FURNACE — BASIC PARAMETERS", 2)

    print(f"  technical specifications: {len(params)} parameters in a table")

    # ── 2. contents, company profile, the client — all ahead of section 3 ────
    # Lifted whole out of the vertical offer so the wording, the borders and
    # the spacing are the same document's, not an imitation of it.
    lead = []

    toc = copy.deepcopy(next(t for t in src.tables
                             if t.rows[0].cells[0].text.strip() == "S. No."
                             and t.rows[0].cells[1].text.strip() == "Section")._tbl)
    refill(Table(toc, doc), CONTENTS)
    ann = copy.deepcopy(next(t for t in src.tables
                             if t.rows[0].cells[0].text.strip() == "Annexure No.")._tbl)
    refill(Table(ann, doc), ANNEXURES)

    toc_head = doc.add_paragraph()
    make_heading(toc_head, "TABLE OF CONTENTS", 1, page_break=True)
    ann_head = doc.add_paragraph()
    make_heading(ann_head, "LIST OF ANNEXURES", 3)
    lead += [toc_head._p, toc, ann_head._p, ann]

    profile = harvest(src, "1. COMPANY PROFILE", "2. ABOUT THE CLIENT")
    client = harvest(src, "2. ABOUT THE CLIENT", "3. TECHNICAL SPECIFICATIONS")
    pics = port_images(profile + client, src, doc)
    lead += profile + client

    for el in lead:
        anchor._p.addprevious(el)
    print(f"  contents, company profile and client details inserted"
          + (f" ({pics} picture{'s' if pics != 1 else ''} carried across)" if pics else ""))

    # ── 3. the scope, given heading levels that mean something ──────────────
    # Everything in the scope was Heading 1, so Word's outline listed each
    # sub-assembly as a section of its own. The vertical offer keeps one
    # Heading 1 per section and everything below it at Heading 2.
    scope = find_para(doc, lambda t: t.upper() == "SCOPE OF SUPPLY")
    if scope is not None:
        make_heading(scope, "4. SCOPE OF SUPPLY", 1, page_break=True)
    demoted = 0
    seen_scope = False
    for par in doc.paragraphs:
        text = " ".join((par.text or "").split())
        if text == "4. SCOPE OF SUPPLY":
            seen_scope = True
            continue
        if not seen_scope or par.style is None or par.style.name != "Heading 1":
            continue
        if text.upper().startswith(("ANNEXURE", "PRICE SCHEDULE")):
            break
        make_heading(par, text, 2)
        demoted += 1
    print(f"  scope of supply: {demoted} sub-headings demoted to Heading 2")

    # ── 4. the annexures ────────────────────────────────────────────────────
    price = find_para(doc, lambda t: t.upper() == "PRICE SCHEDULE")
    if price is not None:
        make_heading(price, "ANNEXURE I — PRICE SCHEDULE", 1, page_break=True)

    # Supervision: sentences become the vertical offer's two-column table, and
    # the sentences themselves go, rather than being said twice.
    sup = find_para(doc, lambda t: t.upper().startswith("SUPERVISION CHARGES"))
    if sup is not None:
        make_heading(sup, "SUPERVISION CHARGES FOR ERECTION & COMMISSIONING", 2)
        node = sup._p.getnext()
        while node is not None and node.tag == f"{W}p":
            text = para_text(node)
            if text.upper().startswith("EXCLUSION") or not text:
                break
            nxt = node.getnext()
            body.remove(node)
            node = nxt
        sup_tbl = copy.deepcopy(next(t for t in src.tables
                                     if t.rows[0].cells[0].text.strip()
                                     .startswith("Supervision Charges"))._tbl)
        sup._p.addnext(sup_tbl)
        refill(Table(sup_tbl, doc), SUPERVISION, header=False)

    excl = find_para(doc, lambda t: t.upper().startswith("EXCLUSION"))
    if excl is not None:
        make_heading(excl, "ANNEXURE II — EXCLUSIONS", 1, page_break=True)

    # Terms: the same terms, as the vertical offer's labelled table.
    tnc = find_para(doc, lambda t: t.upper().startswith("TERMS AND CONDITIONS"))
    if tnc is not None:
        make_heading(tnc, "ANNEXURE III — TERMS & CONDITIONS", 1, page_break=True)
        node = tnc._p.getnext()
        while node is not None and node.tag == f"{W}p":
            nxt = node.getnext()
            body.remove(node)
            node = nxt
        tnc_tbl = copy.deepcopy(next(t for t in src.tables
                                     if t.rows[0].cells[0].text.strip() == "Prices")._tbl)
        tnc._p.addnext(tnc_tbl)
        refill(Table(tnc_tbl, doc), TERMS, header=False)
        print(f"  terms and conditions: {len(TERMS)} rows in a table")

    # The client list and the make list have no BRF counterpart at all; they
    # come across whole. They are the last thing in the vertical offer, so the
    # harvest runs to the end and the trailing section break is dropped.
    tail = harvest(src, "ANNEXURE V — REFERENCE LIST OF CLIENTS")
    port_images(tail, src, doc)
    sect = body.find(qn("w:sectPr"))
    for el in tail:
        if el.tag == f"{W}p" and el.find(f"{W}pPr/{W}sectPr") is not None:
            continue
        (sect.addprevious(el) if sect is not None else body.append(el))
    for el in tail:
        text = para_text(el)
        if text.startswith("ANNEXURE V —"):
            from docx.text.paragraph import Paragraph
            make_heading(Paragraph(el, doc), "ANNEXURE IV — REFERENCE LIST OF CLIENTS",
                         1, page_break=True)
        elif text.startswith("ANNEXURE VI —"):
            from docx.text.paragraph import Paragraph
            make_heading(Paragraph(el, doc), "ANNEXURE V — MAKE LIST", 1, page_break=True)
    print("  reference list of clients and make list appended")

    doc.save(TARGET)
    print(f"{os.path.basename(TARGET)} rebuilt to the vertical offer's structure")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
